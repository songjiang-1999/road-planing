















# import pandas as pd
# import numpy as np
# from scipy.interpolate import griddata
# from sklearn.metrics import mean_squared_error, mean_absolute_error, accuracy_score
# from sklearn.model_selection import train_test_split
# import heapq
# import math
# import plotly.graph_objs as go
# import plotly.io as pio
# from docx import Document
# import tkinter as tk
# from tkinter import filedialog, messagebox
# import threading
#
# # 数据处理类
# class DataProcessor:
#     def __init__(self, dat_path, x_resolution=1, y_resolution=1):
#         self.dat_path = dat_path
#         self.x_resolution = x_resolution
#         self.y_resolution = y_resolution
#         self.df = None
#         self.X = None
#         self.Y = None
#         self.Z = None
#         self.TrafFlag = None
#         self.maxAcc = None
#         self.Slope = None
#         self.StepElev = None
#         self.Roughness = None
#         self.x_range = None
#         self.y_range = None
#         self.X_grid = None
#         self.Y_grid = None
#         self.Z_grid = None
#         self.TrafFlag_grid = None
#         self.Acc_grid = None
#         self.Slope_grid = None
#         self.StepElev_grid = None
#         self.Rough_grid = None
#
#     def load_data(self):
#         try:
#             # 读取 .dat 文件，指定分隔符为逗号
#             self.df = pd.read_csv(self.dat_path, delimiter=',')
#             print("数据文件中的列名：", self.df.columns.tolist())
#         except FileNotFoundError:
#             print(f"文件未找到: {self.dat_path}")
#             return
#         except pd.errors.EmptyDataError:
#             print("CSV 文件为空")
#             return
#         except Exception as e:
#             print(f"读取 CSV 文件时出错: {e}")
#             return
#
#         required_columns = {'X', 'Y', 'Z', 'TrafFlag', 'maxAcc', 'Slope', 'StepElevationDifference', 'Roughness'}
#         if not required_columns.issubset(self.df.columns):
#             print(f"CSV 文件缺少必要的列: {required_columns - set(self.df.columns)}")
#             return
#
#         # 提取所需数据
#         self.X = self.df['X'].values
#         self.Y = self.df['Y'].values
#         self.Z = self.df['Z'].values
#         self.TrafFlag = self.df['TrafFlag'].values
#         self.maxAcc = self.df['maxAcc'].values
#         self.Slope = self.df['Slope'].values
#         self.StepElev = self.df['StepElevationDifference'].values
#         self.Roughness = self.df['Roughness'].values
#
#         # 计算并打印 Slope 和 StepElevationDifference 的最小值和最大值
#         slope_min = np.min(self.Slope)
#         slope_max_val = np.max(self.Slope)
#         step_min = np.min(self.StepElev)
#         step_max = np.max(self.StepElev)
#
#         print(f"坡度 (Slope) 最小值: {slope_min}")
#         print(f"坡度 (Slope) 最大值: {slope_max_val}")
#         print(f"阶跃高程差 (StepElevationDifference) 最小值: {step_min}")
#         print(f"阶跃高程差 (StepElevationDifference) 最大值: {step_max}")
#
#     def interpolate_data(self):
#         # 计算网格范围和生成网格
#         x_min, x_max = self.X.min(), self.X.max()
#         y_min, y_max = self.Y.min(), self.Y.max()
#         self.x_range = np.arange(x_min, x_max + self.x_resolution, self.x_resolution)
#         self.y_range = np.arange(y_min, y_max + self.y_resolution, self.y_resolution)
#         self.X_grid, self.Y_grid = np.meshgrid(self.x_range, self.y_range)
#
#         # 插值
#         self.Z_grid = griddata((self.X, self.Y), self.Z, (self.X_grid, self.Y_grid), method='cubic')
#         self.TrafFlag_grid = griddata((self.X, self.Y), self.TrafFlag, (self.X_grid, self.Y_grid), method='nearest')
#         self.Acc_grid = griddata((self.X, self.Y), self.maxAcc, (self.X_grid, self.Y_grid), method='nearest')
#         self.Slope_grid = griddata((self.X, self.Y), self.Slope, (self.X_grid, self.Y_grid), method='nearest')
#         self.StepElev_grid = griddata((self.X, self.Y), self.StepElev, (self.X_grid, self.Y_grid), method='nearest')
#         self.Rough_grid = griddata((self.X, self.Y), self.Roughness, (self.X_grid, self.Y_grid), method='nearest')
#
#     def fill_nan_values(self):
#         # 填充 NaN 值
#         self.Z_grid = np.nan_to_num(self.Z_grid, nan=np.nanmean(self.Z))
#         self.Acc_grid = np.nan_to_num(self.Acc_grid, nan=0)
#         self.TrafFlag_grid = np.nan_to_num(self.TrafFlag_grid, nan=0)
#         self.Slope_grid = np.nan_to_num(self.Slope_grid, nan=0)
#         self.StepElev_grid = np.nan_to_num(self.StepElev_grid, nan=0)
#         self.Rough_grid = np.nan_to_num(self.Rough_grid, nan=0)
#
#         # 验证网格数据的形状
#         print(f"Z_grid shape: {self.Z_grid.shape}")
#         print(f"TrafFlag_grid shape: {self.TrafFlag_grid.shape}")
#         print(f"Acc_grid shape: {self.Acc_grid.shape}")
#         print(f"Slope_grid shape: {self.Slope_grid.shape}")
#         print(f"StepElev_grid shape: {self.StepElev_grid.shape}")
#         print(f"Rough_grid shape: {self.Rough_grid.shape}")
#
#         # 计算并打印插值后网格数据的最小值和最大值
#         slope_grid_min = np.min(self.Slope_grid)
#         slope_grid_max = np.max(self.Slope_grid)
#         step_grid_min = np.min(self.StepElev_grid)
#         step_grid_max = np.max(self.StepElev_grid)
#
#         print(f"插值后坡度网格 (Slope_grid) 最小值: {slope_grid_min}")
#         print(f"插值后坡度网格 (Slope_grid) 最大值: {slope_grid_max}")
#         print(f"插值后阶跃高程差网格 (StepElev_grid) 最小值: {step_grid_min}")
#         print(f"插值后阶跃高程差网格 (StepElev_grid) 最大值: {step_grid_max}")
#
#     def coord_to_index(self, coord):
#         x_min, y_min = self.x_range.min(), self.y_range.min()
#         col = int((coord[0] - x_min) / self.x_resolution)
#         row = int((coord[1] - y_min) / self.y_resolution)
#         return (row, col)
#
#     def is_valid_position(self, pos):
#         return 0 <= pos[0] < self.TrafFlag_grid.shape[0] and 0 <= pos[1] < self.TrafFlag_grid.shape[1]
#
#     def split_data_and_evaluate(self):
#         # 将所有数据合并
#         data = np.column_stack(
#             (self.X, self.Y, self.Z, self.TrafFlag, self.maxAcc, self.Slope, self.StepElev, self.Roughness))
#
#         # 分割数据集为训练集和测试集
#         test_size = 0.2
#         random_state = 42
#
#         train_data, test_data = train_test_split(data, test_size=test_size, random_state=random_state)
#
#         # 拆分训练集和测试集的数据
#         X_train, Y_train = train_data[:, 0], train_data[:, 1]
#         Z_train = train_data[:, 2]
#         TrafFlag_train = train_data[:, 3]
#         maxAcc_train = train_data[:, 4]
#         Slope_train = train_data[:, 5]
#         StepElev_train = train_data[:, 6]
#         Rough_train = train_data[:, 7]
#
#         X_test, Y_test = test_data[:, 0], test_data[:, 1]
#         Z_test = test_data[:, 2]
#         TrafFlag_test = test_data[:, 3]
#         maxAcc_test = test_data[:, 4]
#         Slope_test = test_data[:, 5]
#         StepElev_test = test_data[:, 6]
#         Rough_test = test_data[:, 7]
#
#         # 使用训练集进行插值
#         Z_grid_train = griddata((X_train, Y_train), Z_train, (self.X_grid, self.Y_grid), method='cubic')
#         TrafFlag_grid_train = griddata((X_train, Y_train), TrafFlag_train, (self.X_grid, self.Y_grid), method='nearest')
#         Acc_grid_train = griddata((X_train, Y_train), maxAcc_train, (self.X_grid, self.Y_grid), method='nearest')
#         Slope_grid_train = griddata((X_train, Y_train), Slope_train, (self.X_grid, self.Y_grid), method='nearest')
#         StepElev_grid_train = griddata((X_train, Y_train), StepElev_train, (self.X_grid, self.Y_grid), method='nearest')
#         Rough_grid_train = griddata((X_train, Y_train), Rough_train, (self.X_grid, self.Y_grid), method='nearest')
#
#         # 在测试集上预测值
#         Z_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Z_grid_train.flatten(), (X_test, Y_test),
#                           method='cubic')
#         TrafFlag_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), TrafFlag_grid_train.flatten(),
#                                  (X_test, Y_test), method='nearest')
#         Acc_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Acc_grid_train.flatten(), (X_test, Y_test),
#                             method='nearest')
#         Slope_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Slope_grid_train.flatten(),
#                               (X_test, Y_test), method='nearest')
#         StepElev_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), StepElev_grid_train.flatten(),
#                                  (X_test, Y_test), method='nearest')
#         Rough_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Rough_grid_train.flatten(),
#                               (X_test, Y_test), method='nearest')
#
#         # 处理 NaN 值并计算误差指标
#
#         # 高程数据误差
#         valid_indices_Z = ~np.isnan(Z_pred) & ~np.isnan(Z_test)
#         if np.any(valid_indices_Z):
#             mse_Z = mean_squared_error(Z_test[valid_indices_Z], Z_pred[valid_indices_Z])
#             mae_Z = mean_absolute_error(Z_test[valid_indices_Z], Z_pred[valid_indices_Z])
#             max_error_Z = np.max(np.abs(Z_test[valid_indices_Z] - Z_pred[valid_indices_Z]))
#             print(f"高程插值 - MSE: {mse_Z}, MAE: {mae_Z}, Max Error: {max_error_Z}")
#         else:
#             print("高程插值预测值中全部为 NaN，无法计算误差。")
#
#         # 通行标志（分类准确率）
#         valid_indices_TrafFlag = ~np.isnan(TrafFlag_pred) & ~np.isnan(TrafFlag_test)
#         if np.any(valid_indices_TrafFlag):
#             accuracy_TrafFlag = accuracy_score(TrafFlag_test[valid_indices_TrafFlag],
#                                                TrafFlag_pred[valid_indices_TrafFlag])
#             print(f"通行标志插值 - 准确率: {accuracy_TrafFlag}")
#         else:
#             print("通行标志插值预测值中全部为 NaN，无法计算准确率。")
#
#         # 最大加速度误差
#         valid_indices_Acc = ~np.isnan(Acc_pred) & ~np.isnan(maxAcc_test)
#         if np.any(valid_indices_Acc):
#             mse_Acc = mean_squared_error(maxAcc_test[valid_indices_Acc], Acc_pred[valid_indices_Acc])
#             mae_Acc = mean_absolute_error(maxAcc_test[valid_indices_Acc], Acc_pred[valid_indices_Acc])
#             max_error_Acc = np.max(np.abs(maxAcc_test[valid_indices_Acc] - Acc_pred[valid_indices_Acc]))
#             print(f"最大加速度插值 - MSE: {mse_Acc}, MAE: {mae_Acc}, Max Error: {max_error_Acc}")
#         else:
#             print("最大加速度插值预测值中全部为 NaN，无法计算误差。")
#
#         # 坡度误差
#         valid_indices_Slope = ~np.isnan(Slope_pred) & ~np.isnan(Slope_test)
#         if np.any(valid_indices_Slope):
#             mse_Slope = mean_squared_error(Slope_test[valid_indices_Slope], Slope_pred[valid_indices_Slope])
#             mae_Slope = mean_absolute_error(Slope_test[valid_indices_Slope], Slope_pred[valid_indices_Slope])
#             max_error_Slope = np.max(np.abs(Slope_test[valid_indices_Slope] - Slope_pred[valid_indices_Slope]))
#             print(f"坡度插值 - MSE: {mse_Slope}, MAE: {mae_Slope}, Max Error: {max_error_Slope}")
#         else:
#             print("坡度插值预测值中全部为 NaN，无法计算误差。")
#
#         # 阶跃高程差误差
#         valid_indices_StepElev = ~np.isnan(StepElev_pred) & ~np.isnan(StepElev_test)
#         if np.any(valid_indices_StepElev):
#             mse_StepElev = mean_squared_error(StepElev_test[valid_indices_StepElev],
#                                               StepElev_pred[valid_indices_StepElev])
#             mae_StepElev = mean_absolute_error(StepElev_test[valid_indices_StepElev],
#                                                StepElev_pred[valid_indices_StepElev])
#             max_error_StepElev = np.max(
#                 np.abs(StepElev_test[valid_indices_StepElev] - StepElev_pred[valid_indices_StepElev]))
#             print(f"阶跃高程差插值 - MSE: {mse_StepElev}, MAE: {mae_StepElev}, Max Error: {max_error_StepElev}")
#         else:
#             print("阶跃高程差插值预测值中全部为 NaN，无法计算误差。")
#
#         # 粗糙度误差
#         valid_indices_Rough = ~np.isnan(Rough_pred) & ~np.isnan(Rough_test)
#         if np.any(valid_indices_Rough):
#             mse_Rough = mean_squared_error(Rough_test[valid_indices_Rough], Rough_pred[valid_indices_Rough])
#             mae_Rough = mean_absolute_error(Rough_test[valid_indices_Rough], Rough_pred[valid_indices_Rough])
#             max_error_Rough = np.max(np.abs(Rough_test[valid_indices_Rough] - Rough_pred[valid_indices_Rough]))
#             print(f"粗糙度插值 - MSE: {mse_Rough}, MAE: {mae_Rough}, Max Error: {max_error_Rough}")
#         else:
#             print("粗糙度插值预测值中全部为 NaN，无法计算误差。")
#
#
# # 节点类和路径规划类
# class Node:
#     def __init__(self, position, parent=None):
#         self.position = position  # (row, col)
#         self.parent = parent
#         self.g = 0  # 累计代价
#         self.h = 0  # 启发代价
#         self.f = 0  # 总代价
#
#     def __lt__(self, other):
#         return self.f < other.f
# class PathPlanner:
#     def __init__(self, grid_data, max_slope=None, max_step_elev=None):
#         self.grid = grid_data['TrafFlag_grid']
#         self.Z_grid = grid_data['Z_grid']
#         self.Acc_grid = grid_data['Acc_grid']
#         self.Slope_grid = grid_data['Slope_grid']
#         self.StepElev_grid = grid_data['StepElev_grid']
#         self.Rough_grid = grid_data['Rough_grid']
#         self.max_slope = max_slope
#         self.max_step_elev = max_step_elev
#
#     def heuristic_distance_3d(self, current_pos, end_pos):
#         d_xy = np.sqrt((current_pos[0] - end_pos[0]) ** 2 +
#                        (current_pos[1] - end_pos[1]) ** 2)
#         z_current = self.Z_grid[current_pos[0], current_pos[1]]
#         z_end = self.Z_grid[end_pos[0], end_pos[1]]
#         d_z = (z_end - z_current) ** 2
#         return np.sqrt(d_xy ** 2 + d_z)
#
#     def astar(self, start, end, mode='distance', alpha_env=0.6, beta_env=0.4):
#         start_node = Node(start)
#         end_node = Node(end)
#
#         open_list = []
#         heapq.heappush(open_list, start_node)
#         open_dict = {start_node.position: start_node}
#
#         closed_set = set()
#
#         while open_list:
#             current_node = heapq.heappop(open_list)
#             del open_dict[current_node.position]
#             closed_set.add(current_node.position)
#
#             # 检查是否到达终点
#             if current_node.position == end_node.position:
#                 path = []
#                 while current_node:
#                     path.append(current_node.position)
#                     current_node = current_node.parent
#                 return path[::-1]
#
#             # 定义8个邻居（包括对角线）
#             neighbors = [(0, 1), (0, -1), (1, 0), (-1, 0),
#                          (1, 1), (1, -1), (-1, 1), (-1, -1)]
#
#             for offset in neighbors:
#                 neighbor_pos = (current_node.position[0] + offset[0],
#                                 current_node.position[1] + offset[1])
#
#                 # 检查邻居是否在网格范围内
#                 if not (0 <= neighbor_pos[0] < self.grid.shape[0] and
#                         0 <= neighbor_pos[1] < self.grid.shape[1]):
#                     continue
#
#                 # 检查 TrafFlag 是否可通行
#                 if self.grid[neighbor_pos] == 0:
#                     continue  # 不可通行区域
#
#                 # 在所有模式下应用坡度和阶跃高程差的限制
#                 if self.max_slope is not None and self.Slope_grid[neighbor_pos] > self.max_slope:
#                     continue  # 坡度超过最大限制
#                 if self.max_step_elev is not None and self.StepElev_grid[neighbor_pos] > self.max_step_elev:
#                     continue  # 阶跃高程差超过最大限制
#
#                 # 检查是否已经在关闭列表中
#                 if neighbor_pos in closed_set:
#                     continue
#
#                 # 处理不同模式
#                 if mode == 'distance':
#                     movement_cost = self.heuristic_distance_3d(current_node.position, neighbor_pos)
#                     g_cost = current_node.g + movement_cost
#                     h_cost = self.heuristic_distance_3d(neighbor_pos, end)
#                 elif mode == 'time':
#                     a_max = self.Acc_grid[neighbor_pos[0], neighbor_pos[1]]
#                     if a_max <= 0:
#                         continue  # 无法加速，跳过
#                     movement_cost = self.heuristic_distance_3d(current_node.position, neighbor_pos)
#                     g_cost = current_node.g + movement_cost / a_max  # 简化为距离除以加速度
#                     h_cost = self.heuristic_distance_3d(neighbor_pos, end)
#                 elif mode == 'env':
#                     movement_uncertainty = (self.Slope_grid[neighbor_pos[0], neighbor_pos[1]] / 45.0 +
#                                             self.StepElev_grid[neighbor_pos[0], neighbor_pos[1]] / 1.5 +
#                                             self.Rough_grid[neighbor_pos[0], neighbor_pos[1]] / 6.0)
#                     g_cost = current_node.g + movement_uncertainty
#                     h_cost = self.heuristic_distance_3d(neighbor_pos, end)
#                 else:
#                     raise ValueError("Invalid mode. Choose from 'distance', 'time', 'env'.")
#
#                 # 创建邻居节点
#                 neighbor_node = Node(neighbor_pos, current_node)
#                 neighbor_node.g = g_cost
#                 neighbor_node.h = h_cost
#                 neighbor_node.f = neighbor_node.g + neighbor_node.h
#
#                 # 检查开放列表中是否有更好的节点
#                 if neighbor_pos in open_dict:
#                     existing_node = open_dict[neighbor_pos]
#                     if existing_node.g <= neighbor_node.g:
#                         continue  # 已有更优节点
#                     else:
#                         # 替换为更优节点
#                         open_list.remove(existing_node)
#                         heapq.heapify(open_list)
#                         heapq.heappush(open_list, neighbor_node)
#                         open_dict[neighbor_pos] = neighbor_node
#                 else:
#                     heapq.heappush(open_list, neighbor_node)
#                     open_dict[neighbor_pos] = neighbor_node
#
#         return None  # 如果找不到路径，返回 None
#
#     def extract_path_coordinates(self, path, x_range, y_range):
#         return [(x_range[p[1]], y_range[p[0]], self.Z_grid[p[0], p[1]]) for p in path]
#
#     def calculate_path_length(self, path_coords):
#         length = 0
#         for i in range(1, len(path_coords)):
#             x1, y1, z1 = path_coords[i - 1]
#             x2, y2, z2 = path_coords[i]
#             length += np.sqrt((x2 - x1) ** 2 + (y2 - y1) ** 2 + (z2 - z1) ** 2)
#         return length
#
#     def compute_traversal_time(self, path_coords, y_range, x_range):
#         total_time = 0.0
#         current_speed = 0.0  # 初始速度为0
#         max_speed = 70.0  # 最大速度70 m/s
#
#         speeds = [current_speed]  # 存储每个点的速度
#         accelerations = []  # 存储每段的加速度
#
#         for i in range(1, len(path_coords)):
#             x1, y1, z1 = path_coords[i - 1]
#             x2, y2, z2 = path_coords[i]
#             # 计算两点之间的距离
#             d = np.sqrt((x2 - x1) ** 2 + (y2 - y1) ** 2 + (z2 - z1) ** 2)
#
#             # 获取当前点的网格索引
#             row = np.argmin(np.abs(y_range - y1))
#             col = np.argmin(np.abs(x_range - x1))
#
#             # 获取当前点的加速度
#             a = self.Acc_grid[row, col]
#             if a <= 0:
#                 # 如果加速度不可用或为0，无法继续加速
#                 a = 0.0
#
#             # 计算需要加速到最大速度的速度增量
#             delta_v = max_speed - current_speed
#
#             if current_speed < max_speed and a > 0:
#                 # 需要加速
#                 t_acc = delta_v / a  # 加速所需时间
#                 d_acc = current_speed * t_acc + 0.5 * a * t_acc ** 2  # 加速所需距离
#
#                 if d_acc >= d:
#                     # 在当前段无法加速到最大速度
#                     # 使用运动学方程计算通过距离d所需的时间
#                     # 解方程: 0.5 * a * t^2 + current_speed * t - d = 0
#                     discriminant = current_speed ** 2 + 2 * a * d
#                     if discriminant < 0:
#                         # 无实数解，无法到达下一个点
#                         return float('inf'), [], []
#                     t = (-current_speed + math.sqrt(discriminant)) / a
#                     total_time += t
#                     current_speed += a * t
#                     accelerations.append(a)
#                     speeds.append(current_speed)
#                 else:
#                     # 在当前段可以加速到最大速度
#                     # 先加速到最大速度
#                     total_time += t_acc
#                     current_speed = max_speed
#                     accelerations.append(a)
#                     speeds.append(current_speed)
#                     # 剩余距离
#                     d_remain = d - d_acc
#                     # 以最大速度匀速行驶剩余距离
#                     t_remain = d_remain / current_speed
#                     total_time += t_remain
#                     accelerations.append(0.0)
#                     speeds.append(current_speed)
#             else:
#                 # 无法加速或已达到最大速度，匀速行驶
#                 if current_speed == 0:
#                     # 无法移动
#                     return float('inf'), [], []
#                 t = d / current_speed
#                 total_time += t
#                 accelerations.append(0.0)
#                 speeds.append(current_speed)
#
#         return total_time, speeds, accelerations
#
#
# # 可视化类
# class Visualizer:
#     def __init__(self, X_grid, Y_grid, Z_grid, TrafFlag_grid):
#         self.X_grid = X_grid
#         self.Y_grid = Y_grid
#         self.Z_grid = Z_grid
#         self.TrafFlag_grid = TrafFlag_grid
#         self.fig = go.Figure()
#
#     def plot_terrain(self):
#         self.fig.add_trace(go.Surface(
#             x=self.X_grid, y=self.Y_grid, z=self.Z_grid,
#             surfacecolor=self.TrafFlag_grid,
#             colorscale='Viridis',
#             opacity=0.8,
#             showscale=True,
#             colorbar=dict(title='TrafFlag')
#         ))
#
#     def add_path(self, coords, speeds, accs, color, name, length, time_taken, path_width=4):
#         if coords:
#             x, y, z = zip(*coords)
#             # 处理不可行路径（时间为inf）的情况
#             if math.isinf(time_taken):
#                 time_str = "无法到达"
#             else:
#                 time_str = f"{time_taken:.2f} 秒"
#
#             # 准备 customdata，将速度和加速度添加进去
#             if accs:
#                 accs = [0.0] + accs  # 与 speeds 对齐
#             else:
#                 accs = [0.0] * len(speeds)
#
#             customdata = np.array([speeds, accs]).T  # 转置使其每行对应一个点的 [speed, acceleration]
#
#             self.fig.add_trace(go.Scatter3d(
#                 x=x, y=y, z=z,
#                 mode='lines+markers',
#                 line=dict(color=color, width=path_width),
#                 marker=dict(size=2, color=color),
#                 name=f"{name} (长度: {length:.2f} 米, 时间: {time_str})",
#                 customdata=customdata,
#                 hovertemplate=(
#                         'X: %{x}<br>' +
#                         'Y: %{y}<br>' +
#                         'Z: %{z}<br>' +
#                         '速度: %{customdata[0]:.2f} m/s<br>' +
#                         '加速度: %{customdata[1]:.2f} m/s²<br>' +
#                         '<extra></extra>'
#                 )
#             ))
#         else:
#             print(f"未找到{name}")
#
#     def add_start_and_end_points(self, start_coord, end_coord):
#         self.fig.add_trace(go.Scatter3d(
#             x=[start_coord[0]],
#             y=[start_coord[1]],
#             z=[start_coord[2]],
#             mode='markers',
#             marker=dict(size=8, color='green'),
#             name='起点'
#         ))
#         self.fig.add_trace(go.Scatter3d(
#             x=[end_coord[0]],
#             y=[end_coord[1]],
#             z=[end_coord[2]],
#             mode='markers',
#             marker=dict(size=8, color='red'),
#             name='终点'
#         ))
#
#     def show(self):
#         self.fig.update_layout(
#             scene=dict(
#                 xaxis_title='X 坐标',
#                 yaxis_title='Y 坐标',
#                 zaxis_title='Z 坐标 (高程)',
#                 aspectmode='data'
#             ),
#             title='3D 地形路径规划（基于距离、时间与最小不确定性）',
#             legend=dict(
#                 x=0,
#                 y=1,
#                 bgcolor='rgba(255, 255, 255, 0)',
#                 bordercolor='rgba(255, 255, 255, 0)'
#             )
#         )
#         pio.show(self.fig)
#
#
# # 报告生成函数
# def write_paths_to_word(document, path_name, path_coords):
#     if not path_coords:
#         return
#     # 添加路径名称和点数量
#     document.add_heading(f"{path_name} (点数: {len(path_coords)})", level=2)
#     # 创建表格，包含三列：X相对坐标、Y相对坐标、Z高程坐标
#     table = document.add_table(rows=1, cols=3)
#     table.style = 'Light Grid'
#     hdr_cells = table.rows[0].cells
#     hdr_cells[0].text = 'X相对坐标'
#     hdr_cells[1].text = 'Y相对坐标'
#     hdr_cells[2].text = 'Z高程坐标'
#     # 填充表格内容
#     for coord in path_coords:
#         row_cells = table.add_row().cells
#         row_cells[0].text = f"{coord[0]}"
#         row_cells[1].text = f"{coord[1]}"
#         row_cells[2].text = f"{coord[2]}"
#     # 添加一个空行
#     document.add_paragraph()
# # 处理函数
# def run_processing(csv_path, start_coord, end_coord, output_path):
#     # 设置路径宽度
#     path_width = 4  # 可以根据需要调整路径宽度
#
#     # 数据处理
#     data_processor = DataProcessor(csv_path)
#     data_processor.load_data()
#     data_processor.interpolate_data()
#     data_processor.fill_nan_values()
#     data_processor.split_data_and_evaluate()
#
#     # 将 start_coord 和 end_coord 转换为网格索引
#     start = data_processor.coord_to_index(start_coord)
#     end = data_processor.coord_to_index(end_coord)
#
#     # 检查起点和终点是否在网格范围内
#     if not data_processor.is_valid_position(start):
#         raise ValueError("起点坐标超出网格范围")
#     if not data_processor.is_valid_position(end):
#         raise ValueError("终点坐标超出网格范围")
#
#     # 确保起点和终点可通行
#     if data_processor.TrafFlag_grid[start] == 0:
#         raise ValueError("起点不可通行")
#     if data_processor.TrafFlag_grid[end] == 0:
#         raise ValueError("终点不可通行")
#
#     # 定义坡度和阶跃高程差的最大允许值（应用于所有模式）
#     max_allowed_slope = 25.0  # 最大坡度（度）
#     max_allowed_step_elev = 10  # 最大阶跃高程差（米）
#
#     # 路径规划
#     grid_data = {
#         'TrafFlag_grid': data_processor.TrafFlag_grid,
#         'Z_grid': data_processor.Z_grid,
#         'Acc_grid': data_processor.Acc_grid,
#         'Slope_grid': data_processor.Slope_grid,
#         'StepElev_grid': data_processor.StepElev_grid,
#         'Rough_grid': data_processor.Rough_grid
#     }
#     path_planner = PathPlanner(grid_data, max_slope=max_allowed_slope, max_step_elev=max_allowed_step_elev)
#
#     print("正在计算基于距离的最短路径...")
#     path_distance = path_planner.astar(start, end, mode='distance')
#     if path_distance:
#         print("基于距离的路径计算完成")
#     else:
#         print("未找到基于距离的路径")
#
#     print("正在计算基于时间的最短耗时路径...")
#     path_time = path_planner.astar(start, end, mode='time')
#     if path_time:
#         print("基于时间的路径计算完成")
#     else:
#         print("未找到基于时间的路径")
#
#     print("正在计算最小不确定性路径...")
#     path_env = path_planner.astar(start, end, mode='env')
#     if path_env:
#         print("最小不确定性路径计算完成")
#     else:
#         print("未找到最小不确定性路径")
#
#     # 提取路径坐标
#     if path_distance:
#         path_distance_coords = path_planner.extract_path_coordinates(path_distance, data_processor.x_range,
#                                                                      data_processor.y_range)
#         length_distance = path_planner.calculate_path_length(path_distance_coords)
#     else:
#         path_distance_coords = None
#         length_distance = None
#     if path_time:
#         path_time_coords = path_planner.extract_path_coordinates(path_time, data_processor.x_range,
#                                                                  data_processor.y_range)
#         length_time = path_planner.calculate_path_length(path_time_coords)
#     else:
#         path_time_coords = None
#         length_time = None
#     if path_env:
#         path_env_coords = path_planner.extract_path_coordinates(path_env, data_processor.x_range,
#                                                                 data_processor.y_range)
#         length_env = path_planner.calculate_path_length(path_env_coords)
#     else:
#         path_env_coords = None
#         length_env = None
#
#     # 计算遍历时间，并获取速度和加速度数据
#     if path_distance:
#         time_distance, speeds_distance, accs_distance = path_planner.compute_traversal_time(
#             path_distance_coords, data_processor.y_range, data_processor.x_range)
#     else:
#         speeds_distance, accs_distance, time_distance = None, None, None
#
#     if path_time:
#         time_time, speeds_time, accs_time = path_planner.compute_traversal_time(
#             path_time_coords, data_processor.y_range, data_processor.x_range)
#     else:
#         speeds_time, accs_time, time_time = None, None, None
#
#     if path_env:
#         time_env, speeds_env, accs_env = path_planner.compute_traversal_time(
#             path_env_coords, data_processor.y_range, data_processor.x_range)
#     else:
#         speeds_env, accs_env, time_env = None, None, None
#
#     # 可视化
#     visualizer = Visualizer(
#         data_processor.X_grid,
#         data_processor.Y_grid,
#         data_processor.Z_grid,
#         data_processor.TrafFlag_grid
#     )
#     visualizer.plot_terrain()
#
#     colors = ['blue', 'orange', 'purple']
#     names = ['MinDistPath 最短距离路径', 'MinTimePath 最短耗时路径', 'MinUncertPath 最小不确定性路径']
#     lengths = [length_distance, length_time, length_env]
#     times = [time_distance, time_time, time_env]
#     coords_list = [path_distance_coords, path_time_coords, path_env_coords]
#     speeds_list = [speeds_distance, speeds_time, speeds_env]
#     accs_list = [accs_distance, accs_time, accs_env]
#
#     for coords, speeds, accs, color, name, length, time_taken in zip(
#             coords_list, speeds_list, accs_list,
#             colors, names, lengths, times):
#         visualizer.add_path(coords, speeds, accs, color, name, length, time_taken, path_width)
#
#     # 标记起点和终点
#     if path_distance or path_time or path_env:
#         selected_path = None
#         if path_distance:
#             selected_path = path_distance_coords
#         elif path_time:
#             selected_path = path_time_coords
#         elif path_env:
#             selected_path = path_env_coords
#
#         if selected_path:
#             # 使用起点坐标
#             start_x, start_y, start_z = selected_path[0]
#             # 使用终点坐标
#             end_x, end_y, end_z = selected_path[-1]
#             visualizer.add_start_and_end_points((start_x, start_y, start_z), (end_x, end_y, end_z))
#
#     visualizer.show()
#
#     # 创建 Word 文档
#     document = Document()
#     document.add_heading('路径规划结果', 0)
#     # 写入路径信息
#     if path_distance or path_time or path_env:
#         if path_distance:
#             write_paths_to_word(document, 'MinDistPath ', path_distance_coords)
#         if path_time:
#             write_paths_to_word(document, 'MinTimePath ', path_time_coords)
#         if path_env:
#             write_paths_to_word(document, 'MinUncertPath ', path_env_coords)
#
#         # 保存文档
#         try:
#             document.save(output_path)
#             print(f"路径坐标已成功保存到 Word 文档: {output_path}")
#         except Exception as e:
#             print(f"保存 Word 文档时出错: {e}")
#     else:
#         print("没有路径可写入 Word 文档")
#
#     # 打印路径长度和遍历时间
#     if path_distance:
#         if math.isinf(time_distance):
#             print(f"最短路径长度: {length_distance:.2f} 米")
#             print(f"最短距离路径遍历时间: 无法到达")
#         else:
#             print(f"最短路径长度: {length_distance:.2f} 米")
#             print(f"最短距离路径遍历时间: {time_distance:.2f} 秒")
#     if path_time:
#         if math.isinf(time_time):
#             print(f"最短耗时路径长度: {length_time:.2f} 米")
#             print(f"最短耗时路径遍历时间: 无法到达")
#         else:
#             print(f"最短耗时路径长度: {length_time:.2f} 米")
#             print(f"最短耗时路径遍历时间: {time_time:.2f} 秒")
#     if path_env:
#         if math.isinf(time_env):
#             print(f"最小不确定性路径长度: {length_env:.2f} 米")
#             print(f"最小不确定性路径遍历时间: 无法到达")
#         else:
#             print(f"最小不确定性路径长度: {length_env:.2f} 米")
#             print(f"最小不确定性路径遍历时间: {time_env:.2f} 秒")
#
# # 主程序
# def main():
#     # 创建主窗口
#     root = tk.Tk()
#     root.title("路径规划软件")
#
#     # 地图数据文件路径
#     tk.Label(root, text="地图数据文件路径:").grid(row=0, column=0, sticky='e')
#     map_data_path_entry = tk.Entry(root, width=50)
#     map_data_path_entry.grid(row=0, column=1)
#     def browse_map_data():
#         filename = filedialog.askopenfilename(filetypes=[("CSV文件", "*.csv")])
#         if filename:
#             map_data_path_entry.delete(0, tk.END)
#             map_data_path_entry.insert(0, filename)
#     tk.Button(root, text="浏览...", command=browse_map_data).grid(row=0, column=2)
#
#     # 起点坐标
#     tk.Label(root, text="起点坐标 (格式: X,Y):").grid(row=1, column=0, sticky='e')
#     start_coord_entry = tk.Entry(root, width=50)
#     start_coord_entry.grid(row=1, column=1)
#
#     # 终点坐标
#     tk.Label(root, text="终点坐标 (格式: X,Y):").grid(row=2, column=0, sticky='e')
#     end_coord_entry = tk.Entry(root, width=50)
#     end_coord_entry.grid(row=2, column=1)
#
#     # 输出路径
#     tk.Label(root, text="输出 Word 文档路径:").grid(row=3, column=0, sticky='e')
#     output_path_entry = tk.Entry(root, width=50)
#     output_path_entry.grid(row=3, column=1)
#     def browse_output_path():
#         filename = filedialog.asksaveasfilename(defaultextension=".docx",
#                                                 filetypes=[("Word文档", "*.docx")])
#         if filename:
#             output_path_entry.delete(0, tk.END)
#             output_path_entry.insert(0, filename)
#     tk.Button(root, text="浏览...", command=browse_output_path).grid(row=3, column=2)
#
#     # 运行按钮
#     def run():
#         # 获取输入值
#         csv_path = map_data_path_entry.get()
#         start_coord_input = start_coord_entry.get()
#         end_coord_input = end_coord_entry.get()
#         output_path = output_path_entry.get()
#
#         # 验证输入
#         if not csv_path:
#             messagebox.showerror("错误", "请提供地图数据文件路径。")
#             return
#         if not output_path:
#             messagebox.showerror("错误", "请提供输出 Word 文档路径。")
#             return
#         if not start_coord_input:
#             messagebox.showerror("错误", "请提供起点坐标。")
#             return
#         if not end_coord_input:
#             messagebox.showerror("错误", "请提供终点坐标。")
#             return
#         try:
#             start_x, start_y = map(float, start_coord_input.strip().split(','))
#             start_coord = (start_x, start_y)
#         except ValueError:
#             messagebox.showerror("错误", "起点坐标格式错误，请使用 'X,Y' 的格式。")
#             return
#
#         try:
#             end_x, end_y = map(float, end_coord_input.strip().split(','))
#             end_coord = (end_x, end_y)
#         except ValueError:
#             messagebox.showerror("错误", "终点坐标格式错误，请使用 'X,Y' 的格式。")
#             return
#
#         # 在后台线程中运行，以避免阻塞GUI
#         def process():
#             try:
#                 # 调用处理函数
#                 run_processing(csv_path, start_coord, end_coord, output_path)
#                 messagebox.showinfo("完成", "路径规划已完成。")
#             except Exception as e:
#                 messagebox.showerror("错误", f"发生错误: {e}")
#
#         threading.Thread(target=process).start()
#
#     tk.Button(root, text="运行", command=run).grid(row=4, column=1)
#
#     root.mainloop()
#
# if __name__ == "__main__":
#     main()









# import pandas as pd
# import numpy as np
# from scipy.interpolate import griddata
# from sklearn.metrics import mean_squared_error, mean_absolute_error, accuracy_score
# from sklearn.model_selection import train_test_split
# import heapq
# import math
# import plotly.graph_objs as go
# import plotly.io as pio
# from docx import Document
# import tkinter as tk
# from tkinter import filedialog, messagebox
# import threading
# import sys
# import os
#
# # 添加 resource_path 函数，用于获取资源文件的正确路径
# def resource_path(relative_path):
#     """获取资源文件的绝对路径，兼容 PyInstaller 打包后的情况"""
#     try:
#         # PyInstaller 创建临时文件夹，并将路径存储在 _MEIPASS 中
#         base_path = sys._MEIPASS
#     except Exception:
#         base_path = os.path.abspath(".")
#     return os.path.join(base_path, relative_path)
#
# # 数据处理类
# class DataProcessor:
#     def __init__(self, dat_path, x_resolution=1, y_resolution=1):
#         self.dat_path = dat_path
#         self.x_resolution = x_resolution
#         self.y_resolution = y_resolution
#         self.df = None
#         self.X = None
#         self.Y = None
#         self.Z = None
#         self.TrafFlag = None
#         self.maxAcc = None
#         self.Slope = None
#         self.StepElev = None
#         self.Roughness = None
#         self.x_range = None
#         self.y_range = None
#         self.X_grid = None
#         self.Y_grid = None
#         self.Z_grid = None
#         self.TrafFlag_grid = None
#         self.Acc_grid = None
#         self.Slope_grid = None
#         self.StepElev_grid = None
#         self.Rough_grid = None
#
#     def load_data(self):
#         try:
#             # 使用 resource_path 函数获取数据文件路径
#             data_path = resource_path(self.dat_path)
#             # 读取 .dat 文件，指定分隔符为逗号
#             self.df = pd.read_csv(data_path, delimiter=',')
#             print("数据文件中的列名：", self.df.columns.tolist())
#         except FileNotFoundError:
#             print(f"文件未找到: {self.dat_path}")
#             return
#         except pd.errors.EmptyDataError:
#             print("CSV 文件为空")
#             return
#         except Exception as e:
#             print(f"读取 CSV 文件时出错: {e}")
#             return
#
#         required_columns = {'X', 'Y', 'Z', 'TrafFlag', 'maxAcc', 'Slope', 'StepElevationDifference', 'Roughness'}
#         if not required_columns.issubset(self.df.columns):
#             print(f"CSV 文件缺少必要的列: {required_columns - set(self.df.columns)}")
#             return
#
#         # 提取所需数据
#         self.X = self.df['X'].values
#         self.Y = self.df['Y'].values
#         self.Z = self.df['Z'].values
#         self.TrafFlag = self.df['TrafFlag'].values
#         self.maxAcc = self.df['maxAcc'].values
#         self.Slope = self.df['Slope'].values
#         self.StepElev = self.df['StepElevationDifference'].values
#         self.Roughness = self.df['Roughness'].values
#
#         # 计算并打印 Slope 和 StepElevationDifference 的最小值和最大值
#         slope_min = np.min(self.Slope)
#         slope_max_val = np.max(self.Slope)
#         step_min = np.min(self.StepElev)
#         step_max = np.max(self.StepElev)
#
#         print(f"坡度 (Slope) 最小值: {slope_min}")
#         print(f"坡度 (Slope) 最大值: {slope_max_val}")
#         print(f"阶跃高程差 (StepElevationDifference) 最小值: {step_min}")
#         print(f"阶跃高程差 (StepElevationDifference) 最大值: {step_max}")
#
#     def interpolate_data(self):
#         # 计算网格范围和生成网格
#         x_min, x_max = self.X.min(), self.X.max()
#         y_min, y_max = self.Y.min(), self.Y.max()
#         self.x_range = np.arange(x_min, x_max + self.x_resolution, self.x_resolution)
#         self.y_range = np.arange(y_min, y_max + self.y_resolution, self.y_resolution)
#         self.X_grid, self.Y_grid = np.meshgrid(self.x_range, self.y_range)
#
#         # 插值
#         self.Z_grid = griddata((self.X, self.Y), self.Z, (self.X_grid, self.Y_grid), method='cubic')
#         self.TrafFlag_grid = griddata((self.X, self.Y), self.TrafFlag, (self.X_grid, self.Y_grid), method='nearest')
#         self.Acc_grid = griddata((self.X, self.Y), self.maxAcc, (self.X_grid, self.Y_grid), method='nearest')
#         self.Slope_grid = griddata((self.X, self.Y), self.Slope, (self.X_grid, self.Y_grid), method='nearest')
#         self.StepElev_grid = griddata((self.X, self.Y), self.StepElev, (self.X_grid, self.Y_grid), method='nearest')
#         self.Rough_grid = griddata((self.X, self.Y), self.Roughness, (self.X_grid, self.Y_grid), method='nearest')
#
#     def fill_nan_values(self):
#         # 填充 NaN 值
#         self.Z_grid = np.nan_to_num(self.Z_grid, nan=np.nanmean(self.Z))
#         self.Acc_grid = np.nan_to_num(self.Acc_grid, nan=0)
#         self.TrafFlag_grid = np.nan_to_num(self.TrafFlag_grid, nan=0)
#         self.Slope_grid = np.nan_to_num(self.Slope_grid, nan=0)
#         self.StepElev_grid = np.nan_to_num(self.StepElev_grid, nan=0)
#         self.Rough_grid = np.nan_to_num(self.Rough_grid, nan=0)
#
#         # 验证网格数据的形状
#         print(f"Z_grid shape: {self.Z_grid.shape}")
#         print(f"TrafFlag_grid shape: {self.TrafFlag_grid.shape}")
#         print(f"Acc_grid shape: {self.Acc_grid.shape}")
#         print(f"Slope_grid shape: {self.Slope_grid.shape}")
#         print(f"StepElev_grid shape: {self.StepElev_grid.shape}")
#         print(f"Rough_grid shape: {self.Rough_grid.shape}")
#
#         # 计算并打印插值后网格数据的最小值和最大值
#         slope_grid_min = np.min(self.Slope_grid)
#         slope_grid_max = np.max(self.Slope_grid)
#         step_grid_min = np.min(self.StepElev_grid)
#         step_grid_max = np.max(self.StepElev_grid)
#
#         print(f"插值后坡度网格 (Slope_grid) 最小值: {slope_grid_min}")
#         print(f"插值后坡度网格 (Slope_grid) 最大值: {slope_grid_max}")
#         print(f"插值后阶跃高程差网格 (StepElev_grid) 最小值: {step_grid_min}")
#         print(f"插值后阶跃高程差网格 (StepElev_grid) 最大值: {step_grid_max}")
#
#     def coord_to_index(self, coord):
#         x_min, y_min = self.x_range.min(), self.y_range.min()
#         col = int((coord[0] - x_min) / self.x_resolution)
#         row = int((coord[1] - y_min) / self.y_resolution)
#         return (row, col)
#
#     def is_valid_position(self, pos):
#         return 0 <= pos[0] < self.TrafFlag_grid.shape[0] and 0 <= pos[1] < self.TrafFlag_grid.shape[1]
#
#     def split_data_and_evaluate(self):
#         # 将所有数据合并
#         data = np.column_stack(
#             (self.X, self.Y, self.Z, self.TrafFlag, self.maxAcc, self.Slope, self.StepElev, self.Roughness))
#
#         # 分割数据集为训练集和测试集
#         test_size = 0.2
#         random_state = 42
#
#         train_data, test_data = train_test_split(data, test_size=test_size, random_state=random_state)
#
#         # 拆分训练集和测试集的数据
#         X_train, Y_train = train_data[:, 0], train_data[:, 1]
#         Z_train = train_data[:, 2]
#         TrafFlag_train = train_data[:, 3]
#         maxAcc_train = train_data[:, 4]
#         Slope_train = train_data[:, 5]
#         StepElev_train = train_data[:, 6]
#         Rough_train = train_data[:, 7]
#
#         X_test, Y_test = test_data[:, 0], test_data[:, 1]
#         Z_test = test_data[:, 2]
#         TrafFlag_test = test_data[:, 3]
#         maxAcc_test = test_data[:, 4]
#         Slope_test = test_data[:, 5]
#         StepElev_test = test_data[:, 6]
#         Rough_test = test_data[:, 7]
#
#         # 使用训练集进行插值
#         Z_grid_train = griddata((X_train, Y_train), Z_train, (self.X_grid, self.Y_grid), method='cubic')
#         TrafFlag_grid_train = griddata((X_train, Y_train), TrafFlag_train, (self.X_grid, self.Y_grid), method='nearest')
#         Acc_grid_train = griddata((X_train, Y_train), maxAcc_train, (self.X_grid, self.Y_grid), method='nearest')
#         Slope_grid_train = griddata((X_train, Y_train), Slope_train, (self.X_grid, self.Y_grid), method='nearest')
#         StepElev_grid_train = griddata((X_train, Y_train), StepElev_train, (self.X_grid, self.Y_grid), method='nearest')
#         Rough_grid_train = griddata((X_train, Y_train), Rough_train, (self.X_grid, self.Y_grid), method='nearest')
#
#         # 在测试集上预测值
#         Z_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Z_grid_train.flatten(), (X_test, Y_test),
#                           method='cubic')
#         TrafFlag_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), TrafFlag_grid_train.flatten(),
#                                  (X_test, Y_test), method='nearest')
#         Acc_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Acc_grid_train.flatten(), (X_test, Y_test),
#                             method='nearest')
#         Slope_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Slope_grid_train.flatten(),
#                               (X_test, Y_test), method='nearest')
#         StepElev_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), StepElev_grid_train.flatten(),
#                                  (X_test, Y_test), method='nearest')
#         Rough_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Rough_grid_train.flatten(),
#                               (X_test, Y_test), method='nearest')
#
#         # 处理 NaN 值并计算误差指标
#
#         # 高程数据误差
#         valid_indices_Z = ~np.isnan(Z_pred) & ~np.isnan(Z_test)
#         if np.any(valid_indices_Z):
#             mse_Z = mean_squared_error(Z_test[valid_indices_Z], Z_pred[valid_indices_Z])
#             mae_Z = mean_absolute_error(Z_test[valid_indices_Z], Z_pred[valid_indices_Z])
#             max_error_Z = np.max(np.abs(Z_test[valid_indices_Z] - Z_pred[valid_indices_Z]))
#             print(f"高程插值 - MSE: {mse_Z}, MAE: {mae_Z}, Max Error: {max_error_Z}")
#         else:
#             print("高程插值预测值中全部为 NaN，无法计算误差。")
#
#         # 通行标志（分类准确率）
#         valid_indices_TrafFlag = ~np.isnan(TrafFlag_pred) & ~np.isnan(TrafFlag_test)
#         if np.any(valid_indices_TrafFlag):
#             accuracy_TrafFlag = accuracy_score(TrafFlag_test[valid_indices_TrafFlag],
#                                                TrafFlag_pred[valid_indices_TrafFlag])
#             print(f"通行标志插值 - 准确率: {accuracy_TrafFlag}")
#         else:
#             print("通行标志插值预测值中全部为 NaN，无法计算准确率。")
#
#         # 最大加速度误差
#         valid_indices_Acc = ~np.isnan(Acc_pred) & ~np.isnan(maxAcc_test)
#         if np.any(valid_indices_Acc):
#             mse_Acc = mean_squared_error(maxAcc_test[valid_indices_Acc], Acc_pred[valid_indices_Acc])
#             mae_Acc = mean_absolute_error(maxAcc_test[valid_indices_Acc], Acc_pred[valid_indices_Acc])
#             max_error_Acc = np.max(np.abs(maxAcc_test[valid_indices_Acc] - Acc_pred[valid_indices_Acc]))
#             print(f"最大加速度插值 - MSE: {mse_Acc}, MAE: {mae_Acc}, Max Error: {max_error_Acc}")
#         else:
#             print("最大加速度插值预测值中全部为 NaN，无法计算误差。")
#
#         # 坡度误差
#         valid_indices_Slope = ~np.isnan(Slope_pred) & ~np.isnan(Slope_test)
#         if np.any(valid_indices_Slope):
#             mse_Slope = mean_squared_error(Slope_test[valid_indices_Slope], Slope_pred[valid_indices_Slope])
#             mae_Slope = mean_absolute_error(Slope_test[valid_indices_Slope], Slope_pred[valid_indices_Slope])
#             max_error_Slope = np.max(np.abs(Slope_test[valid_indices_Slope] - Slope_pred[valid_indices_Slope]))
#             print(f"坡度插值 - MSE: {mse_Slope}, MAE: {mae_Slope}, Max Error: {max_error_Slope}")
#         else:
#             print("坡度插值预测值中全部为 NaN，无法计算误差。")
#
#         # 阶跃高程差误差
#         valid_indices_StepElev = ~np.isnan(StepElev_pred) & ~np.isnan(StepElev_test)
#         if np.any(valid_indices_StepElev):
#             mse_StepElev = mean_squared_error(StepElev_test[valid_indices_StepElev],
#                                               StepElev_pred[valid_indices_StepElev])
#             mae_StepElev = mean_absolute_error(StepElev_test[valid_indices_StepElev],
#                                                StepElev_pred[valid_indices_StepElev])
#             max_error_StepElev = np.max(
#                 np.abs(StepElev_test[valid_indices_StepElev] - StepElev_pred[valid_indices_StepElev]))
#             print(f"阶跃高程差插值 - MSE: {mse_StepElev}, MAE: {mae_StepElev}, Max Error: {max_error_StepElev}")
#         else:
#             print("阶跃高程差插值预测值中全部为 NaN，无法计算误差。")
#
#         # 粗糙度误差
#         valid_indices_Rough = ~np.isnan(Rough_pred) & ~np.isnan(Rough_test)
#         if np.any(valid_indices_Rough):
#             mse_Rough = mean_squared_error(Rough_test[valid_indices_Rough], Rough_pred[valid_indices_Rough])
#             mae_Rough = mean_absolute_error(Rough_test[valid_indices_Rough], Rough_pred[valid_indices_Rough])
#             max_error_Rough = np.max(np.abs(Rough_test[valid_indices_Rough] - Rough_pred[valid_indices_Rough]))
#             print(f"粗糙度插值 - MSE: {mse_Rough}, MAE: {mae_Rough}, Max Error: {max_error_Rough}")
#         else:
#             print("粗糙度插值预测值中全部为 NaN，无法计算误差。")
#
# # 节点类和路径规划类
# class Node:
#     def __init__(self, position, parent=None):
#         self.position = position  # (row, col)
#         self.parent = parent
#         self.g = 0  # 累计代价
#         self.h = 0  # 启发代价
#         self.f = 0  # 总代价
#
#     def __lt__(self, other):
#         return self.f < other.f
#
# class PathPlanner:
#     def __init__(self, grid_data, max_slope=None, max_step_elev=None):
#         self.grid = grid_data['TrafFlag_grid']
#         self.Z_grid = grid_data['Z_grid']
#         self.Acc_grid = grid_data['Acc_grid']
#         self.Slope_grid = grid_data['Slope_grid']
#         self.StepElev_grid = grid_data['StepElev_grid']
#         self.Rough_grid = grid_data['Rough_grid']
#         self.max_slope = max_slope
#         self.max_step_elev = max_step_elev
#
#     def heuristic_distance_3d(self, current_pos, end_pos):
#         d_xy = np.sqrt((current_pos[0] - end_pos[0]) ** 2 +
#                        (current_pos[1] - end_pos[1]) ** 2)
#         z_current = self.Z_grid[current_pos[0], current_pos[1]]
#         z_end = self.Z_grid[end_pos[0], end_pos[1]]
#         d_z = (z_end - z_current) ** 2
#         return np.sqrt(d_xy ** 2 + d_z)
#
#     def astar(self, start, end, mode='distance', alpha_env=0.6, beta_env=0.4):
#         start_node = Node(start)
#         end_node = Node(end)
#
#         open_list = []
#         heapq.heappush(open_list, start_node)
#         open_dict = {start_node.position: start_node}
#
#         closed_set = set()
#
#         while open_list:
#             current_node = heapq.heappop(open_list)
#             del open_dict[current_node.position]
#             closed_set.add(current_node.position)
#
#             # 检查是否到达终点
#             if current_node.position == end_node.position:
#                 path = []
#                 while current_node:
#                     path.append(current_node.position)
#                     current_node = current_node.parent
#                 return path[::-1]
#
#             # 定义8个邻居（包括对角线）
#             neighbors = [(0, 1), (0, -1), (1, 0), (-1, 0),
#                          (1, 1), (1, -1), (-1, 1), (-1, -1)]
#
#             for offset in neighbors:
#                 neighbor_pos = (current_node.position[0] + offset[0],
#                                 current_node.position[1] + offset[1])
#
#                 # 检查邻居是否在网格范围内
#                 if not (0 <= neighbor_pos[0] < self.grid.shape[0] and
#                         0 <= neighbor_pos[1] < self.grid.shape[1]):
#                     continue
#
#                 # 检查 TrafFlag 是否可通行
#                 if self.grid[neighbor_pos] == 0:
#                     continue  # 不可通行区域
#
#                 # 在所有模式下应用坡度和阶跃高程差的限制
#                 if self.max_slope is not None and self.Slope_grid[neighbor_pos] > self.max_slope:
#                     continue  # 坡度超过最大限制
#                 if self.max_step_elev is not None and self.StepElev_grid[neighbor_pos] > self.max_step_elev:
#                     continue  # 阶跃高程差超过最大限制
#
#                 # 检查是否已经在关闭列表中
#                 if neighbor_pos in closed_set:
#                     continue
#
#                 # 处理不同模式
#                 if mode == 'distance':
#                     movement_cost = self.heuristic_distance_3d(current_node.position, neighbor_pos)
#                     g_cost = current_node.g + movement_cost
#                     h_cost = self.heuristic_distance_3d(neighbor_pos, end)
#                 elif mode == 'time':
#                     a_max = self.Acc_grid[neighbor_pos[0], neighbor_pos[1]]
#                     if a_max <= 0:
#                         continue  # 无法加速，跳过
#                     movement_cost = self.heuristic_distance_3d(current_node.position, neighbor_pos)
#                     g_cost = current_node.g + movement_cost / a_max  # 简化为距离除以加速度
#                     h_cost = self.heuristic_distance_3d(neighbor_pos, end)
#                 elif mode == 'env':
#                     movement_uncertainty = (self.Slope_grid[neighbor_pos[0], neighbor_pos[1]] / 45.0 +
#                                             self.StepElev_grid[neighbor_pos[0], neighbor_pos[1]] / 1.5 +
#                                             self.Rough_grid[neighbor_pos[0], neighbor_pos[1]] / 6.0)
#                     g_cost = current_node.g + movement_uncertainty
#                     h_cost = self.heuristic_distance_3d(neighbor_pos, end)
#                 else:
#                     raise ValueError("Invalid mode. Choose from 'distance', 'time', 'env'.")
#
#                 # 创建邻居节点
#                 neighbor_node = Node(neighbor_pos, current_node)
#                 neighbor_node.g = g_cost
#                 neighbor_node.h = h_cost
#                 neighbor_node.f = neighbor_node.g + neighbor_node.h
#
#                 # 检查开放列表中是否有更好的节点
#                 if neighbor_pos in open_dict:
#                     existing_node = open_dict[neighbor_pos]
#                     if existing_node.g <= neighbor_node.g:
#                         continue  # 已有更优节点
#                     else:
#                         # 替换为更优节点
#                         open_list.remove(existing_node)
#                         heapq.heapify(open_list)
#                         heapq.heappush(open_list, neighbor_node)
#                         open_dict[neighbor_pos] = neighbor_node
#                 else:
#                     heapq.heappush(open_list, neighbor_node)
#                     open_dict[neighbor_pos] = neighbor_node
#
#         return None  # 如果找不到路径，返回 None
#
#     def extract_path_coordinates(self, path, x_range, y_range):
#         return [(x_range[p[1]], y_range[p[0]], self.Z_grid[p[0], p[1]]) for p in path]
#
#     def calculate_path_length(self, path_coords):
#         length = 0
#         for i in range(1, len(path_coords)):
#             x1, y1, z1 = path_coords[i - 1]
#             x2, y2, z2 = path_coords[i]
#             length += np.sqrt((x2 - x1) ** 2 + (y2 - y1) ** 2 + (z2 - z1) ** 2)
#         return length
#
#     def compute_traversal_time(self, path_coords, y_range, x_range):
#         total_time = 0.0
#         current_speed = 0.0  # 初始速度为0
#         max_speed = 70.0  # 最大速度70 m/s
#
#         speeds = [current_speed]  # 存储每个点的速度
#         accelerations = []  # 存储每段的加速度
#
#         for i in range(1, len(path_coords)):
#             x1, y1, z1 = path_coords[i - 1]
#             x2, y2, z2 = path_coords[i]
#             # 计算两点之间的距离
#             d = np.sqrt((x2 - x1) ** 2 + (y2 - y1) ** 2 + (z2 - z1) ** 2)
#
#             # 获取当前点的网格索引
#             row = np.argmin(np.abs(y_range - y1))
#             col = np.argmin(np.abs(x_range - x1))
#
#             # 获取当前点的加速度
#             a = self.Acc_grid[row, col]
#             if a <= 0:
#                 # 如果加速度不可用或为0，无法继续加速
#                 a = 0.0
#
#             # 计算需要加速到最大速度的速度增量
#             delta_v = max_speed - current_speed
#
#             if current_speed < max_speed and a > 0:
#                 # 需要加速
#                 t_acc = delta_v / a  # 加速所需时间
#                 d_acc = current_speed * t_acc + 0.5 * a * t_acc ** 2  # 加速所需距离
#
#                 if d_acc >= d:
#                     # 在当前段无法加速到最大速度
#                     # 使用运动学方程计算通过距离d所需的时间
#                     # 解方程: 0.5 * a * t^2 + current_speed * t - d = 0
#                     discriminant = current_speed ** 2 + 2 * a * d
#                     if discriminant < 0:
#                         # 无实数解，无法到达下一个点
#                         return float('inf'), [], []
#                     t = (-current_speed + math.sqrt(discriminant)) / a
#                     total_time += t
#                     current_speed += a * t
#                     accelerations.append(a)
#                     speeds.append(current_speed)
#                 else:
#                     # 在当前段可以加速到最大速度
#                     # 先加速到最大速度
#                     total_time += t_acc
#                     current_speed = max_speed
#                     accelerations.append(a)
#                     speeds.append(current_speed)
#                     # 剩余距离
#                     d_remain = d - d_acc
#                     # 以最大速度匀速行驶剩余距离
#                     t_remain = d_remain / current_speed
#                     total_time += t_remain
#                     accelerations.append(0.0)
#                     speeds.append(current_speed)
#             else:
#                 # 无法加速或已达到最大速度，匀速行驶
#                 if current_speed == 0:
#                     # 无法移动
#                     return float('inf'), [], []
#                 t = d / current_speed
#                 total_time += t
#                 accelerations.append(0.0)
#                 speeds.append(current_speed)
#
#         return total_time, speeds, accelerations
#
# # 可视化类
# class Visualizer:
#     def __init__(self, X_grid, Y_grid, Z_grid, TrafFlag_grid):
#         self.X_grid = X_grid
#         self.Y_grid = Y_grid
#         self.Z_grid = Z_grid
#         self.TrafFlag_grid = TrafFlag_grid
#         self.fig = go.Figure()
#
#     def plot_terrain(self):
#         self.fig.add_trace(go.Surface(
#             x=self.X_grid, y=self.Y_grid, z=self.Z_grid,
#             surfacecolor=self.TrafFlag_grid,
#             colorscale='Viridis',
#             opacity=0.8,
#             showscale=True,
#             colorbar=dict(title='TrafFlag')
#         ))
#
#     def add_path(self, coords, speeds, accs, color, name, length, time_taken, path_width=4):
#         if coords:
#             x, y, z = zip(*coords)
#             # 处理不可行路径（时间为inf）的情况
#             if math.isinf(time_taken):
#                 time_str = "无法到达"
#             else:
#                 time_str = f"{time_taken:.2f} 秒"
#
#             # 准备 customdata，将速度和加速度添加进去
#             if accs:
#                 accs = [0.0] + accs  # 与 speeds 对齐
#             else:
#                 accs = [0.0] * len(speeds)
#
#             customdata = np.array([speeds, accs]).T  # 转置使其每行对应一个点的 [speed, acceleration]
#
#             self.fig.add_trace(go.Scatter3d(
#                 x=x, y=y, z=z,
#                 mode='lines+markers',
#                 line=dict(color=color, width=path_width),
#                 marker=dict(size=2, color=color),
#                 name=f"{name} (长度: {length:.2f} 米, 时间: {time_str})",
#                 customdata=customdata,
#                 hovertemplate=(
#                         'X: %{x}<br>' +
#                         'Y: %{y}<br>' +
#                         'Z: %{z}<br>' +
#                         '速度: %{customdata[0]:.2f} m/s<br>' +
#                         '加速度: %{customdata[1]:.2f} m/s²<br>' +
#                         '<extra></extra>'
#                 )
#             ))
#         else:
#             print(f"未找到{name}")
#
#     def add_start_and_end_points(self, start_coord, end_coord):
#         self.fig.add_trace(go.Scatter3d(
#             x=[start_coord[0]],
#             y=[start_coord[1]],
#             z=[start_coord[2]],
#             mode='markers',
#             marker=dict(size=8, color='green'),
#             name='起点'
#         ))
#         self.fig.add_trace(go.Scatter3d(
#             x=[end_coord[0]],
#             y=[end_coord[1]],
#             z=[end_coord[2]],
#             mode='markers',
#             marker=dict(size=8, color='red'),
#             name='终点'
#         ))
#
#     def show(self):
#         self.fig.update_layout(
#             scene=dict(
#                 xaxis_title='X 坐标',
#                 yaxis_title='Y 坐标',
#                 zaxis_title='Z 坐标 (高程)',
#                 aspectmode='data'
#             ),
#             title='3D 地形路径规划（基于距离、时间与最小不确定性）',
#             legend=dict(
#                 x=0,
#                 y=1,
#                 bgcolor='rgba(255, 255, 255, 0)',
#                 bordercolor='rgba(255, 255, 255, 0)'
#             )
#         )
#         pio.show(self.fig)
#
# # 报告生成函数
# def write_paths_to_word(document, path_name, path_coords):
#     if not path_coords:
#         return
#     # 添加路径名称和点数量
#     document.add_heading(f"{path_name} (点数: {len(path_coords)})", level=2)
#     # 创建表格，包含三列：X相对坐标、Y相对坐标、Z高程坐标
#     table = document.add_table(rows=1, cols=3)
#     table.style = 'Light Grid'
#     hdr_cells = table.rows[0].cells
#     hdr_cells[0].text = 'X相对坐标'
#     hdr_cells[1].text = 'Y相对坐标'
#     hdr_cells[2].text = 'Z高程坐标'
#     # 填充表格内容
#     for coord in path_coords:
#         row_cells = table.add_row().cells
#         row_cells[0].text = f"{coord[0]}"
#         row_cells[1].text = f"{coord[1]}"
#         row_cells[2].text = f"{coord[2]}"
#     # 添加一个空行
#     document.add_paragraph()
#
# # 处理函数
# def run_processing(csv_path, start_coord, end_coord, output_path):
#     # 设置路径宽度
#     path_width = 4  # 可以根据需要调整路径宽度
#
#     # 数据处理
#     data_processor = DataProcessor(csv_path)
#     data_processor.load_data()
#     data_processor.interpolate_data()
#     data_processor.fill_nan_values()
#     data_processor.split_data_and_evaluate()
#
#     # 将 start_coord 和 end_coord 转换为网格索引
#     start = data_processor.coord_to_index(start_coord)
#     end = data_processor.coord_to_index(end_coord)
#
#     # 检查起点和终点是否在网格范围内
#     if not data_processor.is_valid_position(start):
#         raise ValueError("起点坐标超出网格范围")
#     if not data_processor.is_valid_position(end):
#         raise ValueError("终点坐标超出网格范围")
#
#     # 确保起点和终点可通行
#     if data_processor.TrafFlag_grid[start] == 0:
#         raise ValueError("起点不可通行")
#     if data_processor.TrafFlag_grid[end] == 0:
#         raise ValueError("终点不可通行")
#
#     # 定义坡度和阶跃高程差的最大允许值（应用于所有模式）
#     max_allowed_slope = 25.0  # 最大坡度（度）
#     max_allowed_step_elev = 10  # 最大阶跃高程差（米）
#
#     # 路径规划
#     grid_data = {
#         'TrafFlag_grid': data_processor.TrafFlag_grid,
#         'Z_grid': data_processor.Z_grid,
#         'Acc_grid': data_processor.Acc_grid,
#         'Slope_grid': data_processor.Slope_grid,
#         'StepElev_grid': data_processor.StepElev_grid,
#         'Rough_grid': data_processor.Rough_grid
#     }
#     path_planner = PathPlanner(grid_data, max_slope=max_allowed_slope, max_step_elev=max_allowed_step_elev)
#
#     print("正在计算基于距离的最短路径...")
#     path_distance = path_planner.astar(start, end, mode='distance')
#     if path_distance:
#         print("基于距离的路径计算完成")
#     else:
#         print("未找到基于距离的路径")
#
#     print("正在计算基于时间的最短耗时路径...")
#     path_time = path_planner.astar(start, end, mode='time')
#     if path_time:
#         print("基于时间的路径计算完成")
#     else:
#         print("未找到基于时间的路径")
#
#     print("正在计算最小不确定性路径...")
#     path_env = path_planner.astar(start, end, mode='env')
#     if path_env:
#         print("最小不确定性路径计算完成")
#     else:
#         print("未找到最小不确定性路径")
#
#     # 提取路径坐标
#     if path_distance:
#         path_distance_coords = path_planner.extract_path_coordinates(path_distance, data_processor.x_range,
#                                                                      data_processor.y_range)
#         length_distance = path_planner.calculate_path_length(path_distance_coords)
#     else:
#         path_distance_coords = None
#         length_distance = None
#     if path_time:
#         path_time_coords = path_planner.extract_path_coordinates(path_time, data_processor.x_range,
#                                                                  data_processor.y_range)
#         length_time = path_planner.calculate_path_length(path_time_coords)
#     else:
#         path_time_coords = None
#         length_time = None
#     if path_env:
#         path_env_coords = path_planner.extract_path_coordinates(path_env, data_processor.x_range,
#                                                                 data_processor.y_range)
#         length_env = path_planner.calculate_path_length(path_env_coords)
#     else:
#         path_env_coords = None
#         length_env = None
#
#     # 计算遍历时间，并获取速度和加速度数据
#     if path_distance:
#         time_distance, speeds_distance, accs_distance = path_planner.compute_traversal_time(
#             path_distance_coords, data_processor.y_range, data_processor.x_range)
#     else:
#         speeds_distance, accs_distance, time_distance = None, None, None
#
#     if path_time:
#         time_time, speeds_time, accs_time = path_planner.compute_traversal_time(
#             path_time_coords, data_processor.y_range, data_processor.x_range)
#     else:
#         speeds_time, accs_time, time_time = None, None, None
#
#     if path_env:
#         time_env, speeds_env, accs_env = path_planner.compute_traversal_time(
#             path_env_coords, data_processor.y_range, data_processor.x_range)
#     else:
#         speeds_env, accs_env, time_env = None, None, None
#
#     # 可视化
#     visualizer = Visualizer(
#         data_processor.X_grid,
#         data_processor.Y_grid,
#         data_processor.Z_grid,
#         data_processor.TrafFlag_grid
#     )
#     visualizer.plot_terrain()
#
#     colors = ['blue', 'orange', 'purple']
#     names = ['MinDistPath 最短距离路径', 'MinTimePath 最短耗时路径', 'MinUncertPath 最小不确定性路径']
#     lengths = [length_distance, length_time, length_env]
#     times = [time_distance, time_time, time_env]
#     coords_list = [path_distance_coords, path_time_coords, path_env_coords]
#     speeds_list = [speeds_distance, speeds_time, speeds_env]
#     accs_list = [accs_distance, accs_time, accs_env]
#
#     for coords, speeds, accs, color, name, length, time_taken in zip(
#             coords_list, speeds_list, accs_list,
#             colors, names, lengths, times):
#         visualizer.add_path(coords, speeds, accs, color, name, length, time_taken, path_width)
#
#     # 标记起点和终点
#     if path_distance or path_time or path_env:
#         selected_path = None
#         if path_distance:
#             selected_path = path_distance_coords
#         elif path_time:
#             selected_path = path_time_coords
#         elif path_env:
#             selected_path = path_env_coords
#
#         if selected_path:
#             # 使用起点坐标
#             start_x, start_y, start_z = selected_path[0]
#             # 使用终点坐标
#             end_x, end_y, end_z = selected_path[-1]
#             visualizer.add_start_and_end_points((start_x, start_y, start_z), (end_x, end_y, end_z))
#
#     visualizer.show()
#
#     # 创建 Word 文档
#     document = Document()
#     document.add_heading('路径规划结果', 0)
#     # 写入路径信息
#     if path_distance or path_time or path_env:
#         if path_distance:
#             write_paths_to_word(document, 'MinDistPath ', path_distance_coords)
#         if path_time:
#             write_paths_to_word(document, 'MinTimePath ', path_time_coords)
#         if path_env:
#             write_paths_to_word(document, 'MinUncertPath ', path_env_coords)
#
#         # 保存文档
#         try:
#             # 使用 resource_path 函数获取输出文件路径
#             output_file_path = output_path  # 如果 output_path 是用户指定的完整路径，则无需使用 resource_path
#             document.save(output_file_path)
#             print(f"路径坐标已成功保存到 Word 文档: {output_file_path}")
#         except Exception as e:
#             print(f"保存 Word 文档时出错: {e}")
#     else:
#         print("没有路径可写入 Word 文档")
#
#     # 打印路径长度和遍历时间
#     if path_distance:
#         if math.isinf(time_distance):
#             print(f"最短路径长度: {length_distance:.2f} 米")
#             print(f"最短距离路径遍历时间: 无法到达")
#         else:
#             print(f"最短路径长度: {length_distance:.2f} 米")
#             print(f"最短距离路径遍历时间: {time_distance:.2f} 秒")
#     if path_time:
#         if math.isinf(time_time):
#             print(f"最短耗时路径长度: {length_time:.2f} 米")
#             print(f"最短耗时路径遍历时间: 无法到达")
#         else:
#             print(f"最短耗时路径长度: {length_time:.2f} 米")
#             print(f"最短耗时路径遍历时间: {time_time:.2f} 秒")
#     if path_env:
#         if math.isinf(time_env):
#             print(f"最小不确定性路径长度: {length_env:.2f} 米")
#             print(f"最小不确定性路径遍历时间: 无法到达")
#         else:
#             print(f"最小不确定性路径长度: {length_env:.2f} 米")
#             print(f"最小不确定性路径遍历时间: {time_env:.2f} 秒")
#
# # 主程序
# def main():
#     # 创建主窗口
#     root = tk.Tk()
#     root.title("路径规划软件")
#
#     # 地图数据文件路径
#     tk.Label(root, text="地图数据文件路径:").grid(row=0, column=0, sticky='e')
#     map_data_path_entry = tk.Entry(root, width=50)
#     map_data_path_entry.grid(row=0, column=1)
#     def browse_map_data():
#         filename = filedialog.askopenfilename(filetypes=[("CSV文件", "*.csv")])
#         if filename:
#             map_data_path_entry.delete(0, tk.END)
#             map_data_path_entry.insert(0, filename)
#     tk.Button(root, text="浏览...", command=browse_map_data).grid(row=0, column=2)
#
#     # 起点坐标
#     tk.Label(root, text="起点坐标 (格式: X,Y):").grid(row=1, column=0, sticky='e')
#     start_coord_entry = tk.Entry(root, width=50)
#     start_coord_entry.grid(row=1, column=1)
#
#     # 终点坐标
#     tk.Label(root, text="终点坐标 (格式: X,Y):").grid(row=2, column=0, sticky='e')
#     end_coord_entry = tk.Entry(root, width=50)
#     end_coord_entry.grid(row=2, column=1)
#
#     # 输出路径
#     tk.Label(root, text="输出 Word 文档路径:").grid(row=3, column=0, sticky='e')
#     output_path_entry = tk.Entry(root, width=50)
#     output_path_entry.grid(row=3, column=1)
#     def browse_output_path():
#         filename = filedialog.asksaveasfilename(defaultextension=".docx",
#                                                 filetypes=[("Word文档", "*.docx")])
#         if filename:
#             output_path_entry.delete(0, tk.END)
#             output_path_entry.insert(0, filename)
#     tk.Button(root, text="浏览...", command=browse_output_path).grid(row=3, column=2)
#
#     # 运行按钮
#     def run():
#         # 获取输入值
#         csv_path = map_data_path_entry.get()
#         start_coord_input = start_coord_entry.get()
#         end_coord_input = end_coord_entry.get()
#         output_path = output_path_entry.get()
#
#         # 验证输入
#         if not csv_path:
#             messagebox.showerror("错误", "请提供地图数据文件路径。")
#             return
#         if not output_path:
#             messagebox.showerror("错误", "请提供输出 Word 文档路径。")
#             return
#         if not start_coord_input:
#             messagebox.showerror("错误", "请提供起点坐标。")
#             return
#         if not end_coord_input:
#             messagebox.showerror("错误", "请提供终点坐标。")
#             return
#         try:
#             start_x, start_y = map(float, start_coord_input.strip().split(','))
#             start_coord = (start_x, start_y)
#         except ValueError:
#             messagebox.showerror("错误", "起点坐标格式错误，请使用 'X,Y' 的格式。")
#             return
#
#         try:
#             end_x, end_y = map(float, end_coord_input.strip().split(','))
#             end_coord = (end_x, end_y)
#         except ValueError:
#             messagebox.showerror("错误", "终点坐标格式错误，请使用 'X,Y' 的格式。")
#             return
#
#         # 在后台线程中运行，以避免阻塞GUI
#         def process():
#             try:
#                 # 调用处理函数
#                 run_processing(csv_path, start_coord, end_coord, output_path)
#                 messagebox.showinfo("完成", "路径规划已完成。")
#             except Exception as e:
#                 messagebox.showerror("错误", f"发生错误: {e}")
#
#         threading.Thread(target=process).start()
#
#     tk.Button(root, text="运行", command=run).grid(row=4, column=1)
#
#     root.mainloop()
#
# if __name__ == "__main__":
#     main()






# import pandas as pd
# import numpy as np
# from scipy.interpolate import griddata
# from sklearn.metrics import mean_squared_error, mean_absolute_error, accuracy_score
# from sklearn.model_selection import train_test_split
# import heapq
# import math
# import plotly.graph_objs as go
# import plotly.io as pio
# from docx import Document
# import tkinter as tk
# from tkinter import filedialog, messagebox
# import threading
# import sys
# import os
# import logging
#
# # 设置日志记录
# logging.basicConfig(filename='error.log', level=logging.ERROR)
#
# # 设置 Plotly 渲染器为浏览器
# pio.renderers.default = 'browser'
#
# # 添加 resource_path 函数，用于获取资源文件的正确路径
# def resource_path(relative_path):
#     """获取资源文件的绝对路径，兼容 PyInstaller 打包后的情况"""
#     try:
#         # PyInstaller 会把路径存储在 _MEIPASS 中
#         base_path = sys._MEIPASS
#     except AttributeError:
#         base_path = os.path.abspath(".")
#     return os.path.join(base_path, relative_path)
#
# # 数据处理类
# class DataProcessor:
#     def __init__(self, dat_path, x_resolution=1, y_resolution=1):
#         self.dat_path = dat_path
#         self.x_resolution = x_resolution
#         self.y_resolution = y_resolution
#         self.df = None
#         self.X = None
#         self.Y = None
#         self.Z = None
#         self.TrafFlag = None
#         self.maxAcc = None
#         self.Slope = None
#         self.StepElev = None
#         self.Roughness = None
#         self.x_range = None
#         self.y_range = None
#         self.X_grid = None
#         self.Y_grid = None
#         self.Z_grid = None
#         self.TrafFlag_grid = None
#         self.Acc_grid = None
#         self.Slope_grid = None
#         self.StepElev_grid = None
#         self.Rough_grid = None
#
#     def load_data(self):
#         try:
#             # 使用 resource_path 函数获取数据文件路径
#             data_path = resource_path(self.dat_path)
#             # 读取 .csv 文件，指定分隔符为逗号
#             self.df = pd.read_csv(data_path, delimiter=',')
#             print("数据文件中的列名：", self.df.columns.tolist())
#         except FileNotFoundError:
#             print(f"文件未找到: {self.dat_path}")
#             return
#         except pd.errors.EmptyDataError:
#             print("CSV 文件为空")
#             return
#         except Exception as e:
#             print(f"读取 CSV 文件时出错: {e}")
#             logging.exception("读取 CSV 文件时出错")
#             return
#
#         required_columns = {'X', 'Y', 'Z', 'TrafFlag', 'maxAcc', 'Slope', 'StepElevationDifference', 'Roughness'}
#         if not required_columns.issubset(self.df.columns):
#             print(f"CSV 文件缺少必要的列: {required_columns - set(self.df.columns)}")
#             return
#
#         # 提取所需数据
#         self.X = self.df['X'].values
#         self.Y = self.df['Y'].values
#         self.Z = self.df['Z'].values
#         self.TrafFlag = self.df['TrafFlag'].values
#         self.maxAcc = self.df['maxAcc'].values
#         self.Slope = self.df['Slope'].values
#         self.StepElev = self.df['StepElevationDifference'].values
#         self.Roughness = self.df['Roughness'].values
#
#         # 计算并打印 Slope 和 StepElevationDifference 的最小值和最大值
#         slope_min = np.min(self.Slope)
#         slope_max_val = np.max(self.Slope)
#         step_min = np.min(self.StepElev)
#         step_max = np.max(self.StepElev)
#
#         print(f"坡度 (Slope) 最小值: {slope_min}")
#         print(f"坡度 (Slope) 最大值: {slope_max_val}")
#         print(f"阶跃高程差 (StepElevationDifference) 最小值: {step_min}")
#         print(f"阶跃高程差 (StepElevationDifference) 最大值: {step_max}")
#
#     def interpolate_data(self):
#         # 计算网格范围和生成网格
#         x_min, x_max = self.X.min(), self.X.max()
#         y_min, y_max = self.Y.min(), self.Y.max()
#         self.x_range = np.arange(x_min, x_max + self.x_resolution, self.x_resolution)
#         self.y_range = np.arange(y_min, y_max + self.y_resolution, self.y_resolution)
#         self.X_grid, self.Y_grid = np.meshgrid(self.x_range, self.y_range)
#
#         # 插值
#         self.Z_grid = griddata((self.X, self.Y), self.Z, (self.X_grid, self.Y_grid), method='cubic')
#         self.TrafFlag_grid = griddata((self.X, self.Y), self.TrafFlag, (self.X_grid, self.Y_grid), method='nearest')
#         self.Acc_grid = griddata((self.X, self.Y), self.maxAcc, (self.X_grid, self.Y_grid), method='nearest')
#         self.Slope_grid = griddata((self.X, self.Y), self.Slope, (self.X_grid, self.Y_grid), method='nearest')
#         self.StepElev_grid = griddata((self.X, self.Y), self.StepElev, (self.X_grid, self.Y_grid), method='nearest')
#         self.Rough_grid = griddata((self.X, self.Y), self.Roughness, (self.X_grid, self.Y_grid), method='nearest')
#
#     def fill_nan_values(self):
#         # 填充 NaN 值
#         self.Z_grid = np.nan_to_num(self.Z_grid, nan=np.nanmean(self.Z))
#         self.Acc_grid = np.nan_to_num(self.Acc_grid, nan=0)
#         self.TrafFlag_grid = np.nan_to_num(self.TrafFlag_grid, nan=0)
#         self.Slope_grid = np.nan_to_num(self.Slope_grid, nan=0)
#         self.StepElev_grid = np.nan_to_num(self.StepElev_grid, nan=0)
#         self.Rough_grid = np.nan_to_num(self.Rough_grid, nan=0)
#
#         # 验证网格数据的形状
#         print(f"Z_grid shape: {self.Z_grid.shape}")
#         print(f"TrafFlag_grid shape: {self.TrafFlag_grid.shape}")
#         print(f"Acc_grid shape: {self.Acc_grid.shape}")
#         print(f"Slope_grid shape: {self.Slope_grid.shape}")
#         print(f"StepElev_grid shape: {self.StepElev_grid.shape}")
#         print(f"Rough_grid shape: {self.Rough_grid.shape}")
#
#         # 计算并打印插值后网格数据的最小值和最大值
#         slope_grid_min = np.min(self.Slope_grid)
#         slope_grid_max = np.max(self.Slope_grid)
#         step_grid_min = np.min(self.StepElev_grid)
#         step_grid_max = np.max(self.StepElev_grid)
#
#         print(f"插值后坡度网格 (Slope_grid) 最小值: {slope_grid_min}")
#         print(f"插值后坡度网格 (Slope_grid) 最大值: {slope_grid_max}")
#         print(f"插值后阶跃高程差网格 (StepElev_grid) 最小值: {step_grid_min}")
#         print(f"插值后阶跃高程差网格 (StepElev_grid) 最大值: {step_grid_max}")
#
#     def coord_to_index(self, coord):
#         x_min, y_min = self.x_range.min(), self.y_range.min()
#         col = int((coord[0] - x_min) / self.x_resolution)
#         row = int((coord[1] - y_min) / self.y_resolution)
#         return (row, col)
#
#     def is_valid_position(self, pos):
#         return 0 <= pos[0] < self.TrafFlag_grid.shape[0] and 0 <= pos[1] < self.TrafFlag_grid.shape[1]
#
#     def split_data_and_evaluate(self):
#         try:
#             # 将所有数据合并
#             data = np.column_stack(
#                 (self.X, self.Y, self.Z, self.TrafFlag, self.maxAcc, self.Slope, self.StepElev, self.Roughness))
#
#             # 分割数据集为训练集和测试集
#             test_size = 0.2
#             random_state = 42
#
#             train_data, test_data = train_test_split(data, test_size=test_size, random_state=random_state)
#
#             # 拆分训练集和测试集的数据
#             X_train, Y_train = train_data[:, 0], train_data[:, 1]
#             Z_train = train_data[:, 2]
#             TrafFlag_train = train_data[:, 3]
#             maxAcc_train = train_data[:, 4]
#             Slope_train = train_data[:, 5]
#             StepElev_train = train_data[:, 6]
#             Rough_train = train_data[:, 7]
#
#             X_test, Y_test = test_data[:, 0], test_data[:, 1]
#             Z_test = test_data[:, 2]
#             TrafFlag_test = test_data[:, 3]
#             maxAcc_test = test_data[:, 4]
#             Slope_test = test_data[:, 5]
#             StepElev_test = test_data[:, 6]
#             Rough_test = test_data[:, 7]
#
#             # 使用训练集进行插值
#             Z_grid_train = griddata((X_train, Y_train), Z_train, (self.X_grid, self.Y_grid), method='cubic')
#             TrafFlag_grid_train = griddata((X_train, Y_train), TrafFlag_train, (self.X_grid, self.Y_grid), method='nearest')
#             Acc_grid_train = griddata((X_train, Y_train), maxAcc_train, (self.X_grid, self.Y_grid), method='nearest')
#             Slope_grid_train = griddata((X_train, Y_train), Slope_train, (self.X_grid, self.Y_grid), method='nearest')
#             StepElev_grid_train = griddata((X_train, Y_train), StepElev_train, (self.X_grid, self.Y_grid), method='nearest')
#             Rough_grid_train = griddata((X_train, Y_train), Rough_train, (self.X_grid, self.Y_grid), method='nearest')
#
#             # 在测试集上预测值
#             Z_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Z_grid_train.flatten(), (X_test, Y_test),
#                               method='cubic')
#             TrafFlag_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), TrafFlag_grid_train.flatten(),
#                                      (X_test, Y_test), method='nearest')
#             Acc_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Acc_grid_train.flatten(), (X_test, Y_test),
#                                 method='nearest')
#             Slope_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Slope_grid_train.flatten(),
#                                   (X_test, Y_test), method='nearest')
#             StepElev_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), StepElev_grid_train.flatten(),
#                                      (X_test, Y_test), method='nearest')
#             Rough_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Rough_grid_train.flatten(),
#                                   (X_test, Y_test), method='nearest')
#
#             # 处理 NaN 值并计算误差指标
#
#             # 高程数据误差
#             valid_indices_Z = ~np.isnan(Z_pred) & ~np.isnan(Z_test)
#             if np.any(valid_indices_Z):
#                 mse_Z = mean_squared_error(Z_test[valid_indices_Z], Z_pred[valid_indices_Z])
#                 mae_Z = mean_absolute_error(Z_test[valid_indices_Z], Z_pred[valid_indices_Z])
#                 max_error_Z = np.max(np.abs(Z_test[valid_indices_Z] - Z_pred[valid_indices_Z]))
#                 print(f"高程插值 - MSE: {mse_Z}, MAE: {mae_Z}, Max Error: {max_error_Z}")
#             else:
#                 print("高程插值预测值中全部为 NaN，无法计算误差。")
#
#             # 通行标志（分类准确率）
#             valid_indices_TrafFlag = ~np.isnan(TrafFlag_pred) & ~np.isnan(TrafFlag_test)
#             if np.any(valid_indices_TrafFlag):
#                 accuracy_TrafFlag = accuracy_score(TrafFlag_test[valid_indices_TrafFlag],
#                                                    TrafFlag_pred[valid_indices_TrafFlag])
#                 print(f"通行标志插值 - 准确率: {accuracy_TrafFlag}")
#             else:
#                 print("通行标志插值预测值中全部为 NaN，无法计算准确率。")
#
#             # 最大加速度误差
#             valid_indices_Acc = ~np.isnan(Acc_pred) & ~np.isnan(maxAcc_test)
#             if np.any(valid_indices_Acc):
#                 mse_Acc = mean_squared_error(maxAcc_test[valid_indices_Acc], Acc_pred[valid_indices_Acc])
#                 mae_Acc = mean_absolute_error(maxAcc_test[valid_indices_Acc], Acc_pred[valid_indices_Acc])
#                 max_error_Acc = np.max(np.abs(maxAcc_test[valid_indices_Acc] - Acc_pred[valid_indices_Acc]))
#                 print(f"最大加速度插值 - MSE: {mse_Acc}, MAE: {mae_Acc}, Max Error: {max_error_Acc}")
#             else:
#                 print("最大加速度插值预测值中全部为 NaN，无法计算误差。")
#
#             # 坡度误差
#             valid_indices_Slope = ~np.isnan(Slope_pred) & ~np.isnan(Slope_test)
#             if np.any(valid_indices_Slope):
#                 mse_Slope = mean_squared_error(Slope_test[valid_indices_Slope], Slope_pred[valid_indices_Slope])
#                 mae_Slope = mean_absolute_error(Slope_test[valid_indices_Slope], Slope_pred[valid_indices_Slope])
#                 max_error_Slope = np.max(np.abs(Slope_test[valid_indices_Slope] - Slope_pred[valid_indices_Slope]))
#                 print(f"坡度插值 - MSE: {mse_Slope}, MAE: {mae_Slope}, Max Error: {max_error_Slope}")
#             else:
#                 print("坡度插值预测值中全部为 NaN，无法计算误差。")
#
#             # 阶跃高程差误差
#             valid_indices_StepElev = ~np.isnan(StepElev_pred) & ~np.isnan(StepElev_test)
#             if np.any(valid_indices_StepElev):
#                 mse_StepElev = mean_squared_error(StepElev_test[valid_indices_StepElev],
#                                                   StepElev_pred[valid_indices_StepElev])
#                 mae_StepElev = mean_absolute_error(StepElev_test[valid_indices_StepElev],
#                                                    StepElev_pred[valid_indices_StepElev])
#                 max_error_StepElev = np.max(
#                     np.abs(StepElev_test[valid_indices_StepElev] - StepElev_pred[valid_indices_StepElev]))
#                 print(f"阶跃高程差插值 - MSE: {mse_StepElev}, MAE: {mae_StepElev}, Max Error: {max_error_StepElev}")
#             else:
#                 print("阶跃高程差插值预测值中全部为 NaN，无法计算误差。")
#
#             # 粗糙度误差
#             valid_indices_Rough = ~np.isnan(Rough_pred) & ~np.isnan(Rough_test)
#             if np.any(valid_indices_Rough):
#                 mse_Rough = mean_squared_error(Rough_test[valid_indices_Rough], Rough_pred[valid_indices_Rough])
#                 mae_Rough = mean_absolute_error(Rough_test[valid_indices_Rough], Rough_pred[valid_indices_Rough])
#                 max_error_Rough = np.max(np.abs(Rough_test[valid_indices_Rough] - Rough_pred[valid_indices_Rough]))
#                 print(f"粗糙度插值 - MSE: {mse_Rough}, MAE: {mae_Rough}, Max Error: {max_error_Rough}")
#             else:
#                 print("粗糙度插值预测值中全部为 NaN，无法计算误差。")
#         except Exception as e:
#             print(f"数据评估时发生错误: {e}")
#             logging.exception("数据评估时发生错误")
#
# # 节点类和路径规划类
# class Node:
#     def __init__(self, position, parent=None):
#         self.position = position  # (row, col)
#         self.parent = parent
#         self.g = 0  # 累计代价
#         self.h = 0  # 启发代价
#         self.f = 0  # 总代价
#
#     def __lt__(self, other):
#         return self.f < other.f
#
# class PathPlanner:
#     def __init__(self, grid_data, max_slope=None, max_step_elev=None):
#         self.grid = grid_data['TrafFlag_grid']
#         self.Z_grid = grid_data['Z_grid']
#         self.Acc_grid = grid_data['Acc_grid']
#         self.Slope_grid = grid_data['Slope_grid']
#         self.StepElev_grid = grid_data['StepElev_grid']
#         self.Rough_grid = grid_data['Rough_grid']
#         self.max_slope = max_slope
#         self.max_step_elev = max_step_elev
#
#     def heuristic_distance_3d(self, current_pos, end_pos):
#         d_xy = np.sqrt((current_pos[0] - end_pos[0]) ** 2 +
#                        (current_pos[1] - end_pos[1]) ** 2)
#         z_current = self.Z_grid[current_pos[0], current_pos[1]]
#         z_end = self.Z_grid[end_pos[0], end_pos[1]]
#         d_z = (z_end - z_current) ** 2
#         return np.sqrt(d_xy ** 2 + d_z)
#
#     def astar(self, start, end, mode='distance', alpha_env=0.6, beta_env=0.4):
#         start_node = Node(start)
#         end_node = Node(end)
#
#         open_list = []
#         heapq.heappush(open_list, start_node)
#         open_dict = {start_node.position: start_node}
#
#         closed_set = set()
#
#         while open_list:
#             current_node = heapq.heappop(open_list)
#             del open_dict[current_node.position]
#             closed_set.add(current_node.position)
#
#             # 检查是否到达终点
#             if current_node.position == end_node.position:
#                 path = []
#                 while current_node:
#                     path.append(current_node.position)
#                     current_node = current_node.parent
#                 return path[::-1]
#
#             # 定义8个邻居（包括对角线）
#             neighbors = [(0, 1), (0, -1), (1, 0), (-1, 0),
#                          (1, 1), (1, -1), (-1, 1), (-1, -1)]
#
#             for offset in neighbors:
#                 neighbor_pos = (current_node.position[0] + offset[0],
#                                 current_node.position[1] + offset[1])
#
#                 # 检查邻居是否在网格范围内
#                 if not (0 <= neighbor_pos[0] < self.grid.shape[0] and
#                         0 <= neighbor_pos[1] < self.grid.shape[1]):
#                     continue
#
#                 # 检查 TrafFlag 是否可通行
#                 if self.grid[neighbor_pos] == 0:
#                     continue  # 不可通行区域
#
#                 # 在所有模式下应用坡度和阶跃高程差的限制
#                 if self.max_slope is not None and self.Slope_grid[neighbor_pos] > self.max_slope:
#                     continue  # 坡度超过最大限制
#                 if self.max_step_elev is not None and self.StepElev_grid[neighbor_pos] > self.max_step_elev:
#                     continue  # 阶跃高程差超过最大限制
#
#                 # 检查是否已经在关闭列表中
#                 if neighbor_pos in closed_set:
#                     continue
#
#                 # 处理不同模式
#                 if mode == 'distance':
#                     movement_cost = self.heuristic_distance_3d(current_node.position, neighbor_pos)
#                     g_cost = current_node.g + movement_cost
#                     h_cost = self.heuristic_distance_3d(neighbor_pos, end)
#                 elif mode == 'time':
#                     a_max = self.Acc_grid[neighbor_pos[0], neighbor_pos[1]]
#                     if a_max <= 0:
#                         continue  # 无法加速，跳过
#                     movement_cost = self.heuristic_distance_3d(current_node.position, neighbor_pos)
#                     g_cost = current_node.g + movement_cost / a_max  # 简化为距离除以加速度
#                     h_cost = self.heuristic_distance_3d(neighbor_pos, end)
#                 elif mode == 'env':
#                     movement_uncertainty = (self.Slope_grid[neighbor_pos[0], neighbor_pos[1]] / 45.0 +
#                                             self.StepElev_grid[neighbor_pos[0], neighbor_pos[1]] / 1.5 +
#                                             self.Rough_grid[neighbor_pos[0], neighbor_pos[1]] / 6.0)
#                     g_cost = current_node.g + movement_uncertainty
#                     h_cost = self.heuristic_distance_3d(neighbor_pos, end)
#                 else:
#                     raise ValueError("Invalid mode. Choose from 'distance', 'time', 'env'.")
#
#                 # 创建邻居节点
#                 neighbor_node = Node(neighbor_pos, current_node)
#                 neighbor_node.g = g_cost
#                 neighbor_node.h = h_cost
#                 neighbor_node.f = neighbor_node.g + neighbor_node.h
#
#                 # 检查开放列表中是否有更好的节点
#                 if neighbor_pos in open_dict:
#                     existing_node = open_dict[neighbor_pos]
#                     if existing_node.g <= neighbor_node.g:
#                         continue  # 已有更优节点
#                     else:
#                         # 替换为更优节点
#                         open_list.remove(existing_node)
#                         heapq.heapify(open_list)
#                         heapq.heappush(open_list, neighbor_node)
#                         open_dict[neighbor_pos] = neighbor_node
#                 else:
#                     heapq.heappush(open_list, neighbor_node)
#                     open_dict[neighbor_pos] = neighbor_node
#
#         return None  # 如果找不到路径，返回 None
#
#     def extract_path_coordinates(self, path, x_range, y_range):
#         return [(x_range[p[1]], y_range[p[0]], self.Z_grid[p[0], p[1]]) for p in path]
#
#     def calculate_path_length(self, path_coords):
#         length = 0
#         for i in range(1, len(path_coords)):
#             x1, y1, z1 = path_coords[i - 1]
#             x2, y2, z2 = path_coords[i]
#             length += np.sqrt((x2 - x1) ** 2 + (y2 - y1) ** 2 + (z2 - z1) ** 2)
#         return length
#
#     def compute_traversal_time(self, path_coords, y_range, x_range):
#         total_time = 0.0
#         current_speed = 0.0  # 初始速度为0
#         max_speed = 70.0  # 最大速度70 m/s
#
#         speeds = [current_speed]  # 存储每个点的速度
#         accelerations = []  # 存储每段的加速度
#
#         for i in range(1, len(path_coords)):
#             x1, y1, z1 = path_coords[i - 1]
#             x2, y2, z2 = path_coords[i]
#             # 计算两点之间的距离
#             d = np.sqrt((x2 - x1) ** 2 + (y2 - y1) ** 2 + (z2 - z1) ** 2)
#
#             # 获取当前点的网格索引
#             row = np.argmin(np.abs(y_range - y1))
#             col = np.argmin(np.abs(x_range - x1))
#
#             # 获取当前点的加速度
#             a = self.Acc_grid[row, col]
#             if a <= 0:
#                 # 如果加速度不可用或为0，无法继续加速
#                 a = 0.0
#
#             # 计算需要加速到最大速度的速度增量
#             delta_v = max_speed - current_speed
#
#             if current_speed < max_speed and a > 0:
#                 # 需要加速
#                 t_acc = delta_v / a  # 加速所需时间
#                 d_acc = current_speed * t_acc + 0.5 * a * t_acc ** 2  # 加速所需距离
#
#                 if d_acc >= d:
#                     # 在当前段无法加速到最大速度
#                     # 使用运动学方程计算通过距离d所需的时间
#                     # 解方程: 0.5 * a * t^2 + current_speed * t - d = 0
#                     discriminant = current_speed ** 2 + 2 * a * d
#                     if discriminant < 0:
#                         # 无实数解，无法到达下一个点
#                         return float('inf'), [], []
#                     t = (-current_speed + math.sqrt(discriminant)) / a
#                     total_time += t
#                     current_speed += a * t
#                     accelerations.append(a)
#                     speeds.append(current_speed)
#                 else:
#                     # 在当前段可以加速到最大速度
#                     # 先加速到最大速度
#                     total_time += t_acc
#                     current_speed = max_speed
#                     accelerations.append(a)
#                     speeds.append(current_speed)
#                     # 剩余距离
#                     d_remain = d - d_acc
#                     # 以最大速度匀速行驶剩余距离
#                     t_remain = d_remain / current_speed
#                     total_time += t_remain
#                     accelerations.append(0.0)
#                     speeds.append(current_speed)
#             else:
#                 # 无法加速或已达到最大速度，匀速行驶
#                 if current_speed == 0:
#                     # 无法移动
#                     return float('inf'), [], []
#                 t = d / current_speed
#                 total_time += t
#                 accelerations.append(0.0)
#                 speeds.append(current_speed)
#
#         return total_time, speeds, accelerations
#
# # 可视化类
# class Visualizer:
#     def __init__(self, X_grid, Y_grid, Z_grid, TrafFlag_grid):
#         self.X_grid = X_grid
#         self.Y_grid = Y_grid
#         self.Z_grid = Z_grid
#         self.TrafFlag_grid = TrafFlag_grid
#         self.fig = go.Figure()
#
#     def plot_terrain(self):
#         self.fig.add_trace(go.Surface(
#             x=self.X_grid, y=self.Y_grid, z=self.Z_grid,
#             surfacecolor=self.TrafFlag_grid,
#             colorscale='Viridis',
#             opacity=0.8,
#             showscale=True,
#             colorbar=dict(title='TrafFlag')
#         ))
#
#     def add_path(self, coords, speeds, accs, color, name, length, time_taken, path_width=4):
#         if coords:
#             x, y, z = zip(*coords)
#             # 处理不可行路径（时间为inf）的情况
#             if math.isinf(time_taken):
#                 time_str = "无法到达"
#             else:
#                 time_str = f"{time_taken:.2f} 秒"
#
#             # 准备 customdata，将速度和加速度添加进去
#             if accs:
#                 accs = [0.0] + accs  # 与 speeds 对齐
#             else:
#                 accs = [0.0] * len(speeds)
#
#             customdata = np.array([speeds, accs]).T  # 转置使其每行对应一个点的 [speed, acceleration]
#
#             self.fig.add_trace(go.Scatter3d(
#                 x=x, y=y, z=z,
#                 mode='lines+markers',
#                 line=dict(color=color, width=path_width),
#                 marker=dict(size=2, color=color),
#                 name=f"{name} (长度: {length:.2f} 米, 时间: {time_str})",
#                 customdata=customdata,
#                 hovertemplate=(
#                         'X: %{x}<br>' +
#                         'Y: %{y}<br>' +
#                         'Z: %{z}<br>' +
#                         '速度: %{customdata[0]:.2f} m/s<br>' +
#                         '加速度: %{customdata[1]:.2f} m/s²<br>' +
#                         '<extra></extra>'
#                 )
#             ))
#         else:
#             print(f"未找到{name}")
#
#     def add_start_and_end_points(self, start_coord, end_coord):
#         self.fig.add_trace(go.Scatter3d(
#             x=[start_coord[0]],
#             y=[start_coord[1]],
#             z=[start_coord[2]],
#             mode='markers',
#             marker=dict(size=8, color='green'),
#             name='起点'
#         ))
#         self.fig.add_trace(go.Scatter3d(
#             x=[end_coord[0]],
#             y=[end_coord[1]],
#             z=[end_coord[2]],
#             mode='markers',
#             marker=dict(size=8, color='red'),
#             name='终点'
#         ))
#
#     def show(self):
#         self.fig.update_layout(
#             scene=dict(
#                 xaxis_title='X 坐标',
#                 yaxis_title='Y 坐标',
#                 zaxis_title='Z 坐标 (高程)',
#                 aspectmode='data'
#             ),
#             title='3D 地形路径规划（基于距离、时间与最小不确定性）',
#             legend=dict(
#                 x=0,
#                 y=1,
#                 bgcolor='rgba(255, 255, 255, 0)',
#                 bordercolor='rgba(255, 255, 255, 0)'
#             )
#         )
#         pio.show(self.fig)
#
# # 报告生成函数
# def write_paths_to_word(document, path_name, path_coords):
#     if not path_coords:
#         return
#     # 添加路径名称和点数量
#     document.add_heading(f"{path_name} (点数: {len(path_coords)})", level=2)
#     # 创建表格，包含三列：X相对坐标、Y相对坐标、Z高程坐标
#     table = document.add_table(rows=1, cols=3)
#     table.style = 'Light Grid'
#     hdr_cells = table.rows[0].cells
#     hdr_cells[0].text = 'X相对坐标'
#     hdr_cells[1].text = 'Y相对坐标'
#     hdr_cells[2].text = 'Z高程坐标'
#     # 填充表格内容
#     for coord in path_coords:
#         row_cells = table.add_row().cells
#         row_cells[0].text = f"{coord[0]}"
#         row_cells[1].text = f"{coord[1]}"
#         row_cells[2].text = f"{coord[2]}"
#     # 添加一个空行
#     document.add_paragraph()
#
# # 处理函数
# def run_processing(csv_path, start_coord, end_coord, output_path):
#     try:
#         # 设置路径宽度
#         path_width = 4  # 可以根据需要调整路径宽度
#
#         # 数据处理
#         data_processor = DataProcessor(csv_path)
#         data_processor.load_data()
#         data_processor.interpolate_data()
#         data_processor.fill_nan_values()
#         data_processor.split_data_and_evaluate()
#
#         # 将 start_coord 和 end_coord 转换为网格索引
#         start = data_processor.coord_to_index(start_coord)
#         end = data_processor.coord_to_index(end_coord)
#
#         # 检查起点和终点是否在网格范围内
#         if not data_processor.is_valid_position(start):
#             raise ValueError("起点坐标超出网格范围")
#         if not data_processor.is_valid_position(end):
#             raise ValueError("终点坐标超出网格范围")
#
#         # 确保起点和终点可通行
#         if data_processor.TrafFlag_grid[start] == 0:
#             raise ValueError("起点不可通行")
#         if data_processor.TrafFlag_grid[end] == 0:
#             raise ValueError("终点不可通行")
#
#         # 定义坡度和阶跃高程差的最大允许值（应用于所有模式）
#         max_allowed_slope = 25.0  # 最大坡度（度）
#         max_allowed_step_elev = 10  # 最大阶跃高程差（米）
#
#         # 路径规划
#         grid_data = {
#             'TrafFlag_grid': data_processor.TrafFlag_grid,
#             'Z_grid': data_processor.Z_grid,
#             'Acc_grid': data_processor.Acc_grid,
#             'Slope_grid': data_processor.Slope_grid,
#             'StepElev_grid': data_processor.StepElev_grid,
#             'Rough_grid': data_processor.Rough_grid
#         }
#         path_planner = PathPlanner(grid_data, max_slope=max_allowed_slope, max_step_elev=max_allowed_step_elev)
#
#         print("正在计算基于距离的最短路径...")
#         path_distance = path_planner.astar(start, end, mode='distance')
#         if path_distance:
#             print("基于距离的路径计算完成")
#         else:
#             print("未找到基于距离的路径")
#
#         print("正在计算基于时间的最短耗时路径...")
#         path_time = path_planner.astar(start, end, mode='time')
#         if path_time:
#             print("基于时间的路径计算完成")
#         else:
#             print("未找到基于时间的路径")
#
#         print("正在计算最小不确定性路径...")
#         path_env = path_planner.astar(start, end, mode='env')
#         if path_env:
#             print("最小不确定性路径计算完成")
#         else:
#             print("未找到最小不确定性路径")
#
#         # 提取路径坐标
#         if path_distance:
#             path_distance_coords = path_planner.extract_path_coordinates(path_distance, data_processor.x_range,
#                                                                          data_processor.y_range)
#             length_distance = path_planner.calculate_path_length(path_distance_coords)
#         else:
#             path_distance_coords = None
#             length_distance = None
#         if path_time:
#             path_time_coords = path_planner.extract_path_coordinates(path_time, data_processor.x_range,
#                                                                      data_processor.y_range)
#             length_time = path_planner.calculate_path_length(path_time_coords)
#         else:
#             path_time_coords = None
#             length_time = None
#         if path_env:
#             path_env_coords = path_planner.extract_path_coordinates(path_env, data_processor.x_range,
#                                                                     data_processor.y_range)
#             length_env = path_planner.calculate_path_length(path_env_coords)
#         else:
#             path_env_coords = None
#             length_env = None
#
#         # 计算遍历时间，并获取速度和加速度数据
#         if path_distance:
#             time_distance, speeds_distance, accs_distance = path_planner.compute_traversal_time(
#                 path_distance_coords, data_processor.y_range, data_processor.x_range)
#         else:
#             speeds_distance, accs_distance, time_distance = None, None, None
#
#         if path_time:
#             time_time, speeds_time, accs_time = path_planner.compute_traversal_time(
#                 path_time_coords, data_processor.y_range, data_processor.x_range)
#         else:
#             speeds_time, accs_time, time_time = None, None, None
#
#         if path_env:
#             time_env, speeds_env, accs_env = path_planner.compute_traversal_time(
#                 path_env_coords, data_processor.y_range, data_processor.x_range)
#         else:
#             speeds_env, accs_env, time_env = None, None, None
#
#         # 可视化
#         visualizer = Visualizer(
#             data_processor.X_grid,
#             data_processor.Y_grid,
#             data_processor.Z_grid,
#             data_processor.TrafFlag_grid
#         )
#         visualizer.plot_terrain()
#
#         colors = ['blue', 'orange', 'purple']
#         names = ['MinDistPath 最短距离路径', 'MinTimePath 最短耗时路径', 'MinUncertPath 最小不确定性路径']
#         lengths = [length_distance, length_time, length_env]
#         times = [time_distance, time_time, time_env]
#         coords_list = [path_distance_coords, path_time_coords, path_env_coords]
#         speeds_list = [speeds_distance, speeds_time, speeds_env]
#         accs_list = [accs_distance, accs_time, accs_env]
#
#         for coords, speeds, accs, color, name, length, time_taken in zip(
#                 coords_list, speeds_list, accs_list,
#                 colors, names, lengths, times):
#             visualizer.add_path(coords, speeds, accs, color, name, length, time_taken, path_width)
#
#         # 标记起点和终点
#         if path_distance or path_time or path_env:
#             selected_path = None
#             if path_distance:
#                 selected_path = path_distance_coords
#             elif path_time:
#                 selected_path = path_time_coords
#             elif path_env:
#                 selected_path = path_env_coords
#
#             if selected_path:
#                 # 使用起点坐标
#                 start_x, start_y, start_z = selected_path[0]
#                 # 使用终点坐标
#                 end_x, end_y, end_z = selected_path[-1]
#                 visualizer.add_start_and_end_points((start_x, start_y, start_z), (end_x, end_y, end_z))
#
#         visualizer.show()
#
#         # 创建 Word 文档
#         document = Document()
#         document.add_heading('路径规划结果', 0)
#         # 写入路径信息
#         if path_distance or path_time or path_env:
#             if path_distance:
#                 write_paths_to_word(document, 'MinDistPath ', path_distance_coords)
#             if path_time:
#                 write_paths_to_word(document, 'MinTimePath ', path_time_coords)
#             if path_env:
#                 write_paths_to_word(document, 'MinUncertPath ', path_env_coords)
#
#             # 保存文档
#             try:
#                 # 使用 resource_path 函数获取输出文件路径
#                 output_file_path = output_path  # 如果 output_path 是用户指定的完整路径，则无需使用 resource_path
#                 document.save(output_file_path)
#                 print(f"路径坐标已成功保存到 Word 文档: {output_file_path}")
#             except Exception as e:
#                 print(f"保存 Word 文档时出错: {e}")
#                 logging.exception("保存 Word 文档时出错")
#         else:
#             print("没有路径可写入 Word 文档")
#
#         # 打印路径长度和遍历时间
#         if path_distance:
#             if math.isinf(time_distance):
#                 print(f"最短路径长度: {length_distance:.2f} 米")
#                 print(f"最短距离路径遍历时间: 无法到达")
#             else:
#                 print(f"最短路径长度: {length_distance:.2f} 米")
#                 print(f"最短距离路径遍历时间: {time_distance:.2f} 秒")
#         if path_time:
#             if math.isinf(time_time):
#                 print(f"最短耗时路径长度: {length_time:.2f} 米")
#                 print(f"最短耗时路径遍历时间: 无法到达")
#             else:
#                 print(f"最短耗时路径长度: {length_time:.2f} 米")
#                 print(f"最短耗时路径遍历时间: {time_time:.2f} 秒")
#         if path_env:
#             if math.isinf(time_env):
#                 print(f"最小不确定性路径长度: {length_env:.2f} 米")
#                 print(f"最小不确定性路径遍历时间: 无法到达")
#             else:
#                 print(f"最小不确定性路径长度: {length_env:.2f} 米")
#                 print(f"最小不确定性路径遍历时间: {time_env:.2f} 秒")
#     except Exception as e:
#         print(f"运行处理时发生错误: {e}")
#         logging.exception("运行处理时发生错误")
#
# # 主程序
# def main():
#     try:
#         # 创建主窗口
#         root = tk.Tk()
#         root.title("路径规划软件")
#
#         # 地图数据文件路径
#         tk.Label(root, text="地图数据文件路径:").grid(row=0, column=0, sticky='e')
#         map_data_path_entry = tk.Entry(root, width=50)
#         map_data_path_entry.grid(row=0, column=1)
#         def browse_map_data():
#             filename = filedialog.askopenfilename(filetypes=[("CSV文件", "*.csv")])
#             if filename:
#                 map_data_path_entry.delete(0, tk.END)
#                 map_data_path_entry.insert(0, filename)
#         tk.Button(root, text="浏览...", command=browse_map_data).grid(row=0, column=2)
#
#         # 起点坐标
#         tk.Label(root, text="起点坐标 (格式: X,Y):").grid(row=1, column=0, sticky='e')
#         start_coord_entry = tk.Entry(root, width=50)
#         start_coord_entry.grid(row=1, column=1)
#
#         # 终点坐标
#         tk.Label(root, text="终点坐标 (格式: X,Y):").grid(row=2, column=0, sticky='e')
#         end_coord_entry = tk.Entry(root, width=50)
#         end_coord_entry.grid(row=2, column=1)
#
#         # 输出路径
#         tk.Label(root, text="输出 Word 文档路径:").grid(row=3, column=0, sticky='e')
#         output_path_entry = tk.Entry(root, width=50)
#         output_path_entry.grid(row=3, column=1)
#         def browse_output_path():
#             filename = filedialog.asksaveasfilename(defaultextension=".docx",
#                                                     filetypes=[("Word文档", "*.docx")])
#             if filename:
#                 output_path_entry.delete(0, tk.END)
#                 output_path_entry.insert(0, filename)
#         tk.Button(root, text="浏览...", command=browse_output_path).grid(row=3, column=2)
#
#         # 运行按钮
#         def run():
#             # 获取输入值
#             csv_path = map_data_path_entry.get()
#             start_coord_input = start_coord_entry.get()
#             end_coord_input = end_coord_entry.get()
#             output_path = output_path_entry.get()
#
#             # 验证输入
#             if not csv_path:
#                 messagebox.showerror("错误", "请提供地图数据文件路径。")
#                 return
#             if not output_path:
#                 messagebox.showerror("错误", "请提供输出 Word 文档路径。")
#                 return
#             if not start_coord_input:
#                 messagebox.showerror("错误", "请提供起点坐标。")
#                 return
#             if not end_coord_input:
#                 messagebox.showerror("错误", "请提供终点坐标。")
#                 return
#             try:
#                 start_x, start_y = map(float, start_coord_input.strip().split(','))
#                 start_coord = (start_x, start_y)
#             except ValueError:
#                 messagebox.showerror("错误", "起点坐标格式错误，请使用 'X,Y' 的格式。")
#                 return
#
#             try:
#                 end_x, end_y = map(float, end_coord_input.strip().split(','))
#                 end_coord = (end_x, end_y)
#             except ValueError:
#                 messagebox.showerror("错误", "终点坐标格式错误，请使用 'X,Y' 的格式。")
#                 return
#
#             # 在后台线程中运行，以避免阻塞GUI
#             def process():
#                 try:
#                     # 调用处理函数
#                     run_processing(csv_path, start_coord, end_coord, output_path)
#                     messagebox.showinfo("完成", "路径规划已完成。")
#                 except Exception as e:
#                     messagebox.showerror("错误", f"发生错误: {e}")
#                     logging.exception("后台线程运行时发生错误")
#
#             threading.Thread(target=process).start()
#
#         tk.Button(root, text="运行", command=run).grid(row=4, column=1)
#
#         root.mainloop()
#     except Exception as e:
#         print(f"主程序运行时发生错误: {e}")
#         logging.exception("主程序运行时发生错误")
#
# if __name__ == "__main__":
#     main()











# import pandas as pd
# import numpy as np
# from scipy.interpolate import griddata
# from sklearn.metrics import mean_squared_error, mean_absolute_error, accuracy_score
# from sklearn.model_selection import train_test_split
# import heapq
# import math
# import plotly.graph_objs as go
# import plotly.io as pio
# from docx import Document
# import tkinter as tk
# from tkinter import filedialog, messagebox
# import threading
# import sys
# import os
# import logging
#
# # 设置日志记录
# logging.basicConfig(filename='error.log', level=logging.ERROR)
#
# # 设置 Plotly 渲染器为浏览器
# pio.renderers.default = 'browser'
#
# # 添加 resource_path 函数，用于获取资源文件的正确路径
# def resource_path(relative_path):
#     """获取资源文件的绝对路径，兼容 PyInstaller 打包后的情况"""
#     try:
#         # PyInstaller 会把路径存储在 _MEIPASS 中
#         base_path = sys._MEIPASS
#     except AttributeError:
#         base_path = os.path.abspath(".")
#     return os.path.join(base_path, relative_path)
#
# # 数据处理类
# class DataProcessor:
#     def __init__(self, dat_path, x_resolution=1, y_resolution=1):
#         self.dat_path = dat_path
#         self.x_resolution = x_resolution
#         self.y_resolution = y_resolution
#         self.df = None
#         self.X = None
#         self.Y = None
#         self.Z = None
#         self.TrafFlag = None
#         self.maxAcc = None
#         self.Slope = None
#         self.StepElev = None
#         self.Roughness = None
#         self.x_range = None
#         self.y_range = None
#         self.X_grid = None
#         self.Y_grid = None
#         self.Z_grid = None
#         self.TrafFlag_grid = None
#         self.Acc_grid = None
#         self.Slope_grid = None
#         self.StepElev_grid = None
#         self.Rough_grid = None
#
#     def load_data(self):
#         try:
#             # 使用 resource_path 函数获取数据文件路径
#             data_path = resource_path(self.dat_path)
#             # 读取 .csv 文件，指定分隔符为逗号
#             self.df = pd.read_csv(data_path, delimiter=',')
#             print("数据文件中的列名：", self.df.columns.tolist())
#         except FileNotFoundError:
#             print(f"文件未找到: {self.dat_path}")
#             return
#         except pd.errors.EmptyDataError:
#             print("CSV 文件为空")
#             return
#         except Exception as e:
#             print(f"读取 CSV 文件时出错: {e}")
#             logging.exception("读取 CSV 文件时出错")
#             return
#
#         required_columns = {'X', 'Y', 'Z', 'TrafFlag', 'maxAcc', 'Slope', 'StepElevationDifference', 'Roughness'}
#         if not required_columns.issubset(self.df.columns):
#             print(f"CSV 文件缺少必要的列: {required_columns - set(self.df.columns)}")
#             return
#
#         # 提取所需数据
#         self.X = self.df['X'].values
#         self.Y = self.df['Y'].values
#         self.Z = self.df['Z'].values
#         self.TrafFlag = self.df['TrafFlag'].values
#         self.maxAcc = self.df['maxAcc'].values
#         self.Slope = self.df['Slope'].values
#         self.StepElev = self.df['StepElevationDifference'].values
#         self.Roughness = self.df['Roughness'].values
#
#         # 计算并打印 Slope 和 StepElevationDifference 的最小值和最大值
#         slope_min = np.min(self.Slope)
#         slope_max_val = np.max(self.Slope)
#         step_min = np.min(self.StepElev)
#         step_max = np.max(self.StepElev)
#
#         print(f"坡度 (Slope) 最小值: {slope_min}")
#         print(f"坡度 (Slope) 最大值: {slope_max_val}")
#         print(f"阶跃高程差 (StepElevationDifference) 最小值: {step_min}")
#         print(f"阶跃高程差 (StepElevationDifference) 最大值: {step_max}")
#
#     def interpolate_data(self):
#         # 计算网格范围和生成网格
#         x_min, x_max = self.X.min(), self.X.max()
#         y_min, y_max = self.Y.min(), self.Y.max()
#         self.x_range = np.arange(x_min, x_max + self.x_resolution, self.x_resolution)
#         self.y_range = np.arange(y_min, y_max + self.y_resolution, self.y_resolution)
#         self.X_grid, self.Y_grid = np.meshgrid(self.x_range, self.y_range)
#
#         # 插值
#         self.Z_grid = griddata((self.X, self.Y), self.Z, (self.X_grid, self.Y_grid), method='cubic')
#         self.TrafFlag_grid = griddata((self.X, self.Y), self.TrafFlag, (self.X_grid, self.Y_grid), method='nearest')
#         self.Acc_grid = griddata((self.X, self.Y), self.maxAcc, (self.X_grid, self.Y_grid), method='nearest')
#         self.Slope_grid = griddata((self.X, self.Y), self.Slope, (self.X_grid, self.Y_grid), method='nearest')
#         self.StepElev_grid = griddata((self.X, self.Y), self.StepElev, (self.X_grid, self.Y_grid), method='nearest')
#         self.Rough_grid = griddata((self.X, self.Y), self.Roughness, (self.X_grid, self.Y_grid), method='nearest')
#
#     def fill_nan_values(self):
#         # 填充 NaN 值
#         self.Z_grid = np.nan_to_num(self.Z_grid, nan=np.nanmean(self.Z))
#         self.Acc_grid = np.nan_to_num(self.Acc_grid, nan=0)
#         self.TrafFlag_grid = np.nan_to_num(self.TrafFlag_grid, nan=0)
#         self.Slope_grid = np.nan_to_num(self.Slope_grid, nan=0)
#         self.StepElev_grid = np.nan_to_num(self.StepElev_grid, nan=0)
#         self.Rough_grid = np.nan_to_num(self.Rough_grid, nan=0)
#
#         # 验证网格数据的形状
#         print(f"Z_grid shape: {self.Z_grid.shape}")
#         print(f"TrafFlag_grid shape: {self.TrafFlag_grid.shape}")
#         print(f"Acc_grid shape: {self.Acc_grid.shape}")
#         print(f"Slope_grid shape: {self.Slope_grid.shape}")
#         print(f"StepElev_grid shape: {self.StepElev_grid.shape}")
#         print(f"Rough_grid shape: {self.Rough_grid.shape}")
#
#         # 计算并打印插值后网格数据的最小值和最大值
#         slope_grid_min = np.min(self.Slope_grid)
#         slope_grid_max = np.max(self.Slope_grid)
#         step_grid_min = np.min(self.StepElev_grid)
#         step_grid_max = np.max(self.StepElev_grid)
#
#         print(f"插值后坡度网格 (Slope_grid) 最小值: {slope_grid_min}")
#         print(f"插值后坡度网格 (Slope_grid) 最大值: {slope_grid_max}")
#         print(f"插值后阶跃高程差网格 (StepElev_grid) 最小值: {step_grid_min}")
#         print(f"插值后阶跃高程差网格 (StepElev_grid) 最大值: {step_grid_max}")
#
#     def coord_to_index(self, coord):
#         x_min, y_min = self.x_range.min(), self.y_range.min()
#         col = int((coord[0] - x_min) / self.x_resolution)
#         row = int((coord[1] - y_min) / self.y_resolution)
#         return (row, col)
#
#     def is_valid_position(self, pos):
#         return 0 <= pos[0] < self.TrafFlag_grid.shape[0] and 0 <= pos[1] < self.TrafFlag_grid.shape[1]
#
#     def split_data_and_evaluate(self):
#         try:
#             # 将所有数据合并
#             data = np.column_stack(
#                 (self.X, self.Y, self.Z, self.TrafFlag, self.maxAcc, self.Slope, self.StepElev, self.Roughness))
#
#             # 分割数据集为训练集和测试集
#             test_size = 0.2
#             random_state = 42
#
#             train_data, test_data = train_test_split(data, test_size=test_size, random_state=random_state)
#
#             # 拆分训练集和测试集的数据
#             X_train, Y_train = train_data[:, 0], train_data[:, 1]
#             Z_train = train_data[:, 2]
#             TrafFlag_train = train_data[:, 3]
#             maxAcc_train = train_data[:, 4]
#             Slope_train = train_data[:, 5]
#             StepElev_train = train_data[:, 6]
#             Rough_train = train_data[:, 7]
#
#             X_test, Y_test = test_data[:, 0], test_data[:, 1]
#             Z_test = test_data[:, 2]
#             TrafFlag_test = test_data[:, 3]
#             maxAcc_test = test_data[:, 4]
#             Slope_test = test_data[:, 5]
#             StepElev_test = test_data[:, 6]
#             Rough_test = test_data[:, 7]
#
#             # 使用训练集进行插值
#             Z_grid_train = griddata((X_train, Y_train), Z_train, (self.X_grid, self.Y_grid), method='cubic')
#             TrafFlag_grid_train = griddata((X_train, Y_train), TrafFlag_train, (self.X_grid, self.Y_grid), method='nearest')
#             Acc_grid_train = griddata((X_train, Y_train), maxAcc_train, (self.X_grid, self.Y_grid), method='nearest')
#             Slope_grid_train = griddata((X_train, Y_train), Slope_train, (self.X_grid, self.Y_grid), method='nearest')
#             StepElev_grid_train = griddata((X_train, Y_train), StepElev_train, (self.X_grid, self.Y_grid), method='nearest')
#             Rough_grid_train = griddata((X_train, Y_train), Rough_train, (self.X_grid, self.Y_grid), method='nearest')
#
#             # 在测试集上预测值
#             Z_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Z_grid_train.flatten(), (X_test, Y_test),
#                               method='cubic')
#             TrafFlag_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), TrafFlag_grid_train.flatten(),
#                                      (X_test, Y_test), method='nearest')
#             Acc_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Acc_grid_train.flatten(), (X_test, Y_test),
#                                 method='nearest')
#             Slope_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Slope_grid_train.flatten(),
#                                   (X_test, Y_test), method='nearest')
#             StepElev_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), StepElev_grid_train.flatten(),
#                                      (X_test, Y_test), method='nearest')
#             Rough_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Rough_grid_train.flatten(),
#                                   (X_test, Y_test), method='nearest')
#
#             # 处理 NaN 值并计算误差指标
#
#             # 高程数据误差
#             valid_indices_Z = ~np.isnan(Z_pred) & ~np.isnan(Z_test)
#             if np.any(valid_indices_Z):
#                 mse_Z = mean_squared_error(Z_test[valid_indices_Z], Z_pred[valid_indices_Z])
#                 mae_Z = mean_absolute_error(Z_test[valid_indices_Z], Z_pred[valid_indices_Z])
#                 max_error_Z = np.max(np.abs(Z_test[valid_indices_Z] - Z_pred[valid_indices_Z]))
#                 print(f"高程插值 - MSE: {mse_Z}, MAE: {mae_Z}, Max Error: {max_error_Z}")
#             else:
#                 print("高程插值预测值中全部为 NaN，无法计算误差。")
#
#             # 通行标志（分类准确率）
#             valid_indices_TrafFlag = ~np.isnan(TrafFlag_pred) & ~np.isnan(TrafFlag_test)
#             if np.any(valid_indices_TrafFlag):
#                 accuracy_TrafFlag = accuracy_score(TrafFlag_test[valid_indices_TrafFlag],
#                                                    TrafFlag_pred[valid_indices_TrafFlag])
#                 print(f"通行标志插值 - 准确率: {accuracy_TrafFlag}")
#             else:
#                 print("通行标志插值预测值中全部为 NaN，无法计算准确率。")
#
#             # 最大加速度误差
#             valid_indices_Acc = ~np.isnan(Acc_pred) & ~np.isnan(maxAcc_test)
#             if np.any(valid_indices_Acc):
#                 mse_Acc = mean_squared_error(maxAcc_test[valid_indices_Acc], Acc_pred[valid_indices_Acc])
#                 mae_Acc = mean_absolute_error(maxAcc_test[valid_indices_Acc], Acc_pred[valid_indices_Acc])
#                 max_error_Acc = np.max(np.abs(maxAcc_test[valid_indices_Acc] - Acc_pred[valid_indices_Acc]))
#                 print(f"最大加速度插值 - MSE: {mse_Acc}, MAE: {mae_Acc}, Max Error: {max_error_Acc}")
#             else:
#                 print("最大加速度插值预测值中全部为 NaN，无法计算误差。")
#
#             # 坡度误差
#             valid_indices_Slope = ~np.isnan(Slope_pred) & ~np.isnan(Slope_test)
#             if np.any(valid_indices_Slope):
#                 mse_Slope = mean_squared_error(Slope_test[valid_indices_Slope], Slope_pred[valid_indices_Slope])
#                 mae_Slope = mean_absolute_error(Slope_test[valid_indices_Slope], Slope_pred[valid_indices_Slope])
#                 max_error_Slope = np.max(np.abs(Slope_test[valid_indices_Slope] - Slope_pred[valid_indices_Slope]))
#                 print(f"坡度插值 - MSE: {mse_Slope}, MAE: {mae_Slope}, Max Error: {max_error_Slope}")
#             else:
#                 print("坡度插值预测值中全部为 NaN，无法计算误差。")
#
#             # 阶跃高程差误差
#             valid_indices_StepElev = ~np.isnan(StepElev_pred) & ~np.isnan(StepElev_test)
#             if np.any(valid_indices_StepElev):
#                 mse_StepElev = mean_squared_error(StepElev_test[valid_indices_StepElev],
#                                                   StepElev_pred[valid_indices_StepElev])
#                 mae_StepElev = mean_absolute_error(StepElev_test[valid_indices_StepElev],
#                                                    StepElev_pred[valid_indices_StepElev])
#                 max_error_StepElev = np.max(
#                     np.abs(StepElev_test[valid_indices_StepElev] - StepElev_pred[valid_indices_StepElev]))
#                 print(f"阶跃高程差插值 - MSE: {mse_StepElev}, MAE: {mae_StepElev}, Max Error: {max_error_StepElev}")
#             else:
#                 print("阶跃高程差插值预测值中全部为 NaN，无法计算误差。")
#
#             # 粗糙度误差
#             valid_indices_Rough = ~np.isnan(Rough_pred) & ~np.isnan(Rough_test)
#             if np.any(valid_indices_Rough):
#                 mse_Rough = mean_squared_error(Rough_test[valid_indices_Rough], Rough_pred[valid_indices_Rough])
#                 mae_Rough = mean_absolute_error(Rough_test[valid_indices_Rough], Rough_pred[valid_indices_Rough])
#                 max_error_Rough = np.max(np.abs(Rough_test[valid_indices_Rough] - Rough_pred[valid_indices_Rough]))
#                 print(f"粗糙度插值 - MSE: {mse_Rough}, MAE: {mae_Rough}, Max Error: {max_error_Rough}")
#             else:
#                 print("粗糙度插值预测值中全部为 NaN，无法计算误差。")
#         except Exception as e:
#             print(f"数据评估时发生错误: {e}")
#             logging.exception("数据评估时发生错误")
#
# # 节点类和路径规划类
# class Node:
#     def __init__(self, position, parent=None):
#         self.position = position  # (row, col)
#         self.parent = parent
#         self.g = 0  # 累计代价
#         self.h = 0  # 启发代价
#         self.f = 0  # 总代价
#
#     def __lt__(self, other):
#         return self.f < other.f
#
# class PathPlanner:
#     def __init__(self, grid_data, max_slope=None, max_step_elev=None):
#         self.grid = grid_data['TrafFlag_grid']
#         self.Z_grid = grid_data['Z_grid']
#         self.Acc_grid = grid_data['Acc_grid']
#         self.Slope_grid = grid_data['Slope_grid']
#         self.StepElev_grid = grid_data['StepElev_grid']
#         self.Rough_grid = grid_data['Rough_grid']
#         self.max_slope = max_slope
#         self.max_step_elev = max_step_elev
#
#     def heuristic_distance_3d(self, current_pos, end_pos):
#         d_xy = np.sqrt((current_pos[0] - end_pos[0]) ** 2 +
#                        (current_pos[1] - end_pos[1]) ** 2)
#         z_current = self.Z_grid[current_pos[0], current_pos[1]]
#         z_end = self.Z_grid[end_pos[0], end_pos[1]]
#         d_z = (z_end - z_current) ** 2
#         return np.sqrt(d_xy ** 2 + d_z)
#
#     def astar(self, start, end, mode='distance', alpha_env=0.6, beta_env=0.4):
#         start_node = Node(start)
#         end_node = Node(end)
#
#         open_list = []
#         heapq.heappush(open_list, start_node)
#         open_dict = {start_node.position: start_node}
#
#         closed_set = set()
#
#         while open_list:
#             current_node = heapq.heappop(open_list)
#             del open_dict[current_node.position]
#             closed_set.add(current_node.position)
#
#             # 检查是否到达终点
#             if current_node.position == end_node.position:
#                 path = []
#                 while current_node:
#                     path.append(current_node.position)
#                     current_node = current_node.parent
#                 return path[::-1]
#
#             # 定义8个邻居（包括对角线）
#             neighbors = [(0, 1), (0, -1), (1, 0), (-1, 0),
#                          (1, 1), (1, -1), (-1, 1), (-1, -1)]
#
#             for offset in neighbors:
#                 neighbor_pos = (current_node.position[0] + offset[0],
#                                 current_node.position[1] + offset[1])
#
#                 # 检查邻居是否在网格范围内
#                 if not (0 <= neighbor_pos[0] < self.grid.shape[0] and
#                         0 <= neighbor_pos[1] < self.grid.shape[1]):
#                     continue
#
#                 # 检查 TrafFlag 是否可通行
#                 if self.grid[neighbor_pos] == 0:
#                     continue  # 不可通行区域
#
#                 # 在所有模式下应用坡度和阶跃高程差的限制
#                 if self.max_slope is not None and self.Slope_grid[neighbor_pos] > self.max_slope:
#                     continue  # 坡度超过最大限制
#                 if self.max_step_elev is not None and self.StepElev_grid[neighbor_pos] > self.max_step_elev:
#                     continue  # 阶跃高程差超过最大限制
#
#                 # 检查是否已经在关闭列表中
#                 if neighbor_pos in closed_set:
#                     continue
#
#                 # 处理不同模式
#                 if mode == 'distance':
#                     movement_cost = self.heuristic_distance_3d(current_node.position, neighbor_pos)
#                     g_cost = current_node.g + movement_cost
#                     h_cost = self.heuristic_distance_3d(neighbor_pos, end)
#                 elif mode == 'time':
#                     a_max = self.Acc_grid[neighbor_pos[0], neighbor_pos[1]]
#                     if a_max <= 0:
#                         continue  # 无法加速，跳过
#                     movement_cost = self.heuristic_distance_3d(current_node.position, neighbor_pos)
#                     g_cost = current_node.g + movement_cost / a_max  # 简化为距离除以加速度
#                     h_cost = self.heuristic_distance_3d(neighbor_pos, end)
#                 elif mode == 'env':
#                     movement_uncertainty = (self.Slope_grid[neighbor_pos[0], neighbor_pos[1]] / 45.0 +
#                                             self.StepElev_grid[neighbor_pos[0], neighbor_pos[1]] / 1.5 +
#                                             self.Rough_grid[neighbor_pos[0], neighbor_pos[1]] / 6.0)
#                     g_cost = current_node.g + movement_uncertainty
#                     h_cost = self.heuristic_distance_3d(neighbor_pos, end)
#                 else:
#                     raise ValueError("Invalid mode. Choose from 'distance', 'time', 'env'.")
#
#                 # 创建邻居节点
#                 neighbor_node = Node(neighbor_pos, current_node)
#                 neighbor_node.g = g_cost
#                 neighbor_node.h = h_cost
#                 neighbor_node.f = neighbor_node.g + neighbor_node.h
#
#                 # 检查开放列表中是否有更好的节点
#                 if neighbor_pos in open_dict:
#                     existing_node = open_dict[neighbor_pos]
#                     if existing_node.g <= neighbor_node.g:
#                         continue  # 已有更优节点
#                     else:
#                         # 替换为更优节点
#                         open_list.remove(existing_node)
#                         heapq.heapify(open_list)
#                         heapq.heappush(open_list, neighbor_node)
#                         open_dict[neighbor_pos] = neighbor_node
#                 else:
#                     heapq.heappush(open_list, neighbor_node)
#                     open_dict[neighbor_pos] = neighbor_node
#
#         return None  # 如果找不到路径，返回 None
#
#     def extract_path_coordinates(self, path, x_range, y_range):
#         return [(x_range[p[1]], y_range[p[0]], self.Z_grid[p[0], p[1]]) for p in path]
#
#     def calculate_path_length(self, path_coords):
#         length = 0
#         for i in range(1, len(path_coords)):
#             x1, y1, z1 = path_coords[i - 1]
#             x2, y2, z2 = path_coords[i]
#             length += np.sqrt((x2 - x1) ** 2 + (y2 - y1) ** 2 + (z2 - z1) ** 2)
#         return length
#
#     def compute_traversal_time(self, path_coords, y_range, x_range):
#         total_time = 0.0
#         current_speed = 0.0  # 初始速度为0
#         max_speed = 70.0  # 最大速度70 m/s
#
#         speeds = [current_speed]  # 存储每个点的速度
#         accelerations = []  # 存储每段的加速度
#
#         for i in range(1, len(path_coords)):
#             x1, y1, z1 = path_coords[i - 1]
#             x2, y2, z2 = path_coords[i]
#             # 计算两点之间的距离
#             d = np.sqrt((x2 - x1) ** 2 + (y2 - y1) ** 2 + (z2 - z1) ** 2)
#
#             # 获取当前点的网格索引
#             row = np.argmin(np.abs(y_range - y1))
#             col = np.argmin(np.abs(x_range - x1))
#
#             # 获取当前点的加速度
#             a = self.Acc_grid[row, col]
#             if a <= 0:
#                 # 如果加速度不可用或为0，无法继续加速
#                 a = 0.0
#
#             # 计算需要加速到最大速度的速度增量
#             delta_v = max_speed - current_speed
#
#             if current_speed < max_speed and a > 0:
#                 # 需要加速
#                 t_acc = delta_v / a  # 加速所需时间
#                 d_acc = current_speed * t_acc + 0.5 * a * t_acc ** 2  # 加速所需距离
#
#                 if d_acc >= d:
#                     # 在当前段无法加速到最大速度
#                     # 使用运动学方程计算通过距离d所需的时间
#                     # 解方程: 0.5 * a * t^2 + current_speed * t - d = 0
#                     discriminant = current_speed ** 2 + 2 * a * d
#                     if discriminant < 0:
#                         # 无实数解，无法到达下一个点
#                         return float('inf'), [], []
#                     t = (-current_speed + math.sqrt(discriminant)) / a
#                     total_time += t
#                     current_speed += a * t
#                     accelerations.append(a)
#                     speeds.append(current_speed)
#                 else:
#                     # 在当前段可以加速到最大速度
#                     # 先加速到最大速度
#                     total_time += t_acc
#                     current_speed = max_speed
#                     accelerations.append(a)
#                     speeds.append(current_speed)
#                     # 剩余距离
#                     d_remain = d - d_acc
#                     # 以最大速度匀速行驶剩余距离
#                     t_remain = d_remain / current_speed
#                     total_time += t_remain
#                     accelerations.append(0.0)
#                     speeds.append(current_speed)
#             else:
#                 # 无法加速或已达到最大速度，匀速行驶
#                 if current_speed == 0:
#                     # 无法移动
#                     return float('inf'), [], []
#                 t = d / current_speed
#                 total_time += t
#                 accelerations.append(0.0)
#                 speeds.append(current_speed)
#
#         return total_time, speeds, accelerations
#
# # 可视化类
# class Visualizer:
#     def __init__(self, X_grid, Y_grid, Z_grid, TrafFlag_grid):
#         self.X_grid = X_grid
#         self.Y_grid = Y_grid
#         self.Z_grid = Z_grid
#         self.TrafFlag_grid = TrafFlag_grid
#         self.fig = go.Figure()
#
#     def plot_terrain(self):
#         self.fig.add_trace(go.Surface(
#             x=self.X_grid, y=self.Y_grid, z=self.Z_grid,
#             surfacecolor=self.TrafFlag_grid,
#             colorscale='Viridis',
#             opacity=0.8,
#             showscale=True,
#             colorbar=dict(title='TrafFlag')
#         ))
#
#     def add_path(self, coords, speeds, accs, color, name, length, time_taken, path_width=4):
#         if coords:
#             x, y, z = zip(*coords)
#             # 处理不可行路径（时间为inf）的情况
#             if math.isinf(time_taken):
#                 time_str = "无法到达"
#             else:
#                 time_str = f"{time_taken:.2f} 秒"
#
#             # 准备 customdata，将速度和加速度添加进去
#             if accs:
#                 accs = [0.0] + accs  # 与 speeds 对齐
#             else:
#                 accs = [0.0] * len(speeds)
#
#             customdata = np.array([speeds, accs]).T  # 转置使其每行对应一个点的 [speed, acceleration]
#
#             self.fig.add_trace(go.Scatter3d(
#                 x=x, y=y, z=z,
#                 mode='lines+markers',
#                 line=dict(color=color, width=path_width),
#                 marker=dict(size=2, color=color),
#                 name=f"{name} (长度: {length:.2f} 米, 时间: {time_str})",
#                 customdata=customdata,
#                 hovertemplate=(
#                         'X: %{x}<br>' +
#                         'Y: %{y}<br>' +
#                         'Z: %{z}<br>' +
#                         '速度: %{customdata[0]:.2f} m/s<br>' +
#                         '加速度: %{customdata[1]:.2f} m/s²<br>' +
#                         '<extra></extra>'
#                 )
#             ))
#         else:
#             print(f"未找到{name}")
#
#     def add_start_and_end_points(self, start_coord, end_coord):
#         self.fig.add_trace(go.Scatter3d(
#             x=[start_coord[0]],
#             y=[start_coord[1]],
#             z=[start_coord[2]],
#             mode='markers',
#             marker=dict(size=8, color='green'),
#             name='起点'
#         ))
#         self.fig.add_trace(go.Scatter3d(
#             x=[end_coord[0]],
#             y=[end_coord[1]],
#             z=[end_coord[2]],
#             mode='markers',
#             marker=dict(size=8, color='red'),
#             name='终点'
#         ))
#
#     def show(self):
#         self.fig.update_layout(
#             scene=dict(
#                 xaxis_title='X 坐标',
#                 yaxis_title='Y 坐标',
#                 zaxis_title='Z 坐标 (高程)',
#                 aspectmode='data'
#             ),
#             title='3D 地形路径规划（基于距离、时间与最小不确定性）',
#             legend=dict(
#                 x=0,
#                 y=1,
#                 bgcolor='rgba(255, 255, 255, 0)',
#                 bordercolor='rgba(255, 255, 255, 0)'
#             )
#         )
#         pio.show(self.fig)
#
# # 报告生成函数
# def write_paths_to_word(document, path_name, path_coords):
#     if not path_coords:
#         return
#     # 添加路径名称和点数量
#     document.add_heading(f"{path_name} (点数: {len(path_coords)})", level=2)
#     # 创建表格，包含三列：X相对坐标、Y相对坐标、Z高程坐标
#     table = document.add_table(rows=1, cols=3)
#     table.style = 'Light Grid'
#     hdr_cells = table.rows[0].cells
#     hdr_cells[0].text = 'X相对坐标'
#     hdr_cells[1].text = 'Y相对坐标'
#     hdr_cells[2].text = 'Z高程坐标'
#     # 填充表格内容
#     for coord in path_coords:
#         row_cells = table.add_row().cells
#         row_cells[0].text = f"{coord[0]}"
#         row_cells[1].text = f"{coord[1]}"
#         row_cells[2].text = f"{coord[2]}"
#     # 添加一个空行
#     document.add_paragraph()
#
# # 处理函数
# def run_processing(csv_path, start_coord, end_coord, output_path, max_allowed_slope, max_allowed_step_elev):
#     try:
#         # 设置路径宽度
#         path_width = 4  # 可以根据需要调整路径宽度
#
#         # 数据处理
#         data_processor = DataProcessor(csv_path)
#         data_processor.load_data()
#         data_processor.interpolate_data()
#         data_processor.fill_nan_values()
#         data_processor.split_data_and_evaluate()
#
#         # 将 start_coord 和 end_coord 转换为网格索引
#         start = data_processor.coord_to_index(start_coord)
#         end = data_processor.coord_to_index(end_coord)
#
#         # 检查起点和终点是否在网格范围内
#         if not data_processor.is_valid_position(start):
#             raise ValueError("起点坐标超出网格范围")
#         if not data_processor.is_valid_position(end):
#             raise ValueError("终点坐标超出网格范围")
#
#         # 确保起点和终点可通行
#         if data_processor.TrafFlag_grid[start] == 0:
#             raise ValueError("起点不可通行")
#         if data_processor.TrafFlag_grid[end] == 0:
#             raise ValueError("终点不可通行")
#
#         # 路径规划
#         grid_data = {
#             'TrafFlag_grid': data_processor.TrafFlag_grid,
#             'Z_grid': data_processor.Z_grid,
#             'Acc_grid': data_processor.Acc_grid,
#             'Slope_grid': data_processor.Slope_grid,
#             'StepElev_grid': data_processor.StepElev_grid,
#             'Rough_grid': data_processor.Rough_grid
#         }
#         path_planner = PathPlanner(grid_data, max_slope=max_allowed_slope, max_step_elev=max_allowed_step_elev)
#
#         print("正在计算基于距离的最短路径...")
#         path_distance = path_planner.astar(start, end, mode='distance')
#         if path_distance:
#             print("基于距离的路径计算完成")
#         else:
#             print("未找到基于距离的路径")
#
#         print("正在计算基于时间的最短耗时路径...")
#         path_time = path_planner.astar(start, end, mode='time')
#         if path_time:
#             print("基于时间的路径计算完成")
#         else:
#             print("未找到基于时间的路径")
#
#         print("正在计算最小不确定性路径...")
#         path_env = path_planner.astar(start, end, mode='env')
#         if path_env:
#             print("最小不确定性路径计算完成")
#         else:
#             print("未找到最小不确定性路径")
#
#         # 提取路径坐标
#         if path_distance:
#             path_distance_coords = path_planner.extract_path_coordinates(path_distance, data_processor.x_range,
#                                                                          data_processor.y_range)
#             length_distance = path_planner.calculate_path_length(path_distance_coords)
#         else:
#             path_distance_coords = None
#             length_distance = None
#         if path_time:
#             path_time_coords = path_planner.extract_path_coordinates(path_time, data_processor.x_range,
#                                                                      data_processor.y_range)
#             length_time = path_planner.calculate_path_length(path_time_coords)
#         else:
#             path_time_coords = None
#             length_time = None
#         if path_env:
#             path_env_coords = path_planner.extract_path_coordinates(path_env, data_processor.x_range,
#                                                                     data_processor.y_range)
#             length_env = path_planner.calculate_path_length(path_env_coords)
#         else:
#             path_env_coords = None
#             length_env = None
#
#         # 计算遍历时间，并获取速度和加速度数据
#         if path_distance:
#             time_distance, speeds_distance, accs_distance = path_planner.compute_traversal_time(
#                 path_distance_coords, data_processor.y_range, data_processor.x_range)
#         else:
#             speeds_distance, accs_distance, time_distance = None, None, None
#
#         if path_time:
#             time_time, speeds_time, accs_time = path_planner.compute_traversal_time(
#                 path_time_coords, data_processor.y_range, data_processor.x_range)
#         else:
#             speeds_time, accs_time, time_time = None, None, None
#
#         if path_env:
#             time_env, speeds_env, accs_env = path_planner.compute_traversal_time(
#                 path_env_coords, data_processor.y_range, data_processor.x_range)
#         else:
#             speeds_env, accs_env, time_env = None, None, None
#
#         # 可视化
#         visualizer = Visualizer(
#             data_processor.X_grid,
#             data_processor.Y_grid,
#             data_processor.Z_grid,
#             data_processor.TrafFlag_grid
#         )
#         visualizer.plot_terrain()
#
#         colors = ['blue', 'orange', 'purple']
#         names = ['MinDistPath 最短距离路径', 'MinTimePath 最短耗时路径', 'MinUncertPath 最小不确定性路径']
#         lengths = [length_distance, length_time, length_env]
#         times = [time_distance, time_time, time_env]
#         coords_list = [path_distance_coords, path_time_coords, path_env_coords]
#         speeds_list = [speeds_distance, speeds_time, speeds_env]
#         accs_list = [accs_distance, accs_time, accs_env]
#
#         for coords, speeds, accs, color, name, length, time_taken in zip(
#                 coords_list, speeds_list, accs_list,
#                 colors, names, lengths, times):
#             visualizer.add_path(coords, speeds, accs, color, name, length, time_taken, path_width)
#
#         # 标记起点和终点
#         if path_distance or path_time or path_env:
#             selected_path = None
#             if path_distance:
#                 selected_path = path_distance_coords
#             elif path_time:
#                 selected_path = path_time_coords
#             elif path_env:
#                 selected_path = path_env_coords
#
#             if selected_path:
#                 # 使用起点坐标
#                 start_x, start_y, start_z = selected_path[0]
#                 # 使用终点坐标
#                 end_x, end_y, end_z = selected_path[-1]
#                 visualizer.add_start_and_end_points((start_x, start_y, start_z), (end_x, end_y, end_z))
#
#         visualizer.show()
#
#         # 创建 Word 文档
#         document = Document()
#         document.add_heading('路径规划结果', 0)
#         # 写入路径信息
#         if path_distance or path_time or path_env:
#             if path_distance:
#                 write_paths_to_word(document, 'MinDistPath ', path_distance_coords)
#             if path_time:
#                 write_paths_to_word(document, 'MinTimePath ', path_time_coords)
#             if path_env:
#                 write_paths_to_word(document, 'MinUncertPath ', path_env_coords)
#
#             # 保存文档
#             try:
#                 # 使用 resource_path 函数获取输出文件路径
#                 output_file_path = output_path  # 如果 output_path 是用户指定的完整路径，则无需使用 resource_path
#                 document.save(output_file_path)
#                 print(f"路径坐标已成功保存到 Word 文档: {output_file_path}")
#             except Exception as e:
#                 print(f"保存 Word 文档时出错: {e}")
#                 logging.exception("保存 Word 文档时出错")
#         else:
#             print("没有路径可写入 Word 文档")
#
#         # 打印路径长度和遍历时间
#         if path_distance:
#             if math.isinf(time_distance):
#                 print(f"最短路径长度: {length_distance:.2f} 米")
#                 print(f"最短距离路径遍历时间: 无法到达")
#             else:
#                 print(f"最短路径长度: {length_distance:.2f} 米")
#                 print(f"最短距离路径遍历时间: {time_distance:.2f} 秒")
#         if path_time:
#             if math.isinf(time_time):
#                 print(f"最短耗时路径长度: {length_time:.2f} 米")
#                 print(f"最短耗时路径遍历时间: 无法到达")
#             else:
#                 print(f"最短耗时路径长度: {length_time:.2f} 米")
#                 print(f"最短耗时路径遍历时间: {time_time:.2f} 秒")
#         if path_env:
#             if math.isinf(time_env):
#                 print(f"最小不确定性路径长度: {length_env:.2f} 米")
#                 print(f"最小不确定性路径遍历时间: 无法到达")
#             else:
#                 print(f"最小不确定性路径长度: {length_env:.2f} 米")
#                 print(f"最小不确定性路径遍历时间: {time_env:.2f} 秒")
#     except Exception as e:
#         print(f"运行处理时发生错误: {e}")
#         logging.exception("运行处理时发生错误")
#
# # 主程序
# def main():
#     try:
#         # 创建主窗口
#         root = tk.Tk()
#         root.title("路径规划软件")
#
#         # 地图数据文件路径
#         tk.Label(root, text="地图数据文件路径:").grid(row=0, column=0, sticky='e')
#         map_data_path_entry = tk.Entry(root, width=50)
#         map_data_path_entry.grid(row=0, column=1)
#         def browse_map_data():
#             filename = filedialog.askopenfilename(filetypes=[("CSV文件", "*.csv")])
#             if filename:
#                 map_data_path_entry.delete(0, tk.END)
#                 map_data_path_entry.insert(0, filename)
#         tk.Button(root, text="浏览...", command=browse_map_data).grid(row=0, column=2)
#
#         # 起点坐标
#         tk.Label(root, text="起点坐标 (格式: X,Y):").grid(row=1, column=0, sticky='e')
#         start_coord_entry = tk.Entry(root, width=50)
#         start_coord_entry.grid(row=1, column=1)
#
#         # 终点坐标
#         tk.Label(root, text="终点坐标 (格式: X,Y):").grid(row=2, column=0, sticky='e')
#         end_coord_entry = tk.Entry(root, width=50)
#         end_coord_entry.grid(row=2, column=1)
#
#         # 输出路径
#         tk.Label(root, text="输出 Word 文档路径:").grid(row=3, column=0, sticky='e')
#         output_path_entry = tk.Entry(root, width=50)
#         output_path_entry.grid(row=3, column=1)
#         def browse_output_path():
#             filename = filedialog.asksaveasfilename(defaultextension=".docx",
#                                                     filetypes=[("Word文档", "*.docx")])
#             if filename:
#                 output_path_entry.delete(0, tk.END)
#                 output_path_entry.insert(0, filename)
#         tk.Button(root, text="浏览...", command=browse_output_path).grid(row=3, column=2)
#
#         # 最大坡度
#         tk.Label(root, text="最大坡度 (默认25.0度):").grid(row=4, column=0, sticky='e')
#         max_slope_entry = tk.Entry(root, width=50)
#         max_slope_entry.grid(row=4, column=1)
#         max_slope_entry.insert(0, "25.0")  # 设置默认值
#
#         # 最大阶跃高程差
#         tk.Label(root, text="最大阶跃高程差 (默认10.0米):").grid(row=5, column=0, sticky='e')
#         max_step_elev_entry = tk.Entry(root, width=50)
#         max_step_elev_entry.grid(row=5, column=1)
#         max_step_elev_entry.insert(0, "10.0")  # 设置默认值
#
#         # 运行按钮
#         def run():
#             # 获取输入值
#             csv_path = map_data_path_entry.get()
#             start_coord_input = start_coord_entry.get()
#             end_coord_input = end_coord_entry.get()
#             output_path = output_path_entry.get()
#
#             # 验证输入
#             if not csv_path:
#                 messagebox.showerror("错误", "请提供地图数据文件路径。")
#                 return
#             if not output_path:
#                 messagebox.showerror("错误", "请提供输出 Word 文档路径。")
#                 return
#             if not start_coord_input:
#                 messagebox.showerror("错误", "请提供起点坐标。")
#                 return
#             if not end_coord_input:
#                 messagebox.showerror("错误", "请提供终点坐标。")
#                 return
#             try:
#                 start_x, start_y = map(float, start_coord_input.strip().split(','))
#                 start_coord = (start_x, start_y)
#             except ValueError:
#                 messagebox.showerror("错误", "起点坐标格式错误，请使用 'X,Y' 的格式。")
#                 return
#
#             try:
#                 end_x, end_y = map(float, end_coord_input.strip().split(','))
#                 end_coord = (end_x, end_y)
#             except ValueError:
#                 messagebox.showerror("错误", "终点坐标格式错误，请使用 'X,Y' 的格式。")
#                 return
#
#             # 获取最大坡度和最大阶跃高程差
#             max_slope_input = max_slope_entry.get()
#             max_step_elev_input = max_step_elev_entry.get()
#
#             try:
#                 max_allowed_slope = float(max_slope_input)
#             except ValueError:
#                 messagebox.showerror("错误", "最大坡度格式错误，请输入数字。")
#                 return
#
#             try:
#                 max_allowed_step_elev = float(max_step_elev_input)
#             except ValueError:
#                 messagebox.showerror("错误", "最大阶跃高程差格式错误，请输入数字。")
#                 return
#
#             # 在后台线程中运行，以避免阻塞GUI
#             def process():
#                 try:
#                     # 调用处理函数，传入最大坡度和最大阶跃高程差
#                     run_processing(csv_path, start_coord, end_coord, output_path, max_allowed_slope, max_allowed_step_elev)
#                     messagebox.showinfo("完成", "路径规划已完成。")
#                 except Exception as e:
#                     messagebox.showerror("错误", f"发生错误: {e}")
#                     logging.exception("后台线程运行时发生错误")
#
#             threading.Thread(target=process).start()
#
#         tk.Button(root, text="运行", command=run).grid(row=6, column=1)
#
#         root.mainloop()
#     except Exception as e:
#         print(f"主程序运行时发生错误: {e}")
#         logging.exception("主程序运行时发生错误")
#
# if __name__ == "__main__":
#     main()









# import pandas as pd
# import numpy as np
# from scipy.interpolate import griddata
# from sklearn.metrics import mean_squared_error, mean_absolute_error, accuracy_score
# from sklearn.model_selection import train_test_split
# import heapq
# import math
# import plotly.graph_objs as go
# import plotly.io as pio
# from docx import Document
# import tkinter as tk
# from tkinter import filedialog, messagebox
# import threading
# import sys
# import os
# import logging
#
# # 设置日志记录
# logging.basicConfig(filename='error.log', level=logging.ERROR)
#
# # 设置 Plotly 渲染器为浏览器
# pio.renderers.default = 'browser'
#
# # 添加 resource_path 函数，用于获取资源文件的正确路径
# def resource_path(relative_path):
#     """获取资源文件的绝对路径，兼容 PyInstaller 打包后的情况"""
#     try:
#         # PyInstaller 会把路径存储在 _MEIPASS 中
#         base_path = sys._MEIPASS
#     except AttributeError:
#         base_path = os.path.abspath(".")
#     return os.path.join(base_path, relative_path)
#
# # 数据处理类
# class DataProcessor:
#     def __init__(self, dat_path, x_resolution=1, y_resolution=1):
#         self.dat_path = dat_path
#         self.x_resolution = x_resolution
#         self.y_resolution = y_resolution
#         self.df = None
#         self.X = None
#         self.Y = None
#         self.Z = None
#         self.TrafFlag = None
#         self.maxAcc = None
#         self.Slope = None
#         self.StepElev = None
#         self.Roughness = None
#         self.x_range = None
#         self.y_range = None
#         self.X_grid = None
#         self.Y_grid = None
#         self.Z_grid = None
#         self.TrafFlag_grid = None
#         self.Acc_grid = None
#         self.Slope_grid = None
#         self.StepElev_grid = None
#         self.Rough_grid = None
#
#     def load_data(self):
#         try:
#             # 使用 resource_path 函数获取数据文件路径
#             data_path = resource_path(self.dat_path)
#             # 读取 .csv 文件，指定分隔符为逗号
#             self.df = pd.read_csv(data_path, delimiter=',')
#             print("数据文件中的列名：", self.df.columns.tolist())
#         except FileNotFoundError:
#             print(f"文件未找到: {self.dat_path}")
#             return
#         except pd.errors.EmptyDataError:
#             print("CSV 文件为空")
#             return
#         except Exception as e:
#             print(f"读取 CSV 文件时出错: {e}")
#             logging.exception("读取 CSV 文件时出错")
#             return
#
#         required_columns = {'X', 'Y', 'Z', 'TrafFlag', 'maxAcc', 'Slope', 'StepElevationDifference', 'Roughness'}
#         if not required_columns.issubset(self.df.columns):
#             print(f"CSV文件缺少必要的列: {required_columns - set(self.df.columns)}")
#             return
#
#         # 提取所需数据
#         self.X = self.df['X'].values
#         self.Y = self.df['Y'].values
#         self.Z = self.df['Z'].values
#         self.TrafFlag = self.df['TrafFlag'].values
#         self.maxAcc = self.df['maxAcc'].values
#         self.Slope = self.df['Slope'].values
#         self.StepElev = self.df['StepElevationDifference'].values
#         self.Roughness = self.df['Roughness'].values
#
#         # 计算并打印 Slope 和 StepElevationDifference 的最小值和最大值
#         slope_min = np.min(self.Slope)
#         slope_max_val = np.max(self.Slope)
#         step_min = np.min(self.StepElev)
#         step_max = np.max(self.StepElev)
#
#         print(f"坡度 (Slope) 最小值: {slope_min}")
#         print(f"坡度 (Slope) 最大值: {slope_max_val}")
#         print(f"阶跃高程差 (StepElevationDifference) 最小值: {step_min}")
#         print(f"阶跃高程差 (StepElevationDifference) 最大值: {step_max}")
#
#     def interpolate_data(self):
#         # 计算网格范围和生成网格
#         x_min, x_max = self.X.min(), self.X.max()
#         y_min, y_max = self.Y.min(), self.Y.max()
#         self.x_range = np.arange(x_min, x_max + self.x_resolution, self.x_resolution)
#         self.y_range = np.arange(y_min, y_max + self.y_resolution, self.y_resolution)
#         self.X_grid, self.Y_grid = np.meshgrid(self.x_range, self.y_range)
#
#         # 插值
#         self.Z_grid = griddata((self.X, self.Y), self.Z, (self.X_grid, self.Y_grid), method='cubic')
#         self.TrafFlag_grid = griddata((self.X, self.Y), self.TrafFlag, (self.X_grid, self.Y_grid), method='nearest')
#         self.Acc_grid = griddata((self.X, self.Y), self.maxAcc, (self.X_grid, self.Y_grid), method='nearest')
#         self.Slope_grid = griddata((self.X, self.Y), self.Slope, (self.X_grid, self.Y_grid), method='nearest')
#         self.StepElev_grid = griddata((self.X, self.Y), self.StepElev, (self.X_grid, self.Y_grid), method='nearest')
#         self.Rough_grid = griddata((self.X, self.Y), self.Roughness, (self.X_grid, self.Y_grid), method='nearest')
#
#     def fill_nan_values(self):
#         # 填充 NaN 值
#         self.Z_grid = np.nan_to_num(self.Z_grid, nan=np.nanmean(self.Z))
#         self.Acc_grid = np.nan_to_num(self.Acc_grid, nan=0)
#         self.TrafFlag_grid = np.nan_to_num(self.TrafFlag_grid, nan=0)
#         self.Slope_grid = np.nan_to_num(self.Slope_grid, nan=0)
#         self.StepElev_grid = np.nan_to_num(self.StepElev_grid, nan=0)
#         self.Rough_grid = np.nan_to_num(self.Rough_grid, nan=0)
#
#         # 验证网格数据的形状
#         print(f"Z_grid shape: {self.Z_grid.shape}")
#         print(f"TrafFlag_grid shape: {self.TrafFlag_grid.shape}")
#         print(f"Acc_grid shape: {self.Acc_grid.shape}")
#         print(f"Slope_grid shape: {self.Slope_grid.shape}")
#         print(f"StepElev_grid shape: {self.StepElev_grid.shape}")
#         print(f"Rough_grid shape: {self.Rough_grid.shape}")
#
#         # 计算并打印插值后网格数据的最小值和最大值
#         slope_grid_min = np.min(self.Slope_grid)
#         slope_grid_max = np.max(self.Slope_grid)
#         step_grid_min = np.min(self.StepElev_grid)
#         step_grid_max = np.max(self.StepElev_grid)
#
#         print(f"插值后坡度网格 (Slope_grid) 最小值: {slope_grid_min}")
#         print(f"插值后坡度网格 (Slope_grid) 最大值: {slope_grid_max}")
#         print(f"插值后阶跃高程差网格 (StepElev_grid) 最小值: {step_grid_min}")
#         print(f"插值后阶跃高程差网格 (StepElev_grid) 最大值: {step_grid_max}")
#
#     def coord_to_index(self, coord):
#         x_min, y_min = self.x_range.min(), self.y_range.min()
#         col = int((coord[0] - x_min) / self.x_resolution)
#         row = int((coord[1] - y_min) / self.y_resolution)
#         return (row, col)
#
#     def is_valid_position(self, pos):
#         return 0 <= pos[0] < self.TrafFlag_grid.shape[0] and 0 <= pos[1] < self.TrafFlag_grid.shape[1]
#
#     def split_data_and_evaluate(self):
#         try:
#             # 将所有数据合并
#             data = np.column_stack(
#                 (self.X, self.Y, self.Z, self.TrafFlag, self.maxAcc, self.Slope, self.StepElev, self.Roughness))
#
#             # 分割数据集为训练集和测试集
#             test_size = 0.2
#             random_state = 42
#
#             train_data, test_data = train_test_split(data, test_size=test_size, random_state=random_state)
#
#             # 拆分训练集和测试集的数据
#             X_train, Y_train = train_data[:, 0], train_data[:, 1]
#             Z_train = train_data[:, 2]
#             TrafFlag_train = train_data[:, 3]
#             maxAcc_train = train_data[:, 4]
#             Slope_train = train_data[:, 5]
#             StepElev_train = train_data[:, 6]
#             Rough_train = train_data[:, 7]
#
#             X_test, Y_test = test_data[:, 0], test_data[:, 1]
#             Z_test = test_data[:, 2]
#             TrafFlag_test = test_data[:, 3]
#             maxAcc_test = test_data[:, 4]
#             Slope_test = test_data[:, 5]
#             StepElev_test = test_data[:, 6]
#             Rough_test = test_data[:, 7]
#
#             # 使用训练集进行插值
#             Z_grid_train = griddata((X_train, Y_train), Z_train, (self.X_grid, self.Y_grid), method='cubic')
#             TrafFlag_grid_train = griddata((X_train, Y_train), TrafFlag_train, (self.X_grid, self.Y_grid), method='nearest')
#             Acc_grid_train = griddata((X_train, Y_train), maxAcc_train, (self.X_grid, self.Y_grid), method='nearest')
#             Slope_grid_train = griddata((X_train, Y_train), Slope_train, (self.X_grid, self.Y_grid), method='nearest')
#             StepElev_grid_train = griddata((X_train, Y_train), StepElev_train, (self.X_grid, self.Y_grid), method='nearest')
#             Rough_grid_train = griddata((X_train, Y_train), Rough_train, (self.X_grid, self.Y_grid), method='nearest')
#
#             # 在测试集上预测值
#             Z_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Z_grid_train.flatten(), (X_test, Y_test),
#                               method='cubic')
#             TrafFlag_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), TrafFlag_grid_train.flatten(),
#                                      (X_test, Y_test), method='nearest')
#             Acc_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Acc_grid_train.flatten(), (X_test, Y_test),
#                                 method='nearest')
#             Slope_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Slope_grid_train.flatten(),
#                                   (X_test, Y_test), method='nearest')
#             StepElev_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), StepElev_grid_train.flatten(),
#                                      (X_test, Y_test), method='nearest')
#             Rough_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Rough_grid_train.flatten(),
#                                   (X_test, Y_test), method='nearest')
#
#             # 处理 NaN 值并计算误差指标
#
#             # 高程数据误差
#             valid_indices_Z = ~np.isnan(Z_pred) & ~np.isnan(Z_test)
#             if np.any(valid_indices_Z):
#                 mse_Z = mean_squared_error(Z_test[valid_indices_Z], Z_pred[valid_indices_Z])
#                 mae_Z = mean_absolute_error(Z_test[valid_indices_Z], Z_pred[valid_indices_Z])
#                 max_error_Z = np.max(np.abs(Z_test[valid_indices_Z] - Z_pred[valid_indices_Z]))
#                 print(f"高程插值 - MSE: {mse_Z}, MAE: {mae_Z}, Max Error: {max_error_Z}")
#             else:
#                 print("高程插值预测值中全部为 NaN，无法计算误差。")
#
#             # 通行标志（分类准确率）
#             valid_indices_TrafFlag = ~np.isnan(TrafFlag_pred) & ~np.isnan(TrafFlag_test)
#             if np.any(valid_indices_TrafFlag):
#                 accuracy_TrafFlag = accuracy_score(TrafFlag_test[valid_indices_TrafFlag],
#                                                    TrafFlag_pred[valid_indices_TrafFlag])
#                 print(f"通行标志插值 - 准确率: {accuracy_TrafFlag}")
#             else:
#                 print("通行标志插值预测值中全部为 NaN，无法计算准确率。")
#
#             # 最大加速度误差
#             valid_indices_Acc = ~np.isnan(Acc_pred) & ~np.isnan(maxAcc_test)
#             if np.any(valid_indices_Acc):
#                 mse_Acc = mean_squared_error(maxAcc_test[valid_indices_Acc], Acc_pred[valid_indices_Acc])
#                 mae_Acc = mean_absolute_error(maxAcc_test[valid_indices_Acc], Acc_pred[valid_indices_Acc])
#                 max_error_Acc = np.max(np.abs(maxAcc_test[valid_indices_Acc] - Acc_pred[valid_indices_Acc]))
#                 print(f"最大加速度插值 - MSE: {mse_Acc}, MAE: {mae_Acc}, Max Error: {max_error_Acc}")
#             else:
#                 print("最大加速度插值预测值中全部为 NaN，无法计算误差。")
#
#             # 坡度误差
#             valid_indices_Slope = ~np.isnan(Slope_pred) & ~np.isnan(Slope_test)
#             if np.any(valid_indices_Slope):
#                 mse_Slope = mean_squared_error(Slope_test[valid_indices_Slope], Slope_pred[valid_indices_Slope])
#                 mae_Slope = mean_absolute_error(Slope_test[valid_indices_Slope], Slope_pred[valid_indices_Slope])
#                 max_error_Slope = np.max(np.abs(Slope_test[valid_indices_Slope] - Slope_pred[valid_indices_Slope]))
#                 print(f"坡度插值 - MSE: {mse_Slope}, MAE: {mae_Slope}, Max Error: {max_error_Slope}")
#             else:
#                 print("坡度插值预测值中全部为 NaN，无法计算误差。")
#
#             # 阶跃高程差误差
#             valid_indices_StepElev = ~np.isnan(StepElev_pred) & ~np.isnan(StepElev_test)
#             if np.any(valid_indices_StepElev):
#                 mse_StepElev = mean_squared_error(StepElev_test[valid_indices_StepElev],
#                                                   StepElev_pred[valid_indices_StepElev])
#                 mae_StepElev = mean_absolute_error(StepElev_test[valid_indices_StepElev],
#                                                    StepElev_pred[valid_indices_StepElev])
#                 max_error_StepElev = np.max(
#                     np.abs(StepElev_test[valid_indices_StepElev] - StepElev_pred[valid_indices_StepElev]))
#                 print(f"阶跃高程差插值 - MSE: {mse_StepElev}, MAE: {mae_StepElev}, Max Error: {max_error_StepElev}")
#             else:
#                 print("阶跃高程差插值预测值中全部为 NaN，无法计算误差。")
#
#             # 粗糙度误差
#             valid_indices_Rough = ~np.isnan(Rough_pred) & ~np.isnan(Rough_test)
#             if np.any(valid_indices_Rough):
#                 mse_Rough = mean_squared_error(Rough_test[valid_indices_Rough], Rough_pred[valid_indices_Rough])
#                 mae_Rough = mean_absolute_error(Rough_test[valid_indices_Rough], Rough_pred[valid_indices_Rough])
#                 max_error_Rough = np.max(np.abs(Rough_test[valid_indices_Rough] - Rough_pred[valid_indices_Rough]))
#                 print(f"粗糙度插值 - MSE: {mse_Rough}, MAE: {mae_Rough}, Max Error: {max_error_Rough}")
#             else:
#                 print("粗糙度插值预测值中全部为 NaN，无法计算误差。")
#         except Exception as e:
#             print(f"数据评估时发生错误: {e}")
#             logging.exception("数据评估时发生错误")
#
# # 节点类和路径规划类
# class Node:
#     def __init__(self, position, parent=None):
#         self.position = position  # (row, col)
#         self.parent = parent
#         self.g = 0  # 累计代价
#         self.h = 0  # 启发代价
#         self.f = 0  # 总代价
#
#     def __lt__(self, other):
#         return self.f < other.f
#
# class PathPlanner:
#     def __init__(self, grid_data, max_slope=None, max_step_elev=None, max_speed=70.0):
#         self.grid = grid_data['TrafFlag_grid']
#         self.Z_grid = grid_data['Z_grid']
#         self.Acc_grid = grid_data['Acc_grid']
#         self.Slope_grid = grid_data['Slope_grid']
#         self.StepElev_grid = grid_data['StepElev_grid']
#         self.Rough_grid = grid_data['Rough_grid']
#         self.max_slope = max_slope
#         self.max_step_elev = max_step_elev
#         self.max_speed = max_speed
#
#     def heuristic_distance_3d(self, current_pos, end_pos):
#         d_xy = np.sqrt((current_pos[0] - end_pos[0]) ** 2 +
#                        (current_pos[1] - end_pos[1]) ** 2)
#         z_current = self.Z_grid[current_pos[0], current_pos[1]]
#         z_end = self.Z_grid[end_pos[0], end_pos[1]]
#         d_z = (z_end - z_current) ** 2
#         return np.sqrt(d_xy ** 2 + d_z)
#
#     def astar(self, start, end, mode='distance', alpha_env=0.6, beta_env=0.4):
#         start_node = Node(start)
#         end_node = Node(end)
#
#         open_list = []
#         heapq.heappush(open_list, start_node)
#         open_dict = {start_node.position: start_node}
#
#         closed_set = set()
#
#         while open_list:
#             current_node = heapq.heappop(open_list)
#             del open_dict[current_node.position]
#             closed_set.add(current_node.position)
#
#             # 检查是否到达终点
#             if current_node.position == end_node.position:
#                 path = []
#                 while current_node:
#                     path.append(current_node.position)
#                     current_node = current_node.parent
#                 return path[::-1]
#
#             # 定义8个邻居（包括对角线）
#             neighbors = [(0, 1), (0, -1), (1, 0), (-1, 0),
#                          (1, 1), (1, -1), (-1, 1), (-1, -1)]
#
#             for offset in neighbors:
#                 neighbor_pos = (current_node.position[0] + offset[0],
#                                 current_node.position[1] + offset[1])
#
#                 # 检查邻居是否在网格范围内
#                 if not (0 <= neighbor_pos[0] < self.grid.shape[0] and
#                         0 <= neighbor_pos[1] < self.grid.shape[1]):
#                     continue
#
#                 # 检查 TrafFlag 是否可通行
#                 if self.grid[neighbor_pos] == 0:
#                     continue  # 不可通行区域
#
#                 # 在所有模式下应用坡度和阶跃高程差的限制
#                 if self.max_slope is not None and self.Slope_grid[neighbor_pos] > self.max_slope:
#                     continue  # 坡度超过最大限制
#                 if self.max_step_elev is not None and self.StepElev_grid[neighbor_pos] > self.max_step_elev:
#                     continue  # 阶跃高程差超过最大限制
#
#                 # 检查是否已经在关闭列表中
#                 if neighbor_pos in closed_set:
#                     continue
#
#                 # 处理不同模式
#                 if mode == 'distance':
#                     movement_cost = self.heuristic_distance_3d(current_node.position, neighbor_pos)
#                     g_cost = current_node.g + movement_cost
#                     h_cost = self.heuristic_distance_3d(neighbor_pos, end)
#                 elif mode == 'time':
#                     a_max = self.Acc_grid[neighbor_pos[0], neighbor_pos[1]]
#                     if a_max <= 0:
#                         continue  # 无法加速，跳过
#                     movement_cost = self.heuristic_distance_3d(current_node.position, neighbor_pos)
#                     g_cost = current_node.g + movement_cost / a_max  # 简化为距离除以加速度
#                     h_cost = self.heuristic_distance_3d(neighbor_pos, end)
#                 elif mode == 'env':
#                     movement_uncertainty = (self.Slope_grid[neighbor_pos[0], neighbor_pos[1]] / 45.0 +
#                                             self.StepElev_grid[neighbor_pos[0], neighbor_pos[1]] / 1.5 +
#                                             self.Rough_grid[neighbor_pos[0], neighbor_pos[1]] / 6.0)
#                     g_cost = current_node.g + movement_uncertainty
#                     h_cost = self.heuristic_distance_3d(neighbor_pos, end)
#                 else:
#                     raise ValueError("Invalid mode. Choose from 'distance', 'time', 'env'.")
#
#                 # 创建邻居节点
#                 neighbor_node = Node(neighbor_pos, current_node)
#                 neighbor_node.g = g_cost
#                 neighbor_node.h = h_cost
#                 neighbor_node.f = neighbor_node.g + neighbor_node.h
#
#                 # 检查开放列表中是否有更好的节点
#                 if neighbor_pos in open_dict:
#                     existing_node = open_dict[neighbor_pos]
#                     if existing_node.g <= neighbor_node.g:
#                         continue  # 已有更优节点
#                     else:
#                         # 替换为更优节点
#                         open_list.remove(existing_node)
#                         heapq.heapify(open_list)
#                         heapq.heappush(open_list, neighbor_node)
#                         open_dict[neighbor_pos] = neighbor_node
#                 else:
#                     heapq.heappush(open_list, neighbor_node)
#                     open_dict[neighbor_pos] = neighbor_node
#
#         return None  # 如果找不到路径，返回 None
#
#     def extract_path_coordinates(self, path, x_range, y_range):
#         return [(x_range[p[1]], y_range[p[0]], self.Z_grid[p[0], p[1]]) for p in path]
#
#     def calculate_path_length(self, path_coords):
#         length = 0
#         for i in range(1, len(path_coords)):
#             x1, y1, z1 = path_coords[i - 1]
#             x2, y2, z2 = path_coords[i]
#             length += np.sqrt((x2 - x1) ** 2 + (y2 - y1) ** 2 + (z2 - z1) ** 2)
#         return length
#
#     def compute_traversal_time(self, path_coords, y_range, x_range):
#         total_time = 0.0
#         current_speed = 0.0  # 初始速度为0
#         max_speed = self.max_speed  # 使用用户指定的最大速度
#
#         speeds = [current_speed]  # 存储每个点的速度
#         accelerations = []  # 存储每段的加速度
#
#         for i in range(1, len(path_coords)):
#             x1, y1, z1 = path_coords[i - 1]
#             x2, y2, z2 = path_coords[i]
#             # 计算两点之间的距离
#             d = np.sqrt((x2 - x1) ** 2 + (y2 - y1) ** 2 + (z2 - z1) ** 2)
#
#             # 获取当前点的网格索引
#             row = np.argmin(np.abs(y_range - y1))
#             col = np.argmin(np.abs(x_range - x1))
#
#             # 获取当前点的加速度
#             a = self.Acc_grid[row, col]
#             if a <= 0:
#                 # 如果加速度不可用或为0，无法继续加速
#                 a = 0.0
#
#             # 计算需要加速到最大速度的速度增量
#             delta_v = max_speed - current_speed
#
#             if current_speed < max_speed and a > 0:
#                 # 需要加速
#                 t_acc = delta_v / a  # 加速所需时间
#                 d_acc = current_speed * t_acc + 0.5 * a * t_acc ** 2  # 加速所需距离
#
#                 if d_acc >= d:
#                     # 在当前段无法加速到最大速度
#                     # 使用运动学方程计算通过距离d所需的时间
#                     # 解方程: 0.5 * a * t^2 + current_speed * t - d = 0
#                     discriminant = current_speed ** 2 + 2 * a * d
#                     if discriminant < 0:
#                         # 无实数解，无法到达下一个点
#                         return float('inf'), [], []
#                     t = (-current_speed + math.sqrt(discriminant)) / a
#                     total_time += t
#                     current_speed += a * t
#                     accelerations.append(a)
#                     speeds.append(current_speed)
#                 else:
#                     # 在当前段可以加速到最大速度
#                     # 先加速到最大速度
#                     total_time += t_acc
#                     current_speed = max_speed
#                     accelerations.append(a)
#                     speeds.append(current_speed)
#                     # 剩余距离
#                     d_remain = d - d_acc
#                     # 以最大速度匀速行驶剩余距离
#                     t_remain = d_remain / current_speed
#                     total_time += t_remain
#                     accelerations.append(0.0)
#                     speeds.append(current_speed)
#             else:
#                 # 无法加速或已达到最大速度，匀速行驶
#                 if current_speed == 0:
#                     # 无法移动
#                     return float('inf'), [], []
#                 t = d / current_speed
#                 total_time += t
#                 accelerations.append(0.0)
#                 speeds.append(current_speed)
#
#         return total_time, speeds, accelerations
#
# # 可视化类
# class Visualizer:
#     def __init__(self, X_grid, Y_grid, Z_grid, TrafFlag_grid):
#         self.X_grid = X_grid
#         self.Y_grid = Y_grid
#         self.Z_grid = Z_grid
#         self.TrafFlag_grid = TrafFlag_grid
#         self.fig = go.Figure()
#
#     def plot_terrain(self):
#         self.fig.add_trace(go.Surface(
#             x=self.X_grid, y=self.Y_grid, z=self.Z_grid,
#             surfacecolor=self.TrafFlag_grid,
#             colorscale='Viridis',
#             opacity=0.8,
#             showscale=True,
#             colorbar=dict(title='TrafFlag')
#         ))
#
#     def add_path(self, coords, speeds, accs, color, name, length, time_taken, path_width=4):
#         if coords:
#             x, y, z = zip(*coords)
#             # 处理不可行路径（时间为inf）的情况
#             if math.isinf(time_taken):
#                 time_str = "无法到达"
#             else:
#                 time_str = f"{time_taken:.2f} 秒"
#
#             # 准备 customdata，将速度和加速度添加进去
#             if accs:
#                 accs = [0.0] + accs  # 与 speeds 对齐
#             else:
#                 accs = [0.0] * len(speeds)
#
#             customdata = np.array([speeds, accs]).T  # 转置使其每行对应一个点的 [speed, acceleration]
#
#             self.fig.add_trace(go.Scatter3d(
#                 x=x, y=y, z=z,
#                 mode='lines+markers',
#                 line=dict(color=color, width=path_width),
#                 marker=dict(size=2, color=color),
#                 name=f"{name} (长度: {length:.2f} 米, 时间: {time_str})",
#                 customdata=customdata,
#                 hovertemplate=(
#                         'X: %{x}<br>' +
#                         'Y: %{y}<br>' +
#                         'Z: %{z}<br>' +
#                         '速度: %{customdata[0]:.2f} m/s<br>' +
#                         '加速度: %{customdata[1]:.2f} m/s²<br>' +
#                         '<extra></extra>'
#                 )
#             ))
#         else:
#             print(f"未找到{name}")
#
#     def add_start_and_end_points(self, start_coord, end_coord):
#         self.fig.add_trace(go.Scatter3d(
#             x=[start_coord[0]],
#             y=[start_coord[1]],
#             z=[start_coord[2]],
#             mode='markers',
#             marker=dict(size=8, color='green'),
#             name='起点'
#         ))
#         self.fig.add_trace(go.Scatter3d(
#             x=[end_coord[0]],
#             y=[end_coord[1]],
#             z=[end_coord[2]],
#             mode='markers',
#             marker=dict(size=8, color='red'),
#             name='终点'
#         ))
#
#     def show(self):
#         self.fig.update_layout(
#             scene=dict(
#                 xaxis_title='X 坐标',
#                 yaxis_title='Y 坐标',
#                 zaxis_title='Z 坐标 (高程)',
#                 aspectmode='data'
#             ),
#             title='3D 地形路径规划（基于距离、时间与最小不确定性）',
#             legend=dict(
#                 x=0,
#                 y=1,
#                 bgcolor='rgba(255, 255, 255, 0)',
#                 bordercolor='rgba(255, 255, 255, 0)'
#             )
#         )
#         pio.show(self.fig)
#
# # 报告生成函数
# def write_paths_to_word(document, path_name, path_coords):
#     if not path_coords:
#         return
#     # 添加路径名称和点数量
#     document.add_heading(f"{path_name} (点数: {len(path_coords)})", level=2)
#     # 创建表格，包含三列：X相对坐标、Y相对坐标、Z高程坐标
#     table = document.add_table(rows=1, cols=3)
#     table.style = 'Light Grid'
#     hdr_cells = table.rows[0].cells
#     hdr_cells[0].text = 'X相对坐标'
#     hdr_cells[1].text = 'Y相对坐标'
#     hdr_cells[2].text = 'Z高程坐标'
#     # 填充表格内容
#     for coord in path_coords:
#         row_cells = table.add_row().cells
#         row_cells[0].text = f"{coord[0]}"
#         row_cells[1].text = f"{coord[1]}"
#         row_cells[2].text = f"{coord[2]}"
#     # 添加一个空行
#     document.add_paragraph()
#
# # 处理函数
# def run_processing(csv_path, start_coord, end_coord, output_path, max_allowed_slope, max_allowed_step_elev, max_speed):
#     try:
#         # 设置路径宽度
#         path_width = 4  # 可以根据需要调整路径宽度
#
#         # 数据处理
#         data_processor = DataProcessor(csv_path)
#         data_processor.load_data()
#         data_processor.interpolate_data()
#         data_processor.fill_nan_values()
#         data_processor.split_data_and_evaluate()
#
#         # 将 start_coord 和 end_coord 转换为网格索引
#         start = data_processor.coord_to_index(start_coord)
#         end = data_processor.coord_to_index(end_coord)
#
#         # 检查起点和终点是否在网格范围内
#         if not data_processor.is_valid_position(start):
#             raise ValueError("起点坐标超出网格范围")
#         if not data_processor.is_valid_position(end):
#             raise ValueError("终点坐标超出网格范围")
#
#         # 确保起点和终点可通行
#         if data_processor.TrafFlag_grid[start] == 0:
#             raise ValueError("起点不可通行")
#         if data_processor.TrafFlag_grid[end] == 0:
#             raise ValueError("终点不可通行")
#
#         # 路径规划
#         grid_data = {
#             'TrafFlag_grid': data_processor.TrafFlag_grid,
#             'Z_grid': data_processor.Z_grid,
#             'Acc_grid': data_processor.Acc_grid,
#             'Slope_grid': data_processor.Slope_grid,
#             'StepElev_grid': data_processor.StepElev_grid,
#             'Rough_grid': data_processor.Rough_grid
#         }
#         path_planner = PathPlanner(grid_data, max_slope=max_allowed_slope, max_step_elev=max_allowed_step_elev, max_speed=max_speed)
#
#         print("正在计算基于距离的最短路径...")
#         path_distance = path_planner.astar(start, end, mode='distance')
#         if path_distance:
#             print("基于距离的路径计算完成")
#         else:
#             print("未找到基于距离的路径")
#
#         print("正在计算基于时间的最短耗时路径...")
#         path_time = path_planner.astar(start, end, mode='time')
#         if path_time:
#             print("基于时间的路径计算完成")
#         else:
#             print("未找到基于时间的路径")
#
#         print("正在计算最小不确定性路径...")
#         path_env = path_planner.astar(start, end, mode='env')
#         if path_env:
#             print("最小不确定性路径计算完成")
#         else:
#             print("未找到最小不确定性路径")
#
#         # 提取路径坐标
#         if path_distance:
#             path_distance_coords = path_planner.extract_path_coordinates(path_distance, data_processor.x_range,
#                                                                          data_processor.y_range)
#             length_distance = path_planner.calculate_path_length(path_distance_coords)
#         else:
#             path_distance_coords = None
#             length_distance = None
#         if path_time:
#             path_time_coords = path_planner.extract_path_coordinates(path_time, data_processor.x_range,
#                                                                      data_processor.y_range)
#             length_time = path_planner.calculate_path_length(path_time_coords)
#         else:
#             path_time_coords = None
#             length_time = None
#         if path_env:
#             path_env_coords = path_planner.extract_path_coordinates(path_env, data_processor.x_range,
#                                                                     data_processor.y_range)
#             length_env = path_planner.calculate_path_length(path_env_coords)
#         else:
#             path_env_coords = None
#             length_env = None
#
#         # 计算遍历时间，并获取速度和加速度数据
#         if path_distance:
#             time_distance, speeds_distance, accs_distance = path_planner.compute_traversal_time(
#                 path_distance_coords, data_processor.y_range, data_processor.x_range)
#         else:
#             speeds_distance, accs_distance, time_distance = None, None, None
#
#         if path_time:
#             time_time, speeds_time, accs_time = path_planner.compute_traversal_time(
#                 path_time_coords, data_processor.y_range, data_processor.x_range)
#         else:
#             speeds_time, accs_time, time_time = None, None, None
#
#         if path_env:
#             time_env, speeds_env, accs_env = path_planner.compute_traversal_time(
#                 path_env_coords, data_processor.y_range, data_processor.x_range)
#         else:
#             speeds_env, accs_env, time_env = None, None, None
#
#         # 可视化
#         visualizer = Visualizer(
#             data_processor.X_grid,
#             data_processor.Y_grid,
#             data_processor.Z_grid,
#             data_processor.TrafFlag_grid
#         )
#         visualizer.plot_terrain()
#
#         colors = ['blue', 'orange', 'purple']
#         names = ['MinDistPath 最短距离路径', 'MinTimePath 最短耗时路径', 'MinUncertPath 最小不确定性路径']
#         lengths = [length_distance, length_time, length_env]
#         times = [time_distance, time_time, time_env]
#         coords_list = [path_distance_coords, path_time_coords, path_env_coords]
#         speeds_list = [speeds_distance, speeds_time, speeds_env]
#         accs_list = [accs_distance, accs_time, accs_env]
#
#         for coords, speeds, accs, color, name, length, time_taken in zip(
#                 coords_list, speeds_list, accs_list,
#                 colors, names, lengths, times):
#             visualizer.add_path(coords, speeds, accs, color, name, length, time_taken, path_width)
#
#         # 标记起点和终点
#         if path_distance or path_time or path_env:
#             selected_path = None
#             if path_distance:
#                 selected_path = path_distance_coords
#             elif path_time:
#                 selected_path = path_time_coords
#             elif path_env:
#                 selected_path = path_env_coords
#
#             if selected_path:
#                 # 使用起点坐标
#                 start_x, start_y, start_z = selected_path[0]
#                 # 使用终点坐标
#                 end_x, end_y, end_z = selected_path[-1]
#                 visualizer.add_start_and_end_points((start_x, start_y, start_z), (end_x, end_y, end_z))
#
#         visualizer.show()
#
#         # 创建 Word 文档
#         document = Document()
#         document.add_heading('路径规划结果', 0)
#         # 写入路径信息
#         if path_distance or path_time or path_env:
#             if path_distance:
#                 write_paths_to_word(document, 'MinDistPath ', path_distance_coords)
#             if path_time:
#                 write_paths_to_word(document, 'MinTimePath ', path_time_coords)
#             if path_env:
#                 write_paths_to_word(document, 'MinUncertPath ', path_env_coords)
#
#             # 保存文档
#             try:
#                 # 使用 resource_path 函数获取输出文件路径
#                 output_file_path = output_path  # 如果 output_path 是用户指定的完整路径，则无需使用 resource_path
#                 document.save(output_file_path)
#                 print(f"路径坐标已成功保存到 Word 文档: {output_file_path}")
#             except Exception as e:
#                 print(f"保存 Word 文档时出错: {e}")
#                 logging.exception("保存 Word 文档时出错")
#         else:
#             print("没有路径可写入 Word 文档")
#
#         # 打印路径长度和遍历时间
#         if path_distance:
#             if math.isinf(time_distance):
#                 print(f"最短路径长度: {length_distance:.2f} 米")
#                 print(f"最短距离路径遍历时间: 无法到达")
#             else:
#                 print(f"最短路径长度: {length_distance:.2f} 米")
#                 print(f"最短距离路径遍历时间: {time_distance:.2f} 秒")
#         if path_time:
#             if math.isinf(time_time):
#                 print(f"最短耗时路径长度: {length_time:.2f} 米")
#                 print(f"最短耗时路径遍历时间: 无法到达")
#             else:
#                 print(f"最短耗时路径长度: {length_time:.2f} 米")
#                 print(f"最短耗时路径遍历时间: {time_time:.2f} 秒")
#         if path_env:
#             if math.isinf(time_env):
#                 print(f"最小不确定性路径长度: {length_env:.2f} 米")
#                 print(f"最小不确定性路径遍历时间: 无法到达")
#             else:
#                 print(f"最小不确定性路径长度: {length_env:.2f} 米")
#                 print(f"最小不确定性路径遍历时间: {time_env:.2f} 秒")
#     except Exception as e:
#         print(f"运行处理时发生错误: {e}")
#         logging.exception("运行处理时发生错误")
#
# # 主程序
# def main():
#     try:
#         # 创建主窗口
#         root = tk.Tk()
#         root.title("路径规划软件")
#
#         # 地图数据文件路径
#         tk.Label(root, text="地图数据文件路径:").grid(row=0, column=0, sticky='e')
#         map_data_path_entry = tk.Entry(root, width=50)
#         map_data_path_entry.grid(row=0, column=1)
#         def browse_map_data():
#             filename = filedialog.askopenfilename(filetypes=[("CSV文件", "*.csv")])
#             if filename:
#                 map_data_path_entry.delete(0, tk.END)
#                 map_data_path_entry.insert(0, filename)
#         tk.Button(root, text="浏览...", command=browse_map_data).grid(row=0, column=2)
#
#         # 起点坐标
#         tk.Label(root, text="起点坐标 (格式: X,Y):").grid(row=1, column=0, sticky='e')
#         start_coord_entry = tk.Entry(root, width=50)
#         start_coord_entry.grid(row=1, column=1)
#
#         # 终点坐标
#         tk.Label(root, text="终点坐标 (格式: X,Y):").grid(row=2, column=0, sticky='e')
#         end_coord_entry = tk.Entry(root, width=50)
#         end_coord_entry.grid(row=2, column=1)
#
#         # 输出路径
#         tk.Label(root, text="输出路径文档路径:").grid(row=3, column=0, sticky='e')
#         output_path_entry = tk.Entry(root, width=50)
#         output_path_entry.grid(row=3, column=1)
#         def browse_output_path():
#             filename = filedialog.asksaveasfilename(defaultextension=".docx",
#                                                     filetypes=[("Word文档", "*.docx")])
#             if filename:
#                 output_path_entry.delete(0, tk.END)
#                 output_path_entry.insert(0, filename)
#         tk.Button(root, text="浏览...", command=browse_output_path).grid(row=3, column=2)
#
#         # 最大坡度
#         tk.Label(root, text="最大爬坡度 (默认25.0度):").grid(row=4, column=0, sticky='e')
#         max_slope_entry = tk.Entry(root, width=50)
#         max_slope_entry.grid(row=4, column=1)
#         max_slope_entry.insert(0, "25.0")  # 设置默认值
#
#         # 最大阶跃高程差
#         tk.Label(root, text="最大阶跃高程差 (默认0.3米):").grid(row=5, column=0, sticky='e')
#         max_step_elev_entry = tk.Entry(root, width=50)
#         max_step_elev_entry.grid(row=5, column=1)
#         max_step_elev_entry.insert(0, "10.0")  # 设置默认值
#
#         # 最大速度
#         tk.Label(root, text="最大速度 (默认20.0 m/s):").grid(row=6, column=0, sticky='e')
#         max_speed_entry = tk.Entry(root, width=50)
#         max_speed_entry.grid(row=6, column=1)
#         max_speed_entry.insert(0, "20.0")  # 设置默认值
#
#         # 运行按钮
#         def run():
#             # 获取输入值
#             csv_path = map_data_path_entry.get()
#             start_coord_input = start_coord_entry.get()
#             end_coord_input = end_coord_entry.get()
#             output_path = output_path_entry.get()
#
#             # 验证输入
#             if not csv_path:
#                 messagebox.showerror("错误", "请提供地图数据文件路径。")
#                 return
#             if not output_path:
#                 messagebox.showerror("错误", "请提供输出 Word 文档路径。")
#                 return
#             if not start_coord_input:
#                 messagebox.showerror("错误", "请提供起点坐标。")
#                 return
#             if not end_coord_input:
#                 messagebox.showerror("错误", "请提供终点坐标。")
#                 return
#             try:
#                 start_x, start_y = map(float, start_coord_input.strip().split(','))
#                 start_coord = (start_x, start_y)
#             except ValueError:
#                 messagebox.showerror("错误", "起点坐标格式错误，请使用 'X,Y' 的格式。")
#                 return
#
#             try:
#                 end_x, end_y = map(float, end_coord_input.strip().split(','))
#                 end_coord = (end_x, end_y)
#             except ValueError:
#                 messagebox.showerror("错误", "终点坐标格式错误，请使用 'X,Y' 的格式。")
#                 return
#
#             # 获取最大坡度和最大阶跃高程差
#             max_slope_input = max_slope_entry.get()
#             max_step_elev_input = max_step_elev_entry.get()
#
#             try:
#                 max_allowed_slope = float(max_slope_input)
#             except ValueError:
#                 messagebox.showerror("错误", "最大坡度格式错误，请输入数字。")
#                 return
#
#             try:
#                 max_allowed_step_elev = float(max_step_elev_input)
#             except ValueError:
#                 messagebox.showerror("错误", "最大阶跃高程差格式错误，请输入数字。")
#                 return
#
#             # 获取最大速度
#             max_speed_input = max_speed_entry.get()
#
#             try:
#                 max_speed = float(max_speed_input)
#             except ValueError:
#                 messagebox.showerror("错误", "最大速度格式错误，请输入数字。")
#                 return
#
#             # 在后台线程中运行，以避免阻塞GUI
#             def process():
#                 try:
#                     # 调用处理函数，传入最大坡度、最大阶跃高程差和最大速度
#                     run_processing(csv_path, start_coord, end_coord, output_path, max_allowed_slope, max_allowed_step_elev, max_speed)
#                     messagebox.showinfo("完成", "路径规划已完成。")
#                 except Exception as e:
#                     messagebox.showerror("错误", f"发生错误: {e}")
#                     logging.exception("后台线程运行时发生错误")
#
#             threading.Thread(target=process).start()
#
#         tk.Button(root, text="运行", command=run).grid(row=7, column=1)
#
#         root.mainloop()
#     except Exception as e:
#         print(f"主程序运行时发生错误: {e}")
#         logging.exception("主程序运行时发生错误")
#
# if __name__ == "__main__":
#     main()










# import pandas as pd
# import numpy as np
# from scipy.interpolate import griddata
# from sklearn.metrics import mean_squared_error, mean_absolute_error, accuracy_score
# from sklearn.model_selection import train_test_split
# import heapq
# import math
# import plotly.graph_objs as go
# import plotly.io as pio
# from docx import Document
# import tkinter as tk
# from tkinter import filedialog, messagebox
# import threading
# import sys
# import os
# import logging
#
# # 设置日志记录
# logging.basicConfig(filename='error.log', level=logging.ERROR)
#
# # 设置 Plotly 渲染器为浏览器
# pio.renderers.default = 'browser'
#
# # 添加 resource_path 函数，用于获取资源文件的正确路径
# def resource_path(relative_path):
#     """获取资源文件的绝对路径，兼容 PyInstaller 打包后的情况"""
#     try:
#         # PyInstaller 会把路径存储在 _MEIPASS 中
#         base_path = sys._MEIPASS
#     except AttributeError:
#         base_path = os.path.abspath(".")
#     return os.path.join(base_path, relative_path)
#
# # 数据处理类
# class DataProcessor:
#     def __init__(self, dat_path, x_resolution=1, y_resolution=1):
#         self.dat_path = dat_path
#         self.x_resolution = x_resolution
#         self.y_resolution = y_resolution
#         self.df = None
#         self.X = None
#         self.Y = None
#         self.Z = None
#         self.TrafFlag = None
#         self.maxAcc = None
#         self.Slope = None
#         self.StepElev = None
#         self.Roughness = None
#         self.x_range = None
#         self.y_range = None
#         self.X_grid = None
#         self.Y_grid = None
#         self.Z_grid = None
#         self.TrafFlag_grid = None
#         self.Acc_grid = None
#         self.Slope_grid = None
#         self.StepElev_grid = None
#         self.Rough_grid = None
#
#     def load_data(self):
#         try:
#             # 使用 resource_path 函数获取数据文件路径
#             data_path = resource_path(self.dat_path)
#             # 读取 .csv 文件，指定分隔符为逗号
#             self.df = pd.read_csv(data_path, delimiter=',')
#             print("数据文件中的列名：", self.df.columns.tolist())
#         except FileNotFoundError:
#             print(f"文件未找到: {self.dat_path}")
#             return
#         except pd.errors.EmptyDataError:
#             print("CSV 文件为空")
#             return
#         except Exception as e:
#             print(f"读取 CSV 文件时出错: {e}")
#             logging.exception("读取 CSV 文件时出错")
#             return
#
#         required_columns = {'X', 'Y', 'Z', 'TrafFlag', 'maxAcc', 'Slope', 'StepElevationDifference', 'Roughness'}
#         if not required_columns.issubset(self.df.columns):
#             print(f"CSV文件缺少必要的列: {required_columns - set(self.df.columns)}")
#             return
#
#         # 提取所需数据
#         self.X = self.df['X'].values
#         self.Y = self.df['Y'].values
#         self.Z = self.df['Z'].values
#         self.TrafFlag = self.df['TrafFlag'].values
#         self.maxAcc = self.df['maxAcc'].values
#         self.Slope = self.df['Slope'].values
#         self.StepElev = self.df['StepElevationDifference'].values
#         self.Roughness = self.df['Roughness'].values
#
#         # 计算并打印 Slope 和 StepElevationDifference 的最小值和最大值
#         slope_min = np.min(self.Slope)
#         slope_max_val = np.max(self.Slope)
#         step_min = np.min(self.StepElev)
#         step_max = np.max(self.StepElev)
#
#         print(f"坡度 (Slope) 最小值: {slope_min}")
#         print(f"坡度 (Slope) 最大值: {slope_max_val}")
#         print(f"阶跃高程差 (StepElevationDifference) 最小值: {step_min}")
#         print(f"阶跃高程差 (StepElevationDifference) 最大值: {step_max}")
#
#     def interpolate_data(self):
#         # 计算网格范围和生成网格
#         x_min, x_max = self.X.min(), self.X.max()
#         y_min, y_max = self.Y.min(), self.Y.max()
#         self.x_range = np.arange(x_min, x_max + self.x_resolution, self.x_resolution)
#         self.y_range = np.arange(y_min, y_max + self.y_resolution, self.y_resolution)
#         self.X_grid, self.Y_grid = np.meshgrid(self.x_range, self.y_range)
#
#         # 插值
#         self.Z_grid = griddata((self.X, self.Y), self.Z, (self.X_grid, self.Y_grid), method='cubic')
#         self.TrafFlag_grid = griddata((self.X, self.Y), self.TrafFlag, (self.X_grid, self.Y_grid), method='nearest')
#         self.Acc_grid = griddata((self.X, self.Y), self.maxAcc, (self.X_grid, self.Y_grid), method='nearest')
#         self.Slope_grid = griddata((self.X, self.Y), self.Slope, (self.X_grid, self.Y_grid), method='nearest')
#         self.StepElev_grid = griddata((self.X, self.Y), self.StepElev, (self.X_grid, self.Y_grid), method='nearest')
#         self.Rough_grid = griddata((self.X, self.Y), self.Roughness, (self.X_grid, self.Y_grid), method='nearest')
#
#     def fill_nan_values(self):
#         # 填充 NaN 值
#         self.Z_grid = np.nan_to_num(self.Z_grid, nan=np.nanmean(self.Z))
#         self.Acc_grid = np.nan_to_num(self.Acc_grid, nan=0)
#         self.TrafFlag_grid = np.nan_to_num(self.TrafFlag_grid, nan=0)
#         # 将 TrafFlag_grid 的值限制为 0 或 1
#         self.TrafFlag_grid = np.where(self.TrafFlag_grid >= 0.5, 1, 0)
#         self.Slope_grid = np.nan_to_num(self.Slope_grid, nan=0)
#         self.StepElev_grid = np.nan_to_num(self.StepElev_grid, nan=0)
#         self.Rough_grid = np.nan_to_num(self.Rough_grid, nan=0)
#
#         # 验证网格数据的形状
#         print(f"Z_grid shape: {self.Z_grid.shape}")
#         print(f"TrafFlag_grid shape: {self.TrafFlag_grid.shape}")
#         print(f"Acc_grid shape: {self.Acc_grid.shape}")
#         print(f"Slope_grid shape: {self.Slope_grid.shape}")
#         print(f"StepElev_grid shape: {self.StepElev_grid.shape}")
#         print(f"Rough_grid shape: {self.Rough_grid.shape}")
#
#         # 计算并打印插值后网格数据的最小值和最大值
#         slope_grid_min = np.min(self.Slope_grid)
#         slope_grid_max = np.max(self.Slope_grid)
#         step_grid_min = np.min(self.StepElev_grid)
#         step_grid_max = np.max(self.StepElev_grid)
#
#         print(f"插值后坡度网格 (Slope_grid) 最小值: {slope_grid_min}")
#         print(f"插值后坡度网格 (Slope_grid) 最大值: {slope_grid_max}")
#         print(f"插值后阶跃高程差网格 (StepElev_grid) 最小值: {step_grid_min}")
#         print(f"插值后阶跃高程差网格 (StepElev_grid) 最大值: {step_grid_max}")
#
#     def coord_to_index(self, coord):
#         x_min, y_min = self.x_range.min(), self.y_range.min()
#         col = int((coord[0] - x_min) / self.x_resolution)
#         row = int((coord[1] - y_min) / self.y_resolution)
#         return (row, col)
#
#     def is_valid_position(self, pos):
#         return 0 <= pos[0] < self.TrafFlag_grid.shape[0] and 0 <= pos[1] < self.TrafFlag_grid.shape[1]
#
#     def split_data_and_evaluate(self):
#         try:
#             # 将所有数据合并
#             data = np.column_stack(
#                 (self.X, self.Y, self.Z, self.TrafFlag, self.maxAcc, self.Slope, self.StepElev, self.Roughness))
#
#             # 分割数据集为训练集和测试集
#             test_size = 0.2
#             random_state = 42
#
#             train_data, test_data = train_test_split(data, test_size=test_size, random_state=random_state)
#
#             # 拆分训练集和测试集的数据
#             X_train, Y_train = train_data[:, 0], train_data[:, 1]
#             Z_train = train_data[:, 2]
#             TrafFlag_train = train_data[:, 3]
#             maxAcc_train = train_data[:, 4]
#             Slope_train = train_data[:, 5]
#             StepElev_train = train_data[:, 6]
#             Rough_train = train_data[:, 7]
#
#             X_test, Y_test = test_data[:, 0], test_data[:, 1]
#             Z_test = test_data[:, 2]
#             TrafFlag_test = test_data[:, 3]
#             maxAcc_test = test_data[:, 4]
#             Slope_test = test_data[:, 5]
#             StepElev_test = test_data[:, 6]
#             Rough_test = test_data[:, 7]
#
#             # 使用训练集进行插值
#             Z_grid_train = griddata((X_train, Y_train), Z_train, (self.X_grid, self.Y_grid), method='cubic')
#             TrafFlag_grid_train = griddata((X_train, Y_train), TrafFlag_train, (self.X_grid, self.Y_grid), method='nearest')
#             Acc_grid_train = griddata((X_train, Y_train), maxAcc_train, (self.X_grid, self.Y_grid), method='nearest')
#             Slope_grid_train = griddata((X_train, Y_train), Slope_train, (self.X_grid, self.Y_grid), method='nearest')
#             StepElev_grid_train = griddata((X_train, Y_train), StepElev_train, (self.X_grid, self.Y_grid), method='nearest')
#             Rough_grid_train = griddata((X_train, Y_train), Rough_train, (self.X_grid, self.Y_grid), method='nearest')
#
#             # 在测试集上预测值
#             Z_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Z_grid_train.flatten(), (X_test, Y_test),
#                               method='cubic')
#             TrafFlag_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), TrafFlag_grid_train.flatten(),
#                                      (X_test, Y_test), method='nearest')
#             Acc_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Acc_grid_train.flatten(), (X_test, Y_test),
#                                 method='nearest')
#             Slope_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Slope_grid_train.flatten(),
#                                   (X_test, Y_test), method='nearest')
#             StepElev_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), StepElev_grid_train.flatten(),
#                                      (X_test, Y_test), method='nearest')
#             Rough_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Rough_grid_train.flatten(),
#                                   (X_test, Y_test), method='nearest')
#
#             # 处理 NaN 值并计算误差指标
#
#             # 高程数据误差
#             valid_indices_Z = ~np.isnan(Z_pred) & ~np.isnan(Z_test)
#             if np.any(valid_indices_Z):
#                 mse_Z = mean_squared_error(Z_test[valid_indices_Z], Z_pred[valid_indices_Z])
#                 mae_Z = mean_absolute_error(Z_test[valid_indices_Z], Z_pred[valid_indices_Z])
#                 max_error_Z = np.max(np.abs(Z_test[valid_indices_Z] - Z_pred[valid_indices_Z]))
#                 print(f"高程插值 - MSE: {mse_Z}, MAE: {mae_Z}, Max Error: {max_error_Z}")
#             else:
#                 print("高程插值预测值中全部为 NaN，无法计算误差。")
#
#             # 通行标志（分类准确率）
#             valid_indices_TrafFlag = ~np.isnan(TrafFlag_pred) & ~np.isnan(TrafFlag_test)
#             if np.any(valid_indices_TrafFlag):
#                 accuracy_TrafFlag = accuracy_score(TrafFlag_test[valid_indices_TrafFlag],
#                                                    TrafFlag_pred[valid_indices_TrafFlag])
#                 print(f"通行标志插值 - 准确率: {accuracy_TrafFlag}")
#             else:
#                 print("通行标志插值预测值中全部为 NaN，无法计算准确率。")
#
#             # 最大加速度误差
#             valid_indices_Acc = ~np.isnan(Acc_pred) & ~np.isnan(maxAcc_test)
#             if np.any(valid_indices_Acc):
#                 mse_Acc = mean_squared_error(maxAcc_test[valid_indices_Acc], Acc_pred[valid_indices_Acc])
#                 mae_Acc = mean_absolute_error(maxAcc_test[valid_indices_Acc], Acc_pred[valid_indices_Acc])
#                 max_error_Acc = np.max(np.abs(maxAcc_test[valid_indices_Acc] - Acc_pred[valid_indices_Acc]))
#                 print(f"最大加速度插值 - MSE: {mse_Acc}, MAE: {mae_Acc}, Max Error: {max_error_Acc}")
#             else:
#                 print("最大加速度插值预测值中全部为 NaN，无法计算误差。")
#
#             # 坡度误差
#             valid_indices_Slope = ~np.isnan(Slope_pred) & ~np.isnan(Slope_test)
#             if np.any(valid_indices_Slope):
#                 mse_Slope = mean_squared_error(Slope_test[valid_indices_Slope], Slope_pred[valid_indices_Slope])
#                 mae_Slope = mean_absolute_error(Slope_test[valid_indices_Slope], Slope_pred[valid_indices_Slope])
#                 max_error_Slope = np.max(np.abs(Slope_test[valid_indices_Slope] - Slope_pred[valid_indices_Slope]))
#                 print(f"坡度插值 - MSE: {mse_Slope}, MAE: {mae_Slope}, Max Error: {max_error_Slope}")
#             else:
#                 print("坡度插值预测值中全部为 NaN，无法计算误差。")
#
#             # 阶跃高程差误差
#             valid_indices_StepElev = ~np.isnan(StepElev_pred) & ~np.isnan(StepElev_test)
#             if np.any(valid_indices_StepElev):
#                 mse_StepElev = mean_squared_error(StepElev_test[valid_indices_StepElev],
#                                                   StepElev_pred[valid_indices_StepElev])
#                 mae_StepElev = mean_absolute_error(StepElev_test[valid_indices_StepElev],
#                                                    StepElev_pred[valid_indices_StepElev])
#                 max_error_StepElev = np.max(
#                     np.abs(StepElev_test[valid_indices_StepElev] - StepElev_pred[valid_indices_StepElev]))
#                 print(f"阶跃高程差插值 - MSE: {mse_StepElev}, MAE: {mae_StepElev}, Max Error: {max_error_StepElev}")
#             else:
#                 print("阶跃高程差插值预测值中全部为 NaN，无法计算误差。")
#
#             # 粗糙度误差
#             valid_indices_Rough = ~np.isnan(Rough_pred) & ~np.isnan(Rough_test)
#             if np.any(valid_indices_Rough):
#                 mse_Rough = mean_squared_error(Rough_test[valid_indices_Rough], Rough_pred[valid_indices_Rough])
#                 mae_Rough = mean_absolute_error(Rough_test[valid_indices_Rough], Rough_pred[valid_indices_Rough])
#                 max_error_Rough = np.max(np.abs(Rough_test[valid_indices_Rough] - Rough_pred[valid_indices_Rough]))
#                 print(f"粗糙度插值 - MSE: {mse_Rough}, MAE: {mae_Rough}, Max Error: {max_error_Rough}")
#             else:
#                 print("粗糙度插值预测值中全部为 NaN，无法计算误差。")
#         except Exception as e:
#             print(f"数据评估时发生错误: {e}")
#             logging.exception("数据评估时发生错误")
#
# # 节点类和路径规划类
# class Node:
#     def __init__(self, position, parent=None):
#         self.position = position  # (row, col)
#         self.parent = parent
#         self.g = 0  # 累计代价
#         self.h = 0  # 启发代价
#         self.f = 0  # 总代价
#
#     def __lt__(self, other):
#         return self.f < other.f
#
# class PathPlanner:
#     def __init__(self, grid_data, max_slope=None, max_step_elev=None, max_speed=70.0):
#         self.grid = grid_data['TrafFlag_grid']
#         self.Z_grid = grid_data['Z_grid']
#         self.Acc_grid = grid_data['Acc_grid']
#         self.Slope_grid = grid_data['Slope_grid']
#         self.StepElev_grid = grid_data['StepElev_grid']
#         self.Rough_grid = grid_data['Rough_grid']
#         self.max_slope = max_slope
#         self.max_step_elev = max_step_elev
#         self.max_speed = max_speed
#
#     def heuristic_distance_3d(self, current_pos, end_pos):
#         d_xy = np.sqrt((current_pos[0] - end_pos[0]) ** 2 +
#                        (current_pos[1] - end_pos[1]) ** 2)
#         z_current = self.Z_grid[current_pos[0], current_pos[1]]
#         z_end = self.Z_grid[end_pos[0], end_pos[1]]
#         d_z = (z_end - z_current) ** 2
#         return np.sqrt(d_xy ** 2 + d_z)
#
#     def astar(self, start, end, mode='distance', alpha_env=0.6, beta_env=0.4):
#         start_node = Node(start)
#         end_node = Node(end)
#
#         open_list = []
#         heapq.heappush(open_list, start_node)
#         open_dict = {start_node.position: start_node}
#
#         closed_set = set()
#
#         while open_list:
#             current_node = heapq.heappop(open_list)
#             del open_dict[current_node.position]
#             closed_set.add(current_node.position)
#
#             # 检查是否到达终点
#             if current_node.position == end_node.position:
#                 path = []
#                 while current_node:
#                     path.append(current_node.position)
#                     current_node = current_node.parent
#                 return path[::-1]
#
#             # 定义8个邻居（包括对角线）
#             neighbors = [(0, 1), (0, -1), (1, 0), (-1, 0),
#                          (1, 1), (1, -1), (-1, 1), (-1, -1)]
#
#             for offset in neighbors:
#                 neighbor_pos = (current_node.position[0] + offset[0],
#                                 current_node.position[1] + offset[1])
#
#                 # 检查邻居是否在网格范围内
#                 if not (0 <= neighbor_pos[0] < self.grid.shape[0] and
#                         0 <= neighbor_pos[1] < self.grid.shape[1]):
#                     continue
#
#                 # 检查 TrafFlag 是否可通行
#                 if self.grid[neighbor_pos] == 0:
#                     continue  # 不可通行区域
#
#                 # 在所有模式下应用坡度和阶跃高程差的限制
#                 if self.max_slope is not None and self.Slope_grid[neighbor_pos] > self.max_slope:
#                     continue  # 坡度超过最大限制
#                 if self.max_step_elev is not None and self.StepElev_grid[neighbor_pos] > self.max_step_elev:
#                     continue  # 阶跃高程差超过最大限制
#
#                 # 检查是否已经在关闭列表中
#                 if neighbor_pos in closed_set:
#                     continue
#
#                 # 处理不同模式
#                 if mode == 'distance':
#                     movement_cost = self.heuristic_distance_3d(current_node.position, neighbor_pos)
#                     g_cost = current_node.g + movement_cost
#                     h_cost = self.heuristic_distance_3d(neighbor_pos, end)
#                 elif mode == 'time':
#                     a_max = self.Acc_grid[neighbor_pos[0], neighbor_pos[1]]
#                     if a_max <= 0:
#                         continue  # 无法加速，跳过
#                     movement_cost = self.heuristic_distance_3d(current_node.position, neighbor_pos)
#                     g_cost = current_node.g + movement_cost / a_max  # 简化为距离除以加速度
#                     h_cost = self.heuristic_distance_3d(neighbor_pos, end)
#                 elif mode == 'env':
#                     movement_uncertainty = (self.Slope_grid[neighbor_pos[0], neighbor_pos[1]] / 45.0 +
#                                             self.StepElev_grid[neighbor_pos[0], neighbor_pos[1]] / 1.5 +
#                                             self.Rough_grid[neighbor_pos[0], neighbor_pos[1]] / 6.0)
#                     g_cost = current_node.g + movement_uncertainty
#                     h_cost = self.heuristic_distance_3d(neighbor_pos, end)
#                 else:
#                     raise ValueError("Invalid mode. Choose from 'distance', 'time', 'env'.")
#
#                 # 创建邻居节点
#                 neighbor_node = Node(neighbor_pos, current_node)
#                 neighbor_node.g = g_cost
#                 neighbor_node.h = h_cost
#                 neighbor_node.f = neighbor_node.g + neighbor_node.h
#
#                 # 检查开放列表中是否有更好的节点
#                 if neighbor_pos in open_dict:
#                     existing_node = open_dict[neighbor_pos]
#                     if existing_node.g <= neighbor_node.g:
#                         continue  # 已有更优节点
#                     else:
#                         # 替换为更优节点
#                         open_list.remove(existing_node)
#                         heapq.heapify(open_list)
#                         heapq.heappush(open_list, neighbor_node)
#                         open_dict[neighbor_pos] = neighbor_node
#                 else:
#                     heapq.heappush(open_list, neighbor_node)
#                     open_dict[neighbor_pos] = neighbor_node
#
#         return None  # 如果找不到路径，返回 None
#
#     def extract_path_coordinates(self, path, x_range, y_range):
#         return [(x_range[p[1]], y_range[p[0]], self.Z_grid[p[0], p[1]]) for p in path]
#
#     def calculate_path_length(self, path_coords):
#         length = 0
#         for i in range(1, len(path_coords)):
#             x1, y1, z1 = path_coords[i - 1]
#             x2, y2, z2 = path_coords[i]
#             length += np.sqrt((x2 - x1) ** 2 + (y2 - y1) ** 2 + (z2 - z1) ** 2)
#         return length
#
#     def compute_traversal_time(self, path_coords, y_range, x_range):
#         total_time = 0.0
#         current_speed = 0.0  # 初始速度为0
#         max_speed = self.max_speed  # 使用用户指定的最大速度
#
#         speeds = [current_speed]  # 存储每个点的速度
#         accelerations = []  # 存储每段的加速度
#
#         for i in range(1, len(path_coords)):
#             x1, y1, z1 = path_coords[i - 1]
#             x2, y2, z2 = path_coords[i]
#             # 计算两点之间的距离
#             d = np.sqrt((x2 - x1) ** 2 + (y2 - y1) ** 2 + (z2 - z1) ** 2)
#
#             # 获取当前点的网格索引
#             row = np.argmin(np.abs(y_range - y1))
#             col = np.argmin(np.abs(x_range - x1))
#
#             # 获取当前点的加速度
#             a = self.Acc_grid[row, col]
#             if a <= 0:
#                 # 如果加速度不可用或为0，无法继续加速
#                 a = 0.0
#
#             # 计算需要加速到最大速度的速度增量
#             delta_v = max_speed - current_speed
#
#             if current_speed < max_speed and a > 0:
#                 # 需要加速
#                 t_acc = delta_v / a  # 加速所需时间
#                 d_acc = current_speed * t_acc + 0.5 * a * t_acc ** 2  # 加速所需距离
#
#                 if d_acc >= d:
#                     # 在当前段无法加速到最大速度
#                     # 使用运动学方程计算通过距离d所需的时间
#                     # 解方程: 0.5 * a * t^2 + current_speed * t - d = 0
#                     discriminant = current_speed ** 2 + 2 * a * d
#                     if discriminant < 0:
#                         # 无实数解，无法到达下一个点
#                         return float('inf'), [], []
#                     t = (-current_speed + math.sqrt(discriminant)) / a
#                     total_time += t
#                     current_speed += a * t
#                     accelerations.append(a)
#                     speeds.append(current_speed)
#                 else:
#                     # 在当前段可以加速到最大速度
#                     # 先加速到最大速度
#                     total_time += t_acc
#                     current_speed = max_speed
#                     accelerations.append(a)
#                     speeds.append(current_speed)
#                     # 剩余距离
#                     d_remain = d - d_acc
#                     # 以最大速度匀速行驶剩余距离
#                     t_remain = d_remain / current_speed
#                     total_time += t_remain
#                     accelerations.append(0.0)
#                     speeds.append(current_speed)
#             else:
#                 # 无法加速或已达到最大速度，匀速行驶
#                 if current_speed == 0:
#                     # 无法移动
#                     return float('inf'), [], []
#                 t = d / current_speed
#                 total_time += t
#                 accelerations.append(0.0)
#                 speeds.append(current_speed)
#
#         return total_time, speeds, accelerations
#
# # 可视化类
# class Visualizer:
#     def __init__(self, X_grid, Y_grid, Z_grid, TrafFlag_grid):
#         self.X_grid = X_grid
#         self.Y_grid = Y_grid
#         self.Z_grid = Z_grid
#         self.TrafFlag_grid = TrafFlag_grid
#         self.fig = go.Figure()
#
#     def plot_terrain(self):
#         # 定义自定义颜色尺度
#         colorscale = [
#             [0.0, 'gray'],        # TrafFlag == 0 的区域为灰色
#             [1.0, 'lightgreen']   # TrafFlag == 1 的区域为浅绿色
#         ]
#
#         self.fig.add_trace(go.Surface(
#             x=self.X_grid, y=self.Y_grid, z=self.Z_grid,
#             surfacecolor=self.TrafFlag_grid,
#             colorscale=colorscale,
#             cmin=0,
#             cmax=1,
#             opacity=0.8,
#             showscale=True,
#             colorbar=dict(title='TrafFlag')
#         ))
#
#     def add_path(self, coords, speeds, accs, color, name, length, time_taken, path_width=4):
#         if coords:
#             x, y, z = zip(*coords)
#             # 处理不可行路径（时间为inf）的情况
#             if math.isinf(time_taken):
#                 time_str = "无法到达"
#             else:
#                 time_str = f"{time_taken:.2f} 秒"
#
#             # 准备 customdata，将速度和加速度添加进去
#             if accs:
#                 accs = [0.0] + accs  # 与 speeds 对齐
#             else:
#                 accs = [0.0] * len(speeds)
#
#             customdata = np.array([speeds, accs]).T  # 转置使其每行对应一个点的 [speed, acceleration]
#
#             self.fig.add_trace(go.Scatter3d(
#                 x=x, y=y, z=z,
#                 mode='lines+markers',
#                 line=dict(color=color, width=path_width),
#                 marker=dict(size=2, color=color),
#                 name=f"{name} (长度: {length:.2f} 米, 时间: {time_str})",
#                 customdata=customdata,
#                 hovertemplate=(
#                         'X: %{x}<br>' +
#                         'Y: %{y}<br>' +
#                         'Z: %{z}<br>' +
#                         '速度: %{customdata[0]:.2f} m/s<br>' +
#                         '加速度: %{customdata[1]:.2f} m/s²<br>' +
#                         '<extra></extra>'
#                 )
#             ))
#         else:
#             print(f"未找到{name}")
#
#     def add_start_and_end_points(self, start_coord, end_coord):
#         self.fig.add_trace(go.Scatter3d(
#             x=[start_coord[0]],
#             y=[start_coord[1]],
#             z=[start_coord[2]],
#             mode='markers',
#             marker=dict(size=8, color='green'),
#             name='起点'
#         ))
#         self.fig.add_trace(go.Scatter3d(
#             x=[end_coord[0]],
#             y=[end_coord[1]],
#             z=[end_coord[2]],
#             mode='markers',
#             marker=dict(size=8, color='red'),
#             name='终点'
#         ))
#
#     def show(self):
#         self.fig.update_layout(
#             scene=dict(
#                 xaxis_title='X 坐标',
#                 yaxis_title='Y 坐标',
#                 zaxis_title='Z 坐标 (高程)',
#                 aspectmode='data'
#             ),
#             title='3D 地形路径规划（基于距离、时间与最小不确定性）',
#             legend=dict(
#                 x=0,
#                 y=1,
#                 bgcolor='rgba(255, 255, 255, 0)',
#                 bordercolor='rgba(255, 255, 255, 0)'
#             )
#         )
#         pio.show(self.fig)
#
# # 报告生成函数
# def write_paths_to_word(document, path_name, path_coords):
#     if not path_coords:
#         return
#     # 添加路径名称和点数量
#     document.add_heading(f"{path_name} (点数: {len(path_coords)})", level=2)
#     # 创建表格，包含三列：X相对坐标、Y相对坐标、Z高程坐标
#     table = document.add_table(rows=1, cols=3)
#     table.style = 'Light Grid'
#     hdr_cells = table.rows[0].cells
#     hdr_cells[0].text = 'X相对坐标'
#     hdr_cells[1].text = 'Y相对坐标'
#     hdr_cells[2].text = 'Z高程坐标'
#     # 填充表格内容
#     for coord in path_coords:
#         row_cells = table.add_row().cells
#         row_cells[0].text = f"{coord[0]}"
#         row_cells[1].text = f"{coord[1]}"
#         row_cells[2].text = f"{coord[2]}"
#     # 添加一个空行
#     document.add_paragraph()
#
# # 处理函数
# def run_processing(csv_path, start_coord, end_coord, output_path, max_allowed_slope, max_allowed_step_elev, max_speed):
#     try:
#         # 设置路径宽度
#         path_width = 4  # 可以根据需要调整路径宽度
#
#         # 数据处理
#         data_processor = DataProcessor(csv_path)
#         data_processor.load_data()
#         data_processor.interpolate_data()
#         data_processor.fill_nan_values()
#         data_processor.split_data_and_evaluate()
#
#         # 将 start_coord 和 end_coord 转换为网格索引
#         start = data_processor.coord_to_index(start_coord)
#         end = data_processor.coord_to_index(end_coord)
#
#         # 检查起点和终点是否在网格范围内
#         if not data_processor.is_valid_position(start):
#             raise ValueError("起点坐标超出网格范围")
#         if not data_processor.is_valid_position(end):
#             raise ValueError("终点坐标超出网格范围")
#
#         # 确保起点和终点可通行
#         if data_processor.TrafFlag_grid[start] == 0:
#             raise ValueError("起点不可通行")
#         if data_processor.TrafFlag_grid[end] == 0:
#             raise ValueError("终点不可通行")
#
#         # 路径规划
#         grid_data = {
#             'TrafFlag_grid': data_processor.TrafFlag_grid,
#             'Z_grid': data_processor.Z_grid,
#             'Acc_grid': data_processor.Acc_grid,
#             'Slope_grid': data_processor.Slope_grid,
#             'StepElev_grid': data_processor.StepElev_grid,
#             'Rough_grid': data_processor.Rough_grid
#         }
#         path_planner = PathPlanner(grid_data, max_slope=max_allowed_slope, max_step_elev=max_allowed_step_elev, max_speed=max_speed)
#
#         print("正在计算基于距离的最短路径...")
#         path_distance = path_planner.astar(start, end, mode='distance')
#         if path_distance:
#             print("基于距离的路径计算完成")
#         else:
#             print("未找到基于距离的路径")
#
#         print("正在计算基于时间的最短耗时路径...")
#         path_time = path_planner.astar(start, end, mode='time')
#         if path_time:
#             print("基于时间的路径计算完成")
#         else:
#             print("未找到基于时间的路径")
#
#         print("正在计算最小不确定性路径...")
#         path_env = path_planner.astar(start, end, mode='env')
#         if path_env:
#             print("最小不确定性路径计算完成")
#         else:
#             print("未找到最小不确定性路径")
#
#         # 提取路径坐标
#         if path_distance:
#             path_distance_coords = path_planner.extract_path_coordinates(path_distance, data_processor.x_range,
#                                                                          data_processor.y_range)
#             length_distance = path_planner.calculate_path_length(path_distance_coords)
#         else:
#             path_distance_coords = None
#             length_distance = None
#         if path_time:
#             path_time_coords = path_planner.extract_path_coordinates(path_time, data_processor.x_range,
#                                                                      data_processor.y_range)
#             length_time = path_planner.calculate_path_length(path_time_coords)
#         else:
#             path_time_coords = None
#             length_time = None
#         if path_env:
#             path_env_coords = path_planner.extract_path_coordinates(path_env, data_processor.x_range,
#                                                                     data_processor.y_range)
#             length_env = path_planner.calculate_path_length(path_env_coords)
#         else:
#             path_env_coords = None
#             length_env = None
#
#         # 计算遍历时间，并获取速度和加速度数据
#         if path_distance:
#             time_distance, speeds_distance, accs_distance = path_planner.compute_traversal_time(
#                 path_distance_coords, data_processor.y_range, data_processor.x_range)
#         else:
#             speeds_distance, accs_distance, time_distance = None, None, None
#
#         if path_time:
#             time_time, speeds_time, accs_time = path_planner.compute_traversal_time(
#                 path_time_coords, data_processor.y_range, data_processor.x_range)
#         else:
#             speeds_time, accs_time, time_time = None, None, None
#
#         if path_env:
#             time_env, speeds_env, accs_env = path_planner.compute_traversal_time(
#                 path_env_coords, data_processor.y_range, data_processor.x_range)
#         else:
#             speeds_env, accs_env, time_env = None, None, None
#
#         # 可视化
#         visualizer = Visualizer(
#             data_processor.X_grid,
#             data_processor.Y_grid,
#             data_processor.Z_grid,
#             data_processor.TrafFlag_grid
#         )
#         visualizer.plot_terrain()
#
#         colors = ['blue', 'orange', 'purple']
#         names = ['MinDistPath 最短距离路径', 'MinTimePath 最短耗时路径', 'MinUncertPath 最小不确定性路径']
#         lengths = [length_distance, length_time, length_env]
#         times = [time_distance, time_time, time_env]
#         coords_list = [path_distance_coords, path_time_coords, path_env_coords]
#         speeds_list = [speeds_distance, speeds_time, speeds_env]
#         accs_list = [accs_distance, accs_time, accs_env]
#
#         for coords, speeds, accs, color, name, length, time_taken in zip(
#                 coords_list, speeds_list, accs_list,
#                 colors, names, lengths, times):
#             visualizer.add_path(coords, speeds, accs, color, name, length, time_taken, path_width)
#
#         # 标记起点和终点
#         if path_distance or path_time or path_env:
#             selected_path = None
#             if path_distance:
#                 selected_path = path_distance_coords
#             elif path_time:
#                 selected_path = path_time_coords
#             elif path_env:
#                 selected_path = path_env_coords
#
#             if selected_path:
#                 # 使用起点坐标
#                 start_x, start_y, start_z = selected_path[0]
#                 # 使用终点坐标
#                 end_x, end_y, end_z = selected_path[-1]
#                 visualizer.add_start_and_end_points((start_x, start_y, start_z), (end_x, end_y, end_z))
#
#         visualizer.show()
#
#         # 创建 Word 文档
#         document = Document()
#         document.add_heading('路径规划结果', 0)
#         # 写入路径信息
#         if path_distance or path_time or path_env:
#             if path_distance:
#                 write_paths_to_word(document, 'MinDistPath ', path_distance_coords)
#             if path_time:
#                 write_paths_to_word(document, 'MinTimePath ', path_time_coords)
#             if path_env:
#                 write_paths_to_word(document, 'MinUncertPath ', path_env_coords)
#
#             # 保存文档
#             try:
#                 # 使用 resource_path 函数获取输出文件路径
#                 output_file_path = output_path  # 如果 output_path 是用户指定的完整路径，则无需使用 resource_path
#                 document.save(output_file_path)
#                 print(f"路径坐标已成功保存到 Word 文档: {output_file_path}")
#             except Exception as e:
#                 print(f"保存 Word 文档时出错: {e}")
#                 logging.exception("保存 Word 文档时出错")
#         else:
#             print("没有路径可写入 Word 文档")
#
#         # 打印路径长度和遍历时间
#         if path_distance:
#             if math.isinf(time_distance):
#                 print(f"最短路径长度: {length_distance:.2f} 米")
#                 print(f"最短距离路径遍历时间: 无法到达")
#             else:
#                 print(f"最短路径长度: {length_distance:.2f} 米")
#                 print(f"最短距离路径遍历时间: {time_distance:.2f} 秒")
#         if path_time:
#             if math.isinf(time_time):
#                 print(f"最短耗时路径长度: {length_time:.2f} 米")
#                 print(f"最短耗时路径遍历时间: 无法到达")
#             else:
#                 print(f"最短耗时路径长度: {length_time:.2f} 米")
#                 print(f"最短耗时路径遍历时间: {time_time:.2f} 秒")
#         if path_env:
#             if math.isinf(time_env):
#                 print(f"最小不确定性路径长度: {length_env:.2f} 米")
#                 print(f"最小不确定性路径遍历时间: 无法到达")
#             else:
#                 print(f"最小不确定性路径长度: {length_env:.2f} 米")
#                 print(f"最小不确定性路径遍历时间: {time_env:.2f} 秒")
#     except Exception as e:
#         print(f"运行处理时发生错误: {e}")
#         logging.exception("运行处理时发生错误")
#
# # 主程序
# def main():
#     try:
#         # 创建主窗口
#         root = tk.Tk()
#         root.title("路径规划软件")
#
#         # 地图数据文件路径
#         tk.Label(root, text="地图数据文件路径:").grid(row=0, column=0, sticky='e')
#         map_data_path_entry = tk.Entry(root, width=50)
#         map_data_path_entry.grid(row=0, column=1)
#         def browse_map_data():
#             filename = filedialog.askopenfilename(filetypes=[("CSV文件", "*.csv")])
#             if filename:
#                 map_data_path_entry.delete(0, tk.END)
#                 map_data_path_entry.insert(0, filename)
#         tk.Button(root, text="浏览...", command=browse_map_data).grid(row=0, column=2)
#
#         # 起点坐标
#         tk.Label(root, text="起点坐标 (格式: X,Y):").grid(row=1, column=0, sticky='e')
#         start_coord_entry = tk.Entry(root, width=50)
#         start_coord_entry.grid(row=1, column=1)
#
#         # 终点坐标
#         tk.Label(root, text="终点坐标 (格式: X,Y):").grid(row=2, column=0, sticky='e')
#         end_coord_entry = tk.Entry(root, width=50)
#         end_coord_entry.grid(row=2, column=1)
#
#         # 输出路径
#         tk.Label(root, text="输出路径文档路径:").grid(row=3, column=0, sticky='e')
#         output_path_entry = tk.Entry(root, width=50)
#         output_path_entry.grid(row=3, column=1)
#         def browse_output_path():
#             filename = filedialog.asksaveasfilename(defaultextension=".docx",
#                                                     filetypes=[("Word文档", "*.docx")])
#             if filename:
#                 output_path_entry.delete(0, tk.END)
#                 output_path_entry.insert(0, filename)
#         tk.Button(root, text="浏览...", command=browse_output_path).grid(row=3, column=2)
#
#         # 最大坡度
#         tk.Label(root, text="最大爬坡度 (默认25.0度):").grid(row=4, column=0, sticky='e')
#         max_slope_entry = tk.Entry(root, width=50)
#         max_slope_entry.grid(row=4, column=1)
#         max_slope_entry.insert(0, "25.0")  # 设置默认值
#
#         # 最大阶跃高程差
#         tk.Label(root, text="最大阶跃高程差 (默认10.0米):").grid(row=5, column=0, sticky='e')
#         max_step_elev_entry = tk.Entry(root, width=50)
#         max_step_elev_entry.grid(row=5, column=1)
#         max_step_elev_entry.insert(0, "10.0")  # 设置默认值
#
#         # 最大速度
#         tk.Label(root, text="最大速度 (默认20.0 m/s):").grid(row=6, column=0, sticky='e')
#         max_speed_entry = tk.Entry(root, width=50)
#         max_speed_entry.grid(row=6, column=1)
#         max_speed_entry.insert(0, "20.0")  # 设置默认值
#
#         # 运行按钮
#         def run():
#             # 获取输入值
#             csv_path = map_data_path_entry.get()
#             start_coord_input = start_coord_entry.get()
#             end_coord_input = end_coord_entry.get()
#             output_path = output_path_entry.get()
#
#             # 验证输入
#             if not csv_path:
#                 messagebox.showerror("错误", "请提供地图数据文件路径。")
#                 return
#             if not output_path:
#                 messagebox.showerror("错误", "请提供输出 Word 文档路径。")
#                 return
#             if not start_coord_input:
#                 messagebox.showerror("错误", "请提供起点坐标。")
#                 return
#             if not end_coord_input:
#                 messagebox.showerror("错误", "请提供终点坐标。")
#                 return
#             try:
#                 start_x, start_y = map(float, start_coord_input.strip().split(','))
#                 start_coord = (start_x, start_y)
#             except ValueError:
#                 messagebox.showerror("错误", "起点坐标格式错误，请使用 'X,Y' 的格式。")
#                 return
#
#             try:
#                 end_x, end_y = map(float, end_coord_input.strip().split(','))
#                 end_coord = (end_x, end_y)
#             except ValueError:
#                 messagebox.showerror("错误", "终点坐标格式错误，请使用 'X,Y' 的格式。")
#                 return
#
#             # 获取最大坡度和最大阶跃高程差
#             max_slope_input = max_slope_entry.get()
#             max_step_elev_input = max_step_elev_entry.get()
#
#             try:
#                 max_allowed_slope = float(max_slope_input)
#             except ValueError:
#                 messagebox.showerror("错误", "最大坡度格式错误，请输入数字。")
#                 return
#
#             try:
#                 max_allowed_step_elev = float(max_step_elev_input)
#             except ValueError:
#                 messagebox.showerror("错误", "最大阶跃高程差格式错误，请输入数字。")
#                 return
#
#             # 获取最大速度
#             max_speed_input = max_speed_entry.get()
#
#             try:
#                 max_speed = float(max_speed_input)
#             except ValueError:
#                 messagebox.showerror("错误", "最大速度格式错误，请输入数字。")
#                 return
#
#             # 在后台线程中运行，以避免阻塞GUI
#             def process():
#                 try:
#                     # 调用处理函数，传入最大坡度、最大阶跃高程差和最大速度
#                     run_processing(csv_path, start_coord, end_coord, output_path, max_allowed_slope, max_allowed_step_elev, max_speed)
#                     messagebox.showinfo("完成", "路径规划已完成。")
#                 except Exception as e:
#                     messagebox.showerror("错误", f"发生错误: {e}")
#                     logging.exception("后台线程运行时发生错误")
#
#             threading.Thread(target=process).start()
#
#         tk.Button(root, text="运行", command=run).grid(row=7, column=1)
#
#         root.mainloop()
#     except Exception as e:
#         print(f"主程序运行时发生错误: {e}")
#         logging.exception("主程序运行时发生错误")
#
# if __name__ == "__main__":
#     main()

















#
# import pandas as pd
# import numpy as np
# from scipy.interpolate import griddata
# from sklearn.metrics import mean_squared_error, mean_absolute_error, accuracy_score
# from sklearn.model_selection import train_test_split
# import heapq
# import math
# import plotly.graph_objs as go
# import plotly.io as pio
# from docx import Document
# import tkinter as tk
# from tkinter import filedialog, messagebox
# import threading
# import sys
# import os
# import logging
#
# # 设置日志记录
# logging.basicConfig(filename='error.log', level=logging.ERROR)
#
# # 设置 Plotly 渲染器为浏览器
# pio.renderers.default = 'browser'
#
# # 添加 resource_path 函数，用于获取资源文件的正确路径
# def resource_path(relative_path):
#     """获取资源文件的绝对路径，兼容 PyInstaller 打包后的情况"""
#     try:
#         # PyInstaller 会把路径存储在 _MEIPASS 中
#         base_path = sys._MEIPASS
#     except AttributeError:
#         base_path = os.path.abspath(".")
#     return os.path.join(base_path, relative_path)
#
# # 数据处理类
# class DataProcessor:
#     def __init__(self, dat_path, x_resolution=1, y_resolution=1):
#         self.dat_path = dat_path
#         self.x_resolution = x_resolution
#         self.y_resolution = y_resolution
#         self.df = None
#         self.X = None
#         self.Y = None
#         self.Z = None
#         self.TrafFlag = None
#         self.maxAcc = None
#         self.Slope = None
#         self.StepElev = None
#         self.Roughness = None
#         self.x_range = None
#         self.y_range = None
#         self.X_grid = None
#         self.Y_grid = None
#         self.Z_grid = None
#         self.TrafFlag_grid = None
#         self.Acc_grid = None
#         self.Slope_grid = None
#         self.StepElev_grid = None
#         self.Rough_grid = None
#
#     def load_data(self):
#         try:
#             # 使用 resource_path 函数获取数据文件路径
#             data_path = resource_path(self.dat_path)
#             # 读取 .csv 文件，指定分隔符为逗号
#             self.df = pd.read_csv(data_path, delimiter=',')
#             print("数据文件中的列名：", self.df.columns.tolist())
#         except FileNotFoundError:
#             print(f"文件未找到: {self.dat_path}")
#             return
#         except pd.errors.EmptyDataError:
#             print("CSV 文件为空")
#             return
#         except Exception as e:
#             print(f"读取 CSV 文件时出错: {e}")
#             logging.exception("读取 CSV 文件时出错")
#             return
#
#         required_columns = {'X', 'Y', 'Z', 'TrafFlag', 'maxAcc', 'Slope', 'StepElevationDifference', 'Roughness'}
#         if not required_columns.issubset(self.df.columns):
#             print(f"CSV文件缺少必要的列: {required_columns - set(self.df.columns)}")
#             return
#
#         # 提取所需数据
#         self.X = self.df['X'].values
#         self.Y = self.df['Y'].values
#         self.Z = self.df['Z'].values
#         self.TrafFlag = self.df['TrafFlag'].values
#         self.maxAcc = self.df['maxAcc'].values
#         self.Slope = self.df['Slope'].values
#         self.StepElev = self.df['StepElevationDifference'].values
#         self.Roughness = self.df['Roughness'].values
#
#         # 计算并打印 Slope 和 StepElevationDifference 的最小值和最大值
#         slope_min = np.min(self.Slope)
#         slope_max_val = np.max(self.Slope)
#         step_min = np.min(self.StepElev)
#         step_max = np.max(self.StepElev)
#
#         print(f"坡度 (Slope) 最小值: {slope_min}")
#         print(f"坡度 (Slope) 最大值: {slope_max_val}")
#         print(f"阶跃高程差 (StepElevationDifference) 最小值: {step_min}")
#         print(f"阶跃高程差 (StepElevationDifference) 最大值: {step_max}")
#
#     def interpolate_data(self):
#         # 计算网格范围和生成网格
#         x_min, x_max = self.X.min(), self.X.max()
#         y_min, y_max = self.Y.min(), self.Y.max()
#         self.x_range = np.arange(x_min, x_max + self.x_resolution, self.x_resolution)
#         self.y_range = np.arange(y_min, y_max + self.y_resolution, self.y_resolution)
#         self.X_grid, self.Y_grid = np.meshgrid(self.x_range, self.y_range)
#
#         # 插值
#         self.Z_grid = griddata((self.X, self.Y), self.Z, (self.X_grid, self.Y_grid), method='cubic')
#         self.TrafFlag_grid = griddata((self.X, self.Y), self.TrafFlag, (self.X_grid, self.Y_grid), method='nearest')
#         self.Acc_grid = griddata((self.X, self.Y), self.maxAcc, (self.X_grid, self.Y_grid), method='nearest')
#         self.Slope_grid = griddata((self.X, self.Y), self.Slope, (self.X_grid, self.Y_grid), method='nearest')
#         self.StepElev_grid = griddata((self.X, self.Y), self.StepElev, (self.X_grid, self.Y_grid), method='nearest')
#         self.Rough_grid = griddata((self.X, self.Y), self.Roughness, (self.X_grid, self.Y_grid), method='nearest')
#
#     def fill_nan_values(self):
#         # 填充 NaN 值
#         self.Z_grid = np.nan_to_num(self.Z_grid, nan=np.nanmean(self.Z))
#         self.Acc_grid = np.nan_to_num(self.Acc_grid, nan=0)
#         self.TrafFlag_grid = np.nan_to_num(self.TrafFlag_grid, nan=0)
#         # 将 TrafFlag_grid 的值限制为 0 或 1
#         self.TrafFlag_grid = np.where(self.TrafFlag_grid >= 0.5, 1, 0)
#         self.Slope_grid = np.nan_to_num(self.Slope_grid, nan=0)
#         self.StepElev_grid = np.nan_to_num(self.StepElev_grid, nan=0)
#         self.Rough_grid = np.nan_to_num(self.Rough_grid, nan=0)
#
#         # 验证网格数据的形状
#         print(f"Z_grid shape: {self.Z_grid.shape}")
#         print(f"TrafFlag_grid shape: {self.TrafFlag_grid.shape}")
#         print(f"Acc_grid shape: {self.Acc_grid.shape}")
#         print(f"Slope_grid shape: {self.Slope_grid.shape}")
#         print(f"StepElev_grid shape: {self.StepElev_grid.shape}")
#         print(f"Rough_grid shape: {self.Rough_grid.shape}")
#
#         # 计算并打印插值后网格数据的最小值和最大值
#         slope_grid_min = np.min(self.Slope_grid)
#         slope_grid_max = np.max(self.Slope_grid)
#         step_grid_min = np.min(self.StepElev_grid)
#         step_grid_max = np.max(self.StepElev_grid)
#
#         print(f"插值后坡度网格 (Slope_grid) 最小值: {slope_grid_min}")
#         print(f"插值后坡度网格 (Slope_grid) 最大值: {slope_grid_max}")
#         print(f"插值后阶跃高程差网格 (StepElev_grid) 最小值: {step_grid_min}")
#         print(f"插值后阶跃高程差网格 (StepElev_grid) 最大值: {step_grid_max}")
#
#     def coord_to_index(self, coord):
#         x_min, y_min = self.x_range.min(), self.y_range.min()
#         col = int((coord[0] - x_min) / self.x_resolution)
#         row = int((coord[1] - y_min) / self.y_resolution)
#         return (row, col)
#
#     def is_valid_position(self, pos):
#         return 0 <= pos[0] < self.TrafFlag_grid.shape[0] and 0 <= pos[1] < self.TrafFlag_grid.shape[1]
#
#     def split_data_and_evaluate(self):
#         try:
#             # 将所有数据合并
#             data = np.column_stack(
#                 (self.X, self.Y, self.Z, self.TrafFlag, self.maxAcc, self.Slope, self.StepElev, self.Roughness))
#
#             # 分割数据集为训练集和测试集
#             test_size = 0.2
#             random_state = 42
#
#             train_data, test_data = train_test_split(data, test_size=test_size, random_state=random_state)
#
#             # 拆分训练集和测试集的数据
#             X_train, Y_train = train_data[:, 0], train_data[:, 1]
#             Z_train = train_data[:, 2]
#             TrafFlag_train = train_data[:, 3]
#             maxAcc_train = train_data[:, 4]
#             Slope_train = train_data[:, 5]
#             StepElev_train = train_data[:, 6]
#             Rough_train = train_data[:, 7]
#
#             X_test, Y_test = test_data[:, 0], test_data[:, 1]
#             Z_test = test_data[:, 2]
#             TrafFlag_test = test_data[:, 3]
#             maxAcc_test = test_data[:, 4]
#             Slope_test = test_data[:, 5]
#             StepElev_test = test_data[:, 6]
#             Rough_test = test_data[:, 7]
#
#             # 使用训练集进行插值
#             Z_grid_train = griddata((X_train, Y_train), Z_train, (self.X_grid, self.Y_grid), method='cubic')
#             TrafFlag_grid_train = griddata((X_train, Y_train), TrafFlag_train, (self.X_grid, self.Y_grid), method='nearest')
#             Acc_grid_train = griddata((X_train, Y_train), maxAcc_train, (self.X_grid, self.Y_grid), method='nearest')
#             Slope_grid_train = griddata((X_train, Y_train), Slope_train, (self.X_grid, self.Y_grid), method='nearest')
#             StepElev_grid_train = griddata((X_train, Y_train), StepElev_train, (self.X_grid, self.Y_grid), method='nearest')
#             Rough_grid_train = griddata((X_train, Y_train), Rough_train, (self.X_grid, self.Y_grid), method='nearest')
#
#             # 在测试集上预测值
#             Z_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Z_grid_train.flatten(), (X_test, Y_test),
#                               method='cubic')
#             TrafFlag_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), TrafFlag_grid_train.flatten(),
#                                      (X_test, Y_test), method='nearest')
#             Acc_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Acc_grid_train.flatten(), (X_test, Y_test),
#                                 method='nearest')
#             Slope_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Slope_grid_train.flatten(),
#                                   (X_test, Y_test), method='nearest')
#             StepElev_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), StepElev_grid_train.flatten(),
#                                      (X_test, Y_test), method='nearest')
#             Rough_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Rough_grid_train.flatten(),
#                                   (X_test, Y_test), method='nearest')
#
#             # 处理 NaN 值并计算误差指标
#
#             # 高程数据误差
#             valid_indices_Z = ~np.isnan(Z_pred) & ~np.isnan(Z_test)
#             if np.any(valid_indices_Z):
#                 mse_Z = mean_squared_error(Z_test[valid_indices_Z], Z_pred[valid_indices_Z])
#                 mae_Z = mean_absolute_error(Z_test[valid_indices_Z], Z_pred[valid_indices_Z])
#                 max_error_Z = np.max(np.abs(Z_test[valid_indices_Z] - Z_pred[valid_indices_Z]))
#                 print(f"高程插值 - MSE: {mse_Z}, MAE: {mae_Z}, Max Error: {max_error_Z}")
#             else:
#                 print("高程插值预测值中全部为 NaN，无法计算误差。")
#
#             # 通行标志（分类准确率）
#             valid_indices_TrafFlag = ~np.isnan(TrafFlag_pred) & ~np.isnan(TrafFlag_test)
#             if np.any(valid_indices_TrafFlag):
#                 accuracy_TrafFlag = accuracy_score(TrafFlag_test[valid_indices_TrafFlag],
#                                                    TrafFlag_pred[valid_indices_TrafFlag])
#                 print(f"通行标志插值 - 准确率: {accuracy_TrafFlag}")
#             else:
#                 print("通行标志插值预测值中全部为 NaN，无法计算准确率。")
#
#             # 最大加速度误差
#             valid_indices_Acc = ~np.isnan(Acc_pred) & ~np.isnan(maxAcc_test)
#             if np.any(valid_indices_Acc):
#                 mse_Acc = mean_squared_error(maxAcc_test[valid_indices_Acc], Acc_pred[valid_indices_Acc])
#                 mae_Acc = mean_absolute_error(maxAcc_test[valid_indices_Acc], Acc_pred[valid_indices_Acc])
#                 max_error_Acc = np.max(np.abs(maxAcc_test[valid_indices_Acc] - Acc_pred[valid_indices_Acc]))
#                 print(f"最大加速度插值 - MSE: {mse_Acc}, MAE: {mae_Acc}, Max Error: {max_error_Acc}")
#             else:
#                 print("最大加速度插值预测值中全部为 NaN，无法计算误差。")
#
#             # 坡度误差
#             valid_indices_Slope = ~np.isnan(Slope_pred) & ~np.isnan(Slope_test)
#             if np.any(valid_indices_Slope):
#                 mse_Slope = mean_squared_error(Slope_test[valid_indices_Slope], Slope_pred[valid_indices_Slope])
#                 mae_Slope = mean_absolute_error(Slope_test[valid_indices_Slope], Slope_pred[valid_indices_Slope])
#                 max_error_Slope = np.max(np.abs(Slope_test[valid_indices_Slope] - Slope_pred[valid_indices_Slope]))
#                 print(f"坡度插值 - MSE: {mse_Slope}, MAE: {mae_Slope}, Max Error: {max_error_Slope}")
#             else:
#                 print("坡度插值预测值中全部为 NaN，无法计算误差。")
#
#             # 阶跃高程差误差
#             valid_indices_StepElev = ~np.isnan(StepElev_pred) & ~np.isnan(StepElev_test)
#             if np.any(valid_indices_StepElev):
#                 mse_StepElev = mean_squared_error(StepElev_test[valid_indices_StepElev],
#                                                   StepElev_pred[valid_indices_StepElev])
#                 mae_StepElev = mean_absolute_error(StepElev_test[valid_indices_StepElev],
#                                                    StepElev_pred[valid_indices_StepElev])
#                 max_error_StepElev = np.max(
#                     np.abs(StepElev_test[valid_indices_StepElev] - StepElev_pred[valid_indices_StepElev]))
#                 print(f"阶跃高程差插值 - MSE: {mse_StepElev}, MAE: {mae_StepElev}, Max Error: {max_error_StepElev}")
#             else:
#                 print("阶跃高程差插值预测值中全部为 NaN，无法计算误差。")
#
#             # 粗糙度误差
#             valid_indices_Rough = ~np.isnan(Rough_pred) & ~np.isnan(Rough_test)
#             if np.any(valid_indices_Rough):
#                 mse_Rough = mean_squared_error(Rough_test[valid_indices_Rough], Rough_pred[valid_indices_Rough])
#                 mae_Rough = mean_absolute_error(Rough_test[valid_indices_Rough], Rough_pred[valid_indices_Rough])
#                 max_error_Rough = np.max(np.abs(Rough_test[valid_indices_Rough] - Rough_pred[valid_indices_Rough]))
#                 print(f"粗糙度插值 - MSE: {mse_Rough}, MAE: {mae_Rough}, Max Error: {max_error_Rough}")
#             else:
#                 print("粗糙度插值预测值中全部为 NaN，无法计算误差。")
#         except Exception as e:
#             print(f"数据评估时发生错误: {e}")
#             logging.exception("数据评估时发生错误")
#
# # 节点类和路径规划类
# class Node:
#     def __init__(self, position, parent=None):
#         self.position = position  # (row, col)
#         self.parent = parent
#         self.g = 0  # 累计代价
#         self.h = 0  # 启发代价
#         self.f = 0  # 总代价
#
#     def __lt__(self, other):
#         return self.f < other.f
#
# class PathPlanner:
#     def __init__(self, grid_data, max_slope=None, max_step_elev=None, max_speed=70.0):
#         self.grid = grid_data['TrafFlag_grid']
#         self.Z_grid = grid_data['Z_grid']
#         self.Acc_grid = grid_data['Acc_grid']
#         self.Slope_grid = grid_data['Slope_grid']
#         self.StepElev_grid = grid_data['StepElev_grid']
#         self.Rough_grid = grid_data['Rough_grid']
#         self.max_slope = max_slope
#         self.max_step_elev = max_step_elev
#         self.max_speed = max_speed
#
#     def heuristic_distance_3d(self, current_pos, end_pos):
#         d_xy = np.sqrt((current_pos[0] - end_pos[0]) ** 2 +
#                        (current_pos[1] - end_pos[1]) ** 2)
#         z_current = self.Z_grid[current_pos[0], current_pos[1]]
#         z_end = self.Z_grid[end_pos[0], end_pos[1]]
#         d_z = (z_end - z_current) ** 2
#         return np.sqrt(d_xy ** 2 + d_z)
#
#     def astar(self, start, end, mode='distance', alpha_env=0.6, beta_env=0.4):
#         start_node = Node(start)
#         end_node = Node(end)
#
#         open_list = []
#         heapq.heappush(open_list, start_node)
#         open_dict = {start_node.position: start_node}
#
#         closed_set = set()
#
#         while open_list:
#             current_node = heapq.heappop(open_list)
#             del open_dict[current_node.position]
#             closed_set.add(current_node.position)
#
#             # 检查是否到达终点
#             if current_node.position == end_node.position:
#                 path = []
#                 while current_node:
#                     path.append(current_node.position)
#                     current_node = current_node.parent
#                 return path[::-1]
#
#             # 定义8个邻居（包括对角线）
#             neighbors = [(0, 1), (0, -1), (1, 0), (-1, 0),
#                          (1, 1), (1, -1), (-1, 1), (-1, -1)]
#
#             for offset in neighbors:
#                 neighbor_pos = (current_node.position[0] + offset[0],
#                                 current_node.position[1] + offset[1])
#
#                 # 检查邻居是否在网格范围内
#                 if not (0 <= neighbor_pos[0] < self.grid.shape[0] and
#                         0 <= neighbor_pos[1] < self.grid.shape[1]):
#                     continue
#
#                 # 检查 TrafFlag 是否可通行
#                 if self.grid[neighbor_pos] == 0:
#                     continue  # 不可通行区域
#
#                 # 在所有模式下应用坡度和阶跃高程差的限制
#                 if self.max_slope is not None and self.Slope_grid[neighbor_pos] > self.max_slope:
#                     continue  # 坡度超过最大限制
#                 if self.max_step_elev is not None and self.StepElev_grid[neighbor_pos] > self.max_step_elev:
#                     continue  # 阶跃高程差超过最大限制
#
#                 # 检查是否已经在关闭列表中
#                 if neighbor_pos in closed_set:
#                     continue
#
#                 # 处理不同模式
#                 if mode == 'distance':
#                     movement_cost = self.heuristic_distance_3d(current_node.position, neighbor_pos)
#                     g_cost = current_node.g + movement_cost
#                     h_cost = self.heuristic_distance_3d(neighbor_pos, end)
#                 elif mode == 'time':
#                     a_max = self.Acc_grid[neighbor_pos[0], neighbor_pos[1]]
#                     if a_max <= 0:
#                         continue  # 无法加速，跳过
#                     movement_cost = self.heuristic_distance_3d(current_node.position, neighbor_pos)
#                     g_cost = current_node.g + movement_cost / a_max  # 简化为距离除以加速度
#                     h_cost = self.heuristic_distance_3d(neighbor_pos, end)
#                 elif mode == 'env':
#                     movement_uncertainty = (self.Slope_grid[neighbor_pos[0], neighbor_pos[1]] / 45.0 +
#                                             self.StepElev_grid[neighbor_pos[0], neighbor_pos[1]] / 1.5 +
#                                             self.Rough_grid[neighbor_pos[0], neighbor_pos[1]] / 6.0)
#                     g_cost = current_node.g + movement_uncertainty
#                     h_cost = self.heuristic_distance_3d(neighbor_pos, end)
#                 else:
#                     raise ValueError("Invalid mode. Choose from 'distance', 'time', 'env'.")
#
#                 # 创建邻居节点
#                 neighbor_node = Node(neighbor_pos, current_node)
#                 neighbor_node.g = g_cost
#                 neighbor_node.h = h_cost
#                 neighbor_node.f = neighbor_node.g + neighbor_node.h
#
#                 # 检查开放列表中是否有更好的节点
#                 if neighbor_pos in open_dict:
#                     existing_node = open_dict[neighbor_pos]
#                     if existing_node.g <= neighbor_node.g:
#                         continue  # 已有更优节点
#                     else:
#                         # 替换为更优节点
#                         open_list.remove(existing_node)
#                         heapq.heapify(open_list)
#                         heapq.heappush(open_list, neighbor_node)
#                         open_dict[neighbor_pos] = neighbor_node
#                 else:
#                     heapq.heappush(open_list, neighbor_node)
#                     open_dict[neighbor_pos] = neighbor_node
#
#         return None  # 如果找不到路径，返回 None
#
#     def extract_path_coordinates(self, path, x_range, y_range):
#         return [(x_range[p[1]], y_range[p[0]], self.Z_grid[p[0], p[1]]) for p in path]
#
#     def calculate_path_length(self, path_coords):
#         length = 0
#         for i in range(1, len(path_coords)):
#             x1, y1, z1 = path_coords[i - 1]
#             x2, y2, z2 = path_coords[i]
#             length += np.sqrt((x2 - x1) ** 2 + (y2 - y1) ** 2 + (z2 - z1) ** 2)
#         return length
#
#     def compute_traversal_time(self, path_coords, y_range, x_range):
#         total_time = 0.0
#         current_speed = 0.0  # 初始速度为0
#         max_speed = self.max_speed  # 使用用户指定的最大速度
#
#         speeds = [current_speed]  # 存储每个点的速度
#         accelerations = []  # 存储每段的加速度
#
#         for i in range(1, len(path_coords)):
#             x1, y1, z1 = path_coords[i - 1]
#             x2, y2, z2 = path_coords[i]
#             # 计算两点之间的距离
#             d = np.sqrt((x2 - x1) ** 2 + (y2 - y1) ** 2 + (z2 - z1) ** 2)
#
#             # 获取当前点的网格索引
#             row = np.argmin(np.abs(y_range - y1))
#             col = np.argmin(np.abs(x_range - x1))
#
#             # 获取当前点的加速度
#             a = self.Acc_grid[row, col]
#             if a <= 0:
#                 # 如果加速度不可用或为0，无法继续加速
#                 a = 0.0
#
#             # 计算需要加速到最大速度的速度增量
#             delta_v = max_speed - current_speed
#
#             if current_speed < max_speed and a > 0:
#                 # 需要加速
#                 t_acc = delta_v / a  # 加速所需时间
#                 d_acc = current_speed * t_acc + 0.5 * a * t_acc ** 2  # 加速所需距离
#
#                 if d_acc >= d:
#                     # 在当前段无法加速到最大速度
#                     # 使用运动学方程计算通过距离d所需的时间
#                     # 解方程: 0.5 * a * t^2 + current_speed * t - d = 0
#                     discriminant = current_speed ** 2 + 2 * a * d
#                     if discriminant < 0:
#                         # 无实数解，无法到达下一个点
#                         return float('inf'), [], []
#                     t = (-current_speed + math.sqrt(discriminant)) / a
#                     total_time += t
#                     current_speed += a * t
#                     accelerations.append(a)
#                     speeds.append(current_speed)
#                 else:
#                     # 在当前段可以加速到最大速度
#                     # 先加速到最大速度
#                     total_time += t_acc
#                     current_speed = max_speed
#                     accelerations.append(a)
#                     speeds.append(current_speed)
#                     # 剩余距离
#                     d_remain = d - d_acc
#                     # 以最大速度匀速行驶剩余距离
#                     t_remain = d_remain / current_speed
#                     total_time += t_remain
#                     accelerations.append(0.0)
#                     speeds.append(current_speed)
#             else:
#                 # 无法加速或已达到最大速度，匀速行驶
#                 if current_speed == 0:
#                     # 无法移动
#                     return float('inf'), [], []
#                 t = d / current_speed
#                 total_time += t
#                 accelerations.append(0.0)
#                 speeds.append(current_speed)
#
#         return total_time, speeds, accelerations
#
# # 可视化类
# class Visualizer:
#     def __init__(self, X_grid, Y_grid, Z_grid, TrafFlag_grid):
#         self.X_grid = X_grid
#         self.Y_grid = Y_grid
#         self.Z_grid = Z_grid
#         self.TrafFlag_grid = TrafFlag_grid
#         self.fig = go.Figure()
#
#     def plot_terrain(self):
#         # 定义自定义颜色尺度
#         colorscale = [
#             [0.0, 'gray'],        # TrafFlag == 0 的区域为灰色
#             [1.0, 'lightgreen']   # TrafFlag == 1 的区域为浅绿色
#         ]
#
#         self.fig.add_trace(go.Surface(
#             x=self.X_grid, y=self.Y_grid, z=self.Z_grid,
#             surfacecolor=self.TrafFlag_grid,
#             colorscale=colorscale,
#             cmin=0,
#             cmax=1,
#             opacity=0.8,
#             showscale=True,
#             colorbar=dict(title='TrafFlag')
#         ))
#
#     def add_path(self, coords, speeds, accs, color, name, length, time_taken, path_width=4):
#         if coords:
#             x, y, z = zip(*coords)
#             # 处理不可行路径（时间为inf）的情况
#             if math.isinf(time_taken):
#                 time_str = "无法到达"
#             else:
#                 time_str = f"{time_taken:.2f} 秒"
#
#             # 准备 customdata，将速度和加速度添加进去
#             if accs:
#                 accs = [0.0] + accs  # 与 speeds 对齐
#             else:
#                 accs = [0.0] * len(speeds)
#
#             customdata = np.array([speeds, accs]).T  # 转置使其每行对应一个点的 [speed, acceleration]
#
#             self.fig.add_trace(go.Scatter3d(
#                 x=x, y=y, z=z,
#                 mode='lines+markers',
#                 line=dict(color=color, width=path_width),
#                 marker=dict(size=2, color=color),
#                 name=f"{name} (长度: {length:.2f} 米, 时间: {time_str})",
#                 customdata=customdata,
#                 hovertemplate=(
#                         'X: %{x}<br>' +
#                         'Y: %{y}<br>' +
#                         'Z: %{z}<br>' +
#                         '速度: %{customdata[0]:.2f} m/s<br>' +
#                         '加速度: %{customdata[1]:.2f} m/s²<br>' +
#                         '<extra></extra>'
#                 )
#             ))
#         else:
#             print(f"未找到{name}")
#
#     def add_start_and_end_points(self, start_coord, end_coord):
#         self.fig.add_trace(go.Scatter3d(
#             x=[start_coord[0]],
#             y=[start_coord[1]],
#             z=[start_coord[2]],
#             mode='markers',
#             marker=dict(size=8, color='green'),
#             name='起点'
#         ))
#         self.fig.add_trace(go.Scatter3d(
#             x=[end_coord[0]],
#             y=[end_coord[1]],
#             z=[end_coord[2]],
#             mode='markers',
#             marker=dict(size=8, color='red'),
#             name='终点'
#         ))
#
#     def show(self):
#         self.fig.update_layout(
#             scene=dict(
#                 xaxis_title='X 坐标',
#                 yaxis_title='Y 坐标',
#                 zaxis_title='Z 坐标 (高程)',
#                 aspectmode='data'
#             ),
#             title='3D 地形路径规划（基于距离、时间与最小不确定性）',
#             legend=dict(
#                 x=0,
#                 y=1,
#                 bgcolor='rgba(255, 255, 255, 0)',
#                 bordercolor='rgba(255, 255, 255, 0)'
#             )
#         )
#         pio.show(self.fig)
#
# # 报告生成函数
# def write_paths_to_word(document, path_name, path_coords):
#     if not path_coords:
#         return
#     # 添加路径名称和点数量
#     document.add_heading(f"{path_name} (点数: {len(path_coords)})", level=2)
#     # 创建表格，包含三列：X相对坐标、Y相对坐标、Z高程坐标
#     table = document.add_table(rows=1, cols=3)
#     table.style = 'Light Grid'
#     hdr_cells = table.rows[0].cells
#     hdr_cells[0].text = 'X相对坐标'
#     hdr_cells[1].text = 'Y相对坐标'
#     hdr_cells[2].text = 'Z高程坐标'
#     # 填充表格内容
#     for coord in path_coords:
#         row_cells = table.add_row().cells
#         row_cells[0].text = f"{coord[0]}"
#         row_cells[1].text = f"{coord[1]}"
#         row_cells[2].text = f"{coord[2]}"
#     # 添加一个空行
#     document.add_paragraph()
#
# # 路径规划函数
# def run_path_planning(data_processor, start_coord, end_coord, output_path, max_allowed_slope, max_allowed_step_elev, max_speed):
#     try:
#         # 确保已进行数据评估
#         data_processor.split_data_and_evaluate()
#
#         # 将 start_coord 和 end_coord 转换为网格索引
#         start = data_processor.coord_to_index(start_coord)
#         end = data_processor.coord_to_index(end_coord)
#
#         # 检查起点和终点是否在网格范围内
#         if not data_processor.is_valid_position(start):
#             raise ValueError("起点坐标超出网格范围")
#         if not data_processor.is_valid_position(end):
#             raise ValueError("终点坐标超出网格范围")
#
#         # 确保起点和终点可通行
#         if data_processor.TrafFlag_grid[start] == 0:
#             raise ValueError("起点不可通行")
#         if data_processor.TrafFlag_grid[end] == 0:
#             raise ValueError("终点不可通行")
#
#         # 路径规划
#         grid_data = {
#             'TrafFlag_grid': data_processor.TrafFlag_grid,
#             'Z_grid': data_processor.Z_grid,
#             'Acc_grid': data_processor.Acc_grid,
#             'Slope_grid': data_processor.Slope_grid,
#             'StepElev_grid': data_processor.StepElev_grid,
#             'Rough_grid': data_processor.Rough_grid
#         }
#         path_planner = PathPlanner(grid_data, max_slope=max_allowed_slope, max_step_elev=max_allowed_step_elev, max_speed=max_speed)
#
#         print("正在计算基于距离的最短路径...")
#         path_distance = path_planner.astar(start, end, mode='distance')
#         if path_distance:
#             print("基于距离的路径计算完成")
#         else:
#             print("未找到基于距离的路径")
#
#         print("正在计算基于时间的最短耗时路径...")
#         path_time = path_planner.astar(start, end, mode='time')
#         if path_time:
#             print("基于时间的路径计算完成")
#         else:
#             print("未找到基于时间的路径")
#
#         print("正在计算最小不确定性路径...")
#         path_env = path_planner.astar(start, end, mode='env')
#         if path_env:
#             print("最小不确定性路径计算完成")
#         else:
#             print("未找到最小不确定性路径")
#
#         # 提取路径坐标
#         if path_distance:
#             path_distance_coords = path_planner.extract_path_coordinates(path_distance, data_processor.x_range,
#                                                                          data_processor.y_range)
#             length_distance = path_planner.calculate_path_length(path_distance_coords)
#         else:
#             path_distance_coords = None
#             length_distance = None
#         if path_time:
#             path_time_coords = path_planner.extract_path_coordinates(path_time, data_processor.x_range,
#                                                                      data_processor.y_range)
#             length_time = path_planner.calculate_path_length(path_time_coords)
#         else:
#             path_time_coords = None
#             length_time = None
#         if path_env:
#             path_env_coords = path_planner.extract_path_coordinates(path_env, data_processor.x_range,
#                                                                     data_processor.y_range)
#             length_env = path_planner.calculate_path_length(path_env_coords)
#         else:
#             path_env_coords = None
#             length_env = None
#
#         # 计算遍历时间，并获取速度和加速度数据
#         if path_distance:
#             time_distance, speeds_distance, accs_distance = path_planner.compute_traversal_time(
#                 path_distance_coords, data_processor.y_range, data_processor.x_range)
#         else:
#             speeds_distance, accs_distance, time_distance = None, None, None
#
#         if path_time:
#             time_time, speeds_time, accs_time = path_planner.compute_traversal_time(
#                 path_time_coords, data_processor.y_range, data_processor.x_range)
#         else:
#             speeds_time, accs_time, time_time = None, None, None
#
#         if path_env:
#             time_env, speeds_env, accs_env = path_planner.compute_traversal_time(
#                 path_env_coords, data_processor.y_range, data_processor.x_range)
#         else:
#             speeds_env, accs_env, time_env = None, None, None
#
#         # 可视化
#         visualizer = Visualizer(
#             data_processor.X_grid,
#             data_processor.Y_grid,
#             data_processor.Z_grid,
#             data_processor.TrafFlag_grid
#         )
#         visualizer.plot_terrain()
#
#         colors = ['blue', 'orange', 'purple']
#         names = ['MinDistPath 最短距离路径', 'MinTimePath 最短耗时路径', 'MinUncertPath 最小不确定性路径']
#         lengths = [length_distance, length_time, length_env]
#         times = [time_distance, time_time, time_env]
#         coords_list = [path_distance_coords, path_time_coords, path_env_coords]
#         speeds_list = [speeds_distance, speeds_time, speeds_env]
#         accs_list = [accs_distance, accs_time, accs_env]
#
#         for coords, speeds, accs, color, name, length, time_taken in zip(
#                 coords_list, speeds_list, accs_list,
#                 colors, names, lengths, times):
#             visualizer.add_path(coords, speeds, accs, color, name, length, time_taken)
#
#         # 标记起点和终点
#         if path_distance or path_time or path_env:
#             selected_path = None
#             if path_distance:
#                 selected_path = path_distance_coords
#             elif path_time:
#                 selected_path = path_time_coords
#             elif path_env:
#                 selected_path = path_env_coords
#
#             if selected_path:
#                 # 使用起点坐标
#                 start_x, start_y, start_z = selected_path[0]
#                 # 使用终点坐标
#                 end_x, end_y, end_z = selected_path[-1]
#                 visualizer.add_start_and_end_points((start_x, start_y, start_z), (end_x, end_y, end_z))
#
#         visualizer.show()
#
#         # 创建 Word 文档
#         document = Document()
#         document.add_heading('路径规划结果', 0)
#         # 写入路径信息
#         if path_distance or path_time or path_env:
#             if path_distance:
#                 write_paths_to_word(document, 'MinDistPath ', path_distance_coords)
#             if path_time:
#                 write_paths_to_word(document, 'MinTimePath ', path_time_coords)
#             if path_env:
#                 write_paths_to_word(document, 'MinUncertPath ', path_env_coords)
#
#             # 保存文档
#             try:
#                 # 使用 resource_path 函数获取输出文件路径
#                 output_file_path = output_path  # 如果 output_path 是用户指定的完整路径，则无需使用 resource_path
#                 document.save(output_file_path)
#                 print(f"路径坐标已成功保存到 Word 文档: {output_file_path}")
#             except Exception as e:
#                 print(f"保存 Word 文档时出错: {e}")
#                 logging.exception("保存 Word 文档时出错")
#         else:
#             print("没有路径可写入 Word 文档")
#
#         # 打印路径长度和遍历时间
#         if path_distance:
#             if math.isinf(time_distance):
#                 print(f"最短路径长度: {length_distance:.2f} 米")
#                 print(f"最短距离路径遍历时间: 无法到达")
#             else:
#                 print(f"最短路径长度: {length_distance:.2f} 米")
#                 print(f"最短距离路径遍历时间: {time_distance:.2f} 秒")
#         if path_time:
#             if math.isinf(time_time):
#                 print(f"最短耗时路径长度: {length_time:.2f} 米")
#                 print(f"最短耗时路径遍历时间: 无法到达")
#             else:
#                 print(f"最短耗时路径长度: {length_time:.2f} 米")
#                 print(f"最短耗时路径遍历时间: {time_time:.2f} 秒")
#         if path_env:
#             if math.isinf(time_env):
#                 print(f"最小不确定性路径长度: {length_env:.2f} 米")
#                 print(f"最小不确定性路径遍历时间: 无法到达")
#             else:
#                 print(f"最小不确定性路径长度: {length_env:.2f} 米")
#                 print(f"最小不确定性路径遍历时间: {time_env:.2f} 秒")
#     except Exception as e:
#         print(f"运行处理时发生错误: {e}")
#         logging.exception("运行处理时发生错误")
#
# # 主程序
# def main():
#     try:
#         # 创建主窗口
#         root = tk.Tk()
#         root.title("路径规划软件")
#
#         # 地图数据文件路径
#         tk.Label(root, text="地图数据文件路径:").grid(row=0, column=0, sticky='e')
#         map_data_path_entry = tk.Entry(root, width=50)
#         map_data_path_entry.grid(row=0, column=1)
#         def browse_map_data():
#             filename = filedialog.askopenfilename(filetypes=[("CSV文件", "*.csv")])
#             if filename:
#                 map_data_path_entry.delete(0, tk.END)
#                 map_data_path_entry.insert(0, filename)
#         tk.Button(root, text="浏览...", command=browse_map_data).grid(row=0, column=2)
#
#         # 起点坐标（初始禁用）
#         tk.Label(root, text="起点坐标 (格式: X,Y):").grid(row=1, column=0, sticky='e')
#         start_coord_entry = tk.Entry(root, width=50)
#         start_coord_entry.grid(row=1, column=1)
#         # 终点坐标（初始禁用）
#         tk.Label(root, text="终点坐标 (格式: X,Y):").grid(row=2, column=0, sticky='e')
#         end_coord_entry = tk.Entry(root, width=50)
#         end_coord_entry.grid(row=2, column=1)
#
#         # 输出路径
#         tk.Label(root, text="输出路径文档路径:").grid(row=3, column=0, sticky='e')
#         output_path_entry = tk.Entry(root, width=50)
#         output_path_entry.grid(row=3, column=1)
#         def browse_output_path():
#             filename = filedialog.asksaveasfilename(defaultextension=".docx",
#                                                     filetypes=[("Word文档", "*.docx")])
#             if filename:
#                 output_path_entry.delete(0, tk.END)
#                 output_path_entry.insert(0, filename)
#         tk.Button(root, text="浏览...", command=browse_output_path).grid(row=3, column=2)
#
#         # 最大坡度
#         tk.Label(root, text="最大爬坡度 (默认25.0度):").grid(row=4, column=0, sticky='e')
#         max_slope_entry = tk.Entry(root, width=50)
#         max_slope_entry.grid(row=4, column=1)
#         max_slope_entry.insert(0, "25.0")  # 设置默认值
#
#         # 最大阶跃高程差
#         tk.Label(root, text="最大阶跃高程差 (默认10.0米):").grid(row=5, column=0, sticky='e')
#         max_step_elev_entry = tk.Entry(root, width=50)
#         max_step_elev_entry.grid(row=5, column=1)
#         max_step_elev_entry.insert(0, "10.0")  # 设置默认值
#
#         # 最大速度
#         tk.Label(root, text="最大速度 (默认20.0 m/s):").grid(row=6, column=0, sticky='e')
#         max_speed_entry = tk.Entry(root, width=50)
#         max_speed_entry.grid(row=6, column=1)
#         max_speed_entry.insert(0, "20.0")  # 设置默认值
#
#         # 禁用输入字段，直到地图加载完成
#         start_coord_entry.configure(state='disabled')
#         end_coord_entry.configure(state='disabled')
#         output_path_entry.configure(state='disabled')
#         max_slope_entry.configure(state='disabled')
#         max_step_elev_entry.configure(state='disabled')
#         max_speed_entry.configure(state='disabled')
#
#         # 加载地图函数
#         def load_map():
#             csv_path = map_data_path_entry.get()
#             if not csv_path:
#                 messagebox.showerror("错误", "请提供地图数据文件路径。")
#                 return
#
#             def process_map():
#                 try:
#                     data_processor = DataProcessor(csv_path)
#                     data_processor.load_data()
#                     data_processor.interpolate_data()
#                     data_processor.fill_nan_values()
#
#                     # 可视化地图
#                     visualizer = Visualizer(
#                         data_processor.X_grid,
#                         data_processor.Y_grid,
#                         data_processor.Z_grid,
#                         data_processor.TrafFlag_grid
#                     )
#                     visualizer.plot_terrain()
#                     visualizer.show()
#
#                     # 存储 data_processor
#                     root.data_processor = data_processor
#
#                     # 启用输入字段
#                     root.after(0, enable_input_fields)
#                     messagebox.showinfo("完成", "地图已加载并显示。请继续输入起点和终点坐标。")
#                 except Exception as e:
#                     messagebox.showerror("错误", f"加载地图时发生错误: {e}")
#                     logging.exception("加载地图时发生错误")
#
#             threading.Thread(target=process_map).start()
#
#         # 启用输入字段函数
#         def enable_input_fields():
#             start_coord_entry.configure(state='normal')
#             end_coord_entry.configure(state='normal')
#             output_path_entry.configure(state='normal')
#             max_slope_entry.configure(state='normal')
#             max_step_elev_entry.configure(state='normal')
#             max_speed_entry.configure(state='normal')
#             run_button.configure(state='normal')
#
#         # 加载地图按钮
#         tk.Button(root, text="加载地图", command=load_map).grid(row=7, column=1)
#
#         # 运行路径规划函数
#         def run():
#             # 获取输入值
#             start_coord_input = start_coord_entry.get()
#             end_coord_input = end_coord_entry.get()
#             output_path = output_path_entry.get()
#
#             # 验证输入
#             if not output_path:
#                 messagebox.showerror("错误", "请提供输出 Word 文档路径。")
#                 return
#             if not start_coord_input:
#                 messagebox.showerror("错误", "请提供起点坐标。")
#                 return
#             if not end_coord_input:
#                 messagebox.showerror("错误", "请提供终点坐标。")
#                 return
#             try:
#                 start_x, start_y = map(float, start_coord_input.strip().split(','))
#                 start_coord = (start_x, start_y)
#             except ValueError:
#                 messagebox.showerror("错误", "起点坐标格式错误，请使用 'X,Y' 的格式。")
#                 return
#
#             try:
#                 end_x, end_y = map(float, end_coord_input.strip().split(','))
#                 end_coord = (end_x, end_y)
#             except ValueError:
#                 messagebox.showerror("错误", "终点坐标格式错误，请使用 'X,Y' 的格式。")
#                 return
#
#             # 获取最大坡度和最大阶跃高程差
#             max_slope_input = max_slope_entry.get()
#             max_step_elev_input = max_step_elev_entry.get()
#
#             try:
#                 max_allowed_slope = float(max_slope_input)
#             except ValueError:
#                 messagebox.showerror("错误", "最大坡度格式错误，请输入数字。")
#                 return
#
#             try:
#                 max_allowed_step_elev = float(max_step_elev_input)
#             except ValueError:
#                 messagebox.showerror("错误", "最大阶跃高程差格式错误，请输入数字。")
#                 return
#
#             # 获取最大速度
#             max_speed_input = max_speed_entry.get()
#
#             try:
#                 max_speed = float(max_speed_input)
#             except ValueError:
#                 messagebox.showerror("错误", "最大速度格式错误，请输入数字。")
#                 return
#
#             # 在后台线程中运行，以避免阻塞GUI
#             def process():
#                 try:
#                     # 使用已加载的 data_processor
#                     data_processor = root.data_processor
#                     # 调用路径规划函数
#                     run_path_planning(data_processor, start_coord, end_coord, output_path, max_allowed_slope, max_allowed_step_elev, max_speed)
#                     messagebox.showinfo("完成", "路径规划已完成。")
#                 except Exception as e:
#                     messagebox.showerror("错误", f"发生错误: {e}")
#                     logging.exception("后台线程运行时发生错误")
#
#             threading.Thread(target=process).start()
#
#         # 运行按钮（初始禁用）
#         run_button = tk.Button(root, text="运行", command=run)
#         run_button.grid(row=8, column=1)
#         run_button.configure(state='disabled')
#
#         root.mainloop()
#     except Exception as e:
#         print(f"主程序运行时发生错误: {e}")
#         logging.exception("主程序运行时发生错误")
#
# if __name__ == "__main__":
#     main()








# import pandas as pd
# import numpy as np
# from scipy.interpolate import griddata
# from sklearn.metrics import mean_squared_error, mean_absolute_error, accuracy_score
# from sklearn.model_selection import train_test_split
# import heapq
# import math
# import plotly.graph_objs as go
# import plotly.io as pio
# from docx import Document
# import tkinter as tk
# from tkinter import filedialog, messagebox
# import threading
# import sys
# import os
# import logging
#
# # 设置日志记录
# logging.basicConfig(filename='error.log', level=logging.ERROR)
#
# # 设置 Plotly 渲染器为浏览器
# pio.renderers.default = 'browser'
#
# # 添加 resource_path 函数，用于获取资源文件的正确路径
# def resource_path(relative_path):
#     """获取资源文件的绝对路径，兼容 PyInstaller 打包后的情况"""
#     try:
#         # PyInstaller 会把路径存储在 _MEIPASS 中
#         base_path = sys._MEIPASS
#     except AttributeError:
#         base_path = os.path.abspath(".")
#     return os.path.join(base_path, relative_path)
#
# # 数据处理类
# class DataProcessor:
#     def __init__(self, dat_path, x_resolution=1, y_resolution=1):
#         self.dat_path = dat_path
#         self.x_resolution = x_resolution
#         self.y_resolution = y_resolution
#         self.df = None
#         self.X = None
#         self.Y = None
#         self.Z = None
#         self.TrafFlag = None
#         self.maxAcc = None
#         self.Slope = None
#         self.StepElev = None
#         self.Roughness = None
#         self.x_range = None
#         self.y_range = None
#         self.X_grid = None
#         self.Y_grid = None
#         self.Z_grid = None
#         self.TrafFlag_grid = None
#         self.Acc_grid = None
#         self.Slope_grid = None
#         self.StepElev_grid = None
#         self.Rough_grid = None
#
#     def load_data(self):
#         try:
#             # 使用 resource_path 函数获取数据文件路径
#             data_path = resource_path(self.dat_path)
#             # 读取 .csv 文件，指定分隔符为逗号
#             self.df = pd.read_csv(data_path, delimiter=',')
#             print("数据文件中的列名：", self.df.columns.tolist())
#         except FileNotFoundError:
#             print(f"文件未找到: {self.dat_path}")
#             return
#         except pd.errors.EmptyDataError:
#             print("CSV 文件为空")
#             return
#         except Exception as e:
#             print(f"读取 CSV 文件时出错: {e}")
#             logging.exception("读取 CSV 文件时出错")
#             return
#
#         required_columns = {'X', 'Y', 'Z', 'TrafFlag', 'maxAcc', 'Slope', 'StepElevationDifference', 'Roughness'}
#         if not required_columns.issubset(self.df.columns):
#             print(f"CSV文件缺少必要的列: {required_columns - set(self.df.columns)}")
#             return
#
#         # 提取所需数据
#         self.X = self.df['X'].values
#         self.Y = self.df['Y'].values
#         self.Z = self.df['Z'].values
#         self.TrafFlag = self.df['TrafFlag'].values
#         self.maxAcc = self.df['maxAcc'].values
#         self.Slope = self.df['Slope'].values
#         self.StepElev = self.df['StepElevationDifference'].values
#         self.Roughness = self.df['Roughness'].values
#
#         # 计算并打印 Slope 和 StepElevationDifference 的最小值和最大值
#         slope_min = np.min(self.Slope)
#         slope_max_val = np.max(self.Slope)
#         step_min = np.min(self.StepElev)
#         step_max = np.max(self.StepElev)
#
#         print(f"坡度 (Slope) 最小值: {slope_min}")
#         print(f"坡度 (Slope) 最大值: {slope_max_val}")
#         print(f"阶跃高程差 (StepElevationDifference) 最小值: {step_min}")
#         print(f"阶跃高程差 (StepElevationDifference) 最大值: {step_max}")
#
#     def interpolate_data(self):
#         # 计算网格范围和生成网格
#         x_min, x_max = self.X.min(), self.X.max()
#         y_min, y_max = self.Y.min(), self.Y.max()
#         self.x_range = np.arange(x_min, x_max + self.x_resolution, self.x_resolution)
#         self.y_range = np.arange(y_min, y_max + self.y_resolution, self.y_resolution)
#         self.X_grid, self.Y_grid = np.meshgrid(self.x_range, self.y_range)
#
#         # 插值
#         self.Z_grid = griddata((self.X, self.Y), self.Z, (self.X_grid, self.Y_grid), method='cubic')
#         self.TrafFlag_grid = griddata((self.X, self.Y), self.TrafFlag, (self.X_grid, self.Y_grid), method='nearest')
#         self.Acc_grid = griddata((self.X, self.Y), self.maxAcc, (self.X_grid, self.Y_grid), method='nearest')
#         self.Slope_grid = griddata((self.X, self.Y), self.Slope, (self.X_grid, self.Y_grid), method='nearest')
#         self.StepElev_grid = griddata((self.X, self.Y), self.StepElev, (self.X_grid, self.Y_grid), method='nearest')
#         self.Rough_grid = griddata((self.X, self.Y), self.Roughness, (self.X_grid, self.Y_grid), method='nearest')
#
#     def fill_nan_values(self):
#         # 填充 NaN 值
#         self.Z_grid = np.nan_to_num(self.Z_grid, nan=np.nanmean(self.Z))
#         self.Acc_grid = np.nan_to_num(self.Acc_grid, nan=0)
#         self.TrafFlag_grid = np.nan_to_num(self.TrafFlag_grid, nan=0)
#         # 将 TrafFlag_grid 的值限制为 0 或 1
#         self.TrafFlag_grid = np.where(self.TrafFlag_grid >= 0.5, 1, 0)
#         self.Slope_grid = np.nan_to_num(self.Slope_grid, nan=0)
#         self.StepElev_grid = np.nan_to_num(self.StepElev_grid, nan=0)
#         self.Rough_grid = np.nan_to_num(self.Rough_grid, nan=0)
#
#         # 验证网格数据的形状
#         print(f"Z_grid shape: {self.Z_grid.shape}")
#         print(f"TrafFlag_grid shape: {self.TrafFlag_grid.shape}")
#         print(f"Acc_grid shape: {self.Acc_grid.shape}")
#         print(f"Slope_grid shape: {self.Slope_grid.shape}")
#         print(f"StepElev_grid shape: {self.StepElev_grid.shape}")
#         print(f"Rough_grid shape: {self.Rough_grid.shape}")
#
#         # 计算并打印插值后网格数据的最小值和最大值
#         slope_grid_min = np.min(self.Slope_grid)
#         slope_grid_max = np.max(self.Slope_grid)
#         step_grid_min = np.min(self.StepElev_grid)
#         step_grid_max = np.max(self.StepElev_grid)
#
#         print(f"插值后坡度网格 (Slope_grid) 最小值: {slope_grid_min}")
#         print(f"插值后坡度网格 (Slope_grid) 最大值: {slope_grid_max}")
#         print(f"插值后阶跃高程差网格 (StepElev_grid) 最小值: {step_grid_min}")
#         print(f"插值后阶跃高程差网格 (StepElev_grid) 最大值: {step_grid_max}")
#
#     def coord_to_index(self, coord):
#         x_min, y_min = self.x_range.min(), self.y_range.min()
#         col = int((coord[0] - x_min) / self.x_resolution)
#         row = int((coord[1] - y_min) / self.y_resolution)
#         return (row, col)
#
#     def is_valid_position(self, pos):
#         return 0 <= pos[0] < self.TrafFlag_grid.shape[0] and 0 <= pos[1] < self.TrafFlag_grid.shape[1]
#
#     def split_data_and_evaluate(self):
#         try:
#             # 将所有数据合并
#             data = np.column_stack(
#                 (self.X, self.Y, self.Z, self.TrafFlag, self.maxAcc, self.Slope, self.StepElev, self.Roughness))
#
#             # 分割数据集为训练集和测试集
#             test_size = 0.2
#             random_state = 42
#
#             train_data, test_data = train_test_split(data, test_size=test_size, random_state=random_state)
#
#             # 拆分训练集和测试集的数据
#             X_train, Y_train = train_data[:, 0], train_data[:, 1]
#             Z_train = train_data[:, 2]
#             TrafFlag_train = train_data[:, 3]
#             maxAcc_train = train_data[:, 4]
#             Slope_train = train_data[:, 5]
#             StepElev_train = train_data[:, 6]
#             Rough_train = train_data[:, 7]
#
#             X_test, Y_test = test_data[:, 0], test_data[:, 1]
#             Z_test = test_data[:, 2]
#             TrafFlag_test = test_data[:, 3]
#             maxAcc_test = test_data[:, 4]
#             Slope_test = test_data[:, 5]
#             StepElev_test = test_data[:, 6]
#             Rough_test = test_data[:, 7]
#
#             # 使用训练集进行插值
#             Z_grid_train = griddata((X_train, Y_train), Z_train, (self.X_grid, self.Y_grid), method='cubic')
#             TrafFlag_grid_train = griddata((X_train, Y_train), TrafFlag_train, (self.X_grid, self.Y_grid), method='nearest')
#             Acc_grid_train = griddata((X_train, Y_train), maxAcc_train, (self.X_grid, self.Y_grid), method='nearest')
#             Slope_grid_train = griddata((X_train, Y_train), Slope_train, (self.X_grid, self.Y_grid), method='nearest')
#             StepElev_grid_train = griddata((X_train, Y_train), StepElev_train, (self.X_grid, self.Y_grid), method='nearest')
#             Rough_grid_train = griddata((X_train, Y_train), Rough_train, (self.X_grid, self.Y_grid), method='nearest')
#
#             # 在测试集上预测值
#             Z_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Z_grid_train.flatten(), (X_test, Y_test),
#                               method='cubic')
#             TrafFlag_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), TrafFlag_grid_train.flatten(),
#                                      (X_test, Y_test), method='nearest')
#             Acc_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Acc_grid_train.flatten(), (X_test, Y_test),
#                                 method='nearest')
#             Slope_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Slope_grid_train.flatten(),
#                                   (X_test, Y_test), method='nearest')
#             StepElev_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), StepElev_grid_train.flatten(),
#                                      (X_test, Y_test), method='nearest')
#             Rough_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Rough_grid_train.flatten(),
#                                   (X_test, Y_test), method='nearest')
#
#             # 处理 NaN 值并计算误差指标
#
#             # 高程数据误差
#             valid_indices_Z = ~np.isnan(Z_pred) & ~np.isnan(Z_test)
#             if np.any(valid_indices_Z):
#                 mse_Z = mean_squared_error(Z_test[valid_indices_Z], Z_pred[valid_indices_Z])
#                 mae_Z = mean_absolute_error(Z_test[valid_indices_Z], Z_pred[valid_indices_Z])
#                 max_error_Z = np.max(np.abs(Z_test[valid_indices_Z] - Z_pred[valid_indices_Z]))
#                 print(f"高程插值 - MSE: {mse_Z}, MAE: {mae_Z}, Max Error: {max_error_Z}")
#             else:
#                 print("高程插值预测值中全部为 NaN，无法计算误差。")
#
#             # 通行标志（分类准确率）
#             valid_indices_TrafFlag = ~np.isnan(TrafFlag_pred) & ~np.isnan(TrafFlag_test)
#             if np.any(valid_indices_TrafFlag):
#                 accuracy_TrafFlag = accuracy_score(TrafFlag_test[valid_indices_TrafFlag],
#                                                    TrafFlag_pred[valid_indices_TrafFlag])
#                 print(f"通行标志插值 - 准确率: {accuracy_TrafFlag}")
#             else:
#                 print("通行标志插值预测值中全部为 NaN，无法计算准确率。")
#
#             # 最大加速度误差
#             valid_indices_Acc = ~np.isnan(Acc_pred) & ~np.isnan(maxAcc_test)
#             if np.any(valid_indices_Acc):
#                 mse_Acc = mean_squared_error(maxAcc_test[valid_indices_Acc], Acc_pred[valid_indices_Acc])
#                 mae_Acc = mean_absolute_error(maxAcc_test[valid_indices_Acc], Acc_pred[valid_indices_Acc])
#                 max_error_Acc = np.max(np.abs(maxAcc_test[valid_indices_Acc] - Acc_pred[valid_indices_Acc]))
#                 print(f"最大加速度插值 - MSE: {mse_Acc}, MAE: {mae_Acc}, Max Error: {max_error_Acc}")
#             else:
#                 print("最大加速度插值预测值中全部为 NaN，无法计算误差。")
#
#             # 坡度误差
#             valid_indices_Slope = ~np.isnan(Slope_pred) & ~np.isnan(Slope_test)
#             if np.any(valid_indices_Slope):
#                 mse_Slope = mean_squared_error(Slope_test[valid_indices_Slope], Slope_pred[valid_indices_Slope])
#                 mae_Slope = mean_absolute_error(Slope_test[valid_indices_Slope], Slope_pred[valid_indices_Slope])
#                 max_error_Slope = np.max(np.abs(Slope_test[valid_indices_Slope] - Slope_pred[valid_indices_Slope]))
#                 print(f"坡度插值 - MSE: {mse_Slope}, MAE: {mae_Slope}, Max Error: {max_error_Slope}")
#             else:
#                 print("坡度插值预测值中全部为 NaN，无法计算误差。")
#
#             # 阶跃高程差误差
#             valid_indices_StepElev = ~np.isnan(StepElev_pred) & ~np.isnan(StepElev_test)
#             if np.any(valid_indices_StepElev):
#                 mse_StepElev = mean_squared_error(StepElev_test[valid_indices_StepElev],
#                                                   StepElev_pred[valid_indices_StepElev])
#                 mae_StepElev = mean_absolute_error(StepElev_test[valid_indices_StepElev],
#                                                    StepElev_pred[valid_indices_StepElev])
#                 max_error_StepElev = np.max(
#                     np.abs(StepElev_test[valid_indices_StepElev] - StepElev_pred[valid_indices_StepElev]))
#                 print(f"阶跃高程差插值 - MSE: {mse_StepElev}, MAE: {mae_StepElev}, Max Error: {max_error_StepElev}")
#             else:
#                 print("阶跃高程差插值预测值中全部为 NaN，无法计算误差。")
#
#             # 粗糙度误差
#             valid_indices_Rough = ~np.isnan(Rough_pred) & ~np.isnan(Rough_test)
#             if np.any(valid_indices_Rough):
#                 mse_Rough = mean_squared_error(Rough_test[valid_indices_Rough], Rough_pred[valid_indices_Rough])
#                 mae_Rough = mean_absolute_error(Rough_test[valid_indices_Rough], Rough_pred[valid_indices_Rough])
#                 max_error_Rough = np.max(np.abs(Rough_test[valid_indices_Rough] - Rough_pred[valid_indices_Rough]))
#                 print(f"粗糙度插值 - MSE: {mse_Rough}, MAE: {mae_Rough}, Max Error: {max_error_Rough}")
#             else:
#                 print("粗糙度插值预测值中全部为 NaN，无法计算误差。")
#         except Exception as e:
#             print(f"数据评估时发生错误: {e}")
#             logging.exception("数据评估时发生错误")
#
# # 节点类和路径规划类
# class Node:
#     def __init__(self, position, parent=None):
#         self.position = position  # (row, col)
#         self.parent = parent
#         self.g = 0  # 累计代价
#         self.h = 0  # 启发代价
#         self.f = 0  # 总代价
#
#     def __lt__(self, other):
#         return self.f < other.f
#
# class PathPlanner:
#     def __init__(self, grid_data, max_slope=None, max_step_elev=None, max_roughness=None, max_speed=70.0):
#         self.grid = grid_data['TrafFlag_grid']
#         self.Z_grid = grid_data['Z_grid']
#         self.Acc_grid = grid_data['Acc_grid']
#         self.Slope_grid = grid_data['Slope_grid']
#         self.StepElev_grid = grid_data['StepElev_grid']
#         self.Rough_grid = grid_data['Rough_grid']
#         self.max_slope = max_slope
#         self.max_step_elev = max_step_elev
#         self.max_roughness = max_roughness  # 新增最大粗糙度参数
#         self.max_speed = max_speed
#
#     def heuristic_distance_3d(self, current_pos, end_pos):
#         d_xy = np.sqrt((current_pos[0] - end_pos[0]) ** 2 +
#                        (current_pos[1] - end_pos[1]) ** 2)
#         z_current = self.Z_grid[current_pos[0], current_pos[1]]
#         z_end = self.Z_grid[end_pos[0], end_pos[1]]
#         d_z = (z_end - z_current) ** 2
#         return np.sqrt(d_xy ** 2 + d_z)
#
#     def astar(self, start, end, mode='distance', alpha_env=0.6, beta_env=0.4):
#         start_node = Node(start)
#         end_node = Node(end)
#
#         open_list = []
#         heapq.heappush(open_list, start_node)
#         open_dict = {start_node.position: start_node}
#
#         closed_set = set()
#
#         while open_list:
#             current_node = heapq.heappop(open_list)
#             del open_dict[current_node.position]
#             closed_set.add(current_node.position)
#
#             # 检查是否到达终点
#             if current_node.position == end_node.position:
#                 path = []
#                 while current_node:
#                     path.append(current_node.position)
#                     current_node = current_node.parent
#                 return path[::-1]
#
#             # 定义8个邻居（包括对角线）
#             neighbors = [(0, 1), (0, -1), (1, 0), (-1, 0),
#                          (1, 1), (1, -1), (-1, 1), (-1, -1)]
#
#             for offset in neighbors:
#                 neighbor_pos = (current_node.position[0] + offset[0],
#                                 current_node.position[1] + offset[1])
#
#                 # 检查邻居是否在网格范围内
#                 if not (0 <= neighbor_pos[0] < self.grid.shape[0] and
#                         0 <= neighbor_pos[1] < self.grid.shape[1]):
#                     continue
#
#                 # 检查 TrafFlag 是否可通行
#                 if self.grid[neighbor_pos] == 0:
#                     continue  # 不可通行区域
#
#                 # 在所有模式下应用坡度和阶跃高程差的限制
#                 if self.max_slope is not None and self.Slope_grid[neighbor_pos] > self.max_slope:
#                     continue  # 坡度超过最大限制
#                 if self.max_step_elev is not None and self.StepElev_grid[neighbor_pos] > self.max_step_elev:
#                     continue  # 阶跃高程差超过最大限制
#
#                 # 新增对最大粗糙度的限制
#                 if self.max_roughness is not None and self.Rough_grid[neighbor_pos] > self.max_roughness:
#                     continue  # 粗糙度超过最大限制
#
#                 # 检查是否已经在关闭列表中
#                 if neighbor_pos in closed_set:
#                     continue
#
#                 # 处理不同模式
#                 if mode == 'distance':
#                     movement_cost = self.heuristic_distance_3d(current_node.position, neighbor_pos)
#                     g_cost = current_node.g + movement_cost
#                     h_cost = self.heuristic_distance_3d(neighbor_pos, end)
#                 elif mode == 'time':
#                     a_max = self.Acc_grid[neighbor_pos[0], neighbor_pos[1]]
#                     if a_max <= 0:
#                         continue  # 无法加速，跳过
#                     movement_cost = self.heuristic_distance_3d(current_node.position, neighbor_pos)
#                     g_cost = current_node.g + movement_cost / a_max  # 简化为距离除以加速度
#                     h_cost = self.heuristic_distance_3d(neighbor_pos, end)
#                 elif mode == 'env':
#                     movement_uncertainty = (self.Slope_grid[neighbor_pos[0], neighbor_pos[1]] / 45.0 +
#                                             self.StepElev_grid[neighbor_pos[0], neighbor_pos[1]] / 1.5 +
#                                             self.Rough_grid[neighbor_pos[0], neighbor_pos[1]] / 6.0)
#                     g_cost = current_node.g + movement_uncertainty
#                     h_cost = self.heuristic_distance_3d(neighbor_pos, end)
#                 else:
#                     raise ValueError("Invalid mode. Choose from 'distance', 'time', 'env'.")
#
#                 # 创建邻居节点
#                 neighbor_node = Node(neighbor_pos, current_node)
#                 neighbor_node.g = g_cost
#                 neighbor_node.h = h_cost
#                 neighbor_node.f = neighbor_node.g + neighbor_node.h
#
#                 # 检查开放列表中是否有更好的节点
#                 if neighbor_pos in open_dict:
#                     existing_node = open_dict[neighbor_pos]
#                     if existing_node.g <= neighbor_node.g:
#                         continue  # 已有更优节点
#                     else:
#                         # 替换为更优节点
#                         open_list.remove(existing_node)
#                         heapq.heapify(open_list)
#                         heapq.heappush(open_list, neighbor_node)
#                         open_dict[neighbor_pos] = neighbor_node
#                 else:
#                     heapq.heappush(open_list, neighbor_node)
#                     open_dict[neighbor_pos] = neighbor_node
#
#         return None  # 如果找不到路径，返回 None
#
#     def extract_path_coordinates(self, path, x_range, y_range):
#         return [(x_range[p[1]], y_range[p[0]], self.Z_grid[p[0], p[1]]) for p in path]
#
#     def calculate_path_length(self, path_coords):
#         length = 0
#         for i in range(1, len(path_coords)):
#             x1, y1, z1 = path_coords[i - 1]
#             x2, y2, z2 = path_coords[i]
#             length += np.sqrt((x2 - x1) ** 2 + (y2 - y1) ** 2 + (z2 - z1) ** 2)
#         return length
#
#     def compute_traversal_time(self, path_coords, y_range, x_range):
#         total_time = 0.0
#         current_speed = 0.0  # 初始速度为0
#         max_speed = self.max_speed  # 使用用户指定的最大速度
#
#         speeds = [current_speed]  # 存储每个点的速度
#         accelerations = []  # 存储每段的加速度
#
#         for i in range(1, len(path_coords)):
#             x1, y1, z1 = path_coords[i - 1]
#             x2, y2, z2 = path_coords[i]
#             # 计算两点之间的距离
#             d = np.sqrt((x2 - x1) ** 2 + (y2 - y1) ** 2 + (z2 - z1) ** 2)
#
#             # 获取当前点的网格索引
#             row = np.argmin(np.abs(y_range - y1))
#             col = np.argmin(np.abs(x_range - x1))
#
#             # 获取当前点的加速度
#             a = self.Acc_grid[row, col]
#             if a <= 0:
#                 # 如果加速度不可用或为0，无法继续加速
#                 a = 0.0
#
#             # 计算需要加速到最大速度的速度增量
#             delta_v = max_speed - current_speed
#
#             if current_speed < max_speed and a > 0:
#                 # 需要加速
#                 t_acc = delta_v / a  # 加速所需时间
#                 d_acc = current_speed * t_acc + 0.5 * a * t_acc ** 2  # 加速所需距离
#
#                 if d_acc >= d:
#                     # 在当前段无法加速到最大速度
#                     # 使用运动学方程计算通过距离d所需的时间
#                     # 解方程: 0.5 * a * t^2 + current_speed * t - d = 0
#                     discriminant = current_speed ** 2 + 2 * a * d
#                     if discriminant < 0:
#                         # 无实数解，无法到达下一个点
#                         return float('inf'), [], []
#                     t = (-current_speed + math.sqrt(discriminant)) / a
#                     total_time += t
#                     current_speed += a * t
#                     accelerations.append(a)
#                     speeds.append(current_speed)
#                 else:
#                     # 在当前段可以加速到最大速度
#                     # 先加速到最大速度
#                     total_time += t_acc
#                     current_speed = max_speed
#                     accelerations.append(a)
#                     speeds.append(current_speed)
#                     # 剩余距离
#                     d_remain = d - d_acc
#                     # 以最大速度匀速行驶剩余距离
#                     t_remain = d_remain / current_speed
#                     total_time += t_remain
#                     accelerations.append(0.0)
#                     speeds.append(current_speed)
#             else:
#                 # 无法加速或已达到最大速度，匀速行驶
#                 if current_speed == 0:
#                     # 无法移动
#                     return float('inf'), [], []
#                 t = d / current_speed
#                 total_time += t
#                 accelerations.append(0.0)
#                 speeds.append(current_speed)
#
#         return total_time, speeds, accelerations
#
# # 可视化类
# class Visualizer:
#     def __init__(self, X_grid, Y_grid, Z_grid, TrafFlag_grid):
#         self.X_grid = X_grid
#         self.Y_grid = Y_grid
#         self.Z_grid = Z_grid
#         self.TrafFlag_grid = TrafFlag_grid
#         self.fig = go.Figure()
#
#     def plot_terrain(self):
#         # 定义自定义颜色尺度
#         colorscale = [
#             [0.0, 'gray'],        # TrafFlag == 0 的区域为灰色
#             [1.0, 'lightgreen']   # TrafFlag == 1 的区域为浅绿色
#         ]
#
#         self.fig.add_trace(go.Surface(
#             x=self.X_grid, y=self.Y_grid, z=self.Z_grid,
#             surfacecolor=self.TrafFlag_grid,
#             colorscale=colorscale,
#             cmin=0,
#             cmax=1,
#             opacity=0.8,
#             showscale=True,
#             colorbar=dict(title='TrafFlag')
#         ))
#
#     def add_path(self, coords, speeds, accs, color, name, length, time_taken, path_width=4):
#         if coords:
#             x, y, z = zip(*coords)
#             # 处理不可行路径（时间为inf）的情况
#             if math.isinf(time_taken):
#                 time_str = "无法到达"
#             else:
#                 time_str = f"{time_taken:.2f} 秒"
#
#             # 准备 customdata，将速度和加速度添加进去
#             if accs:
#                 accs = [0.0] + accs  # 与 speeds 对齐
#             else:
#                 accs = [0.0] * len(speeds)
#
#             customdata = np.array([speeds, accs]).T  # 转置使其每行对应一个点的 [speed, acceleration]
#
#             self.fig.add_trace(go.Scatter3d(
#                 x=x, y=y, z=z,
#                 mode='lines+markers',
#                 line=dict(color=color, width=path_width),
#                 marker=dict(size=2, color=color),
#                 name=f"{name} (长度: {length:.2f} 米, 时间: {time_str})",
#                 customdata=customdata,
#                 hovertemplate=(
#                         'X: %{x}<br>' +
#                         'Y: %{y}<br>' +
#                         'Z: %{z}<br>' +
#                         '速度: %{customdata[0]:.2f} m/s<br>' +
#                         '加速度: %{customdata[1]:.2f} m/s²<br>' +
#                         '<extra></extra>'
#                 )
#             ))
#         else:
#             print(f"未找到{name}")
#
#     def add_start_and_end_points(self, start_coord, end_coord):
#         self.fig.add_trace(go.Scatter3d(
#             x=[start_coord[0]],
#             y=[start_coord[1]],
#             z=[start_coord[2]],
#             mode='markers',
#             marker=dict(size=8, color='green'),
#             name='起点'
#         ))
#         self.fig.add_trace(go.Scatter3d(
#             x=[end_coord[0]],
#             y=[end_coord[1]],
#             z=[end_coord[2]],
#             mode='markers',
#             marker=dict(size=8, color='red'),
#             name='终点'
#         ))
#
#     def show(self):
#         self.fig.update_layout(
#             scene=dict(
#                 xaxis_title='X 坐标',
#                 yaxis_title='Y 坐标',
#                 zaxis_title='Z 坐标 (高程)',
#                 aspectmode='data'
#             ),
#             title='3D 地形路径规划（基于距离、时间与最小不确定性）',
#             legend=dict(
#                 x=0,
#                 y=1,
#                 bgcolor='rgba(255, 255, 255, 0)',
#                 bordercolor='rgba(255, 255, 255, 0)'
#             )
#         )
#         pio.show(self.fig)
#
# # 报告生成函数
# def write_paths_to_word(document, path_name, path_coords):
#     if not path_coords:
#         return
#     # 添加路径名称和点数量
#     document.add_heading(f"{path_name} (点数: {len(path_coords)})", level=2)
#     # 创建表格，包含三列：X相对坐标、Y相对坐标、Z高程坐标
#     table = document.add_table(rows=1, cols=3)
#     table.style = 'Light Grid'
#     hdr_cells = table.rows[0].cells
#     hdr_cells[0].text = 'X相对坐标'
#     hdr_cells[1].text = 'Y相对坐标'
#     hdr_cells[2].text = 'Z高程坐标'
#     # 填充表格内容
#     for coord in path_coords:
#         row_cells = table.add_row().cells
#         row_cells[0].text = f"{coord[0]}"
#         row_cells[1].text = f"{coord[1]}"
#         row_cells[2].text = f"{coord[2]}"
#     # 添加一个空行
#     document.add_paragraph()
#
# # 路径规划函数
# def run_path_planning(data_processor, start_coord, end_coord, output_path, max_allowed_slope, max_allowed_step_elev, max_allowed_roughness, max_speed):
#     try:
#         # 确保已进行数据评估
#         data_processor.split_data_and_evaluate()
#
#         # 将 start_coord 和 end_coord 转换为网格索引
#         start = data_processor.coord_to_index(start_coord)
#         end = data_processor.coord_to_index(end_coord)
#
#         # 检查起点和终点是否在网格范围内
#         if not data_processor.is_valid_position(start):
#             raise ValueError("起点坐标超出网格范围")
#         if not data_processor.is_valid_position(end):
#             raise ValueError("终点坐标超出网格范围")
#
#         # 确保起点和终点可通行
#         if data_processor.TrafFlag_grid[start] == 0:
#             raise ValueError("起点不可通行")
#         if data_processor.TrafFlag_grid[end] == 0:
#             raise ValueError("终点不可通行")
#
#         # 路径规划
#         grid_data = {
#             'TrafFlag_grid': data_processor.TrafFlag_grid,
#             'Z_grid': data_processor.Z_grid,
#             'Acc_grid': data_processor.Acc_grid,
#             'Slope_grid': data_processor.Slope_grid,
#             'StepElev_grid': data_processor.StepElev_grid,
#             'Rough_grid': data_processor.Rough_grid
#         }
#         path_planner = PathPlanner(grid_data, max_slope=max_allowed_slope, max_step_elev=max_allowed_step_elev, max_roughness=max_allowed_roughness, max_speed=max_speed)
#
#         print("正在计算基于距离的最短路径...")
#         path_distance = path_planner.astar(start, end, mode='distance')
#         if path_distance:
#             print("基于距离的路径计算完成")
#         else:
#             print("未找到基于距离的路径")
#
#         print("正在计算基于时间的最短耗时路径...")
#         path_time = path_planner.astar(start, end, mode='time')
#         if path_time:
#             print("基于时间的路径计算完成")
#         else:
#             print("未找到基于时间的路径")
#
#         print("正在计算最小不确定性路径...")
#         path_env = path_planner.astar(start, end, mode='env')
#         if path_env:
#             print("最小不确定性路径计算完成")
#         else:
#             print("未找到最小不确定性路径")
#
#         # 提取路径坐标
#         if path_distance:
#             path_distance_coords = path_planner.extract_path_coordinates(path_distance, data_processor.x_range,
#                                                                          data_processor.y_range)
#             length_distance = path_planner.calculate_path_length(path_distance_coords)
#         else:
#             path_distance_coords = None
#             length_distance = None
#         if path_time:
#             path_time_coords = path_planner.extract_path_coordinates(path_time, data_processor.x_range,
#                                                                      data_processor.y_range)
#             length_time = path_planner.calculate_path_length(path_time_coords)
#         else:
#             path_time_coords = None
#             length_time = None
#         if path_env:
#             path_env_coords = path_planner.extract_path_coordinates(path_env, data_processor.x_range,
#                                                                     data_processor.y_range)
#             length_env = path_planner.calculate_path_length(path_env_coords)
#         else:
#             path_env_coords = None
#             length_env = None
#
#         # 计算遍历时间，并获取速度和加速度数据
#         if path_distance:
#             time_distance, speeds_distance, accs_distance = path_planner.compute_traversal_time(
#                 path_distance_coords, data_processor.y_range, data_processor.x_range)
#         else:
#             speeds_distance, accs_distance, time_distance = None, None, None
#
#         if path_time:
#             time_time, speeds_time, accs_time = path_planner.compute_traversal_time(
#                 path_time_coords, data_processor.y_range, data_processor.x_range)
#         else:
#             speeds_time, accs_time, time_time = None, None, None
#
#         if path_env:
#             time_env, speeds_env, accs_env = path_planner.compute_traversal_time(
#                 path_env_coords, data_processor.y_range, data_processor.x_range)
#         else:
#             speeds_env, accs_env, time_env = None, None, None
#
#         # 可视化
#         visualizer = Visualizer(
#             data_processor.X_grid,
#             data_processor.Y_grid,
#             data_processor.Z_grid,
#             data_processor.TrafFlag_grid
#         )
#         visualizer.plot_terrain()
#
#         colors = ['blue', 'orange', 'purple']
#         names = ['MinDistPath 最短距离路径', 'MinTimePath 最短耗时路径', 'MinUncertPath 最小不确定性路径']
#         lengths = [length_distance, length_time, length_env]
#         times = [time_distance, time_time, time_env]
#         coords_list = [path_distance_coords, path_time_coords, path_env_coords]
#         speeds_list = [speeds_distance, speeds_time, speeds_env]
#         accs_list = [accs_distance, accs_time, accs_env]
#
#         for coords, speeds, accs, color, name, length, time_taken in zip(
#                 coords_list, speeds_list, accs_list,
#                 colors, names, lengths, times):
#             visualizer.add_path(coords, speeds, accs, color, name, length, time_taken)
#
#         # 标记起点和终点
#         if path_distance or path_time or path_env:
#             selected_path = None
#             if path_distance:
#                 selected_path = path_distance_coords
#             elif path_time:
#                 selected_path = path_time_coords
#             elif path_env:
#                 selected_path = path_env_coords
#
#             if selected_path:
#                 # 使用起点坐标
#                 start_x, start_y, start_z = selected_path[0]
#                 # 使用终点坐标
#                 end_x, end_y, end_z = selected_path[-1]
#                 visualizer.add_start_and_end_points((start_x, start_y, start_z), (end_x, end_y, end_z))
#
#         visualizer.show()
#
#         # 创建 Word 文档
#         document = Document()
#         document.add_heading('路径规划结果', 0)
#         # 写入路径信息
#         if path_distance or path_time or path_env:
#             if path_distance:
#                 write_paths_to_word(document, 'MinDistPath ', path_distance_coords)
#             if path_time:
#                 write_paths_to_word(document, 'MinTimePath ', path_time_coords)
#             if path_env:
#                 write_paths_to_word(document, 'MinUncertPath ', path_env_coords)
#
#             # 保存文档
#             try:
#                 # 使用 resource_path 函数获取输出文件路径
#                 output_file_path = output_path  # 如果 output_path 是用户指定的完整路径，则无需使用 resource_path
#                 document.save(output_file_path)
#                 print(f"路径坐标已成功保存到 Word 文档: {output_file_path}")
#             except Exception as e:
#                 print(f"保存 Word 文档时出错: {e}")
#                 logging.exception("保存 Word 文档时出错")
#         else:
#             print("没有路径可写入 Word 文档")
#
#         # 打印路径长度和遍历时间
#         if path_distance:
#             if math.isinf(time_distance):
#                 print(f"最短路径长度: {length_distance:.2f} 米")
#                 print(f"最短距离路径遍历时间: 无法到达")
#             else:
#                 print(f"最短路径长度: {length_distance:.2f} 米")
#                 print(f"最短距离路径遍历时间: {time_distance:.2f} 秒")
#         if path_time:
#             if math.isinf(time_time):
#                 print(f"最短耗时路径长度: {length_time:.2f} 米")
#                 print(f"最短耗时路径遍历时间: 无法到达")
#             else:
#                 print(f"最短耗时路径长度: {length_time:.2f} 米")
#                 print(f"最短耗时路径遍历时间: {time_time:.2f} 秒")
#         if path_env:
#             if math.isinf(time_env):
#                 print(f"最小不确定性路径长度: {length_env:.2f} 米")
#                 print(f"最小不确定性路径遍历时间: 无法到达")
#             else:
#                 print(f"最小不确定性路径长度: {length_env:.2f} 米")
#                 print(f"最小不确定性路径遍历时间: {time_env:.2f} 秒")
#     except Exception as e:
#         print(f"运行处理时发生错误: {e}")
#         logging.exception("运行处理时发生错误")
#
# # 主程序
# def main():
#     try:
#         # 创建主窗口
#         root = tk.Tk()
#         root.title("路径规划软件")
#
#         # 地图数据文件路径
#         tk.Label(root, text="地图数据文件路径:").grid(row=0, column=0, sticky='e')
#         map_data_path_entry = tk.Entry(root, width=50)
#         map_data_path_entry.grid(row=0, column=1)
#         def browse_map_data():
#             filename = filedialog.askopenfilename(filetypes=[("CSV文件", "*.csv")])
#             if filename:
#                 map_data_path_entry.delete(0, tk.END)
#                 map_data_path_entry.insert(0, filename)
#         tk.Button(root, text="浏览...", command=browse_map_data).grid(row=0, column=2)
#
#         # 起点坐标（初始禁用）
#         tk.Label(root, text="起点坐标 (格式: X,Y):").grid(row=1, column=0, sticky='e')
#         start_coord_entry = tk.Entry(root, width=50)
#         start_coord_entry.grid(row=1, column=1)
#         # 终点坐标（初始禁用）
#         tk.Label(root, text="终点坐标 (格式: X,Y):").grid(row=2, column=0, sticky='e')
#         end_coord_entry = tk.Entry(root, width=50)
#         end_coord_entry.grid(row=2, column=1)
#
#         # 输出路径
#         tk.Label(root, text="输出路径文档路径:").grid(row=3, column=0, sticky='e')
#         output_path_entry = tk.Entry(root, width=50)
#         output_path_entry.grid(row=3, column=1)
#         def browse_output_path():
#             filename = filedialog.asksaveasfilename(defaultextension=".docx",
#                                                     filetypes=[("Word文档", "*.docx")])
#             if filename:
#                 output_path_entry.delete(0, tk.END)
#                 output_path_entry.insert(0, filename)
#         tk.Button(root, text="浏览...", command=browse_output_path).grid(row=3, column=2)
#
#         # 最大坡度
#         tk.Label(root, text="最大爬坡度 (默认25.0度):").grid(row=4, column=0, sticky='e')
#         max_slope_entry = tk.Entry(root, width=50)
#         max_slope_entry.grid(row=4, column=1)
#         max_slope_entry.insert(0, "25.0")  # 设置默认值
#
#         # 最大阶跃高程差
#         tk.Label(root, text="最大阶跃高程差 (默认10.0米):").grid(row=5, column=0, sticky='e')
#         max_step_elev_entry = tk.Entry(root, width=50)
#         max_step_elev_entry.grid(row=5, column=1)
#         max_step_elev_entry.insert(0, "10.0")  # 设置默认值
#
#         # 添加最大粗糙度输入字段
#         tk.Label(root, text="最大粗糙度 (默认6.0):").grid(row=6, column=0, sticky='e')
#         max_roughness_entry = tk.Entry(root, width=50)
#         max_roughness_entry.grid(row=6, column=1)
#         max_roughness_entry.insert(0, "6.0")  # 设置默认值
#
#         # 最大速度
#         tk.Label(root, text="最大速度 (默认20.0 m/s):").grid(row=7, column=0, sticky='e')
#         max_speed_entry = tk.Entry(root, width=50)
#         max_speed_entry.grid(row=7, column=1)
#         max_speed_entry.insert(0, "20.0")  # 设置默认值
#
#         # 禁用输入字段，直到地图加载完成
#         start_coord_entry.configure(state='disabled')
#         end_coord_entry.configure(state='disabled')
#         output_path_entry.configure(state='disabled')
#         max_slope_entry.configure(state='disabled')
#         max_step_elev_entry.configure(state='disabled')
#         max_roughness_entry.configure(state='disabled')
#         max_speed_entry.configure(state='disabled')
#
#         # 加载地图函数
#         def load_map():
#             csv_path = map_data_path_entry.get()
#             if not csv_path:
#                 messagebox.showerror("错误", "请提供地图数据文件路径。")
#                 return
#
#             def process_map():
#                 try:
#                     data_processor = DataProcessor(csv_path)
#                     data_processor.load_data()
#                     data_processor.interpolate_data()
#                     data_processor.fill_nan_values()
#
#                     # 可视化地图
#                     visualizer = Visualizer(
#                         data_processor.X_grid,
#                         data_processor.Y_grid,
#                         data_processor.Z_grid,
#                         data_processor.TrafFlag_grid
#                     )
#                     visualizer.plot_terrain()
#                     visualizer.show()
#
#                     # 存储 data_processor
#                     root.data_processor = data_processor
#
#                     # 启用输入字段
#                     root.after(0, enable_input_fields)
#                     messagebox.showinfo("完成", "地图已加载并显示。请继续输入起点和终点坐标。")
#                 except Exception as e:
#                     messagebox.showerror("错误", f"加载地图时发生错误: {e}")
#                     logging.exception("加载地图时发生错误")
#
#             threading.Thread(target=process_map).start()
#
#         # 启用输入字段函数
#         def enable_input_fields():
#             start_coord_entry.configure(state='normal')
#             end_coord_entry.configure(state='normal')
#             output_path_entry.configure(state='normal')
#             max_slope_entry.configure(state='normal')
#             max_step_elev_entry.configure(state='normal')
#             max_roughness_entry.configure(state='normal')
#             max_speed_entry.configure(state='normal')
#             run_button.configure(state='normal')
#
#         # 加载地图按钮
#         tk.Button(root, text="加载地图", command=load_map).grid(row=8, column=1)
#
#         # 运行路径规划函数
#         def run():
#             # 获取输入值
#             start_coord_input = start_coord_entry.get()
#             end_coord_input = end_coord_entry.get()
#             output_path = output_path_entry.get()
#
#             # 验证输入
#             if not output_path:
#                 messagebox.showerror("错误", "请提供输出 Word 文档路径。")
#                 return
#             if not start_coord_input:
#                 messagebox.showerror("错误", "请提供起点坐标。")
#                 return
#             if not end_coord_input:
#                 messagebox.showerror("错误", "请提供终点坐标。")
#                 return
#             try:
#                 start_x, start_y = map(float, start_coord_input.strip().split(','))
#                 start_coord = (start_x, start_y)
#             except ValueError:
#                 messagebox.showerror("错误", "起点坐标格式错误，请使用 'X,Y' 的格式。")
#                 return
#
#             try:
#                 end_x, end_y = map(float, end_coord_input.strip().split(','))
#                 end_coord = (end_x, end_y)
#             except ValueError:
#                 messagebox.showerror("错误", "终点坐标格式错误，请使用 'X,Y' 的格式。")
#                 return
#
#             # 获取最大坡度和最大阶跃高程差
#             max_slope_input = max_slope_entry.get()
#             max_step_elev_input = max_step_elev_entry.get()
#
#             try:
#                 max_allowed_slope = float(max_slope_input)
#             except ValueError:
#                 messagebox.showerror("错误", "最大坡度格式错误，请输入数字。")
#                 return
#
#             try:
#                 max_allowed_step_elev = float(max_step_elev_input)
#             except ValueError:
#                 messagebox.showerror("错误", "最大阶跃高程差格式错误，请输入数字。")
#                 return
#
#             # 获取最大粗糙度
#             max_roughness_input = max_roughness_entry.get()
#             try:
#                 max_allowed_roughness = float(max_roughness_input)
#             except ValueError:
#                 messagebox.showerror("错误", "最大粗糙度格式错误，请输入数字。")
#                 return
#
#             # 获取最大速度
#             max_speed_input = max_speed_entry.get()
#
#             try:
#                 max_speed = float(max_speed_input)
#             except ValueError:
#                 messagebox.showerror("错误", "最大速度格式错误，请输入数字。")
#                 return
#
#             # 在后台线程中运行，以避免阻塞GUI
#             def process():
#                 try:
#                     # 使用已加载的 data_processor
#                     data_processor = root.data_processor
#                     # 调用路径规划函数
#                     run_path_planning(data_processor, start_coord, end_coord, output_path, max_allowed_slope, max_allowed_step_elev, max_allowed_roughness, max_speed)
#                     messagebox.showinfo("完成", "路径规划已完成。")
#                 except Exception as e:
#                     messagebox.showerror("错误", f"发生错误: {e}")
#                     logging.exception("后台线程运行时发生错误")
#
#             threading.Thread(target=process).start()
#
#         # 运行按钮（初始禁用）
#         run_button = tk.Button(root, text="运行", command=run)
#         run_button.grid(row=9, column=1)
#         run_button.configure(state='disabled')
#
#         root.mainloop()
#     except Exception as e:
#         print(f"主程序运行时发生错误: {e}")
#         logging.exception("主程序运行时发生错误")
#
# if __name__ == "__main__":
#     main()








# import pandas as pd
# import numpy as np
# from scipy.interpolate import griddata
# from sklearn.metrics import mean_squared_error, mean_absolute_error, accuracy_score
# from sklearn.model_selection import train_test_split
# import heapq
# import math
# import plotly.graph_objs as go
# import plotly.io as pio
# from docx import Document
# import tkinter as tk
# from tkinter import filedialog, messagebox
# import threading
# import sys
# import os
# import logging
#
# # 设置日志记录
# logging.basicConfig(filename='error.log', level=logging.ERROR)
#
# # 设置 Plotly 渲染器为浏览器
# pio.renderers.default = 'browser'
#
# # 添加 resource_path 函数，用于获取资源文件的正确路径
# def resource_path(relative_path):
#     """获取资源文件的绝对路径，兼容 PyInstaller 打包后的情况"""
#     try:
#         # PyInstaller 会把路径存储在 _MEIPASS 中
#         base_path = sys._MEIPASS
#     except AttributeError:
#         base_path = os.path.abspath(".")
#     return os.path.join(base_path, relative_path)
#
# # 数据处理类
# class DataProcessor:
#     def __init__(self, dat_path, x_resolution=1, y_resolution=1):
#         self.dat_path = dat_path
#         self.x_resolution = x_resolution
#         self.y_resolution = y_resolution
#         self.df = None
#         self.X = None
#         self.Y = None
#         self.Z = None
#         self.TrafFlag = None
#         self.maxAcc = None
#         self.Slope = None
#         self.StepElev = None
#         self.Roughness = None
#         self.x_range = None
#         self.y_range = None
#         self.X_grid = None
#         self.Y_grid = None
#         self.Z_grid = None
#         self.TrafFlag_grid = None
#         self.Acc_grid = None
#         self.Slope_grid = None
#         self.StepElev_grid = None
#         self.Rough_grid = None
#
#     def load_data(self):
#         try:
#             # 使用 resource_path 函数获取数据文件路径
#             data_path = resource_path(self.dat_path)
#             # 读取 .csv 文件，指定分隔符为逗号
#             self.df = pd.read_csv(data_path, delimiter=',')
#             print("数据文件中的列名：", self.df.columns.tolist())
#         except FileNotFoundError:
#             print(f"文件未找到: {self.dat_path}")
#             return
#         except pd.errors.EmptyDataError:
#             print("CSV 文件为空")
#             return
#         except Exception as e:
#             print(f"读取 CSV 文件时出错: {e}")
#             logging.exception("读取 CSV 文件时出错")
#             return
#
#         required_columns = {'X', 'Y', 'Z', 'TrafFlag', 'maxAcc', 'Slope', 'StepElevationDifference', 'Roughness'}
#         if not required_columns.issubset(self.df.columns):
#             print(f"CSV文件缺少必要的列: {required_columns - set(self.df.columns)}")
#             return
#
#         # 提取所需数据
#         self.X = self.df['X'].values
#         self.Y = self.df['Y'].values
#         self.Z = self.df['Z'].values
#         self.TrafFlag = self.df['TrafFlag'].values
#         self.maxAcc = self.df['maxAcc'].values
#         self.Slope = self.df['Slope'].values
#         self.StepElev = self.df['StepElevationDifference'].values
#         self.Roughness = self.df['Roughness'].values
#
#         # 计算并打印 Slope 和 StepElevationDifference 的最小值和最大值
#         slope_min = np.min(self.Slope)
#         slope_max_val = np.max(self.Slope)
#         step_min = np.min(self.StepElev)
#         step_max = np.max(self.StepElev)
#
#         print(f"坡度 (Slope) 最小值: {slope_min}")
#         print(f"坡度 (Slope) 最大值: {slope_max_val}")
#         print(f"阶跃高程差 (StepElevationDifference) 最小值: {step_min}")
#         print(f"阶跃高程差 (StepElevationDifference) 最大值: {step_max}")
#
#     def interpolate_data(self):
#         # 计算网格范围和生成网格
#         x_min, x_max = self.X.min(), self.X.max()
#         y_min, y_max = self.Y.min(), self.Y.max()
#         self.x_range = np.arange(x_min, x_max + self.x_resolution, self.x_resolution)
#         self.y_range = np.arange(y_min, y_max + self.y_resolution, self.y_resolution)
#         self.X_grid, self.Y_grid = np.meshgrid(self.x_range, self.y_range)
#
#         # 插值
#         self.Z_grid = griddata((self.X, self.Y), self.Z, (self.X_grid, self.Y_grid), method='cubic')
#         self.TrafFlag_grid = griddata((self.X, self.Y), self.TrafFlag, (self.X_grid, self.Y_grid), method='nearest')
#         self.Acc_grid = griddata((self.X, self.Y), self.maxAcc, (self.X_grid, self.Y_grid), method='nearest')
#         self.Slope_grid = griddata((self.X, self.Y), self.Slope, (self.X_grid, self.Y_grid), method='nearest')
#         self.StepElev_grid = griddata((self.X, self.Y), self.StepElev, (self.X_grid, self.Y_grid), method='nearest')
#         self.Rough_grid = griddata((self.X, self.Y), self.Roughness, (self.X_grid, self.Y_grid), method='nearest')
#
#     def fill_nan_values(self):
#         # 填充 NaN 值
#         self.Z_grid = np.nan_to_num(self.Z_grid, nan=np.nanmean(self.Z))
#         self.Acc_grid = np.nan_to_num(self.Acc_grid, nan=0)
#         self.TrafFlag_grid = np.nan_to_num(self.TrafFlag_grid, nan=0)
#         # 将 TrafFlag_grid 的值限制为 0 或 1
#         self.TrafFlag_grid = np.where(self.TrafFlag_grid >= 0.5, 1, 0)
#         self.Slope_grid = np.nan_to_num(self.Slope_grid, nan=0)
#         self.StepElev_grid = np.nan_to_num(self.StepElev_grid, nan=0)
#         self.Rough_grid = np.nan_to_num(self.Rough_grid, nan=0)
#
#         # 验证网格数据的形状
#         print(f"Z_grid shape: {self.Z_grid.shape}")
#         print(f"TrafFlag_grid shape: {self.TrafFlag_grid.shape}")
#         print(f"Acc_grid shape: {self.Acc_grid.shape}")
#         print(f"Slope_grid shape: {self.Slope_grid.shape}")
#         print(f"StepElev_grid shape: {self.StepElev_grid.shape}")
#         print(f"Rough_grid shape: {self.Rough_grid.shape}")
#
#         # 计算并打印插值后网格数据的最小值和最大值
#         slope_grid_min = np.min(self.Slope_grid)
#         slope_grid_max = np.max(self.Slope_grid)
#         step_grid_min = np.min(self.StepElev_grid)
#         step_grid_max = np.max(self.StepElev_grid)
#
#         print(f"插值后坡度网格 (Slope_grid) 最小值: {slope_grid_min}")
#         print(f"插值后坡度网格 (Slope_grid) 最大值: {slope_grid_max}")
#         print(f"插值后阶跃高程差网格 (StepElev_grid) 最小值: {step_grid_min}")
#         print(f"插值后阶跃高程差网格 (StepElev_grid) 最大值: {step_grid_max}")
#
#     def coord_to_index(self, coord):
#         x_min, y_min = self.x_range.min(), self.y_range.min()
#         col = int((coord[0] - x_min) / self.x_resolution)
#         row = int((coord[1] - y_min) / self.y_resolution)
#         return (row, col)
#
#     def is_valid_position(self, pos):
#         return 0 <= pos[0] < self.TrafFlag_grid.shape[0] and 0 <= pos[1] < self.TrafFlag_grid.shape[1]
#
#     def split_data_and_evaluate(self):
#         try:
#             # 将所有数据合并
#             data = np.column_stack(
#                 (self.X, self.Y, self.Z, self.TrafFlag, self.maxAcc, self.Slope, self.StepElev, self.Roughness))
#
#             # 分割数据集为训练集和测试集
#             test_size = 0.2
#             random_state = 42
#
#             train_data, test_data = train_test_split(data, test_size=test_size, random_state=random_state)
#
#             # 拆分训练集和测试集的数据
#             X_train, Y_train = train_data[:, 0], train_data[:, 1]
#             Z_train = train_data[:, 2]
#             TrafFlag_train = train_data[:, 3]
#             maxAcc_train = train_data[:, 4]
#             Slope_train = train_data[:, 5]
#             StepElev_train = train_data[:, 6]
#             Rough_train = train_data[:, 7]
#
#             X_test, Y_test = test_data[:, 0], test_data[:, 1]
#             Z_test = test_data[:, 2]
#             TrafFlag_test = test_data[:, 3]
#             maxAcc_test = test_data[:, 4]
#             Slope_test = test_data[:, 5]
#             StepElev_test = test_data[:, 6]
#             Rough_test = test_data[:, 7]
#
#             # 使用训练集进行插值
#             Z_grid_train = griddata((X_train, Y_train), Z_train, (self.X_grid, self.Y_grid), method='cubic')
#             TrafFlag_grid_train = griddata((X_train, Y_train), TrafFlag_train, (self.X_grid, self.Y_grid), method='nearest')
#             Acc_grid_train = griddata((X_train, Y_train), maxAcc_train, (self.X_grid, self.Y_grid), method='nearest')
#             Slope_grid_train = griddata((X_train, Y_train), Slope_train, (self.X_grid, self.Y_grid), method='nearest')
#             StepElev_grid_train = griddata((X_train, Y_train), StepElev_train, (self.X_grid, self.Y_grid), method='nearest')
#             Rough_grid_train = griddata((X_train, Y_train), Rough_train, (self.X_grid, self.Y_grid), method='nearest')
#
#             # 在测试集上预测值
#             Z_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Z_grid_train.flatten(), (X_test, Y_test),
#                               method='cubic')
#             TrafFlag_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), TrafFlag_grid_train.flatten(),
#                                      (X_test, Y_test), method='nearest')
#             Acc_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Acc_grid_train.flatten(), (X_test, Y_test),
#                                 method='nearest')
#             Slope_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Slope_grid_train.flatten(),
#                                   (X_test, Y_test), method='nearest')
#             StepElev_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), StepElev_grid_train.flatten(),
#                                      (X_test, Y_test), method='nearest')
#             Rough_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Rough_grid_train.flatten(),
#                                   (X_test, Y_test), method='nearest')
#
#             # 处理 NaN 值并计算误差指标
#
#             # 高程数据误差
#             valid_indices_Z = ~np.isnan(Z_pred) & ~np.isnan(Z_test)
#             if np.any(valid_indices_Z):
#                 mse_Z = mean_squared_error(Z_test[valid_indices_Z], Z_pred[valid_indices_Z])
#                 mae_Z = mean_absolute_error(Z_test[valid_indices_Z], Z_pred[valid_indices_Z])
#                 max_error_Z = np.max(np.abs(Z_test[valid_indices_Z] - Z_pred[valid_indices_Z]))
#                 print(f"高程插值 - MSE: {mse_Z}, MAE: {mae_Z}, Max Error: {max_error_Z}")
#             else:
#                 print("高程插值预测值中全部为 NaN，无法计算误差。")
#
#             # 通行标志（分类准确率）
#             valid_indices_TrafFlag = ~np.isnan(TrafFlag_pred) & ~np.isnan(TrafFlag_test)
#             if np.any(valid_indices_TrafFlag):
#                 accuracy_TrafFlag = accuracy_score(TrafFlag_test[valid_indices_TrafFlag],
#                                                    TrafFlag_pred[valid_indices_TrafFlag])
#                 print(f"通行标志插值 - 准确率: {accuracy_TrafFlag}")
#             else:
#                 print("通行标志插值预测值中全部为 NaN，无法计算准确率。")
#
#             # 最大加速度误差
#             valid_indices_Acc = ~np.isnan(Acc_pred) & ~np.isnan(maxAcc_test)
#             if np.any(valid_indices_Acc):
#                 mse_Acc = mean_squared_error(maxAcc_test[valid_indices_Acc], Acc_pred[valid_indices_Acc])
#                 mae_Acc = mean_absolute_error(maxAcc_test[valid_indices_Acc], Acc_pred[valid_indices_Acc])
#                 max_error_Acc = np.max(np.abs(maxAcc_test[valid_indices_Acc] - Acc_pred[valid_indices_Acc]))
#                 print(f"最大加速度插值 - MSE: {mse_Acc}, MAE: {mae_Acc}, Max Error: {max_error_Acc}")
#             else:
#                 print("最大加速度插值预测值中全部为 NaN，无法计算误差。")
#
#             # 坡度误差
#             valid_indices_Slope = ~np.isnan(Slope_pred) & ~np.isnan(Slope_test)
#             if np.any(valid_indices_Slope):
#                 mse_Slope = mean_squared_error(Slope_test[valid_indices_Slope], Slope_pred[valid_indices_Slope])
#                 mae_Slope = mean_absolute_error(Slope_test[valid_indices_Slope], Slope_pred[valid_indices_Slope])
#                 max_error_Slope = np.max(np.abs(Slope_test[valid_indices_Slope] - Slope_pred[valid_indices_Slope]))
#                 print(f"坡度插值 - MSE: {mse_Slope}, MAE: {mae_Slope}, Max Error: {max_error_Slope}")
#             else:
#                 print("坡度插值预测值中全部为 NaN，无法计算误差。")
#
#             # 阶跃高程差误差
#             valid_indices_StepElev = ~np.isnan(StepElev_pred) & ~np.isnan(StepElev_test)
#             if np.any(valid_indices_StepElev):
#                 mse_StepElev = mean_squared_error(StepElev_test[valid_indices_StepElev],
#                                                   StepElev_pred[valid_indices_StepElev])
#                 mae_StepElev = mean_absolute_error(StepElev_test[valid_indices_StepElev],
#                                                    StepElev_pred[valid_indices_StepElev])
#                 max_error_StepElev = np.max(
#                     np.abs(StepElev_test[valid_indices_StepElev] - StepElev_pred[valid_indices_StepElev]))
#                 print(f"阶跃高程差插值 - MSE: {mse_StepElev}, MAE: {mae_StepElev}, Max Error: {max_error_StepElev}")
#             else:
#                 print("阶跃高程差插值预测值中全部为 NaN，无法计算误差。")
#
#             # 粗糙度误差
#             valid_indices_Rough = ~np.isnan(Rough_pred) & ~np.isnan(Rough_test)
#             if np.any(valid_indices_Rough):
#                 mse_Rough = mean_squared_error(Rough_test[valid_indices_Rough], Rough_pred[valid_indices_Rough])
#                 mae_Rough = mean_absolute_error(Rough_test[valid_indices_Rough], Rough_pred[valid_indices_Rough])
#                 max_error_Rough = np.max(np.abs(Rough_test[valid_indices_Rough] - Rough_pred[valid_indices_Rough]))
#                 print(f"粗糙度插值 - MSE: {mse_Rough}, MAE: {mae_Rough}, Max Error: {max_error_Rough}")
#             else:
#                 print("粗糙度插值预测值中全部为 NaN，无法计算误差。")
#         except Exception as e:
#             print(f"数据评估时发生错误: {e}")
#             logging.exception("数据评估时发生错误")
#
# # 节点类和路径规划类
# class Node:
#     def __init__(self, position, parent=None):
#         self.position = position  # (row, col)
#         self.parent = parent
#         self.g = 0  # 累计代价
#         self.h = 0  # 启发代价
#         self.f = 0  # 总代价
#
#     def __lt__(self, other):
#         return self.f < other.f
#
# class PathPlanner:
#     def __init__(self, grid_data, max_slope=None, max_step_elev=None, max_roughness=None, max_speed=70.0):
#         self.grid = grid_data['TrafFlag_grid']
#         self.Z_grid = grid_data['Z_grid']
#         self.Acc_grid = grid_data['Acc_grid']
#         self.Slope_grid = grid_data['Slope_grid']
#         self.StepElev_grid = grid_data['StepElev_grid']
#         self.Rough_grid = grid_data['Rough_grid']
#         self.max_slope = max_slope
#         self.max_step_elev = max_step_elev
#         self.max_roughness = max_roughness  # 新增最大粗糙度参数
#         self.max_speed = max_speed
#
#     def heuristic_distance_3d(self, current_pos, end_pos):
#         d_xy = np.sqrt((current_pos[0] - end_pos[0]) ** 2 +
#                        (current_pos[1] - end_pos[1]) ** 2)
#         z_current = self.Z_grid[current_pos[0], current_pos[1]]
#         z_end = self.Z_grid[end_pos[0], end_pos[1]]
#         d_z = (z_end - z_current) ** 2
#         return np.sqrt(d_xy ** 2 + d_z)
#
#     def astar(self, start, end, mode='distance', alpha_env=0.6, beta_env=0.4):
#         start_node = Node(start)
#         end_node = Node(end)
#
#         open_list = []
#         heapq.heappush(open_list, start_node)
#         open_dict = {start_node.position: start_node}
#
#         closed_set = set()
#
#         while open_list:
#             current_node = heapq.heappop(open_list)
#             del open_dict[current_node.position]
#             closed_set.add(current_node.position)
#
#             # 检查是否到达终点
#             if current_node.position == end_node.position:
#                 path = []
#                 while current_node:
#                     path.append(current_node.position)
#                     current_node = current_node.parent
#                 return path[::-1]
#
#             # 定义8个邻居（包括对角线）
#             neighbors = [(0, 1), (0, -1), (1, 0), (-1, 0),
#                          (1, 1), (1, -1), (-1, 1), (-1, -1)]
#
#             for offset in neighbors:
#                 neighbor_pos = (current_node.position[0] + offset[0],
#                                 current_node.position[1] + offset[1])
#
#                 # 检查邻居是否在网格范围内
#                 if not (0 <= neighbor_pos[0] < self.grid.shape[0] and
#                         0 <= neighbor_pos[1] < self.grid.shape[1]):
#                     continue
#
#                 # 检查 TrafFlag 是否可通行
#                 if self.grid[neighbor_pos] == 0:
#                     continue  # 不可通行区域
#
#                 # 在所有模式下应用坡度和阶跃高程差的限制
#                 if self.max_slope is not None and self.Slope_grid[neighbor_pos] > self.max_slope:
#                     continue  # 坡度超过最大限制
#                 if self.max_step_elev is not None and self.StepElev_grid[neighbor_pos] > self.max_step_elev:
#                     continue  # 阶跃高程差超过最大限制
#
#                 # 新增对最大粗糙度的限制
#                 if self.max_roughness is not None and self.Rough_grid[neighbor_pos] > self.max_roughness:
#                     continue  # 粗糙度超过最大限制
#
#                 # 检查是否已经在关闭列表中
#                 if neighbor_pos in closed_set:
#                     continue
#
#                 # 处理不同模式
#                 if mode == 'distance':
#                     movement_cost = self.heuristic_distance_3d(current_node.position, neighbor_pos)
#                     g_cost = current_node.g + movement_cost
#                     h_cost = self.heuristic_distance_3d(neighbor_pos, end)
#                 elif mode == 'time':
#                     a_max = self.Acc_grid[neighbor_pos[0], neighbor_pos[1]]
#                     if a_max <= 0:
#                         continue  # 无法加速，跳过
#                     movement_cost = self.heuristic_distance_3d(current_node.position, neighbor_pos)
#                     g_cost = current_node.g + movement_cost / a_max  # 简化为距离除以加速度
#                     h_cost = self.heuristic_distance_3d(neighbor_pos, end)
#                 elif mode == 'env':
#                     movement_uncertainty = (self.Slope_grid[neighbor_pos[0], neighbor_pos[1]] / 45.0 +
#                                             self.StepElev_grid[neighbor_pos[0], neighbor_pos[1]] / 1.5 +
#                                             self.Rough_grid[neighbor_pos[0], neighbor_pos[1]] / 6.0)
#                     g_cost = current_node.g + movement_uncertainty
#                     h_cost = self.heuristic_distance_3d(neighbor_pos, end)
#                 else:
#                     raise ValueError("Invalid mode. Choose from 'distance', 'time', 'env'.")
#
#                 # 创建邻居节点
#                 neighbor_node = Node(neighbor_pos, current_node)
#                 neighbor_node.g = g_cost
#                 neighbor_node.h = h_cost
#                 neighbor_node.f = neighbor_node.g + neighbor_node.h
#
#                 # 检查开放列表中是否有更好的节点
#                 if neighbor_pos in open_dict:
#                     existing_node = open_dict[neighbor_pos]
#                     if existing_node.g <= neighbor_node.g:
#                         continue  # 已有更优节点
#                     else:
#                         # 替换为更优节点
#                         open_list.remove(existing_node)
#                         heapq.heapify(open_list)
#                         heapq.heappush(open_list, neighbor_node)
#                         open_dict[neighbor_pos] = neighbor_node
#                 else:
#                     heapq.heappush(open_list, neighbor_node)
#                     open_dict[neighbor_pos] = neighbor_node
#
#         return None  # 如果找不到路径，返回 None
#
#     def extract_path_coordinates(self, path, x_range, y_range):
#         return [(x_range[p[1]], y_range[p[0]], self.Z_grid[p[0], p[1]]) for p in path]
#
#     def calculate_path_length(self, path_coords):
#         length = 0
#         for i in range(1, len(path_coords)):
#             x1, y1, z1 = path_coords[i - 1]
#             x2, y2, z2 = path_coords[i]
#             length += np.sqrt((x2 - x1) ** 2 + (y2 - y1) ** 2 + (z2 - z1) ** 2)
#         return length
#
#     def compute_traversal_time(self, path_coords, y_range, x_range):
#         total_time = 0.0
#         current_speed = 0.0  # 初始速度为0
#         max_speed = self.max_speed  # 使用用户指定的最大速度
#
#         speeds = [current_speed]  # 存储每个点的速度
#         accelerations = []  # 存储每段的加速度
#
#         for i in range(1, len(path_coords)):
#             x1, y1, z1 = path_coords[i - 1]
#             x2, y2, z2 = path_coords[i]
#             # 计算两点之间的距离
#             d = np.sqrt((x2 - x1) ** 2 + (y2 - y1) ** 2 + (z2 - z1) ** 2)
#
#             # 获取当前点的网格索引
#             row = np.argmin(np.abs(y_range - y1))
#             col = np.argmin(np.abs(x_range - x1))
#
#             # 获取当前点的加速度
#             a = self.Acc_grid[row, col]
#             if a <= 0:
#                 # 如果加速度不可用或为0，无法继续加速
#                 a = 0.0
#
#             # 计算需要加速到最大速度的速度增量
#             delta_v = max_speed - current_speed
#
#             if current_speed < max_speed and a > 0:
#                 # 需要加速
#                 t_acc = delta_v / a  # 加速所需时间
#                 d_acc = current_speed * t_acc + 0.5 * a * t_acc ** 2  # 加速所需距离
#
#                 if d_acc >= d:
#                     # 在当前段无法加速到最大速度
#                     # 使用运动学方程计算通过距离d所需的时间
#                     # 解方程: 0.5 * a * t^2 + current_speed * t - d = 0
#                     discriminant = current_speed ** 2 + 2 * a * d
#                     if discriminant < 0:
#                         # 无实数解，无法到达下一个点
#                         return float('inf'), [], []
#                     t = (-current_speed + math.sqrt(discriminant)) / a
#                     total_time += t
#                     current_speed += a * t
#                     accelerations.append(a)
#                     speeds.append(current_speed)
#                 else:
#                     # 在当前段可以加速到最大速度
#                     # 先加速到最大速度
#                     total_time += t_acc
#                     current_speed = max_speed
#                     accelerations.append(a)
#                     speeds.append(current_speed)
#                     # 剩余距离
#                     d_remain = d - d_acc
#                     # 以最大速度匀速行驶剩余距离
#                     t_remain = d_remain / current_speed
#                     total_time += t_remain
#                     accelerations.append(0.0)
#                     speeds.append(current_speed)
#             else:
#                 # 无法加速或已达到最大速度，匀速行驶
#                 if current_speed == 0:
#                     # 无法移动
#                     return float('inf'), [], []
#                 t = d / current_speed
#                 total_time += t
#                 accelerations.append(0.0)
#                 speeds.append(current_speed)
#
#         return total_time, speeds, accelerations
#
# # 可视化类
# class Visualizer:
#     def __init__(self, X_grid, Y_grid, Z_grid, TrafFlag_grid):
#         self.X_grid = X_grid
#         self.Y_grid = Y_grid
#         self.Z_grid = Z_grid
#         self.TrafFlag_grid = TrafFlag_grid
#         self.fig = go.Figure()
#
#     def plot_terrain(self):
#         # 定义自定义颜色尺度
#         colorscale = [
#             [0.0, 'gray'],        # TrafFlag == 0 的区域为灰色
#             [1.0, 'lightgreen']   # TrafFlag == 1 的区域为浅绿色
#         ]
#
#         self.fig.add_trace(go.Surface(
#             x=self.X_grid, y=self.Y_grid, z=self.Z_grid,
#             surfacecolor=self.TrafFlag_grid,
#             colorscale=colorscale,
#             cmin=0,
#             cmax=1,
#             opacity=0.8,
#             showscale=True,
#             colorbar=dict(title='TrafFlag')
#         ))
#
#     def add_path(self, coords, speeds, accs, color, name, length, time_taken, path_width=4):
#         if coords:
#             x, y, z = zip(*coords)
#             # 处理不可行路径（时间为inf）的情况
#             if math.isinf(time_taken):
#                 time_str = "无法到达"
#             else:
#                 time_str = f"{time_taken:.2f} 秒"
#
#             # 准备 customdata，将速度和加速度添加进去
#             if accs:
#                 accs = [0.0] + accs  # 与 speeds 对齐
#             else:
#                 accs = [0.0] * len(speeds)
#
#             customdata = np.array([speeds, accs]).T  # 转置使其每行对应一个点的 [speed, acceleration]
#
#             self.fig.add_trace(go.Scatter3d(
#                 x=x, y=y, z=z,
#                 mode='lines+markers',
#                 line=dict(color=color, width=path_width),
#                 marker=dict(size=2, color=color),
#                 name=f"{name} (长度: {length:.2f} 米, 时间: {time_str})",
#                 customdata=customdata,
#                 hovertemplate=(
#                         'X: %{x}<br>' +
#                         'Y: %{y}<br>' +
#                         'Z: %{z}<br>' +
#                         '速度: %{customdata[0]:.2f} m/s<br>' +
#                         '加速度: %{customdata[1]:.2f} m/s²<br>' +
#                         '<extra></extra>'
#                 )
#             ))
#         else:
#             print(f"未找到{name}")
#
#     def add_start_and_end_points(self, start_coord, end_coord):
#         self.fig.add_trace(go.Scatter3d(
#             x=[start_coord[0]],
#             y=[start_coord[1]],
#             z=[start_coord[2]],
#             mode='markers',
#             marker=dict(size=8, color='green'),
#             name='起点'
#         ))
#         self.fig.add_trace(go.Scatter3d(
#             x=[end_coord[0]],
#             y=[end_coord[1]],
#             z=[end_coord[2]],
#             mode='markers',
#             marker=dict(size=8, color='red'),
#             name='终点'
#         ))
#
#     def show(self):
#         self.fig.update_layout(
#             scene=dict(
#                 xaxis_title='X 坐标',
#                 yaxis_title='Y 坐标',
#                 zaxis_title='Z 坐标 (高程)',
#                 aspectmode='data'
#             ),
#             title='3D 地形路径规划（基于距离、时间与最小不确定性）',
#             legend=dict(
#                 x=0,
#                 y=1,
#                 bgcolor='rgba(255, 255, 255, 0)',
#                 bordercolor='rgba(255, 255, 255, 0)'
#             )
#         )
#         pio.show(self.fig)
#
# # 报告生成函数
# def write_paths_to_word(document, path_name, path_coords):
#     if not path_coords:
#         return
#     # 添加路径名称和点数量
#     document.add_heading(f"{path_name} ( {len(path_coords)})", level=2)
#     # 创建表格，包含三列：X相对坐标、Y相对坐标、Z高程坐标
#     table = document.add_table(rows=1, cols=3)
#     table.style = 'Light Grid'
#     hdr_cells = table.rows[0].cells
#     hdr_cells[0].text = 'X'
#     hdr_cells[1].text = 'Y'
#     hdr_cells[2].text = 'Z'
#     # 填充表格内容
#     for coord in path_coords:
#         row_cells = table.add_row().cells
#         row_cells[0].text = f"{coord[0]}"
#         row_cells[1].text = f"{coord[1]}"
#         row_cells[2].text = f"{coord[2]}"
#     # 添加一个空行
#     document.add_paragraph()
#
# # 路径规划函数
# def run_path_planning(data_processor, start_coord, end_coord, output_path, max_allowed_slope, max_allowed_step_elev, max_allowed_roughness, max_speed):
#     try:
#         # 确保已进行数据评估
#         data_processor.split_data_and_evaluate()
#
#         # 将 start_coord 和 end_coord 转换为网格索引
#         start = data_processor.coord_to_index(start_coord)
#         end = data_processor.coord_to_index(end_coord)
#
#         # 检查起点和终点是否在网格范围内
#         if not data_processor.is_valid_position(start):
#             raise ValueError("起点坐标超出网格范围")
#         if not data_processor.is_valid_position(end):
#             raise ValueError("终点坐标超出网格范围")
#
#         # 确保起点和终点可通行
#         if data_processor.TrafFlag_grid[start] == 0:
#             raise ValueError("起点不可通行")
#         if data_processor.TrafFlag_grid[end] == 0:
#             raise ValueError("终点不可通行")
#
#         # 路径规划
#         grid_data = {
#             'TrafFlag_grid': data_processor.TrafFlag_grid,
#             'Z_grid': data_processor.Z_grid,
#             'Acc_grid': data_processor.Acc_grid,
#             'Slope_grid': data_processor.Slope_grid,
#             'StepElev_grid': data_processor.StepElev_grid,
#             'Rough_grid': data_processor.Rough_grid
#         }
#         path_planner = PathPlanner(grid_data, max_slope=max_allowed_slope, max_step_elev=max_allowed_step_elev, max_roughness=max_allowed_roughness, max_speed=max_speed)
#
#         print("正在计算基于距离的最短路径...")
#         path_distance = path_planner.astar(start, end, mode='distance')
#         if path_distance:
#             print("基于距离的路径计算完成")
#         else:
#             print("未找到基于距离的路径")
#
#         print("正在计算基于时间的最短耗时路径...")
#         path_time = path_planner.astar(start, end, mode='time')
#         if path_time:
#             print("基于时间的路径计算完成")
#         else:
#             print("未找到基于时间的路径")
#
#         print("正在计算最小不确定性路径...")
#         path_env = path_planner.astar(start, end, mode='env')
#         if path_env:
#             print("最小不确定性路径计算完成")
#         else:
#             print("未找到最小不确定性路径")
#
#         # 提取路径坐标
#         if path_distance:
#             path_distance_coords = path_planner.extract_path_coordinates(path_distance, data_processor.x_range,
#                                                                          data_processor.y_range)
#             length_distance = path_planner.calculate_path_length(path_distance_coords)
#         else:
#             path_distance_coords = None
#             length_distance = None
#         if path_time:
#             path_time_coords = path_planner.extract_path_coordinates(path_time, data_processor.x_range,
#                                                                      data_processor.y_range)
#             length_time = path_planner.calculate_path_length(path_time_coords)
#         else:
#             path_time_coords = None
#             length_time = None
#         if path_env:
#             path_env_coords = path_planner.extract_path_coordinates(path_env, data_processor.x_range,
#                                                                     data_processor.y_range)
#             length_env = path_planner.calculate_path_length(path_env_coords)
#         else:
#             path_env_coords = None
#             length_env = None
#
#         # 计算遍历时间，并获取速度和加速度数据
#         if path_distance:
#             time_distance, speeds_distance, accs_distance = path_planner.compute_traversal_time(
#                 path_distance_coords, data_processor.y_range, data_processor.x_range)
#         else:
#             speeds_distance, accs_distance, time_distance = None, None, None
#
#         if path_time:
#             time_time, speeds_time, accs_time = path_planner.compute_traversal_time(
#                 path_time_coords, data_processor.y_range, data_processor.x_range)
#         else:
#             speeds_time, accs_time, time_time = None, None, None
#
#         if path_env:
#             time_env, speeds_env, accs_env = path_planner.compute_traversal_time(
#                 path_env_coords, data_processor.y_range, data_processor.x_range)
#         else:
#             speeds_env, accs_env, time_env = None, None, None
#
#         # 可视化
#         visualizer = Visualizer(
#             data_processor.X_grid,
#             data_processor.Y_grid,
#             data_processor.Z_grid,
#             data_processor.TrafFlag_grid
#         )
#         visualizer.plot_terrain()
#
#         colors = ['blue', 'orange', 'purple']
#         names = ['MinDistPath 最短距离路径', 'MinTimePath 最短耗时路径', 'MinUncertPath 最小不确定性路径']
#         lengths = [length_distance, length_time, length_env]
#         times = [time_distance, time_time, time_env]
#         coords_list = [path_distance_coords, path_time_coords, path_env_coords]
#         speeds_list = [speeds_distance, speeds_time, speeds_env]
#         accs_list = [accs_distance, accs_time, accs_env]
#
#         for coords, speeds, accs, color, name, length, time_taken in zip(
#                 coords_list, speeds_list, accs_list,
#                 colors, names, lengths, times):
#             visualizer.add_path(coords, speeds, accs, color, name, length, time_taken)
#
#         # 标记起点和终点
#         if path_distance or path_time or path_env:
#             selected_path = None
#             if path_distance:
#                 selected_path = path_distance_coords
#             elif path_time:
#                 selected_path = path_time_coords
#             elif path_env:
#                 selected_path = path_env_coords
#
#             if selected_path:
#                 # 使用起点坐标
#                 start_x, start_y, start_z = selected_path[0]
#                 # 使用终点坐标
#                 end_x, end_y, end_z = selected_path[-1]
#                 visualizer.add_start_and_end_points((start_x, start_y, start_z), (end_x, end_y, end_z))
#
#         visualizer.show()
#
#         # 创建 Word 文档
#         document = Document()
#         document.add_heading('Route planning results', 0)
#         # 写入路径信息
#         if path_distance or path_time or path_env:
#             if path_distance:
#                 write_paths_to_word(document, 'MinDistPath ', path_distance_coords)
#             if path_time:
#                 write_paths_to_word(document, 'MinTimePath ', path_time_coords)
#             if path_env:
#                 write_paths_to_word(document, 'MinUncertPath ', path_env_coords)
#
#             # 保存文档
#             try:
#                 # 使用 resource_path 函数获取输出文件路径
#                 output_file_path = output_path  # 如果 output_path 是用户指定的完整路径，则无需使用 resource_path
#                 document.save(output_file_path)
#                 print(f"路径坐标已成功保存到 Word 文档: {output_file_path}")
#             except Exception as e:
#                 print(f"保存 Word 文档时出错: {e}")
#                 logging.exception("保存 Word 文档时出错")
#         else:
#             print("没有路径可写入 Word 文档")
#
#         # 打印路径长度和遍历时间
#         if path_distance:
#             if math.isinf(time_distance):
#                 print(f"最短路径长度: {length_distance:.2f} 米")
#                 print(f"最短距离路径遍历时间: 无法到达")
#             else:
#                 print(f"最短路径长度: {length_distance:.2f} 米")
#                 print(f"最短距离路径遍历时间: {time_distance:.2f} 秒")
#         if path_time:
#             if math.isinf(time_time):
#                 print(f"最短耗时路径长度: {length_time:.2f} 米")
#                 print(f"最短耗时路径遍历时间: 无法到达")
#             else:
#                 print(f"最短耗时路径长度: {length_time:.2f} 米")
#                 print(f"最短耗时路径遍历时间: {time_time:.2f} 秒")
#         if path_env:
#             if math.isinf(time_env):
#                 print(f"最小不确定性路径长度: {length_env:.2f} 米")
#                 print(f"最小不确定性路径遍历时间: 无法到达")
#             else:
#                 print(f"最小不确定性路径长度: {length_env:.2f} 米")
#                 print(f"最小不确定性路径遍历时间: {time_env:.2f} 秒")
#     except Exception as e:
#         print(f"运行处理时发生错误: {e}")
#         logging.exception("运行处理时发生错误")
#
# # 主程序
# def main():
#     try:
#         # 创建主窗口
#         root = tk.Tk()
#         root.title("路径规划软件")
#
#         # 地图数据文件路径
#         tk.Label(root, text="地图数据文件路径:").grid(row=0, column=0, sticky='e')
#         map_data_path_entry = tk.Entry(root, width=50)
#         map_data_path_entry.grid(row=0, column=1)
#         def browse_map_data():
#             filename = filedialog.askopenfilename(filetypes=[("CSV文件", "*.csv")])
#             if filename:
#                 map_data_path_entry.delete(0, tk.END)
#                 map_data_path_entry.insert(0, filename)
#         tk.Button(root, text="浏览...", command=browse_map_data).grid(row=0, column=2)
#
#         # 起点坐标（初始禁用）
#         tk.Label(root, text="起点坐标 (格式: X,Y):").grid(row=1, column=0, sticky='e')
#         start_coord_entry = tk.Entry(root, width=50)
#         start_coord_entry.grid(row=1, column=1)
#         # 终点坐标（初始禁用）
#         tk.Label(root, text="终点坐标 (格式: X,Y):").grid(row=2, column=0, sticky='e')
#         end_coord_entry = tk.Entry(root, width=50)
#         end_coord_entry.grid(row=2, column=1)
#
#         # 输出路径
#         tk.Label(root, text="输出路径文档路径:").grid(row=3, column=0, sticky='e')
#         output_path_entry = tk.Entry(root, width=50)
#         output_path_entry.grid(row=3, column=1)
#         def browse_output_path():
#             filename = filedialog.asksaveasfilename(defaultextension=".docx",
#                                                     filetypes=[("Word文档", "*.docx")])
#             if filename:
#                 output_path_entry.delete(0, tk.END)
#                 output_path_entry.insert(0, filename)
#         tk.Button(root, text="浏览...", command=browse_output_path).grid(row=3, column=2)
#
#         # 最大坡度
#         tk.Label(root, text="最大爬坡度 (默认25.0度):").grid(row=4, column=0, sticky='e')
#         max_slope_entry = tk.Entry(root, width=50)
#         max_slope_entry.grid(row=4, column=1)
#         max_slope_entry.insert(0, "25.0")  # 设置默认值
#
#         # 最大阶跃高程差
#         tk.Label(root, text="最大阶跃高程差 (默认0.3米):").grid(row=5, column=0, sticky='e')
#         max_step_elev_entry = tk.Entry(root, width=50)
#         max_step_elev_entry.grid(row=5, column=1)
#         max_step_elev_entry.insert(0, "0.3")  # 设置默认值为 0.3
#
#         # 添加最大粗糙度输入字段
#         tk.Label(root, text="最大粗糙度 (默认6.0):").grid(row=6, column=0, sticky='e')
#         max_roughness_entry = tk.Entry(root, width=50)
#         max_roughness_entry.grid(row=6, column=1)
#         max_roughness_entry.insert(0, "6.0")  # 设置默认值
#
#         # 最大速度
#         tk.Label(root, text="最大速度 (默认20.0 m/s):").grid(row=7, column=0, sticky='e')
#         max_speed_entry = tk.Entry(root, width=50)
#         max_speed_entry.grid(row=7, column=1)
#         max_speed_entry.insert(0, "20.0")  # 设置默认值
#
#         # 禁用输入字段，直到地图加载完成
#         start_coord_entry.configure(state='disabled')
#         end_coord_entry.configure(state='disabled')
#         output_path_entry.configure(state='disabled')
#         max_slope_entry.configure(state='disabled')
#         max_step_elev_entry.configure(state='disabled')
#         max_roughness_entry.configure(state='disabled')
#         max_speed_entry.configure(state='disabled')
#
#         # 加载地图函数
#         def load_map():
#             csv_path = map_data_path_entry.get()
#             if not csv_path:
#                 messagebox.showerror("错误", "请提供地图数据文件路径。")
#                 return
#
#             def process_map():
#                 try:
#                     data_processor = DataProcessor(csv_path)
#                     data_processor.load_data()
#                     data_processor.interpolate_data()
#                     data_processor.fill_nan_values()
#
#                     # 可视化地图
#                     visualizer = Visualizer(
#                         data_processor.X_grid,
#                         data_processor.Y_grid,
#                         data_processor.Z_grid,
#                         data_processor.TrafFlag_grid
#                     )
#                     visualizer.plot_terrain()
#                     visualizer.show()
#
#                     # 存储 data_processor
#                     root.data_processor = data_processor
#
#                     # 启用输入字段
#                     root.after(0, enable_input_fields)
#                     messagebox.showinfo("完成", "地图已加载并显示。请继续输入起点和终点坐标。")
#                 except Exception as e:
#                     messagebox.showerror("错误", f"加载地图时发生错误: {e}")
#                     logging.exception("加载地图时发生错误")
#
#             threading.Thread(target=process_map).start()
#
#         # 启用输入字段函数
#         def enable_input_fields():
#             start_coord_entry.configure(state='normal')
#             end_coord_entry.configure(state='normal')
#             output_path_entry.configure(state='normal')
#             max_slope_entry.configure(state='normal')
#             max_step_elev_entry.configure(state='normal')
#             max_roughness_entry.configure(state='normal')
#             max_speed_entry.configure(state='normal')
#             run_button.configure(state='normal')
#
#         # 加载地图按钮
#         tk.Button(root, text="加载地图", command=load_map).grid(row=8, column=1)
#
#         # 运行路径规划函数
#         def run():
#             # 获取输入值
#             start_coord_input = start_coord_entry.get()
#             end_coord_input = end_coord_entry.get()
#             output_path = output_path_entry.get()
#
#             # 验证输入
#             if not output_path:
#                 messagebox.showerror("错误", "请提供输出 Word 文档路径。")
#                 return
#             if not start_coord_input:
#                 messagebox.showerror("错误", "请提供起点坐标。")
#                 return
#             if not end_coord_input:
#                 messagebox.showerror("错误", "请提供终点坐标。")
#                 return
#             try:
#                 start_x, start_y = map(float, start_coord_input.strip().split(','))
#                 start_coord = (start_x, start_y)
#             except ValueError:
#                 messagebox.showerror("错误", "起点坐标格式错误，请使用 'X,Y' 的格式。")
#                 return
#
#             try:
#                 end_x, end_y = map(float, end_coord_input.strip().split(','))
#                 end_coord = (end_x, end_y)
#             except ValueError:
#                 messagebox.showerror("错误", "终点坐标格式错误，请使用 'X,Y' 的格式。")
#                 return
#
#             # 获取最大坡度和最大阶跃高程差
#             max_slope_input = max_slope_entry.get()
#             max_step_elev_input = max_step_elev_entry.get()
#
#             try:
#                 max_allowed_slope = float(max_slope_input)
#             except ValueError:
#                 messagebox.showerror("错误", "最大坡度格式错误，请输入数字。")
#                 return
#
#             try:
#                 max_allowed_step_elev = float(max_step_elev_input)
#             except ValueError:
#                 messagebox.showerror("错误", "最大阶跃高程差格式错误，请输入数字。")
#                 return
#
#             # 获取最大粗糙度
#             max_roughness_input = max_roughness_entry.get()
#             try:
#                 max_allowed_roughness = float(max_roughness_input)
#             except ValueError:
#                 messagebox.showerror("错误", "最大粗糙度格式错误，请输入数字。")
#                 return
#
#             # 获取最大速度
#             max_speed_input = max_speed_entry.get()
#
#             try:
#                 max_speed = float(max_speed_input)
#             except ValueError:
#                 messagebox.showerror("错误", "最大速度格式错误，请输入数字。")
#                 return
#
#             # 在后台线程中运行，以避免阻塞GUI
#             def process():
#                 try:
#                     # 使用已加载的 data_processor
#                     data_processor = root.data_processor
#                     # 调用路径规划函数
#                     run_path_planning(data_processor, start_coord, end_coord, output_path, max_allowed_slope, max_allowed_step_elev, max_allowed_roughness, max_speed)
#                     messagebox.showinfo("完成", "路径规划已完成。")
#                 except Exception as e:
#                     messagebox.showerror("错误", f"发生错误: {e}")
#                     logging.exception("后台线程运行时发生错误")
#
#             threading.Thread(target=process).start()
#
#         # 运行按钮（初始禁用）
#         run_button = tk.Button(root, text="运行", command=run)
#         run_button.grid(row=9, column=1)
#         run_button.configure(state='disabled')
#
#         root.mainloop()
#     except Exception as e:
#         print(f"主程序运行时发生错误: {e}")
#         logging.exception("主程序运行时发生错误")
#
# if __name__ == "__main__":
#     main()









import pandas as pd
import numpy as np
from scipy.interpolate import griddata
from sklearn.metrics import mean_squared_error, mean_absolute_error, accuracy_score
from sklearn.model_selection import train_test_split
import heapq
import math
import plotly.graph_objs as go
import plotly.io as pio
from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox
import threading
import sys
import os
import logging

# 设置日志记录
logging.basicConfig(filename='error.log', level=logging.ERROR)

# 设置 Plotly 渲染器为浏览器
pio.renderers.default = 'browser'

# 添加 resource_path 函数，用于获取资源文件的正确路径
def resource_path(relative_path):
    """获取资源文件的绝对路径，兼容 PyInstaller 打包后的情况"""
    try:
        # PyInstaller 会把路径存储在 _MEIPASS 中
        base_path = sys._MEIPASS
    except AttributeError:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

# 数据处理类
class DataProcessor:
    def __init__(self, dat_path, x_resolution=1, y_resolution=1):
        self.dat_path = dat_path
        self.x_resolution = x_resolution
        self.y_resolution = y_resolution
        self.df = None
        self.X = None
        self.Y = None
        self.Z = None
        self.TrafFlag = None
        self.maxAcc = None
        self.Slope = None
        self.StepElev = None
        self.Roughness = None
        self.x_range = None
        self.y_range = None
        self.X_grid = None
        self.Y_grid = None
        self.Z_grid = None
        self.TrafFlag_grid = None
        self.Acc_grid = None
        self.Slope_grid = None
        self.StepElev_grid = None
        self.Rough_grid = None

    def load_data(self):
        try:
            # 使用 resource_path 函数获取数据文件路径
            data_path = resource_path(self.dat_path)
            # 读取 .csv 文件，指定分隔符为逗号
            self.df = pd.read_csv(data_path, delimiter=',')
            print("数据文件中的列名：", self.df.columns.tolist())
        except FileNotFoundError:
            print(f"文件未找到: {self.dat_path}")
            return
        except pd.errors.EmptyDataError:
            print("CSV 文件为空")
            return
        except Exception as e:
            print(f"读取 CSV 文件时出错: {e}")
            logging.exception("读取 CSV 文件时出错")
            return

        required_columns = {'X', 'Y', 'Z', 'TrafFlag', 'maxAcc', 'Slope', 'StepElevationDifference', 'Roughness'}
        if not required_columns.issubset(self.df.columns):
            print(f"CSV文件缺少必要的列: {required_columns - set(self.df.columns)}")
            return

        # 提取所需数据
        self.X = self.df['X'].values
        self.Y = self.df['Y'].values
        self.Z = self.df['Z'].values
        self.TrafFlag = self.df['TrafFlag'].values
        self.maxAcc = self.df['maxAcc'].values
        self.Slope = self.df['Slope'].values
        self.StepElev = self.df['StepElevationDifference'].values
        self.Roughness = self.df['Roughness'].values

        # 计算并打印 Slope 和 StepElevationDifference 的最小值和最大值
        slope_min = np.min(self.Slope)
        slope_max_val = np.max(self.Slope)
        step_min = np.min(self.StepElev)
        step_max = np.max(self.StepElev)

        print(f"坡度 (Slope) 最小值: {slope_min}")
        print(f"坡度 (Slope) 最大值: {slope_max_val}")
        print(f"阶跃高程差 (StepElevationDifference) 最小值: {step_min}")
        print(f"阶跃高程差 (StepElevationDifference) 最大值: {step_max}")

    def interpolate_data(self):
        # 计算网格范围和生成网格
        x_min, x_max = self.X.min(), self.X.max()
        y_min, y_max = self.Y.min(), self.Y.max()
        self.x_range = np.arange(x_min, x_max + self.x_resolution, self.x_resolution)
        self.y_range = np.arange(y_min, y_max + self.y_resolution, self.y_resolution)
        self.X_grid, self.Y_grid = np.meshgrid(self.x_range, self.y_range)

        # 插值
        self.Z_grid = griddata((self.X, self.Y), self.Z, (self.X_grid, self.Y_grid), method='cubic')
        self.TrafFlag_grid = griddata((self.X, self.Y), self.TrafFlag, (self.X_grid, self.Y_grid), method='nearest')
        self.Acc_grid = griddata((self.X, self.Y), self.maxAcc, (self.X_grid, self.Y_grid), method='nearest')
        self.Slope_grid = griddata((self.X, self.Y), self.Slope, (self.X_grid, self.Y_grid), method='nearest')
        self.StepElev_grid = griddata((self.X, self.Y), self.StepElev, (self.X_grid, self.Y_grid), method='nearest')
        self.Rough_grid = griddata((self.X, self.Y), self.Roughness, (self.X_grid, self.Y_grid), method='nearest')

    def fill_nan_values(self):
        # 填充 NaN 值
        self.Z_grid = np.nan_to_num(self.Z_grid, nan=np.nanmean(self.Z))
        self.Acc_grid = np.nan_to_num(self.Acc_grid, nan=0)
        self.TrafFlag_grid = np.nan_to_num(self.TrafFlag_grid, nan=0)
        # 将 TrafFlag_grid 的值限制为 0 或 1
        self.TrafFlag_grid = np.where(self.TrafFlag_grid >= 0.5, 1, 0)
        self.Slope_grid = np.nan_to_num(self.Slope_grid, nan=0)
        self.StepElev_grid = np.nan_to_num(self.StepElev_grid, nan=0)
        self.Rough_grid = np.nan_to_num(self.Rough_grid, nan=0)

        # 验证网格数据的形状
        print(f"Z_grid shape: {self.Z_grid.shape}")
        print(f"TrafFlag_grid shape: {self.TrafFlag_grid.shape}")
        print(f"Acc_grid shape: {self.Acc_grid.shape}")
        print(f"Slope_grid shape: {self.Slope_grid.shape}")
        print(f"StepElev_grid shape: {self.StepElev_grid.shape}")
        print(f"Rough_grid shape: {self.Rough_grid.shape}")

        # 计算并打印插值后网格数据的最小值和最大值
        slope_grid_min = np.min(self.Slope_grid)
        slope_grid_max = np.max(self.Slope_grid)
        step_grid_min = np.min(self.StepElev_grid)
        step_grid_max = np.max(self.StepElev_grid)

        print(f"插值后坡度网格 (Slope_grid) 最小值: {slope_grid_min}")
        print(f"插值后坡度网格 (Slope_grid) 最大值: {slope_grid_max}")
        print(f"插值后阶跃高程差网格 (StepElev_grid) 最小值: {step_grid_min}")
        print(f"插值后阶跃高程差网格 (StepElev_grid) 最大值: {step_grid_max}")

    def coord_to_index(self, coord):
        x_min, y_min = self.x_range.min(), self.y_range.min()
        col = int((coord[0] - x_min) / self.x_resolution)
        row = int((coord[1] - y_min) / self.y_resolution)
        return (row, col)

    def is_valid_position(self, pos):
        return 0 <= pos[0] < self.TrafFlag_grid.shape[0] and 0 <= pos[1] < self.TrafFlag_grid.shape[1]

    def split_data_and_evaluate(self):
        try:
            # 将所有数据合并
            data = np.column_stack(
                (self.X, self.Y, self.Z, self.TrafFlag, self.maxAcc, self.Slope, self.StepElev, self.Roughness))

            # 分割数据集为训练集和测试集
            test_size = 0.2
            random_state = 42

            train_data, test_data = train_test_split(data, test_size=test_size, random_state=random_state)

            # 拆分训练集和测试集的数据
            X_train, Y_train = train_data[:, 0], train_data[:, 1]
            Z_train = train_data[:, 2]
            TrafFlag_train = train_data[:, 3]
            maxAcc_train = train_data[:, 4]
            Slope_train = train_data[:, 5]
            StepElev_train = train_data[:, 6]
            Rough_train = train_data[:, 7]

            X_test, Y_test = test_data[:, 0], test_data[:, 1]
            Z_test = test_data[:, 2]
            TrafFlag_test = test_data[:, 3]
            maxAcc_test = test_data[:, 4]
            Slope_test = test_data[:, 5]
            StepElev_test = test_data[:, 6]
            Rough_test = test_data[:, 7]

            # 使用训练集进行插值
            Z_grid_train = griddata((X_train, Y_train), Z_train, (self.X_grid, self.Y_grid), method='cubic')
            TrafFlag_grid_train = griddata((X_train, Y_train), TrafFlag_train, (self.X_grid, self.Y_grid), method='nearest')
            Acc_grid_train = griddata((X_train, Y_train), maxAcc_train, (self.X_grid, self.Y_grid), method='nearest')
            Slope_grid_train = griddata((X_train, Y_train), Slope_train, (self.X_grid, self.Y_grid), method='nearest')
            StepElev_grid_train = griddata((X_train, Y_train), StepElev_train, (self.X_grid, self.Y_grid), method='nearest')
            Rough_grid_train = griddata((X_train, Y_train), Rough_train, (self.X_grid, self.Y_grid), method='nearest')

            # 在测试集上预测值
            Z_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Z_grid_train.flatten(), (X_test, Y_test),
                              method='cubic')
            TrafFlag_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), TrafFlag_grid_train.flatten(),
                                     (X_test, Y_test), method='nearest')
            Acc_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Acc_grid_train.flatten(), (X_test, Y_test),
                                method='nearest')
            Slope_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Slope_grid_train.flatten(),
                                  (X_test, Y_test), method='nearest')
            StepElev_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), StepElev_grid_train.flatten(),
                                     (X_test, Y_test), method='nearest')
            Rough_pred = griddata((self.X_grid.flatten(), self.Y_grid.flatten()), Rough_grid_train.flatten(),
                                  (X_test, Y_test), method='nearest')

            # 处理 NaN 值并计算误差指标

            # 高程数据误差
            valid_indices_Z = ~np.isnan(Z_pred) & ~np.isnan(Z_test)
            if np.any(valid_indices_Z):
                mse_Z = mean_squared_error(Z_test[valid_indices_Z], Z_pred[valid_indices_Z])
                mae_Z = mean_absolute_error(Z_test[valid_indices_Z], Z_pred[valid_indices_Z])
                max_error_Z = np.max(np.abs(Z_test[valid_indices_Z] - Z_pred[valid_indices_Z]))
                print(f"高程插值 - MSE: {mse_Z}, MAE: {mae_Z}, Max Error: {max_error_Z}")
            else:
                print("高程插值预测值中全部为 NaN，无法计算误差。")

            # 通行标志（分类准确率）
            valid_indices_TrafFlag = ~np.isnan(TrafFlag_pred) & ~np.isnan(TrafFlag_test)
            if np.any(valid_indices_TrafFlag):
                accuracy_TrafFlag = accuracy_score(TrafFlag_test[valid_indices_TrafFlag],
                                                   TrafFlag_pred[valid_indices_TrafFlag])
                print(f"通行标志插值 - 准确率: {accuracy_TrafFlag}")
            else:
                print("通行标志插值预测值中全部为 NaN，无法计算准确率。")

            # 最大加速度误差
            valid_indices_Acc = ~np.isnan(Acc_pred) & ~np.isnan(maxAcc_test)
            if np.any(valid_indices_Acc):
                mse_Acc = mean_squared_error(maxAcc_test[valid_indices_Acc], Acc_pred[valid_indices_Acc])
                mae_Acc = mean_absolute_error(maxAcc_test[valid_indices_Acc], Acc_pred[valid_indices_Acc])
                max_error_Acc = np.max(np.abs(maxAcc_test[valid_indices_Acc] - Acc_pred[valid_indices_Acc]))
                print(f"最大加速度插值 - MSE: {mse_Acc}, MAE: {mae_Acc}, Max Error: {max_error_Acc}")
            else:
                print("最大加速度插值预测值中全部为 NaN，无法计算误差。")

            # 坡度误差
            valid_indices_Slope = ~np.isnan(Slope_pred) & ~np.isnan(Slope_test)
            if np.any(valid_indices_Slope):
                mse_Slope = mean_squared_error(Slope_test[valid_indices_Slope], Slope_pred[valid_indices_Slope])
                mae_Slope = mean_absolute_error(Slope_test[valid_indices_Slope], Slope_pred[valid_indices_Slope])
                max_error_Slope = np.max(np.abs(Slope_test[valid_indices_Slope] - Slope_pred[valid_indices_Slope]))
                print(f"坡度插值 - MSE: {mse_Slope}, MAE: {mae_Slope}, Max Error: {max_error_Slope}")
            else:
                print("坡度插值预测值中全部为 NaN，无法计算误差。")

            # 阶跃高程差误差
            valid_indices_StepElev = ~np.isnan(StepElev_pred) & ~np.isnan(StepElev_test)
            if np.any(valid_indices_StepElev):
                mse_StepElev = mean_squared_error(StepElev_test[valid_indices_StepElev],
                                                  StepElev_pred[valid_indices_StepElev])
                mae_StepElev = mean_absolute_error(StepElev_test[valid_indices_StepElev],
                                                   StepElev_pred[valid_indices_StepElev])
                max_error_StepElev = np.max(
                    np.abs(StepElev_test[valid_indices_StepElev] - StepElev_pred[valid_indices_StepElev]))
                print(f"阶跃高程差插值 - MSE: {mse_StepElev}, MAE: {mae_StepElev}, Max Error: {max_error_StepElev}")
            else:
                print("阶跃高程差插值预测值中全部为 NaN，无法计算误差。")

            # 粗糙度误差
            valid_indices_Rough = ~np.isnan(Rough_pred) & ~np.isnan(Rough_test)
            if np.any(valid_indices_Rough):
                mse_Rough = mean_squared_error(Rough_test[valid_indices_Rough], Rough_pred[valid_indices_Rough])
                mae_Rough = mean_absolute_error(Rough_test[valid_indices_Rough], Rough_pred[valid_indices_Rough])
                max_error_Rough = np.max(np.abs(Rough_test[valid_indices_Rough] - Rough_pred[valid_indices_Rough]))
                print(f"粗糙度插值 - MSE: {mse_Rough}, MAE: {mae_Rough}, Max Error: {max_error_Rough}")
            else:
                print("粗糙度插值预测值中全部为 NaN，无法计算误差。")
        except Exception as e:
            print(f"数据评估时发生错误: {e}")
            logging.exception("数据评估时发生错误")

# 节点类和路径规划类
class Node:
    def __init__(self, position, parent=None):
        self.position = position  # (row, col)
        self.parent = parent
        self.g = 0  # 累计代价
        self.h = 0  # 启发代价
        self.f = 0  # 总代价

    def __lt__(self, other):
        return self.f < other.f

class PathPlanner:
    def __init__(self, grid_data, max_slope=None, max_step_elev=None, max_roughness=None, max_speed=70.0):
        self.grid = grid_data['TrafFlag_grid']
        self.Z_grid = grid_data['Z_grid']
        self.Acc_grid = grid_data['Acc_grid']
        self.Slope_grid = grid_data['Slope_grid']
        self.StepElev_grid = grid_data['StepElev_grid']
        self.Rough_grid = grid_data['Rough_grid']
        self.max_slope = max_slope
        self.max_step_elev = max_step_elev
        self.max_roughness = max_roughness  # 新增最大粗糙度参数
        self.max_speed = max_speed

    def heuristic_distance_3d(self, current_pos, end_pos):
        d_xy = np.sqrt((current_pos[0] - end_pos[0]) ** 2 +
                       (current_pos[1] - end_pos[1]) ** 2)
        z_current = self.Z_grid[current_pos[0], current_pos[1]]
        z_end = self.Z_grid[end_pos[0], end_pos[1]]
        d_z = (z_end - z_current) ** 2
        return np.sqrt(d_xy ** 2 + d_z)

    def astar(self, start, end, mode='distance', alpha_env=0.6, beta_env=0.4):
        start_node = Node(start)
        end_node = Node(end)

        open_list = []
        heapq.heappush(open_list, start_node)
        open_dict = {start_node.position: start_node}

        closed_set = set()

        while open_list:
            current_node = heapq.heappop(open_list)
            del open_dict[current_node.position]
            closed_set.add(current_node.position)

            # 检查是否到达终点
            if current_node.position == end_node.position:
                path = []
                while current_node:
                    path.append(current_node.position)
                    current_node = current_node.parent
                return path[::-1]

            # 定义8个邻居（包括对角线）
            neighbors = [(0, 1), (0, -1), (1, 0), (-1, 0),
                         (1, 1), (1, -1), (-1, 1), (-1, -1)]

            for offset in neighbors:
                neighbor_pos = (current_node.position[0] + offset[0],
                                current_node.position[1] + offset[1])

                # 检查邻居是否在网格范围内
                if not (0 <= neighbor_pos[0] < self.grid.shape[0] and
                        0 <= neighbor_pos[1] < self.grid.shape[1]):
                    continue

                # 检查 TrafFlag 是否可通行
                if self.grid[neighbor_pos] == 0:
                    continue  # 不可通行区域

                # 在所有模式下应用坡度和阶跃高程差的限制
                if self.max_slope is not None and self.Slope_grid[neighbor_pos] > self.max_slope:
                    continue  # 坡度超过最大限制
                if self.max_step_elev is not None and self.StepElev_grid[neighbor_pos] > self.max_step_elev:
                    continue  # 阶跃高程差超过最大限制

                # 新增对最大粗糙度的限制
                if self.max_roughness is not None and self.Rough_grid[neighbor_pos] > self.max_roughness:
                    continue  # 粗糙度超过最大限制

                # 检查是否已经在关闭列表中
                if neighbor_pos in closed_set:
                    continue

                # 处理不同模式
                if mode == 'distance':
                    movement_cost = self.heuristic_distance_3d(current_node.position, neighbor_pos)
                    g_cost = current_node.g + movement_cost
                    h_cost = self.heuristic_distance_3d(neighbor_pos, end)
                elif mode == 'time':
                    a_max = self.Acc_grid[neighbor_pos[0], neighbor_pos[1]]
                    if a_max <= 0:
                        continue  # 无法加速，跳过
                    movement_cost = self.heuristic_distance_3d(current_node.position, neighbor_pos)
                    g_cost = current_node.g + movement_cost / a_max  # 简化为距离除以加速度
                    h_cost = self.heuristic_distance_3d(neighbor_pos, end)
                elif mode == 'env':
                    movement_uncertainty = (self.Slope_grid[neighbor_pos[0], neighbor_pos[1]] / 45.0 +
                                            self.StepElev_grid[neighbor_pos[0], neighbor_pos[1]] / 1.5 +
                                            self.Rough_grid[neighbor_pos[0], neighbor_pos[1]] / 6.0)
                    g_cost = current_node.g + movement_uncertainty
                    h_cost = self.heuristic_distance_3d(neighbor_pos, end)
                else:
                    raise ValueError("Invalid mode. Choose from 'distance', 'time', 'env'.")

                # 创建邻居节点
                neighbor_node = Node(neighbor_pos, current_node)
                neighbor_node.g = g_cost
                neighbor_node.h = h_cost
                neighbor_node.f = neighbor_node.g + neighbor_node.h

                # 检查开放列表中是否有更好的节点
                if neighbor_pos in open_dict:
                    existing_node = open_dict[neighbor_pos]
                    if existing_node.g <= neighbor_node.g:
                        continue  # 已有更优节点
                    else:
                        # 替换为更优节点
                        open_list.remove(existing_node)
                        heapq.heapify(open_list)
                        heapq.heappush(open_list, neighbor_node)
                        open_dict[neighbor_pos] = neighbor_node
                else:
                    heapq.heappush(open_list, neighbor_node)
                    open_dict[neighbor_pos] = neighbor_node

        return None  # 如果找不到路径，返回 None

    def extract_path_coordinates(self, path, x_range, y_range):
        return [(x_range[p[1]], y_range[p[0]], self.Z_grid[p[0], p[1]]) for p in path]

    def calculate_path_length(self, path_coords):
        length = 0
        for i in range(1, len(path_coords)):
            x1, y1, z1 = path_coords[i - 1]
            x2, y2, z2 = path_coords[i]
            length += np.sqrt((x2 - x1) ** 2 + (y2 - y1) ** 2 + (z2 - z1) ** 2)
        return length

    def compute_traversal_time(self, path_coords, y_range, x_range):
        total_time = 0.0
        current_speed = 0.0  # 初始速度为0
        max_speed = self.max_speed  # 使用用户指定的最大速度

        speeds = [current_speed]  # 存储每个点的速度
        accelerations = []  # 存储每段的加速度

        for i in range(1, len(path_coords)):
            x1, y1, z1 = path_coords[i - 1]
            x2, y2, z2 = path_coords[i]
            # 计算两点之间的距离
            d = np.sqrt((x2 - x1) ** 2 + (y2 - y1) ** 2 + (z2 - z1) ** 2)

            # 获取当前点的网格索引
            row = np.argmin(np.abs(y_range - y1))
            col = np.argmin(np.abs(x_range - x1))

            # 获取当前点的加速度
            a = self.Acc_grid[row, col]
            if a <= 0:
                # 如果加速度不可用或为0，无法继续加速
                a = 0.0

            # 计算需要加速到最大速度的速度增量
            delta_v = max_speed - current_speed

            if current_speed < max_speed and a > 0:
                # 需要加速
                t_acc = delta_v / a  # 加速所需时间
                d_acc = current_speed * t_acc + 0.5 * a * t_acc ** 2  # 加速所需距离

                if d_acc >= d:
                    # 在当前段无法加速到最大速度
                    # 使用运动学方程计算通过距离d所需的时间
                    # 解方程: 0.5 * a * t^2 + current_speed * t - d = 0
                    discriminant = current_speed ** 2 + 2 * a * d
                    if discriminant < 0:
                        # 无实数解，无法到达下一个点
                        return float('inf'), [], []
                    t = (-current_speed + math.sqrt(discriminant)) / a
                    total_time += t
                    current_speed += a * t
                    accelerations.append(a)
                    speeds.append(current_speed)
                else:
                    # 在当前段可以加速到最大速度
                    # 先加速到最大速度
                    total_time += t_acc
                    current_speed = max_speed
                    accelerations.append(a)
                    speeds.append(current_speed)
                    # 剩余距离
                    d_remain = d - d_acc
                    # 以最大速度匀速行驶剩余距离
                    t_remain = d_remain / current_speed
                    total_time += t_remain
                    accelerations.append(0.0)
                    speeds.append(current_speed)
            else:
                # 无法加速或已达到最大速度，匀速行驶
                if current_speed == 0:
                    # 无法移动
                    return float('inf'), [], []
                t = d / current_speed
                total_time += t
                accelerations.append(0.0)
                speeds.append(current_speed)

        return total_time, speeds, accelerations

# 可视化类
class Visualizer:
    def __init__(self, X_grid, Y_grid, Z_grid, TrafFlag_grid, Acc_grid):
        self.X_grid = X_grid
        self.Y_grid = Y_grid
        self.Z_grid = Z_grid
        self.TrafFlag_grid = TrafFlag_grid
        self.Acc_grid = Acc_grid
        self.fig = go.Figure()

    def plot_terrain(self):
        # 使用您提供的颜色刻度
        colorscale = [
            [0.0, 'rgb(0,0,0)'],
            [0.0667, 'rgb(105,105,105)'],
            [0.1333, 'rgb(169,169,169)'],
            [0.2, 'rgb(160,82,45)'],
            [0.2667, 'rgb(34,139,34)'],
            [0.3333, 'rgb(188,143,143)'],
            [0.4, 'rgb(112,128,144)'],
            [0.4667, 'rgb(139,69,19)'],
            [0.5333, 'rgb(210,180,140)'],
            [0.6, 'rgb(139,69,19)'],
            [0.6667, 'rgb(139,69,19)'],
            [0.7333, 'rgb(139,69,19)'],
            [0.8, 'rgb(139,69,19)'],
            [0.8667, 'rgb(139,69,19)'],
            [0.9333, 'rgb(105,105,105)'],
            [1.0, 'rgb(70,130,180)']
        ]

        # 获取 Acc_grid 的最小值和最大值
        acc_min = np.nanmin(self.Acc_grid)
        acc_max = np.nanmax(self.Acc_grid)

        # 将 Acc_grid 归一化到 [0,1] 范围，用于颜色映射
        normalized_acc = (self.Acc_grid - acc_min) / (acc_max - acc_min) if acc_max > acc_min else self.Acc_grid

        # 添加地形图，使用 Acc_grid 进行着色
        self.fig.add_trace(go.Surface(
            x=self.X_grid,
            y=self.Y_grid,
            z=self.Z_grid,
            surfacecolor=normalized_acc,
            colorscale=colorscale,
            cmin=0,
            cmax=1,
            opacity=0.8,
            showscale=True,
            colorbar=dict(
                title='最大加速度',
                tickvals=[v[0] for v in colorscale],
                ticktext=[f"{acc_min + v[0] * (acc_max - acc_min):.2f}" for v in colorscale]
            )
        ))

    def add_path(self, coords, speeds, accs, color, name, length, time_taken, path_width=4):
        if coords:
            x, y, z = zip(*coords)
            # 处理不可行路径（时间为inf）的情况
            if math.isinf(time_taken):
                time_str = "无法到达"
            else:
                time_str = f"{time_taken:.2f} 秒"

            # 准备 customdata，将速度和加速度添加进去
            if accs:
                accs = [0.0] + accs  # 与 speeds 对齐
            else:
                accs = [0.0] * len(speeds)

            customdata = np.array([speeds, accs]).T  # 转置使其每行对应一个点的 [speed, acceleration]

            self.fig.add_trace(go.Scatter3d(
                x=x, y=y, z=z,
                mode='lines+markers',
                line=dict(color=color, width=path_width),
                marker=dict(size=2, color=color),
                name=f"{name} (长度: {length:.2f} 米, 时间: {time_str})",
                customdata=customdata,
                hovertemplate=(
                        'X: %{x}<br>' +
                        'Y: %{y}<br>' +
                        'Z: %{z}<br>' +
                        '速度: %{customdata[0]:.2f} m/s<br>' +
                        '加速度: %{customdata[1]:.2f} m/s²<br>' +
                        '<extra></extra>'
                )
            ))
        else:
            print(f"未找到{name}")

    def add_start_and_end_points(self, start_coord, end_coord):
        self.fig.add_trace(go.Scatter3d(
            x=[start_coord[0]],
            y=[start_coord[1]],
            z=[start_coord[2]],
            mode='markers',
            marker=dict(size=8, color='green'),
            name='起点'
        ))
        self.fig.add_trace(go.Scatter3d(
            x=[end_coord[0]],
            y=[end_coord[1]],
            z=[end_coord[2]],
            mode='markers',
            marker=dict(size=8, color='red'),
            name='终点'
        ))

    def show(self):
        self.fig.update_layout(
            scene=dict(
                xaxis_title='X 坐标',
                yaxis_title='Y 坐标',
                zaxis_title='Z 坐标 (高程)',
                aspectmode='data'
            ),
            title='3D 地形路径规划（基于距离、时间与最小不确定性）',
            legend=dict(
                x=0,
                y=1,
                bgcolor='rgba(255, 255, 255, 0)',
                bordercolor='rgba(255, 255, 255, 0)'
            )
        )
        pio.show(self.fig)

# 报告生成函数
def write_paths_to_word(document, path_name, path_coords):
    if not path_coords:
        return
    # 添加路径名称和点数量
    document.add_heading(f"{path_name} ( {len(path_coords)})", level=2)
    # 创建表格，包含三列：X相对坐标、Y相对坐标、Z高程坐标
    table = document.add_table(rows=1, cols=3)
    table.style = 'Light Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'X'
    hdr_cells[1].text = 'Y'
    hdr_cells[2].text = 'Z'
    # 填充表格内容
    for coord in path_coords:
        row_cells = table.add_row().cells
        row_cells[0].text = f"{coord[0]}"
        row_cells[1].text = f"{coord[1]}"
        row_cells[2].text = f"{coord[2]}"
    # 添加一个空行
    document.add_paragraph()

# 路径规划函数
def run_path_planning(data_processor, start_coord, end_coord, output_path, max_allowed_slope, max_allowed_step_elev, max_allowed_roughness, max_speed):
    try:
        # 确保已进行数据评估
        data_processor.split_data_and_evaluate()

        # 将 start_coord 和 end_coord 转换为网格索引
        start = data_processor.coord_to_index(start_coord)
        end = data_processor.coord_to_index(end_coord)

        # 检查起点和终点是否在网格范围内
        if not data_processor.is_valid_position(start):
            raise ValueError("起点坐标超出网格范围")
        if not data_processor.is_valid_position(end):
            raise ValueError("终点坐标超出网格范围")

        # 确保起点和终点可通行
        if data_processor.TrafFlag_grid[start] == 0:
            raise ValueError("起点不可通行")
        if data_processor.TrafFlag_grid[end] == 0:
            raise ValueError("终点不可通行")

        # 路径规划
        grid_data = {
            'TrafFlag_grid': data_processor.TrafFlag_grid,
            'Z_grid': data_processor.Z_grid,
            'Acc_grid': data_processor.Acc_grid,
            'Slope_grid': data_processor.Slope_grid,
            'StepElev_grid': data_processor.StepElev_grid,
            'Rough_grid': data_processor.Rough_grid
        }
        path_planner = PathPlanner(grid_data, max_slope=max_allowed_slope, max_step_elev=max_allowed_step_elev, max_roughness=max_allowed_roughness, max_speed=max_speed)

        print("正在计算基于距离的最短路径...")
        path_distance = path_planner.astar(start, end, mode='distance')
        if path_distance:
            print("基于距离的路径计算完成")
        else:
            print("未找到基于距离的路径")

        print("正在计算基于时间的最短耗时路径...")
        path_time = path_planner.astar(start, end, mode='time')
        if path_time:
            print("基于时间的路径计算完成")
        else:
            print("未找到基于时间的路径")

        print("正在计算最小不确定性路径...")
        path_env = path_planner.astar(start, end, mode='env')
        if path_env:
            print("最小不确定性路径计算完成")
        else:
            print("未找到最小不确定性路径")

        # 提取路径坐标
        if path_distance:
            path_distance_coords = path_planner.extract_path_coordinates(path_distance, data_processor.x_range,
                                                                         data_processor.y_range)
            length_distance = path_planner.calculate_path_length(path_distance_coords)
        else:
            path_distance_coords = None
            length_distance = None
        if path_time:
            path_time_coords = path_planner.extract_path_coordinates(path_time, data_processor.x_range,
                                                                     data_processor.y_range)
            length_time = path_planner.calculate_path_length(path_time_coords)
        else:
            path_time_coords = None
            length_time = None
        if path_env:
            path_env_coords = path_planner.extract_path_coordinates(path_env, data_processor.x_range,
                                                                    data_processor.y_range)
            length_env = path_planner.calculate_path_length(path_env_coords)
        else:
            path_env_coords = None
            length_env = None

        # 计算遍历时间，并获取速度和加速度数据
        if path_distance:
            time_distance, speeds_distance, accs_distance = path_planner.compute_traversal_time(
                path_distance_coords, data_processor.y_range, data_processor.x_range)
        else:
            speeds_distance, accs_distance, time_distance = None, None, None

        if path_time:
            time_time, speeds_time, accs_time = path_planner.compute_traversal_time(
                path_time_coords, data_processor.y_range, data_processor.x_range)
        else:
            speeds_time, accs_time, time_time = None, None, None

        if path_env:
            time_env, speeds_env, accs_env = path_planner.compute_traversal_time(
                path_env_coords, data_processor.y_range, data_processor.x_range)
        else:
            speeds_env, accs_env, time_env = None, None, None

        # 可视化
        visualizer = Visualizer(
            data_processor.X_grid,
            data_processor.Y_grid,
            data_processor.Z_grid,
            data_processor.TrafFlag_grid,
            data_processor.Acc_grid  # 传入最大加速度网格
        )
        visualizer.plot_terrain()

        colors = ['blue', 'orange', 'purple']
        names = ['MinDistPath 最短距离路径', 'MinTimePath 最短耗时路径', 'MinUncertPath 最小不确定性路径']
        lengths = [length_distance, length_time, length_env]
        times = [time_distance, time_time, time_env]
        coords_list = [path_distance_coords, path_time_coords, path_env_coords]
        speeds_list = [speeds_distance, speeds_time, speeds_env]
        accs_list = [accs_distance, accs_time, accs_env]

        for coords, speeds, accs, color, name, length, time_taken in zip(
                coords_list, speeds_list, accs_list,
                colors, names, lengths, times):
            visualizer.add_path(coords, speeds, accs, color, name, length, time_taken)

        # 标记起点和终点
        if path_distance or path_time or path_env:
            selected_path = None
            if path_distance:
                selected_path = path_distance_coords
            elif path_time:
                selected_path = path_time_coords
            elif path_env:
                selected_path = path_env_coords

            if selected_path:
                # 使用起点坐标
                start_x, start_y, start_z = selected_path[0]
                # 使用终点坐标
                end_x, end_y, end_z = selected_path[-1]
                visualizer.add_start_and_end_points((start_x, start_y, start_z), (end_x, end_y, end_z))

        visualizer.show()

        # 创建 Word 文档
        document = Document()
        document.add_heading('Route planning results', 0)
        # 写入路径信息
        if path_distance or path_time or path_env:
            if path_distance:
                write_paths_to_word(document, 'MinDistPath ', path_distance_coords)
            if path_time:
                write_paths_to_word(document, 'MinTimePath ', path_time_coords)
            if path_env:
                write_paths_to_word(document, 'MinUncertPath ', path_env_coords)

            # 保存文档
            try:
                # 使用 resource_path 函数获取输出文件路径
                output_file_path = output_path  # 如果 output_path 是用户指定的完整路径，则无需使用 resource_path
                document.save(output_file_path)
                print(f"路径坐标已成功保存到 Word 文档: {output_file_path}")
            except Exception as e:
                print(f"保存 Word 文档时出错: {e}")
                logging.exception("保存 Word 文档时出错")
        else:
            print("没有路径可写入 Word 文档")

        # 打印路径长度和遍历时间
        if path_distance:
            if math.isinf(time_distance):
                print(f"最短路径长度: {length_distance:.2f} 米")
                print(f"最短距离路径遍历时间: 无法到达")
            else:
                print(f"最短路径长度: {length_distance:.2f} 米")
                print(f"最短距离路径遍历时间: {time_distance:.2f} 秒")
        if path_time:
            if math.isinf(time_time):
                print(f"最短耗时路径长度: {length_time:.2f} 米")
                print(f"最短耗时路径遍历时间: 无法到达")
            else:
                print(f"最短耗时路径长度: {length_time:.2f} 米")
                print(f"最短耗时路径遍历时间: {time_time:.2f} 秒")
        if path_env:
            if math.isinf(time_env):
                print(f"最小不确定性路径长度: {length_env:.2f} 米")
                print(f"最小不确定性路径遍历时间: 无法到达")
            else:
                print(f"最小不确定性路径长度: {length_env:.2f} 米")
                print(f"最小不确定性路径遍历时间: {time_env:.2f} 秒")
    except Exception as e:
        print(f"运行处理时发生错误: {e}")
        logging.exception("运行处理时发生错误")

# 主程序
def main():
    try:
        # 创建主窗口
        root = tk.Tk()
        root.title("路径规划软件")

        # 地图数据文件路径
        tk.Label(root, text="地图数据文件路径:").grid(row=0, column=0, sticky='e')
        map_data_path_entry = tk.Entry(root, width=50)
        map_data_path_entry.grid(row=0, column=1)
        def browse_map_data():
            filename = filedialog.askopenfilename(filetypes=[("CSV文件", "*.csv")])
            if filename:
                map_data_path_entry.delete(0, tk.END)
                map_data_path_entry.insert(0, filename)
        tk.Button(root, text="浏览...", command=browse_map_data).grid(row=0, column=2)

        # 起点坐标（初始禁用）
        tk.Label(root, text="起点坐标 (格式: X,Y):").grid(row=1, column=0, sticky='e')
        start_coord_entry = tk.Entry(root, width=50)
        start_coord_entry.grid(row=1, column=1)
        # 终点坐标（初始禁用）
        tk.Label(root, text="终点坐标 (格式: X,Y):").grid(row=2, column=0, sticky='e')
        end_coord_entry = tk.Entry(root, width=50)
        end_coord_entry.grid(row=2, column=1)

        # 输出路径
        tk.Label(root, text="输出路径文档路径:").grid(row=3, column=0, sticky='e')
        output_path_entry = tk.Entry(root, width=50)
        output_path_entry.grid(row=3, column=1)
        def browse_output_path():
            filename = filedialog.asksaveasfilename(defaultextension=".docx",
                                                    filetypes=[("Word文档", "*.docx")])
            if filename:
                output_path_entry.delete(0, tk.END)
                output_path_entry.insert(0, filename)
        tk.Button(root, text="浏览...", command=browse_output_path).grid(row=3, column=2)

        # 最大坡度
        tk.Label(root, text="最大爬坡度 (默认35.0度):").grid(row=4, column=0, sticky='e')
        max_slope_entry = tk.Entry(root, width=50)
        max_slope_entry.grid(row=4, column=1)
        max_slope_entry.insert(0, "35.0")  # 设置默认值

        # 最大阶跃高程差
        tk.Label(root, text="最大阶跃高程差 (默认0.5米):").grid(row=5, column=0, sticky='e')
        max_step_elev_entry = tk.Entry(root, width=50)
        max_step_elev_entry.grid(row=5, column=1)
        max_step_elev_entry.insert(0, "0.5")  # 设置默认值为 0.3

        # 添加最大粗糙度输入字段
        tk.Label(root, text="最大粗糙度 (默认6.0):").grid(row=6, column=0, sticky='e')
        max_roughness_entry = tk.Entry(root, width=50)
        max_roughness_entry.grid(row=6, column=1)
        max_roughness_entry.insert(0, "6.0")  # 设置默认值

        # 最大速度
        tk.Label(root, text="最大速度 (默认20.0 m/s):").grid(row=7, column=0, sticky='e')
        max_speed_entry = tk.Entry(root, width=50)
        max_speed_entry.grid(row=7, column=1)
        max_speed_entry.insert(0, "20.0")  # 设置默认值

        # 禁用输入字段，直到地图加载完成
        start_coord_entry.configure(state='disabled')
        end_coord_entry.configure(state='disabled')
        output_path_entry.configure(state='disabled')
        max_slope_entry.configure(state='disabled')
        max_step_elev_entry.configure(state='disabled')
        max_roughness_entry.configure(state='disabled')
        max_speed_entry.configure(state='disabled')

        # 加载地图函数
        def load_map():
            csv_path = map_data_path_entry.get()
            if not csv_path:
                messagebox.showerror("错误", "请提供地图数据文件路径。")
                return

            def process_map():
                try:
                    data_processor = DataProcessor(csv_path)
                    data_processor.load_data()
                    data_processor.interpolate_data()
                    data_processor.fill_nan_values()

                    # 可视化地图
                    visualizer = Visualizer(
                        data_processor.X_grid,
                        data_processor.Y_grid,
                        data_processor.Z_grid,
                        data_processor.TrafFlag_grid,
                        data_processor.Acc_grid  # 传入最大加速度网格
                    )
                    visualizer.plot_terrain()
                    visualizer.show()

                    # 存储 data_processor
                    root.data_processor = data_processor

                    # 启用输入字段
                    root.after(0, enable_input_fields)
                    messagebox.showinfo("完成", "地图已加载并显示。请继续输入起点和终点坐标。")
                except Exception as e:
                    messagebox.showerror("错误", f"加载地图时发生错误: {e}")
                    logging.exception("加载地图时发生错误")

            threading.Thread(target=process_map).start()

        # 启用输入字段函数
        def enable_input_fields():
            start_coord_entry.configure(state='normal')
            end_coord_entry.configure(state='normal')
            output_path_entry.configure(state='normal')
            max_slope_entry.configure(state='normal')
            max_step_elev_entry.configure(state='normal')
            max_roughness_entry.configure(state='normal')
            max_speed_entry.configure(state='normal')
            run_button.configure(state='normal')

        # 加载地图按钮
        tk.Button(root, text="加载地图", command=load_map).grid(row=8, column=1)

        # 运行路径规划函数
        def run():
            # 获取输入值
            start_coord_input = start_coord_entry.get()
            end_coord_input = end_coord_entry.get()
            output_path = output_path_entry.get()

            # 验证输入
            if not output_path:
                messagebox.showerror("错误", "请提供输出 Word 文档路径。")
                return
            if not start_coord_input:
                messagebox.showerror("错误", "请提供起点坐标。")
                return
            if not end_coord_input:
                messagebox.showerror("错误", "请提供终点坐标。")
                return
            try:
                start_x, start_y = map(float, start_coord_input.strip().split(','))
                start_coord = (start_x, start_y)
            except ValueError:
                messagebox.showerror("错误", "起点坐标格式错误，请使用 'X,Y' 的格式。")
                return

            try:
                end_x, end_y = map(float, end_coord_input.strip().split(','))
                end_coord = (end_x, end_y)
            except ValueError:
                messagebox.showerror("错误", "终点坐标格式错误，请使用 'X,Y' 的格式。")
                return

            # 获取最大坡度和最大阶跃高程差
            max_slope_input = max_slope_entry.get()
            max_step_elev_input = max_step_elev_entry.get()

            try:
                max_allowed_slope = float(max_slope_input)
            except ValueError:
                messagebox.showerror("错误", "最大坡度格式错误，请输入数字。")
                return

            try:
                max_allowed_step_elev = float(max_step_elev_input)
            except ValueError:
                messagebox.showerror("错误", "最大阶跃高程差格式错误，请输入数字。")
                return

            # 获取最大粗糙度
            max_roughness_input = max_roughness_entry.get()
            try:
                max_allowed_roughness = float(max_roughness_input)
            except ValueError:
                messagebox.showerror("错误", "最大粗糙度格式错误，请输入数字。")
                return

            # 获取最大速度
            max_speed_input = max_speed_entry.get()

            try:
                max_speed = float(max_speed_input)
            except ValueError:
                messagebox.showerror("错误", "最大速度格式错误，请输入数字。")
                return

            # 在后台线程中运行，以避免阻塞GUI
            def process():
                try:
                    # 使用已加载的 data_processor
                    data_processor = root.data_processor
                    # 调用路径规划函数
                    run_path_planning(data_processor, start_coord, end_coord, output_path, max_allowed_slope, max_allowed_step_elev, max_allowed_roughness, max_speed)
                    messagebox.showinfo("完成", "路径规划已完成。")
                except Exception as e:
                    messagebox.showerror("错误", f"发生错误: {e}")
                    logging.exception("后台线程运行时发生错误")

            threading.Thread(target=process).start()

        # 运行按钮（初始禁用）
        run_button = tk.Button(root, text="运行", command=run)
        run_button.grid(row=9, column=1)
        run_button.configure(state='disabled')

        root.mainloop()
    except Exception as e:
        print(f"主程序运行时发生错误: {e}")
        logging.exception("主程序运行时发生错误")

if __name__ == "__main__":
    main()
